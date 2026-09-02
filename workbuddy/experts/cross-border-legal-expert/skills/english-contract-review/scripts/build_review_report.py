#!/usr/bin/env python3
"""Build a Chinese contract review report DOCX."""

from __future__ import annotations

import argparse
import json
import re
import unicodedata
from pathlib import Path

from emoji_text import sanitize_data
from skill_paths import generated_path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


DISCLAIMER = (
    "本文档由 AI 辅助生成，仅供参考，不构成正式法律意见，"
    "不能替代具有执业资格的律师。"
)

# Richee output standard colors
RICHEE_BLACK = "1a1a1a"
RICHEE_WHITE = "ffffff"
RICHEE_LINE = "e2e5ea"
RICHEE_BG = "f7f7f7"
REPORT_FONT_SIZES = {
    "body": 12,
    "title": 22,
    "heading1": 18,
    "heading2": 16,
    "heading3": 14,
    "disclaimer": 10.5,
    "table": 10.5,
}
INDEX_HEADER_RE = re.compile(r"^(?:序号|编号|no\.?|id)$", re.IGNORECASE)
COMPACT_HEADER_MARKERS = (
    "等级",
    "标识",
    "重要性",
    "优先级",
    "状态",
    "是否",
    "结果",
)


def set_font(
    run,
    size: float = REPORT_FONT_SIZES["body"],
    bold: bool = False,
    color_hex: str | None = None,
) -> None:
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "PingFang SC"
    if color_hex:
        run.font.color.rgb = RGBColor(
            int(color_hex[0:2], 16),
            int(color_hex[2:4], 16),
            int(color_hex[4:6], 16),
        )
    r_pr = run._element.get_or_add_rPr()
    fonts = r_pr.find(qn("w:rFonts"))
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, fonts)
    fonts.set(qn("w:ascii"), "Arial")
    fonts.set(qn("w:hAnsi"), "Arial")
    fonts.set(qn("w:eastAsia"), "PingFang SC")


def set_cell_shading(cell, color_hex: str) -> None:
    """Set background shading on a table cell (Richee compliance)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), color_hex)


def style_header_row(row, widths: list[int]) -> None:
    """Apply Richee black-header-white-text style to a header row."""
    for idx, cell in enumerate(row.cells):
        set_cell_shading(cell, RICHEE_BLACK)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.bold = True


def configure(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.3)
    section.bottom_margin = Cm(2.3)
    section.left_margin = Cm(2.6)
    section.right_margin = Cm(2.6)
    normal = doc.styles["Normal"]
    normal.font.name = "PingFang SC"
    normal.font.size = Pt(REPORT_FONT_SIZES["body"])
    # 中文 PingFang SC、西文 Arial——ascii/hAnsi 必须显式设 Arial，
    # 否则英文落入 PingFang SC，造成"字体不统一"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(5)
    # Heading 4-9 也必须显式设黑：validate_review_outputs.py 的 check_black_heading_styles
    # 检查全部 heading 样式，python-docx 默认给 Heading 4-9 蓝色，不显式覆盖则交付闸门必失败。
    for name, size in (
        ("Title", REPORT_FONT_SIZES["title"]),
        ("Heading 1", REPORT_FONT_SIZES["heading1"]),
        ("Heading 2", REPORT_FONT_SIZES["heading2"]),
        ("Heading 3", REPORT_FONT_SIZES["heading3"]),
        ("Heading 4", REPORT_FONT_SIZES["heading3"]),
        ("Heading 5", REPORT_FONT_SIZES["heading3"]),
        ("Heading 6", REPORT_FONT_SIZES["heading3"]),
        ("Heading 7", REPORT_FONT_SIZES["heading3"]),
        ("Heading 8", REPORT_FONT_SIZES["heading3"]),
        ("Heading 9", REPORT_FONT_SIZES["heading3"]),
    ):
        style = doc.styles[name]
        style.font.name = "PingFang SC"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.paragraph_format.keep_with_next = True


def set_table_width(table, widths: list[int]) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _display_units(text: str) -> float:
    """Estimate rendered text demand for mixed Chinese and Latin content."""
    units = 0.0
    for character in re.sub(r"\s+", " ", text.strip()):
        if character.isspace():
            units += 0.35
        elif unicodedata.east_asian_width(character) in {"W", "F"}:
            units += 1.0
        elif character.isupper():
            units += 0.65
        else:
            units += 0.55
    return units


def _column_kind(header: str) -> str:
    normalized = re.sub(r"[\s：:()\uff08\uff09【】\[\]]+", "", header).lower()
    if INDEX_HEADER_RE.fullmatch(normalized):
        return "index"
    if any(marker in normalized for marker in COMPACT_HEADER_MARKERS):
        return "compact"
    return "narrative"


def allocate_content_widths(table, total_width: int) -> list[int]:
    """Allocate exact DXA widths by semantics and real cell text demand."""
    headers = [cell.text.strip() for cell in table.rows[0].cells]
    columns = [
        [row.cells[index].text.strip() for row in table.rows]
        for index in range(len(headers))
    ]
    count = len(headers)
    base_min_ratio = min(0.12, 0.68 / count)
    demands: list[float] = []
    minimums: list[int] = []
    maximums: list[int] = []
    for header, values in zip(headers, columns):
        kind = _column_kind(header)
        measured = [min(max(_display_units(value), 1.0), 48.0) for value in values]
        peak = max(measured, default=1.0)
        average = sum(measured) / len(measured) if measured else 1.0
        demands.append(
            max(_display_units(header) + 1.5, peak * 0.68 + average * 0.32, 1.0)
        )
        if kind == "index":
            minimums.append(round(total_width * 0.065))
            maximums.append(round(total_width * 0.09))
        elif kind == "compact":
            minimums.append(round(total_width * min(base_min_ratio, 0.11)))
            maximums.append(round(total_width * 0.20))
        else:
            minimums.append(round(total_width * base_min_ratio))
            maximums.append(round(total_width * (0.80 if count <= 3 else 0.65)))

    while sum(minimums) > total_width:
        minimums[max(range(count), key=lambda item: minimums[item])] -= 1
    widths = minimums[:]
    remaining = total_width - sum(widths)
    active = {index for index in range(count) if widths[index] < maximums[index]}
    while remaining > 0 and active:
        total_demand = sum(demands[index] for index in active)
        planned = []
        for index in active:
            raw = remaining * demands[index] / total_demand
            planned.append((index, int(raw), raw - int(raw)))
        assigned = 0
        for index, amount, _ in planned:
            amount = min(amount, maximums[index] - widths[index])
            widths[index] += amount
            assigned += amount
        leftover = remaining - assigned
        for index, _, _ in sorted(planned, key=lambda item: item[2], reverse=True):
            if leftover <= 0:
                break
            if widths[index] < maximums[index]:
                widths[index] += 1
                assigned += 1
                leftover -= 1
        if assigned == 0:
            break
        remaining -= assigned
        active = {index for index in active if widths[index] < maximums[index]}

    # All caps are soft when their sum is smaller than the physical table.
    order = sorted(range(count), key=lambda item: demands[item], reverse=True)
    while remaining > 0:
        for index in order:
            if remaining <= 0:
                break
            widths[index] += 1
            remaining -= 1
    return widths


def rebalance_table_widths(table, total_width: int) -> list[int]:
    widths = allocate_content_widths(table, total_width)
    set_table_width(table, widths)
    return widths


def configure_table_pagination(table) -> None:
    for index, row in enumerate(table.rows):
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))
        if index == 0 and tr_pr.find(qn("w:tblHeader")) is None:
            header = OxmlElement("w:tblHeader")
            header.set(qn("w:val"), "true")
            tr_pr.append(header)


def add_bilingual_block(
    doc: Document,
    en: str,
    zh: str,
    en_size: float = REPORT_FONT_SIZES["body"],
    zh_size: float = REPORT_FONT_SIZES["body"],
    bullet: bool = False,
) -> None:
    """zh-primary：中文报告默认只渲染中文段，仅当中文缺失时回退英文。

    （修复人工评测 B2：en+zh 双段全渲染导致整段英文嵌入中文报告、严重冗余。
    数据层 bilingual dict 保持不变，双语版交付物另行复用 en。）
    """
    text = zh or en
    if not text:
        return
    add_plain_para(doc, text, zh_size, bullet=bullet)


def _is_bilingual_dict(content) -> bool:
    """True when content is a dict with at least one of 'en'/'zh' keys (and not a simple item dict)."""
    if not isinstance(content, dict):
        return False
    return ("en" in content or "zh" in content) and not any(
        k in content for k in ("title", "issue", "item", "target")
    )


# 模型在 report.json 中夹带的 Markdown 标记（## 标题、**加粗**、`code`、- 列表）
# 不会被 docx 渲染，会以字面符号泄漏到报告里。脚本层兜底转换：
# **…** → 真实加粗 run；行首 #… → 加粗小节标题段；其余记号剥离。
MD_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
MD_HEADING_RE = re.compile(r"^\s{0,3}#{1,6}\s+")
MD_BULLET_RE = re.compile(r"^\s{0,3}[-*•]\s+")


def strip_inline_md(text: str) -> str:
    """表格单元格等无富文本场景：剥离 Markdown 记号只留文字。"""
    if not text:
        return text
    text = MD_HEADING_RE.sub("", text)
    text = MD_BOLD_RE.sub(r"\1", text)
    return text.replace("`", "")


def add_md_runs(paragraph, text: str, size: float, bold: bool = False) -> None:
    """把含 **…** 的文本拆成普通/加粗 run，全部经 set_font。"""
    pos = 0
    for match in MD_BOLD_RE.finditer(text):
        if match.start() > pos:
            set_font(paragraph.add_run(text[pos:match.start()].replace("`", "")), size, bold)
        set_font(paragraph.add_run(match.group(1).replace("`", "")), size, True)
        pos = match.end()
    if pos < len(text):
        set_font(paragraph.add_run(text[pos:].replace("`", "")), size, bold)


def add_plain_para(
    doc: Document,
    text: str,
    size: float = REPORT_FONT_SIZES["body"],
    bullet: bool = False,
) -> None:
    """所有正文段落统一经 set_font（修复 B11 字体不统一），并做 Markdown 兜底转换。"""
    text = text or ""
    if "\n" in text:
        # 多行字符串（模型整段 Markdown 直出）逐行渲染，行首 # 各自成小节标题
        for line in text.split("\n"):
            line = line.strip()
            if line:
                add_plain_para(doc, line, size, bullet=bullet)
        return
    heading = MD_HEADING_RE.match(text)
    if heading:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        add_md_runs(p, text[heading.end():].strip(), size, bold=True)
        return
    if bullet:
        text = MD_BULLET_RE.sub("", text)
    p = doc.add_paragraph(style="List Bullet" if bullet else None)
    add_md_runs(p, text, size)


def add_section(doc: Document, title: str, content) -> None:
    """Enhanced section renderer supporting:
    - bilingual dict: {"en": "...", "zh": "...", "items": [...]} — zh-primary 渲染
    - explained empty: {"empty_reason": "..."} shows reason instead of 无。
    - backward-compatible: plain string, list of dicts, list of strings

    章节永不整体跳过（修复 B1 编号跳跃）：旧版 {"omit_if_empty": true}
    兼容为渲染"本节不适用。"一行，编号保持连续。
    """
    if isinstance(content, dict) and content.get("omit_if_empty"):
        content = {"empty_reason": content.get("empty_reason", "本节不适用。")}

    if isinstance(content, dict) and "empty_reason" in content and not content.get("en") and not content.get("zh"):
        doc.add_heading(title, level=1)
        add_plain_para(doc, content["empty_reason"])
        return

    doc.add_heading(title, level=1)

    if not content:
        add_plain_para(doc, "无。")
        return

    # --- Bilingual dict format（zh-primary）---
    if _is_bilingual_dict(content):
        add_bilingual_block(doc, content.get("en", ""), content.get("zh", ""))
        _render_item_list(doc, content.get("items", []) or content.get("checklist", []))
        return

    # --- Plain string ---
    if isinstance(content, str):
        add_plain_para(doc, content)
        return

    # --- Generic dict（非 bilingual / 非 empty_reason）---
    # 修复 #6：旧版对此类 dict 落入 `for item in content` 会迭代「键名」，
    # 把 `items` 等键当正文渲染。改为：有 items/checklist/list 键则渲染其列表，
    # 否则渲染 dict 的「值」，绝不渲染键名。
    if isinstance(content, dict):
        listed = content.get("items") or content.get("checklist") or content.get("list")
        if isinstance(listed, list):
            _render_item_list(doc, listed)
        else:
            for value in content.values():
                if isinstance(value, (list, dict)):
                    add_section_body(doc, value)
                elif _nonempty(value):
                    add_plain_para(doc, str(value), REPORT_FONT_SIZES["body"], bullet=True)
        return

    # --- List of items ---
    _render_item_list(doc, content)


def add_section_body(doc: Document, content) -> None:
    """渲染一段分节内容但不带标题（供 generic dict 的嵌套值复用）。"""
    if _is_bilingual_dict(content):
        add_bilingual_block(doc, content.get("en", ""), content.get("zh", ""), bullet=True)
        _render_item_list(doc, content.get("items", []) or content.get("checklist", []))
    elif isinstance(content, dict):
        listed = content.get("items") or content.get("checklist") or content.get("list")
        if isinstance(listed, list):
            _render_item_list(doc, listed)
        else:
            for value in content.values():
                if _nonempty(value):
                    add_plain_para(doc, str(value), REPORT_FONT_SIZES["body"], bullet=True)
    elif isinstance(content, list):
        _render_item_list(doc, content)
    elif _nonempty(content):
        add_plain_para(doc, str(content), REPORT_FONT_SIZES["body"], bullet=True)


def _render_item_list(doc: Document, items) -> None:
    """渲染条目列表（dict 条目走 bilingual/title-detail，字符串直渲），统一脱键名。"""
    if not isinstance(items, list):
        return
    for item in items:
        if isinstance(item, dict):
            if _is_bilingual_dict(item):
                add_bilingual_block(doc, item.get("en", ""), item.get("zh", ""), bullet=True)
                continue
            heading = item.get("title") or item.get("issue") or item.get("item") or "事项"
            detail = item.get("detail") or item.get("text") or item.get("status") or ""
            basis = item.get("basis_tag", "")
            add_plain_para(
                doc,
                f"{basis} {heading}：{detail}".strip(),
                REPORT_FONT_SIZES["body"],
                bullet=True,
            )
        else:
            add_plain_para(doc, str(item), REPORT_FONT_SIZES["body"], bullet=True)


def _fill_basis_column(rows: list[dict]) -> list[dict]:
    """结构参数硬检查表的「依据」列需要分类标签 + 详细依据说明，而非裸标签。

    原渲染只输出 `basis_tag`（如 `[惯例]`），看不到具体依据。此函数把每行的
    `basis_tag` 与详细说明（`basis_detail` / `basis`）合并到 `basis` 列；详细说明
    缺失时降级为标签 + 提示，促使补全而非留裸标签（与 warn_structural_basis 告警互补）。
    """
    filled = []
    for row in rows:
        if not isinstance(row, dict):
            filled.append(row)
            continue
        row = dict(row)
        tag = str(row.get("basis_tag", "")).strip()
        detail = str(row.get("basis_detail") or row.get("basis") or "").strip()
        if detail and tag and not detail.startswith(tag):
            row["basis"] = f"{tag} {detail}"
        elif detail:
            row["basis"] = detail
        elif tag:
            row["basis"] = f"{tag} 详见对应条款审查意见"
        else:
            row["basis"] = "基于合理判断，建议专业律师确认"
        filled.append(row)
    return filled


def add_matrix(
    doc: Document,
    title: str,
    rows: list[dict],
    headers: list[str],
    keys: list[str],
    widths: list[int],
    empty_reason: str = "无。",
) -> None:
    """Matrix renderer。rows 为空时不再跳节（修复 B1 编号跳跃），
    而是渲染 empty_reason 一行向读者交代不适用原因。"""
    # 入参容错（真机 17e2c418）：模型常把 ip_analysis 等写成 {"items":[...]}
    # （该形状在 add_section 里是合法的），矩阵渲染需要 list，此处先解包，
    # 避免 build 崩溃、逼主 Agent 手改 report.json 兜底。
    if isinstance(rows, dict):
        rows = rows.get("items") or rows.get("checklist") or rows.get("list") or []
    if not rows:
        doc.add_heading(title, level=1)
        add_plain_para(doc, empty_reason)
        return
    doc.add_heading(title, level=1)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths)
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
        for run in cell.paragraphs[0].runs:
            set_font(run, REPORT_FONT_SIZES["table"], True)
    style_header_row(table.rows[0], widths)
    for item in rows:
        cells = table.add_row().cells
        for cell, key in zip(cells, keys):
            cell.text = strip_inline_md(str(item.get(key, ""))).strip()
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.line_spacing = 1.5
                for run in paragraph.runs:
                    set_font(run, REPORT_FONT_SIZES["table"])
    rebalance_table_widths(table, sum(widths))
    configure_table_pagination(table)


def add_risk_table(doc: Document, title: str, risks: list[dict]) -> None:
    doc.add_heading(title, level=1)
    if not risks:
        add_plain_para(doc, "未识别到需要单列的风险，但仍需完成专业复核。")
        return
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    widths = [680, 1350, 1700, 1700, 2150, 1220]
    set_table_width(table, widths)
    headers = ["编号/等级", "位置", "问题", "影响与可能性", "建议及备选", "依据"]
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
        for run in cell.paragraphs[0].runs:
            set_font(run, REPORT_FONT_SIZES["table"], True)
    style_header_row(table.rows[0], widths)
    for item in risks:
        values = [
            "\n".join(
                value
                for value in (
                    str(item.get("issue_id", "")).strip(),
                    str(item.get("level", "待定")).strip(),
                )
                if value
            ),
            item.get("location", ""),
            item.get("issue", ""),
            str(item.get("impact_likelihood", "")).strip() or "—",
            item.get("recommendation", ""),
            item.get("basis_tag", "[要点]"),
        ]
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            cell.text = strip_inline_md(str(value)).strip()
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.line_spacing = 1.5
                for run in paragraph.runs:
                    set_font(run, REPORT_FONT_SIZES["table"])
    rebalance_table_widths(table, sum(widths))
    configure_table_pagination(table)


REPORT_CORE_FIELDS = ("risks", "scope", "executive_summary", "structural_parameters")
# 「影响与可能性」别名键：模型常用 impact / 影响与可能性 等，导致该列空白
IMPACT_ALIASES = (
    "impact_and_likelihood", "impact", "likelihood",
    "影响与可能性", "影响可能性", "影响及可能性",
)
ID_REF_RE = re.compile(r"(?:ISS|RISK|ISSUE)[-\s_]?\d+|R-\d+", re.IGNORECASE)


def _deliverables_substantive(deliverables) -> bool:
    """deliverables 仅当含中文实质说明（交付范围/验证结论）才渲染成章；
    纯文件名 token（report/redline/decision_state 等）或空值视为无内容而省略——
    文件清单由对话交付消息负责，报告内不重复裸键名。"""
    if not _nonempty(deliverables):
        return False
    text = json.dumps(deliverables, ensure_ascii=False)
    return bool(re.search(r"[一-鿿]", text))


def _nonempty(value) -> bool:
    if value is None:
        return False
    if isinstance(value, str):
        return bool(value.strip())
    if isinstance(value, (list, dict)):
        return len(value) > 0
    return bool(value)


def validate_report_input(data: dict) -> None:
    """拒绝错误输入而非渲染空壳（修复 Bug1）：
    - 误把 operations.json 传给报告构建器；
    - report.json 核心字段全空。"""
    if not isinstance(data, dict):
        raise ValueError("report 输入必须是 JSON 对象")
    if "operations" in data and not any(_nonempty(data.get(f)) for f in REPORT_CORE_FIELDS):
        raise ValueError(
            "输入疑似 operations.json（仅含 operations 数组，供 review_docx.py apply 用），"
            "而 build_review_report.py 需要 report.json（含 risks / scope / "
            "executive_summary 等顶级字段）。请改用 report.json，二者不可混用。"
        )
    if not any(_nonempty(data.get(f)) for f in REPORT_CORE_FIELDS):
        raise ValueError(
            "report.json 核心字段（risks / scope / executive_summary / "
            "structural_parameters）全部为空，将渲染空壳报告。请填充内容后重试。"
        )


# 必填章节（对应审查步骤，缺一即审查未完成）；symmetry/ip/coverage/missing/pending/
# deliverables 允许"本节不适用/无/empty_reason"，不在此表。
REQUIRED_SECTIONS = {
    "scope": "一·审查范围与已知事实",
    "facts": "二·事实、假设与用户立场",
    "playbook_status": "三·客户规则与审查立场校准",
    "executive_summary": "四·执行摘要",
    "structural_parameters": "五·结构参数硬检查",
    "risks": "六·风险清单",
    "verification": "十二·准据法核验状态",
}


def validate_report_completeness(data: dict) -> None:
    """报告必填章节逐节核验：任一为空 = 审查未完成（修复"报告仍有空缺"）。
    强制按审查步骤把每节填实，而非渲染空节。"""
    empty = [name for key, name in REQUIRED_SECTIONS.items() if not _nonempty(data.get(key))]
    risks = data.get("risks") or []
    has_hi_mid = any(
        str(r.get("level", "")).strip().lower() in ("高", "中", "high", "medium")
        for r in risks
    )
    if has_hi_mid and not _nonempty(data.get("recommendations")):
        empty.append("十一·建议文本与谈判备选（存在高/中风险时必填）")
    if empty:
        raise ValueError(
            "审查报告以下必填章节为空，审查未完成（请按审查步骤补全后重新生成）：\n  - "
            + "\n  - ".join(empty)
        )


def validate_risk_completeness(risks: list[dict]) -> None:
    """每条风险的关键列必须填实，避免风险清单出现空单元格（内容空缺/错漏）。"""
    fields = [("location", "位置"), ("issue", "问题"), ("recommendation", "建议及备选")]
    bad = []
    for item in risks:
        rid = str(item.get("issue_id", "?"))
        miss = [label for key, label in fields if not str(item.get(key, "")).strip()]
        if not str(item.get("basis_tag", "")).strip():
            miss.append("依据")
        if miss:
            bad.append(f"{rid}: 缺 {'/'.join(miss)}")
    if bad:
        raise ValueError(
            "风险清单存在内容空缺（每条风险须填齐 位置/问题/建议/依据）：\n  - "
            + "\n  - ".join(bad)
        )


def normalize_risk_fields(risks: list[dict]) -> None:
    """统一 impact_likelihood 别名键（修复 Bug2：模型用 impact / 影响与可能性
    等别名，导致「影响与可能性」列空白）。"""
    for item in risks:
        if not str(item.get("impact_likelihood", "")).strip():
            for alt in IMPACT_ALIASES:
                if str(item.get(alt, "")).strip():
                    item["impact_likelihood"] = item[alt]
                    break


def require_impact_likelihood(risks: list[dict]) -> None:
    """高/中风险必须给出影响×可能性（核心定级方法 D1-S3；修复 Bug2 列空白）。"""
    missing = [
        str(item.get("issue_id", "?"))
        for item in risks
        if str(item.get("level", "")).strip().lower() in ("高", "中", "high", "medium")
        and not str(item.get("impact_likelihood", "")).strip()
    ]
    if missing:
        raise ValueError(
            "以下高/中风险缺少「影响与可能性」(impact_likelihood)，"
            "须按影响×可能性给出（如「影响:高 可能性:中」）: " + ", ".join(missing)
        )


def validate_risk_ids(risks: list[dict]) -> None:
    issue_ids = [str(item.get("issue_id", "")).strip() for item in risks]
    if any(not issue_id for issue_id in issue_ids):
        raise ValueError("every report risk requires issue_id")
    if len(issue_ids) != len(set(issue_ids)):
        raise ValueError("report risk issue_id values must be unique")


def renumber_risks(risks: list[dict]) -> dict:
    """按报告最终排序重编号 R-001…R-N（修复 B7）。返回 旧ID→新ID 映射，
    供 propagate_id_rename 同步全文交叉引用。"""
    mapping: dict[str, str] = {}
    for index, item in enumerate(risks, 1):
        original = str(item.get("issue_id", "")).strip()
        new_id = f"R-{index:03d}"
        if original and original != new_id:
            print(f"risk renumbered: {original} -> {new_id}")
            mapping[original] = new_id
        item["issue_id"] = new_id
    return mapping


def propagate_id_rename(data: dict, mapping: dict) -> None:
    """把风险重编号同步到全文交叉引用（修复 Bug3：风险表用 R-，
    正文执行摘要/建议等仍引 ISS- 旧号，前后不一致）。
    两段式替换（旧号→哨兵→新号）避免新旧 R- 编号互相覆盖。"""
    if not mapping:
        return
    olds = sorted(mapping, key=len, reverse=True)  # 长者优先，避免 ISS-1 误伤 ISS-12
    sentinels = {old: f"\x00{idx}\x00" for idx, old in enumerate(olds)}

    def rep(text: str) -> str:
        for old in olds:
            text = text.replace(old, sentinels[old])
        for old, sent in sentinels.items():
            text = text.replace(sent, mapping[old])
        return text

    def walk(obj):
        if isinstance(obj, str):
            return rep(obj)
        if isinstance(obj, list):
            return [walk(x) for x in obj]
        if isinstance(obj, dict):
            # 不动 issue_id（已是新值），其余字段同步交叉引用
            return {k: (v if k == "issue_id" else walk(v)) for k, v in obj.items()}
        return obj

    for key in list(data.keys()):
        data[key] = walk(data[key])


def warn_dangling_refs(data: dict) -> None:
    """全文扫描风险编号引用，警告未对应到风险清单的悬挂引用（Bug3 安全网）。"""
    valid = {str(r.get("issue_id", "")).strip() for r in data.get("risks", [])}
    chunks = [json.dumps({k: v for k, v in data.items() if k != "risks"}, ensure_ascii=False)]
    for r in data.get("risks", []):
        chunks.append(f"{r.get('issue','')} {r.get('recommendation','')} {r.get('impact_likelihood','')}")
    text = " ".join(chunks)
    bad = set()
    for m in ID_REF_RE.finditer(text):
        token = m.group(0)
        upper = token.upper().replace(" ", "").replace("_", "")
        if upper.startswith("R-"):
            num = re.sub(r"\D", "", upper)
            if num and f"R-{int(num):03d}" not in valid:
                bad.add(token)
        else:
            bad.add(token)  # ISS-/RISK-/ISSUE- 在传播后不应残留
    if bad:
        print(
            "[警告] 报告存在未对应到风险清单的编号引用（疑似交叉引用失配）: "
            + ", ".join(sorted(bad))
        )


_LEVEL_KEYS = ("level", "risk_level", "级别", "风险等级", "severity", "risk")
_LEVEL_EXACT = {
    "high": "高", "medium": "中", "low": "低", "高": "高", "中": "中", "低": "低",
    "h": "高", "m": "中", "l": "低",
}


def _norm_level(value) -> str | None:
    """把各种等级写法归一到 高/中/低；识别不了返回 None。

    真机 655abff3：模型把 level 写成 `高风险`/`H`/`risk_level` 等变体，原严格匹配
    （键必须是 `level`、值必须精确等于 6 个之一）一个都数不上 → 渲染 0/0/0。此处
    放宽键与值：先精确别名，再子串匹配（含 高/high/严重 → 高，依此类推）。
    """
    v = str(value or "").strip().lower()
    if not v:
        return None
    if v in _LEVEL_EXACT:
        return _LEVEL_EXACT[v]
    if any(k in v for k in ("高", "high", "严重", "重大", "critical", "severe")):
        return "高"
    if any(k in v for k in ("中", "medium", "moderate", "中等")):
        return "中"
    if any(k in v for k in ("低", "low", "minor", "轻微")):
        return "低"
    return None


def risk_level_counts(risks: list[dict]) -> dict:
    """从风险清单按等级计算权威数量（修复 #2：执行摘要手数与正文不符）。

    键取 `_LEVEL_KEYS` 中第一个非空者，值经 `_norm_level` 归一，兼容大小写与写法变体。
    """
    counts = {"高": 0, "中": 0, "低": 0}
    for r in risks:
        if not isinstance(r, dict):
            continue
        raw = next((r.get(k) for k in _LEVEL_KEYS if r.get(k)), "")
        level = _norm_level(raw)
        if level:
            counts[level] += 1
    return counts


def guard_risk_counts(risks: list[dict], counts: dict) -> None:
    """风险清单非空但三档计数全 0 = level 字段缺失或取值无法识别，属自相矛盾。

    在生成侧即失败（真机 655abff3：0/0/0 与正文 30 项矛盾漏到下游核验才发现，
    触发一整轮 verify→修复往返）。此守卫把它变成生成侧 0 秒拦截。
    """
    n = len([r for r in risks if isinstance(r, dict)])
    if n > 0 and sum(counts.values()) == 0:
        raise ValueError(
            f"风险清单含 {n} 项但无一项能识别风险等级——每条风险须有 level 字段，"
            "取值为 高/中/低（或 high/medium/low）。请修正各风险条目的 level 后重新生成，"
            "不得渲染与正文矛盾的 0/0/0 风险统计。"
        )


SUMMARY_COUNT_RE = {
    "高": re.compile(r"高风险\s*(\d+)\s*项"),
    "中": re.compile(r"中风险\s*(\d+)\s*项"),
    "低": re.compile(r"低风险\s*(\d+)\s*项"),
}


def lint_summary_counts(executive_summary, counts: dict) -> None:
    """执行摘要声称的风险数量须与实际清单一致（修复 #2，确定性失败）。"""
    text = json.dumps(executive_summary, ensure_ascii=False) if executive_summary else ""
    mismatches = []
    for level, rx in SUMMARY_COUNT_RE.items():
        m = rx.search(text)
        if m and int(m.group(1)) != counts[level]:
            mismatches.append(f"{level}风险 声称{m.group(1)}项 / 实际{counts[level]}项")
    if mismatches:
        raise ValueError(
            "执行摘要风险数量与风险清单不符（请改摘要数字或核对清单）：\n  - "
            + "\n  - ".join(mismatches)
        )


def warn_structural_basis(structural_parameters: list[dict]) -> None:
    """结构参数偏差/缺失项的依据须具体（修复 #1，启发式告警）。"""
    flag_status = ("deviates", "below", "reversed", "absent", "偏差", "低于", "反转", "缺失")
    bare = re.compile(r"^\s*(\[(用规|要点|法规|惯例)\]\s*)+$")
    weak = []
    for p in structural_parameters or []:
        status = str(p.get("status", "")).lower()
        if not any(f in status for f in flag_status):
            continue
        basis = str(p.get("basis_tag", "")).strip()
        if not basis or bare.match(basis):
            weak.append(str(p.get("parameter", "?")))
    if weak:
        print(
            "[警告] 以下结构参数为偏差/缺失但依据为空或仅裸标签，请补具体依据: "
            + ", ".join(weak)
        )


def aggregate_pending(data: dict) -> None:
    """交叉校验"待核查"一致性（修复 B5：十二节列待核查项、
    十三节却称"无"的前后矛盾）。pending 为空时自动聚合全文待核查条目。"""
    found: list[str] = []

    def scan(obj) -> None:
        if isinstance(obj, str):
            if "待核查" in obj or "待核实" in obj:
                found.append(obj if len(obj) <= 120 else obj[:117] + "…")
        elif isinstance(obj, list):
            for entry in obj:
                scan(entry)
        elif isinstance(obj, dict):
            for value in obj.values():
                scan(value)

    for key in ("verification", "risks", "structural_parameters", "missing_terms", "executive_summary"):
        scan(data.get(key))

    pending = data.get("pending")
    is_empty = (
        not pending
        or (isinstance(pending, str) and pending.strip() in {"无", "无。", "暂无"})
        or (isinstance(pending, list) and not pending)
    )
    if found and is_empty:
        deduped = list(dict.fromkeys(found))
        data["pending"] = deduped
        print(
            f"warning: pending was empty but {len(deduped)} 待核查 item(s) found "
            "elsewhere in the report — auto-aggregated into 待核查事项"
        )


IDENTIFIER_RE_KEYS = ("item", "direction", "status")


def validate_coverage(coverage) -> None:
    """coverage 结构校验（修复 B4：模型把 playbook 内部英文键名
    如 universal_rules 直接塞入第七节、渲染为裸标签无展开）。"""
    import re as _re

    bare_identifier = _re.compile(r"^[A-Za-z0-9_\-]+$")
    has_cjk = _re.compile(r"[一-鿿]")

    def fail(detail: str) -> None:
        raise ValueError(
            f"coverage 数据不合格：{detail}。"
            "请使用结构化条目 {item(中文), direction, status, location}，"
            "或无审查原则时使用 {\"empty_reason\": \"…\"}（见 input-schema.md）"
        )

    if coverage is None or isinstance(coverage, str):
        return
    if isinstance(coverage, dict):
        if "empty_reason" in coverage or _is_bilingual_dict(coverage):
            return
        bare = [k for k in coverage if bare_identifier.match(str(k))]
        if bare:
            fail(f"字典键为裸英文标识符 {bare}，疑似 playbook 内部键名直出")
        return
    if isinstance(coverage, list):
        for entry in coverage:
            if isinstance(entry, str) and bare_identifier.match(entry):
                fail(f"列表项为裸英文标识符 '{entry}'")
            if isinstance(entry, dict):
                if not entry.get("item") or not entry.get("status"):
                    fail(f"条目缺少 item/status 字段：{entry}")
                if not has_cjk.search(str(entry.get("item", ""))):
                    fail(f"item 须为中文描述，得到：{entry.get('item')}")


class ChapterNumberer:
    """动态章节编号（修复 B1）：序号按实际渲染顺序分配，
    任何章节的有无都不会造成编号跳跃。"""

    CN = ["一", "二", "三", "四", "五", "六", "七", "八", "九", "十",
          "十一", "十二", "十三", "十四", "十五", "十六"]

    def __init__(self) -> None:
        self.index = 0

    def title(self, name: str) -> str:
        self.index += 1
        return f"{self.CN[self.index - 1]}、{name}"


# --- 立场一致性（修复人工评测：单条风险在同段内并列买卖双方风险，立场摇摆）---
# 一份报告只有一个声明立场；风险清单只从该方视角评估不利后果，
# 双方对比只允许出现在「权利对称性检查」节。
ROLE_PAIRS = [
    ("被许可方", "许可方"),   # 含子串，必须先于「许可方」匹配
    ("买方", "卖方"),
    ("买受人", "出卖人"),
    ("承租方", "出租方"),
    ("承包方", "发包方"),
    ("受让方", "转让方"),
    ("甲方", "乙方"),
]
HEDGE_WORDS = ("反之", "另一方面", "相反", "反过来")
# 担险动词簇：对方被框定为「承受风险的主体」才算串立场；
# 仅提到对方的损失（如卖方视角下的「买方间接损失」=卖方赔付敞口）不算。
RISK_VERBS = "面临|承担|遭受|蒙受|受损|担责|无法获赔|无法充分获赔"


def _role_of(text: str) -> str | None:
    """从立场描述中识别角色 token（被许可方先于许可方）。"""
    for a, b in ROLE_PAIRS:
        if a in text:
            return a
        if b in text:
            return b
    return None


def _opposite_role(role: str) -> str | None:
    for a, b in ROLE_PAIRS:
        if role == a:
            return b
        if role == b:
            return a
    return None


def resolve_position(data: dict) -> dict | None:
    """确定报告的唯一声明立场。优先结构化 review_position，
    否则从 facts 文本推断「基于X方…立场」。返回 {party, basis, role, opposite}。"""
    pos = data.get("review_position")
    party = ""
    basis = ""
    if isinstance(pos, dict):
        party = str(pos.get("party", "")).strip()
        basis = str(pos.get("basis", "")).strip()
    elif isinstance(pos, str):
        party = pos.strip()
    if not party:
        facts = data.get("facts")
        facts_text = json.dumps(facts, ensure_ascii=False) if facts else ""
        m = re.search(r"基于([^，。；,]{1,20}?)立场", facts_text)
        if m:
            party = m.group(1).strip()
            basis = basis or "（自 facts 推断）"
    if not party:
        return None
    role = _role_of(party)
    return {
        "party": party,
        "basis": basis,
        "role": role,
        "opposite": _opposite_role(role) if role else None,
    }


def lint_stance_consistency(risks: list[dict], position: dict | None) -> None:
    """逐条检查风险清单是否串入对方视角。
    硬失败：同一条目出现转折连接词 + 对方被描述为担险（截图原型）。
    告警：对方被描述为担险但无转折词（可能合理，提示复核）。"""
    if not position or not position.get("opposite"):
        return
    opp = position["opposite"]
    declared = position["role"]
    # 对方作为「担险主体」：对方 token 后 6 字内出现担险动词，或「对<对方>…不利」
    subject_at_risk = re.compile(
        rf"{re.escape(opp)}[^。；，,]{{0,6}}?(?:{RISK_VERBS})"
        rf"|对{re.escape(opp)}[^。；，,]{{0,8}}?(?:不利|无法|受损)"
    )
    errors: list[str] = []
    for r in risks:
        rid = r.get("issue_id", "?")
        text = " ".join(
            str(r.get(k, ""))
            for k in ("issue", "impact_likelihood", "recommendation")
        )
        if not subject_at_risk.search(text):
            continue
        hedged = [h for h in HEDGE_WORDS if h in text]
        if hedged:
            errors.append(
                f"{rid}: 风险条目内用「{hedged[0]}」并列了对方（{opp}）"
                "作为担险主体，立场摇摆"
            )
        else:
            print(
                f"[警告] 风险 {rid} 把对方（{opp}）描述为担险主体——"
                f"本报告立场为{position['party']}（{declared or '声明方'}），"
                "风险清单应只评估己方不利后果；双方对比请移入「权利对称性检查」节。"
            )
    if errors:
        raise ValueError(
            "立场一致性校验失败（风险清单必须只从声明立场一方视角评估，"
            "双方对比仅允许出现在「权利对称性检查」节）：\n  - "
            + "\n  - ".join(errors)
        )


def add_position_banner(doc: Document, position: dict | None) -> None:
    """在免责声明下方渲染统一立场声明，作为全文遵循基准。"""
    if not position:
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    basis = f"，依据：{position['basis']}" if position.get("basis") else ""
    set_font(
        p.add_run(
            f"审查立场：{position['party']}{basis}。"
            "本报告风险清单均从该方视角评估不利后果；双方立场对比见「权利对称性检查」节。"
        ),
        REPORT_FONT_SIZES["body"],
        True,
    )


def warn_missing_basis_quote(risks: list[dict]) -> None:
    """高/中风险的 [法规]/[用规]/[要点] 依据须带「」原文摘录；缺失则警告（非致命）。"""
    need = ("[法规]", "[用规]", "[要点]")
    quote = re.compile(r"[「『\"“][^」』\"”]{2,}")
    hits = [f"{r.get('issue_id','?')}({r.get('location','')})"
            for r in risks
            if str(r.get("level", "")).strip().lower() in ("高", "中", "high", "medium")
            and any(t in str(r.get("basis_tag", "")) for t in need)
            and not quote.search(str(r.get("basis_tag", "")))]
    if hits:
        print("[警告] 以下高/中风险依据缺「」原文摘录，请补全原文或降级为待核查: " + "; ".join(hits))


def build(data: dict, output: Path) -> None:
    output = generated_path(output, "review report")
    validate_report_input(data)
    validate_report_completeness(data)
    validate_risk_completeness(data.get("risks", []))
    normalize_risk_fields(data.get("risks", []))
    validate_risk_ids(data.get("risks", []))
    require_impact_likelihood(data.get("risks", []))
    mapping = renumber_risks(data.get("risks", []))
    propagate_id_rename(data, mapping)
    warn_dangling_refs(data)
    validate_coverage(data.get("coverage"))
    warn_missing_basis_quote(data.get("risks", []))
    warn_structural_basis(data.get("structural_parameters", []))
    risk_counts = risk_level_counts(data.get("risks", []))
    guard_risk_counts(data.get("risks", []), risk_counts)
    lint_summary_counts(data.get("executive_summary"), risk_counts)
    position = resolve_position(data)
    lint_stance_consistency(data.get("risks", []), position)
    aggregate_pending(data)
    doc = Document()
    configure(doc)
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(
        title.add_run(data.get("report_title", "合同审查报告")),
        REPORT_FONT_SIZES["title"],
        True,
    )
    first = doc.add_paragraph()
    first.paragraph_format.space_before = Pt(8)
    set_font(first.add_run(DISCLAIMER), REPORT_FONT_SIZES["disclaimer"], True)
    add_position_banner(doc, position)

    ch = ChapterNumberer()
    add_section(doc, ch.title("审查范围与已知事实"), data.get("scope"))
    add_section(doc, ch.title("事实、假设与用户立场"), data.get("facts"))
    add_section(doc, ch.title("客户规则与审查立场校准"), data.get("playbook_status"))
    add_section(doc, ch.title("执行摘要"), data.get("executive_summary"))
    add_matrix(
        doc,
        ch.title("结构参数硬检查"),
        _fill_basis_column(data.get("structural_parameters", [])),
        ["参数", "首选", "底线", "合同实际", "状态", "依据"],
        ["parameter", "preferred", "bottom_line", "actual", "status", "basis"],
        [1150, 1450, 1350, 1750, 1400, 1700],
    )
    add_risk_table(doc, ch.title("风险清单"), data.get("risks", []))
    # 权威风险计数（修复 #2：让读者以脚本计算的数量为准，不依赖摘要手数）
    add_plain_para(
        doc,
        f"风险统计（按实际清单计）：高风险 {risk_counts['高']} 项 / "
        f"中风险 {risk_counts['中']} 项 / 低风险 {risk_counts['低']} 项。",
    )
    coverage = data.get("coverage")
    if isinstance(coverage, list) and coverage and all(isinstance(e, dict) for e in coverage):
        add_matrix(
            doc,
            ch.title("双向审查覆盖状态"),
            coverage,
            ["审查项", "方向", "状态", "位置"],
            ["item", "direction", "status", "location"],
            [3200, 1200, 1600, 2800],
        )
    else:
        add_section(doc, ch.title("双向审查覆盖状态"), coverage)
    add_section(doc, ch.title("缺失保护与系统性不一致"), data.get("missing_terms"))
    add_matrix(
        doc,
        ch.title("权利对称性检查"),
        data.get("symmetry", []),
        ["权利", "客户立场", "相对方立场", "评估"],
        ["right", "client_position", "counterparty_position", "assessment"],
        [1600, 2100, 2100, 3000],
        empty_reason="本节不适用：合同未涉及需对称性比较的实质性权利分配条款。",
    )
    add_matrix(
        doc,
        ch.title("知识产权三层检查"),
        data.get("ip_analysis", []),
        ["对象", "归属", "申请或控制", "维权"],
        ["category", "ownership", "control", "enforcement"],
        [1500, 2350, 2400, 2550],
        empty_reason="本节不适用：合同未涉及知识产权归属、许可或维权条款。",
    )
    add_section(doc, ch.title("建议文本与谈判备选"), data.get("recommendations"))
    add_section(doc, ch.title("准据法核验状态"), data.get("verification"))
    add_section(doc, ch.title("待核查事项"), data.get("pending"))
    # 交付文件清单已在对话交付消息中给出；报告内仅当 deliverables 是一段
    # 实质交付/验证说明（含中文）时才保留该章，纯文件名 token（report/redline/
    # decision_state）或空值一律省略，避免无意义的裸键名章节。
    if _deliverables_substantive(data.get("deliverables")):
        add_section(doc, ch.title("交付与验证说明"), data.get("deliverables"))

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    print(f"created {output}")


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("input", type=Path)
    parser.add_argument("output", type=Path)
    args = parser.parse_args()
    data = json.loads(args.input.read_text(encoding="utf-8"))
    # DOCX hard rule (D3-S2): no emoji. Model-authored status markers map to
    # bracket labels; anything else in emoji ranges is stripped.
    build(sanitize_data(data), args.output)


if __name__ == "__main__":
    main()
