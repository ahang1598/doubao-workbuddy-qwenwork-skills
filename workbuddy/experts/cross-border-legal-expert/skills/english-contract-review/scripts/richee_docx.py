#!/usr/bin/env python3
"""Shared Richee DOCX styling helpers."""

from __future__ import annotations

from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


RICHEE_BLACK = "1A1A1A"
RICHEE_TEXT = "1A1A1A"
RICHEE_MUTED = "6B6B6B"
RICHEE_LINE = "E2E5EA"
RICHEE_BG = "F7F7F7"
RICHEE_WHITE = "FFFFFF"
RICHEE_ACCENT = "32D583"
RICHEE_ACCENT_DARK = "039855"
RICHEE_RED = "D92D20"
RICHEE_RED_BG = "FEF3F2"
RICHEE_AMBER = "B54708"
RICHEE_AMBER_BG = "FFFAEB"
RICHEE_GREEN = "039855"
RICHEE_GREEN_BG = "ECFDF3"
RICHEE_BLUE = "175CD3"
RICHEE_BLUE_BG = "EFF8FF"

FONT_CJK = "PingFang SC"
FONT_LATIN = "Arial"
CONTENT_WIDTH_DXA = 8800


def rgb(color_hex: str) -> RGBColor:
    value = color_hex.lstrip("#")
    return RGBColor(
        int(value[0:2], 16),
        int(value[2:4], 16),
        int(value[4:6], 16),
    )


def set_run_font(
    run,
    *,
    language: str = "zh",
    size: float = 10.5,
    bold: bool = False,
    color_hex: str = RICHEE_TEXT,
) -> None:
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = FONT_LATIN if language == "en" else FONT_CJK
    run.font.color.rgb = rgb(color_hex)
    r_pr = run._element.get_or_add_rPr()
    fonts = r_pr.find(qn("w:rFonts"))
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, fonts)
    fonts.set(qn("w:ascii"), FONT_LATIN)
    fonts.set(qn("w:hAnsi"), FONT_LATIN)
    fonts.set(qn("w:eastAsia"), FONT_CJK)


def set_style_font(
    style,
    *,
    size: float,
    bold: bool = False,
    color_hex: str = RICHEE_TEXT,
    line_spacing: float | None = None,
    space_before: float | None = None,
    space_after: float | None = None,
) -> None:
    style.font.name = FONT_CJK
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = rgb(color_hex)
    r_pr = style._element.get_or_add_rPr()
    fonts = r_pr.find(qn("w:rFonts"))
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, fonts)
    fonts.set(qn("w:ascii"), FONT_LATIN)
    fonts.set(qn("w:hAnsi"), FONT_LATIN)
    fonts.set(qn("w:eastAsia"), FONT_CJK)
    if line_spacing is not None:
        style.paragraph_format.line_spacing = line_spacing
    if space_before is not None:
        style.paragraph_format.space_before = Pt(space_before)
    if space_after is not None:
        style.paragraph_format.space_after = Pt(space_after)


def configure_document(doc, *, report: bool) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2 if report else 2.3)
    section.bottom_margin = Cm(2.2 if report else 2.3)
    section.left_margin = Cm(2.6)
    section.right_margin = Cm(2.6)

    set_style_font(
        doc.styles["Normal"],
        size=10.5,
        line_spacing=1.5 if report else 1.18,
        space_after=5 if report else 4,
    )
    set_style_font(
        doc.styles["Title"],
        size=20 if report else 18,
        bold=True,
        space_after=6,
    )
    set_style_font(
        doc.styles["Heading 1"],
        size=13,
        bold=True,
        space_before=12,
        space_after=6,
    )
    set_style_font(
        doc.styles["Heading 2"],
        size=11.5,
        bold=True,
        space_before=9,
        space_after=4,
    )
    set_style_font(
        doc.styles["List Bullet"],
        size=10.5,
        line_spacing=1.5 if report else 1.18,
        space_after=3,
    )
    for name in ("Heading 1", "Heading 2"):
        doc.styles[name].paragraph_format.keep_with_next = True


def set_paragraph_border(
    paragraph,
    *,
    side: str,
    color_hex: str,
    size: int = 12,
    space: int = 4,
) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    edge = p_bdr.find(qn(f"w:{side}"))
    if edge is None:
        edge = OxmlElement(f"w:{side}")
        p_bdr.append(edge)
    edge.set(qn("w:val"), "single")
    edge.set(qn("w:sz"), str(size))
    edge.set(qn("w:space"), str(space))
    edge.set(qn("w:color"), color_hex)


def set_paragraph_shading(paragraph, color_hex: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shading = p_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        p_pr.append(shading)
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), color_hex)


def style_report_heading(paragraph) -> None:
    set_paragraph_border(
        paragraph,
        side="left",
        color_hex=RICHEE_ACCENT,
        size=18,
        space=5,
    )
    paragraph.paragraph_format.left_indent = Cm(0.18)


def style_title(paragraph, *, report: bool) -> None:
    set_paragraph_border(
        paragraph,
        side="bottom",
        color_hex=RICHEE_ACCENT,
        size=14 if report else 10,
        space=6,
    )


def style_disclaimer(paragraph) -> None:
    set_paragraph_shading(paragraph, RICHEE_AMBER_BG)
    set_paragraph_border(
        paragraph,
        side="left",
        color_hex=RICHEE_AMBER,
        size=18,
        space=5,
    )
    paragraph.paragraph_format.left_indent = Cm(0.18)
    paragraph.paragraph_format.right_indent = Cm(0.12)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(9)
    paragraph.paragraph_format.line_spacing = 1.25


def set_cell_shading(cell, color_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), color_hex)


def set_cell_margins(
    cell,
    *,
    top: int = 100,
    start: int = 120,
    bottom: int = 100,
    end: int = 120,
) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_borders(cell, color_hex: str = RICHEE_LINE, size: int = 4) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge_name in ("top", "start", "bottom", "end", "insideH", "insideV"):
        edge = borders.find(qn(f"w:{edge_name}"))
        if edge is None:
            edge = OxmlElement(f"w:{edge_name}")
            borders.append(edge)
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), str(size))
        edge.set(qn("w:space"), "0")
        edge.set(qn("w:color"), color_hex)


def set_cell_width(cell, width: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width(table, widths: list[int], *, bordered: bool = True) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    indent = tbl_pr.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        tbl_pr.append(indent)
    indent.set(qn("w:w"), "0")
    indent.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if bordered:
                set_cell_borders(cell)


def remove_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge = borders.find(qn(f"w:{edge_name}"))
        if edge is None:
            edge = OxmlElement(f"w:{edge_name}")
            borders.append(edge)
        edge.set(qn("w:val"), "nil")


def configure_table_pagination(table) -> None:
    for index, row in enumerate(table.rows):
        prevent_row_split(row)
        tr_pr = row._tr.get_or_add_trPr()
        if index == 0 and tr_pr.find(qn("w:tblHeader")) is None:
            header = OxmlElement("w:tblHeader")
            header.set(qn("w:val"), "true")
            tr_pr.append(header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def style_header_row(row) -> None:
    for cell in row.cells:
        set_cell_shading(cell, RICHEE_BLACK)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.15
            for run in paragraph.runs:
                set_run_font(
                    run,
                    size=9,
                    bold=True,
                    color_hex=RICHEE_WHITE,
                )


def style_body_rows(table) -> None:
    for index, row in enumerate(table.rows[1:], start=1):
        if index % 2 == 0:
            for cell in row.cells:
                tc_pr = cell._tc.get_or_add_tcPr()
                if tc_pr.find(qn("w:shd")) is None:
                    set_cell_shading(cell, RICHEE_BG)
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.25


def style_status_cell(cell, value: str) -> None:
    normalized = value.strip().lower()
    if any(token in normalized for token in ("high", "高", "reversed", "below bottom line", "absent")):
        fill, color = RICHEE_RED_BG, RICHEE_RED
    elif any(token in normalized for token in ("medium", "mid", "中", "deviates", "待核查", "待定")):
        fill, color = RICHEE_AMBER_BG, RICHEE_AMBER
    elif any(token in normalized for token in ("low", "低", "conforms", "已核验", "通过")):
        fill, color = RICHEE_GREEN_BG, RICHEE_GREEN
    elif any(token in normalized for token in ("not specified", "info", "未说明")):
        fill, color = RICHEE_BLUE_BG, RICHEE_BLUE
    else:
        return
    set_cell_shading(cell, fill)
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            set_run_font(run, size=9, bold=True, color_hex=color)


def add_page_number_footer(doc, label: str) -> None:
    footer = doc.sections[0].footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_before = Pt(5)
    set_paragraph_border(
        paragraph,
        side="top",
        color_hex=RICHEE_LINE,
        size=6,
        space=4,
    )
    set_run_font(
        paragraph.add_run(f"{label}  |  "),
        size=8.5,
        color_hex=RICHEE_MUTED,
    )

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    for element in (begin, instruction, separate):
        run = paragraph.add_run()
        set_run_font(run, size=8.5, color_hex=RICHEE_MUTED)
        run._r.append(element)

    result = paragraph.add_run("1")
    set_run_font(result, size=8.5, color_hex=RICHEE_MUTED)

    end_run = paragraph.add_run()
    set_run_font(end_run, size=8.5, color_hex=RICHEE_MUTED)
    end_run._r.append(end)
