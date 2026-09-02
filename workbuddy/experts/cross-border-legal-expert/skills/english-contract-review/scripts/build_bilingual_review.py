#!/usr/bin/env python3
"""Build a stacked bilingual reviewed-contract DOCX."""

from __future__ import annotations

import argparse
import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from richee_docx import (
    CONTENT_WIDTH_DXA,
    RICHEE_BLACK,
    add_page_number_footer,
    configure_document,
    configure_table_pagination,
    set_run_font,
    set_table_width,
)
from skill_paths import generated_path


LANGUAGE_MODES = {"en_zh", "zh_en"}


def language_order(mode: str) -> tuple[str, str]:
    return ("zh", "en") if mode == "zh_en" else ("en", "zh")


def validate_input(data: dict) -> str:
    mode = str(data.get("language_mode", "en_zh"))
    if mode not in LANGUAGE_MODES:
        raise ValueError(f"language_mode must be one of {sorted(LANGUAGE_MODES)}")
    for key in (
        "title_en",
        "title_zh",
        "language_priority_en",
        "language_priority_zh",
    ):
        if not str(data.get(key, "")).strip():
            raise ValueError(f"bilingual review requires {key}")
    for index, item in enumerate(data.get("paragraphs", []), start=1):
        if item.get("kind") == "table":
            for row_index, row in enumerate(item.get("rows", []), start=1):
                for cell_index, cell in enumerate(row.get("cells", []), start=1):
                    if not str(cell.get("en", "")).strip() or not str(
                        cell.get("zh", "")
                    ).strip():
                        raise ValueError(
                            f"table {index} row {row_index} cell {cell_index} "
                            "requires en and zh"
                        )
        elif not str(item.get("en", "")).strip() or not str(
            item.get("zh", "")
        ).strip():
            raise ValueError(f"paragraph {index} requires en and zh")
    return mode


def add_pair(
    doc: Document,
    mode: str,
    number: str,
    en: str,
    zh: str,
    kind: str,
) -> None:
    heading = kind in {"title", "heading", "article"}
    values = {"en": en, "zh": zh}
    for position, language in enumerate(language_order(mode)):
        paragraph = doc.add_paragraph(
            style="Heading 1" if heading and position == 0 else None
        )
        prefix = f"{number} " if number else ""
        set_run_font(
            paragraph.add_run(prefix + values[language]),
            language=language,
            size=12.5 if heading and position == 0 else 11.5 if heading else 10.5,
            bold=heading,
            color_hex=RICHEE_BLACK,
        )
        paragraph.paragraph_format.keep_with_next = (
            heading or position < len(language_order(mode)) - 1
        )
        paragraph.paragraph_format.space_before = Pt(
            8 if heading and position == 0 else 0
        )
        paragraph.paragraph_format.space_after = Pt(
            2 if position == 0 else 6
        )


def add_table_pair(doc: Document, mode: str, headers: list, rows: list[dict]) -> None:
    column_count = max(len(headers), 1)
    table = doc.add_table(rows=1, cols=column_count)
    table.style = "Table Grid"
    set_table_width(table, [CONTENT_WIDTH_DXA // column_count] * column_count)

    for cell, header in zip(table.rows[0].cells, headers):
        cell.text = ""
        if isinstance(header, dict):
            values = {
                "en": str(header.get("en", "")),
                "zh": str(header.get("zh", "")),
            }
        else:
            values = {"en": str(header), "zh": str(header)}
        for position, language in enumerate(language_order(mode)):
            paragraph = cell.paragraphs[0] if position == 0 else cell.add_paragraph()
            set_run_font(
                paragraph.add_run(values[language]),
                language=language,
                size=9,
                bold=True,
                color_hex=RICHEE_BLACK,
            )
            paragraph.paragraph_format.space_after = Pt(1)

    for row_data in rows:
        cells = table.add_row().cells
        for cell, cell_data in zip(cells, row_data.get("cells", [])):
            cell.text = ""
            for position, language in enumerate(language_order(mode)):
                paragraph = cell.paragraphs[0] if position == 0 else cell.add_paragraph()
                set_run_font(
                    paragraph.add_run(str(cell_data.get(language, ""))),
                    language=language,
                    size=9,
                    color_hex=RICHEE_BLACK,
                )
                paragraph.paragraph_format.space_after = Pt(1)
                paragraph.paragraph_format.keep_together = True
    configure_table_pagination(table)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)


def build(data: dict, output: Path) -> None:
    mode = validate_input(data)
    output = generated_path(output, "bilingual DOCX")
    doc = Document()
    configure_document(doc, report=False)
    doc.core_properties.title = data["title_en"]
    doc.core_properties.subject = f"Reviewed contract; language_mode={mode}"
    doc.core_properties.author = "AI-assisted review"

    titles = {"en": data["title_en"], "zh": data["title_zh"]}
    for position, language in enumerate(language_order(mode)):
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.keep_with_next = True
        paragraph.paragraph_format.space_after = Pt(2 if position == 0 else 12)
        set_run_font(
            paragraph.add_run(titles[language]),
            language=language,
            size=18 if language == "en" else 16,
            bold=True,
            color_hex=RICHEE_BLACK,
        )

    for item in data.get("paragraphs", []):
        kind = item.get("kind", "clause")
        if kind == "table":
            add_table_pair(doc, mode, item.get("headers", []), item.get("rows", []))
        else:
            add_pair(
                doc,
                mode,
                str(item.get("number", "")),
                str(item.get("en", "")),
                str(item.get("zh", "")),
                kind,
            )

    add_pair(
        doc,
        mode,
        "",
        f"Language priority: {data['language_priority_en']}.",
        f"语言优先规则：{data['language_priority_zh']}。",
        "clause",
    )
    add_page_number_footer(doc, data["title_en"])
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    print(f"created {output} ({mode})")


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("input", type=Path)
    parser.add_argument("output", type=Path)
    args = parser.parse_args()
    build(json.loads(args.input.read_text(encoding="utf-8")), args.output)


if __name__ == "__main__":
    main()

