#!/usr/bin/env python3
"""PDF 合同入库转换器（共享内核 · SSOT）。

把 PDF 合同（数字版 / 扫描件 / 分栏等特殊版式）转成标准 Word，并在审查前过
**文字准确性闸门**：高/中/低三档，低档熔断（不对错误转换做法律判断）。

不写死处理工具：每个能力（PDF读取/渲染/OCR/PDF转docx/docx构建）维护一个**优先级
提供方注册表**，运行时探测环境、择优选用、缺则回退；某能力全无提供方时清晰报错而非崩溃。
  python3 pdf_intake.py --probe                      # 查看本环境探测到的提供方
  python3 pdf_intake.py <in.pdf> --out <std.docx> [--lang auto] [--manifest m.json]
  python3 pdf_intake.py <in.pdf> --out o.docx --prefer ocr=tesseract   # 强制某提供方
  PDF_INTAKE_DISABLE=ocrmypdf python3 pdf_intake.py ...                 # 禁用某提供方(测回退)
退出码：0 高/中档可审查（中档带待核对），3 低档熔断，4 缺必要能力。
"""
from __future__ import annotations
import argparse, importlib, io, json, os, re, shutil, statistics, subprocess, sys, tempfile
from pathlib import Path

CJK = re.compile(r"[一-鿿]")
KEY_TOKENS = ["合同", "协议", "甲方", "乙方", "违约", "Agreement", "Party",
              "Liability", "Termination", "Confidential"]

# ---------- 能力提供方注册表（按优先级；不硬编码到单一工具）----------
def _mod(name):
    try:
        importlib.import_module(name); return True
    except Exception:
        return False
def _cli(name): return shutil.which(name) is not None

# 每项：(provider_name, kind, available_fn)。dispatch 时按序取第一个可用且已实现的。
REGISTRY = {
    "pdf_read":   [("pymupdf", "py", lambda: _mod("fitz")),
                   ("pdfplumber", "py", lambda: _mod("pdfplumber")),
                   ("poppler", "cli", lambda: _cli("pdftotext"))],
    "pdf_render": [("pymupdf", "py", lambda: _mod("fitz")),
                   ("poppler", "cli", lambda: _cli("pdftoppm")),
                   ("pdf2image", "py", lambda: _mod("pdf2image"))],
    "ocr":        [("ocrmypdf", "cli", lambda: _cli("ocrmypdf")),
                   ("tesseract", "cli/py", lambda: _cli("tesseract") or _mod("pytesseract")),
                   ("easyocr", "py", lambda: _mod("easyocr"))],
    "pdf2docx":   [("pdf2docx", "py", lambda: _mod("pdf2docx"))],
    "docx_build": [("python-docx", "py", lambda: _mod("docx")),
                   ("ooxml-min", "builtin", lambda: True)],
}
_DISABLED = set(filter(None, os.environ.get("PDF_INTAKE_DISABLE", "").split(",")))
_PREFER: dict[str, str] = {}

def resolve(cap: str) -> str | None:
    pref = _PREFER.get(cap)
    cands = REGISTRY[cap]
    if pref:
        cands = [c for c in cands if c[0] == pref] + [c for c in cands if c[0] != pref]
    for name, _kind, ok in cands:
        if name in _DISABLED:
            continue
        try:
            if ok():
                return name
        except Exception:
            continue
    return None

def probe() -> dict:
    out = {}
    for cap, provs in REGISTRY.items():
        avail = [n for n, _k, ok in provs if n not in _DISABLED and _safe(ok)]
        out[cap] = {"selected": resolve(cap), "available": avail}
    return out
def _safe(fn):
    try: return fn()
    except Exception: return False

# ---------- 能力操作（dispatch 到所选提供方）----------
def read_text(pdf: Path) -> str:
    p = resolve("pdf_read")
    if p == "pymupdf":
        import fitz
        return "".join(pg.get_text() for pg in fitz.open(pdf))
    if p == "pdfplumber":
        import pdfplumber
        with pdfplumber.open(pdf) as d:
            return "\n".join((pg.extract_text() or "") for pg in d.pages)
    if p == "poppler":
        r = subprocess.run(["pdftotext", "-layout", str(pdf), "-"], capture_output=True, text=True)
        return r.stdout
    raise SystemExit("CAP_MISSING:pdf_read")

def doc_stats(pdf: Path) -> dict:
    """页数 / 是否含图像 / 段落(阅读序，支持分栏) —— 用所选 pdf_read 提供方。"""
    p = resolve("pdf_read")
    if p == "pymupdf":
        import fitz
        doc = fitz.open(pdf); pages = len(doc)
        has_images = any(pg.get_images() for pg in doc)
        # 栏数判定（修复：旧版 left>3 and right>3 会把"含双列表格的单栏合同"误判为双栏，
        # 从而走纯文字重排路径、丢掉所有表格）。真双栏的标志是「整宽块极少」+「无表格」。
        left = right = full = total = 0
        has_tables = False
        for pg in doc:
            w = pg.rect.width
            for b in pg.get_text("blocks"):
                if not b[4].strip():
                    continue
                x0, x1, c = b[0], b[2], (b[0] + b[2]) / 2
                total += 1
                if x0 < w * 0.4 and x1 > w * 0.6:      # 跨中线的整宽块
                    full += 1
                elif c < w * 0.45:
                    left += 1
                elif c > w * 0.55:
                    right += 1
            if not has_tables:
                try:
                    if len(pg.find_tables().tables) > 0:
                        has_tables = True
                except Exception:
                    pass
        frac_full = (full / total) if total else 0.0
        # 双栏 = 左右块都多 且 整宽块很少（<35%）且 无表格；否则按单栏（走 pdf2docx 保表）
        cols = 2 if (left > 3 and right > 3 and frac_full < 0.35 and not has_tables) else 1
        return {"pages": pages, "has_images": has_images, "columns": cols,
                "has_tables": has_tables}
    if p == "pdfplumber":
        import pdfplumber
        with pdfplumber.open(pdf) as d:
            pages = len(d.pages); has_images = any(pg.images for pg in d.pages)
        return {"pages": pages, "has_images": has_images, "columns": 1}
    # poppler-only：无版面块信息，栏数未知
    r = subprocess.run(["pdfinfo", str(pdf)], capture_output=True, text=True)
    m = re.search(r"Pages:\s+(\d+)", r.stdout)
    return {"pages": int(m.group(1)) if m else 1, "has_images": True, "columns": 1}

def reading_order_paras(pdf: Path, two_col: bool) -> list[str]:
    p = resolve("pdf_read")
    if p == "pymupdf":
        import fitz
        paras = []
        for pg in fitz.open(pdf):
            blocks = [b for b in pg.get_text("blocks") if b[4].strip()]
            if two_col:
                mid = pg.rect.width / 2
                ordered = sorted([b for b in blocks if (b[0]+b[2])/2 < mid], key=lambda b: b[1]) + \
                          sorted([b for b in blocks if (b[0]+b[2])/2 >= mid], key=lambda b: b[1])
            else:
                ordered = sorted(blocks, key=lambda b: b[1])
            for b in ordered:
                t = re.sub(r"[ \t]*\n[ \t]*", " ", b[4].strip())
                if t: paras.append(t)
        return paras
    # 其他读取器无块坐标：按文本行切段
    return [ln.strip() for ln in read_text(pdf).splitlines() if ln.strip()]

def render_pages(pdf: Path, dpi: int):
    """yield PIL.Image —— 用所选 pdf_render 提供方。"""
    from PIL import Image
    p = resolve("pdf_render")
    if p == "pymupdf":
        import fitz
        for pg in fitz.open(pdf):
            yield Image.open(io.BytesIO(pg.get_pixmap(dpi=dpi).tobytes("png")))
    elif p == "poppler":
        with tempfile.TemporaryDirectory() as tmp:
            subprocess.run(["pdftoppm", "-png", "-r", str(dpi), str(pdf), str(Path(tmp)/"p")], check=True)
            for f in sorted(Path(tmp).glob("p*.png")):
                yield Image.open(f).copy()
    elif p == "pdf2image":
        from pdf2image import convert_from_path
        for im in convert_from_path(str(pdf), dpi=dpi):
            yield im
    else:
        raise SystemExit("CAP_MISSING:pdf_render")

def ocr_to_paras(pdf: Path, lang: str, two_col: bool) -> list[str]:
    p = resolve("ocr")
    if p is None:
        raise SystemExit("CAP_MISSING:ocr")
    if p == "ocrmypdf":
        with tempfile.TemporaryDirectory() as tmp:
            ocred = Path(tmp) / "o.pdf"
            r = subprocess.run(["ocrmypdf", "--language", lang, "--output-type", "pdf",
                                "--optimize", "0", "--force-ocr", str(pdf), str(ocred)],
                               capture_output=True, text=True)
            if r.returncode != 0:
                raise SystemExit("OCR_FAIL:" + (r.stderr.strip().splitlines()[-1] if r.stderr else "ocrmypdf"))
            return reading_order_paras(ocred, two_col)
    if p == "tesseract":
        paras = []
        use_py = _mod("pytesseract") and not _cli("tesseract")
        for img in render_pages(pdf, 200):
            if _mod("pytesseract"):
                import pytesseract
                txt = pytesseract.image_to_string(img, lang=lang)
            else:
                with tempfile.TemporaryDirectory() as tmp:
                    ip = Path(tmp)/"p.png"; img.save(ip)
                    r = subprocess.run(["tesseract", str(ip), "stdout", "-l", lang], capture_output=True, text=True)
                    txt = r.stdout
            paras += [ln.strip() for ln in txt.splitlines() if ln.strip()]
        return paras
    if p == "easyocr":
        import easyocr, numpy as np
        reader = easyocr.Reader([x for x in ("ch_sim", "en") if True])
        paras = []
        for img in render_pages(pdf, 200):
            for line in reader.readtext(np.array(img), detail=0):
                if line.strip(): paras.append(line.strip())
        return paras
    raise SystemExit("CAP_MISSING:ocr")

def ocr_confidence(pdf: Path, lang: str) -> dict | None:
    if resolve("ocr") not in ("ocrmypdf", "tesseract") or not _mod("pytesseract"):
        return None  # 仅 tesseract 系提供词级置信；其他引擎跳过，靠覆盖率/关键词判档
    import pytesseract
    confs = []
    for img in render_pages(pdf, 200):
        data = pytesseract.image_to_data(img, lang=lang, output_type=pytesseract.Output.DICT)
        confs += [int(c) for c in data["conf"] if c not in ("-1", "") and int(c) >= 0]
    if not confs: return None
    return {"mean": round(statistics.mean(confs), 1), "min": min(confs),
            "low_ratio": round(sum(1 for c in confs if c < 60) / len(confs), 3), "words": len(confs)}

def build_docx(paras: list[str], out: Path, lang: str) -> None:
    if resolve("docx_build") == "python-docx":
        from docx import Document
        from docx.shared import Pt
        d = Document(); st = d.styles["Normal"]
        st.font.name = "SimSun" if "chi" in lang else "Times New Roman"; st.font.size = Pt(11)
        for t in paras: d.add_paragraph(t)
        d.save(out); return
    _build_docx_minimal(paras, out)

def _build_docx_minimal(paras: list[str], out: Path) -> None:
    """无 python-docx 时的最小 OOXML 兜底。"""
    import zipfile
    def esc(s): return s.replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")
    body = "".join(f"<w:p><w:r><w:t xml:space='preserve'>{esc(t)}</w:t></w:r></w:p>" for t in paras)
    doc = ("<?xml version='1.0' encoding='UTF-8' standalone='yes'?>"
           "<w:document xmlns:w='http://schemas.openxmlformats.org/wordprocessingml/2006/main'>"
           f"<w:body>{body}</w:body></w:document>")
    ct = ("<?xml version='1.0' encoding='UTF-8' standalone='yes'?>"
          "<Types xmlns='http://schemas.openxmlformats.org/package/2006/content-types'>"
          "<Default Extension='rels' ContentType='application/vnd.openxmlformats-package.relationships+xml'/>"
          "<Override PartName='/word/document.xml' ContentType='application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml'/></Types>")
    rels = ("<?xml version='1.0' encoding='UTF-8' standalone='yes'?>"
            "<Relationships xmlns='http://schemas.openxmlformats.org/package/2006/relationships'>"
            "<Relationship Id='rId1' Type='http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument' Target='word/document.xml'/></Relationships>")
    with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as z:
        z.writestr("[Content_Types].xml", ct); z.writestr("_rels/.rels", rels); z.writestr("word/document.xml", doc)

# ---------- 主流程 ----------
def classify(pdf: Path) -> dict:
    s = doc_stats(pdf)
    text = read_text(pdf); chars = len(text.strip())
    digital = chars >= max(40 * s["pages"], 100)
    kind = "digital" if digital else ("scanned" if s["has_images"] else "empty")
    lang = ("chi_sim+eng" if len(CJK.findall(text)) > 20 else "eng") if digital else "chi_sim+eng"
    return {"pages": s["pages"], "kind": kind, "digital_chars": chars,
            "lang": lang, "columns": s["columns"]}

_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
# 条款编号标记：5. / 5.1. / 2.6.3. / (i) / (a) 等
_CLAUSE_NUM = re.compile(r"^\(?[0-9]+(\.[0-9]+)*\.?\)?$|^\([ivxlcdm]+\)$|^\([a-z]\)$", re.I)


def flatten_layout_tables(out: Path) -> list[str]:
    """pdf2docx 把悬挂缩进的编号条款（5. / 5.1. / 2.6.3.）逐条重建成**布局表格**，
    使合同结构面目全非、且条款文字被拆进多个单元格→破坏 redline 文本匹配。
    本函数把这类"布局表"还原为流式编号段落（保留加粗/下划线标题与缩进层级），
    但**保留真正的数据表**（有底纹的 Key Terms/附件表、含签署标记的签名块）。
    返回处理说明。"""
    from lxml import etree

    def q(t: str) -> str:
        return f"{{{_W_NS}}}{t}"

    import zipfile
    try:
        with zipfile.ZipFile(out) as z:
            names = z.namelist()
            blobs = {n: z.read(n) for n in names}
        root = etree.fromstring(blobs["word/document.xml"])
    except Exception:
        return []

    def cell_text(tc) -> str:
        return "".join(tc.xpath(".//w:t/text()", namespaces={"w": _W_NS})).strip()

    def has_shading(tbl) -> bool:
        for shd in tbl.iter(q("shd")):
            fill = shd.get(q("fill"))
            if fill and fill.upper() not in ("AUTO", "FFFFFF", "FFFFFFFF"):
                return True
        return False

    def is_genuine(tbl) -> bool:
        # 有底纹（表头灰底）的数据表，或签名块，一律保留
        if has_shading(tbl):
            return True
        text = "".join(tbl.xpath(".//w:t/text()", namespaces={"w": _W_NS})).upper()
        if "SIGNATURE" in text or "PRINT NAME" in text or "签名" in text or "盖章" in text:
            return True
        return False

    def first_cell_text(tbl) -> str:
        tr = tbl.find(q("tr"))
        if tr is None:
            return ""
        tc = tr.find(q("tc"))
        return cell_text(tc) if tc is not None else ""

    flattened = 0
    body = root.find(q("body"))
    # 只处理 body 的直接子级表格（顶层条款表；嵌套表不动）
    for tbl in list(body) if body is not None else []:
        if tbl.tag != q("tbl"):
            continue
        if is_genuine(tbl):
            continue
        rows = tbl.findall(q("tr"))
        head = first_cell_text(tbl)
        clause_like = bool(_CLAUSE_NUM.match(head))
        # 布局表判据：首格是条款编号，且（表很小 ≤2 行 —— 条款悬挂缩进的典型形态，
        # 或 某单元格含长正文 ≥80 字 —— 条款正文特征）。
        # 这样首列为数字但**多行短单元格的真数据表**（如 N 行 编号/产品/金额）不会被
        # 误展平，表格结构得以保留——避免把真表当布局表拆掉。
        has_long_cell = any(
            len("".join(tc.xpath(".//w:t/text()", namespaces={"w": _W_NS})).strip()) >= 80
            for tr in rows for tc in tr.findall(q("tc"))
        )
        # 续行碎片表：单行、非编号、整表拼起来是长正文（≥100 字 = 散文而非数据表头）。
        # 真数据表的表头行都很短，不会命中，故安全。
        joined = "".join(tbl.xpath(".//w:t/text()", namespaces={"w": _W_NS})).strip()
        fragment = len(rows) == 1 and not clause_like and len(joined) >= 100
        if not ((clause_like and (len(rows) <= 2 or has_long_cell)) or fragment):
            continue

        new_paras = []
        for tr in rows:
            tcs = tr.findall(q("tc"))
            num = cell_text(tcs[0]) if tcs else ""
            depth = num.count(".") if _CLAUSE_NUM.match(num) else 0
            p = etree.Element(q("p"))
            ppr = etree.SubElement(p, q("pPr"))
            ind = etree.SubElement(ppr, q("ind"))
            ind.set(q("left"), str(min(depth, 4) * 480 + 480))
            ind.set(q("hanging"), "480")
            # 按单元格顺序搬运 runs，单元格间补一个 tab，保留加粗/下划线
            first = True
            for tc in tcs:
                runs = tc.xpath(".//w:r", namespaces={"w": _W_NS})
                runs = [r for r in runs if "".join(r.xpath(".//w:t/text() | .//w:delText/text()", namespaces={"w": _W_NS})).strip()
                        or r.xpath(".//w:tab|.//w:br", namespaces={"w": _W_NS})]
                if not runs:
                    continue
                if not first:
                    tab_r = etree.SubElement(p, q("r"))
                    etree.SubElement(tab_r, q("tab"))
                for r in runs:
                    p.append(etree.fromstring(etree.tostring(r)))
                first = False
            if len(p) > 1:  # 除 pPr 外有内容
                new_paras.append(p)

        idx = list(body).index(tbl)
        for offset, p in enumerate(new_paras):
            body.insert(idx + offset, p)
        body.remove(tbl)
        flattened += 1

    # —— 续行并回父条款 + 删空段（消除多余间距，提升版式还原）——
    XMLSPACE = "{http://www.w3.org/XML/1998/namespace}space"
    clause_pat = re.compile(r"^\(?[0-9]+(\.[0-9]+)*\.|^\([ivxlcdm]+\)|^\([a-z]\)", re.I)
    indent_pat = re.compile(r"^\(?([0-9]+(?:\.[0-9]+)*)\.")
    reflowed = 0

    def _ptext(p) -> str:
        return "".join(p.xpath(".//w:t/text()", namespaces={"w": _W_NS})).strip()

    def _is_center(p) -> bool:
        jc = p.find(f"{q('pPr')}/{q('jc')}")
        return jc is not None and jc.get(q("val")) == "center"

    if body is not None:
        current = None  # 当前"活跃"父条款段落
        for child in list(body):
            if child.tag == q("tbl"):
                current = None
                continue
            if child.tag != q("p"):
                continue
            txt = _ptext(child)
            if not txt:                       # 空段→删（消除多余间距）
                body.remove(child); reflowed += 1; continue
            if _is_center(child):             # 居中标题（Exhibit X / 标题）→独立，不并
                current = None; continue
            if clause_pat.match(txt):         # 新条款编号→成为新的父条款
                current = child; continue
            if current is not None:           # 续行→并回父条款（加一个空格分隔，搬运全部内容）
                sp = etree.SubElement(current, q("r"))
                t = etree.SubElement(sp, q("t")); t.text = " "; t.set(XMLSPACE, "preserve")
                for el in list(child):
                    if el.tag == q("pPr"):
                        continue
                    current.append(etree.fromstring(etree.tostring(el)))
                body.remove(child); reflowed += 1
            else:
                current = None               # 独立非条款段（如引言）→不吸纳后续

        # 正文段落里的定位制表符（pdf2docx 用 tab 还原 PDF 横向间距→"be⟶constrained"间距）
        # 转单空格；每段保留首个 tab（编号→标题的悬挂制表符，维持对齐），其余转空格；
        # 再合并多余空格。不动数据表内的 tab。
        for p in body.findall(q("p")):
            tabs = p.findall(f".//{q('tab')}")
            for i, tabel in enumerate(tabs):
                if i == 0:
                    continue
                par = tabel.getparent()
                if par is None:
                    continue
                sp = etree.Element(q("t")); sp.text = " "; sp.set(XMLSPACE, "preserve")
                par.replace(tabel, sp)
            for t in p.iter(q("t")):
                if t.text and "  " in t.text:
                    t.text = re.sub(r" {2,}", " ", t.text)

        # 统一编号条款缩进（原生 + 展平 + 合并后混排→缩进一致）
        for p in body.findall(q("p")):
            m = indent_pat.match(_ptext(p))
            if not m:
                continue
            depth = m.group(1).count(".")
            ppr = p.find(q("pPr"))
            if ppr is None:
                ppr = etree.Element(q("pPr")); p.insert(0, ppr)
            ind = ppr.find(q("ind"))
            if ind is None:
                ind = etree.SubElement(ppr, q("ind"))
            ind.set(q("left"), str(min(depth, 4) * 480 + 480))
            ind.set(q("hanging"), "480")

    if not (flattened or reflowed):
        return []
    blobs["word/document.xml"] = etree.tostring(
        root, xml_declaration=True, encoding="UTF-8", standalone=True
    )
    with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as z:
        for n in names:
            z.writestr(n, blobs[n])
    return [f"转换后 {flattened} 个布局表已还原为编号段落、{reflowed} 处续行/空段已并回父条款，数据表保留"]


def fit_tables_to_page(out: Path) -> list[str]:
    """pdf2docx 按 PDF 像素宽度给表格绝对列宽，常超出 A4 内容区→边框出界。
    转换后归一：A4 纵向 + 1 英寸页边距（内容区 9026 DXA）；每个超宽表格按
    内容区**等比缩放**列宽与单元格宽、设固定布局、清零过大表缩进。
    仅改版式不动文字，返回处理说明。"""
    from lxml import etree

    def q(t: str) -> str:
        return f"{{{_W_NS}}}{t}"

    def _int(value) -> int:
        # pdf2docx 常把宽度写成浮点字符串（如 "251.9999"），int() 会崩
        try:
            return int(round(float(value)))
        except (TypeError, ValueError):
            return 0

    A4_W, A4_H, MARGIN = 11906, 16838, 1440
    content_w = A4_W - 2 * MARGIN  # 9026
    try:
        import zipfile
        with zipfile.ZipFile(out) as z:
            names = z.namelist()
            blobs = {n: z.read(n) for n in names}
        root = etree.fromstring(blobs["word/document.xml"])
    except Exception:
        return []

    # 1) A4 纵向 + 标准页边距（左右各 1 英寸）
    for sect in root.iter(q("sectPr")):
        pg = sect.find(q("pgSz"))
        if pg is None:
            pg = etree.SubElement(sect, q("pgSz"))
        pg.set(q("w"), str(A4_W)); pg.set(q("h"), str(A4_H)); pg.attrib.pop(q("orient"), None)
        mar = sect.find(q("pgMar"))
        if mar is None:
            mar = etree.SubElement(sect, q("pgMar"))
        for side in ("left", "right"):
            mar.set(q(side), str(MARGIN))
        for side in ("top", "bottom"):
            if not mar.get(q(side)):
                mar.set(q(side), str(MARGIN))

    # 2) 表格等比缩放到内容区
    scaled = 0
    for tbl in root.iter(q("tbl")):
        grid = tbl.find(q("tblGrid"))
        if grid is None:
            continue
        cols = grid.findall(q("gridCol"))
        widths = [_int(c.get(q("w"), "0")) for c in cols]
        total = sum(widths)
        tbl_pr = tbl.find(q("tblPr"))
        # 清零过大表缩进（也会把表推出右边界）
        if tbl_pr is not None:
            ind = tbl_pr.find(q("tblInd"))
            if ind is not None and _int(ind.get(q("w"), "0")) > 0:
                ind.set(q("w"), "0")
        if total <= content_w or total == 0:
            continue
        factor = content_w / total
        new_w = [max(1, int(w * factor)) for w in widths]
        new_w[-1] += content_w - sum(new_w)  # 抹平取整误差，确保总和=内容区
        for c, nw in zip(cols, new_w):
            c.set(q("w"), str(nw))
        # 单元格宽同比缩放（多列合并单元格的 tcW≈所跨列之和，同因子保持一致）
        for tc in tbl.iter(q("tc")):
            tc_pr = tc.find(q("tcPr"))
            if tc_pr is None:
                continue
            tcw = tc_pr.find(q("tcW"))
            if tcw is not None and (tcw.get(q("type")) in (None, "dxa")):
                w = _int(tcw.get(q("w"), "0"))
                if w > 0:
                    tcw.set(q("w"), str(max(1, int(w * factor))))
        # 固定布局 + 表宽=内容区
        if tbl_pr is None:
            tbl_pr = etree.Element(q("tblPr")); tbl.insert(0, tbl_pr)
        tw = tbl_pr.find(q("tblW"))
        if tw is None:
            tw = etree.SubElement(tbl_pr, q("tblW"))
        tw.set(q("type"), "dxa"); tw.set(q("w"), str(content_w))
        layout = tbl_pr.find(q("tblLayout"))
        if layout is None:
            layout = etree.SubElement(tbl_pr, q("tblLayout"))
        layout.set(q("type"), "fixed")
        scaled += 1

    if not scaled:
        return []
    blobs["word/document.xml"] = etree.tostring(
        root, xml_declaration=True, encoding="UTF-8", standalone=True
    )
    import zipfile
    with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as z:
        for n in names:
            z.writestr(n, blobs[n])
    return [f"转换后 {scaled} 个超宽表格已按 A4 内容区等比缩放，表格边框收进页内"]


def convert(pdf: Path, out: Path, info: dict) -> dict:
    kind, lang, cols = info["kind"], info["lang"], info["columns"]
    notes = []
    if kind == "digital" and cols == 1 and resolve("pdf2docx"):
        from pdf2docx import Converter
        c = Converter(str(pdf)); c.convert(str(out)); c.close()
        method = "pdf2docx"
        notes += flatten_layout_tables(out)   # 先还原布局表为编号段落
        notes += fit_tables_to_page(out)      # 再把剩余真数据表收进内容区
    elif kind == "digital":
        build_docx(reading_order_paras(pdf, cols == 2), out, lang)
        method = f"reflow({resolve('pdf_read')})"
        if cols == 2: notes.append("多栏版式已按左→右阅读序重排，请核对跨栏段落顺序")
    elif kind == "scanned":
        build_docx(ocr_to_paras(pdf, lang, cols == 2), out, lang)
        method = f"ocr({resolve('ocr')})+reflow"
        notes.append("扫描件经 OCR 识别，请对照原件核对关键数字/金额/期限")
    else:
        raise SystemExit("空白或无法识别的 PDF（无文字层且无图像）")
    return {"method": method, "notes": notes,
            "providers": {c: resolve(c) for c in REGISTRY}}

def verify(pdf: Path, out: Path, info: dict) -> dict:
    if resolve("docx_build") == "python-docx":
        from docx import Document
        d = Document(out)
        # 表格密集型合同的正文多在单元格里；`d.paragraphs` 不含表格单元格文字，
        # 只数它会把表格合同误判为"空文档"→错误熔断。这里把表格文字一并计入。
        cell_texts = []
        for tbl in d.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        cell_texts.append(cell.text)
        text = "\n".join([p.text for p in d.paragraphs] + cell_texts)
        paras = sum(1 for p in d.paragraphs if p.text.strip()) + len(cell_texts)
    else:
        import zipfile, re as _re
        x = zipfile.ZipFile(out).read("word/document.xml").decode()
        text = " ".join(_re.findall(r"<w:t[^>]*>([^<]*)</w:t>", x)); paras = text.count("  ") + 1
    out_chars = len(text.strip())
    tokens = [t for t in KEY_TOKENS if t.lower() in text.lower()]
    conf = ocr_confidence(pdf, info["lang"]) if info["kind"] == "scanned" else None
    coverage = round(out_chars / max(info["digital_chars"], 1), 2) if info["kind"] == "digital" else None

    tier, reasons = "high", []
    if paras == 0 or out_chars < 50:
        tier, reasons = "low", ["转换后无有效文字（空文档）"]
    elif info["kind"] == "scanned":
        if conf is None:
            tier = "medium"; reasons.append("OCR 引擎不提供置信度，按中档处理：关键数字/金额/期限需人工核对")
        elif conf["mean"] < 75 or conf["low_ratio"] > 0.15:
            tier = "low"; reasons.append(f"OCR 置信偏低 mean={conf['mean']} low_ratio={conf['low_ratio']}")
        else:
            tier = "medium"; reasons.append("扫描件 OCR 结果建议抽查核对")
    elif info["kind"] == "digital":
        if coverage is not None and coverage < 0.6:
            tier = "low"; reasons.append(f"文字覆盖率过低 {coverage}")
        elif info["columns"] == 2 or (coverage is not None and coverage < 0.92):
            tier = "medium"; reasons.append("多栏重排或覆盖率不足，需核对段落顺序与完整性")
    if not tokens:
        tier = "low"; reasons.append("未检出任何合同关键词，疑似转换错误")
    return {"out_chars": out_chars, "paragraphs": paras, "coverage": coverage,
            "key_tokens": tokens, "ocr_confidence": conf, "tier": tier, "reasons": reasons}

def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("pdf", nargs="?", type=Path)
    ap.add_argument("--out", type=Path)
    ap.add_argument("--lang", default="auto")
    ap.add_argument("--manifest", type=Path)
    ap.add_argument("--prefer", action="append", default=[], help="cap=provider，强制某能力的提供方")
    ap.add_argument("--probe", action="store_true", help="只打印探测到的能力提供方")
    args = ap.parse_args()
    for kv in args.prefer:
        if "=" in kv: c, v = kv.split("=", 1); _PREFER[c.strip()] = v.strip()

    if args.probe:
        print(json.dumps(probe(), ensure_ascii=False, indent=2)); return 0
    if not args.pdf or not args.out:
        ap.error("需要 <pdf> 与 --out（或用 --probe）")
    if not args.pdf.exists():
        raise SystemExit(f"missing input: {args.pdf}")

    # 必要能力检查（不硬编码工具，只要求“某个”提供方存在）
    for cap in ("pdf_read", "docx_build"):
        if resolve(cap) is None:
            print(f"[缺能力] 无可用 {cap} 提供方，无法入库。请在环境中安装/接入相应工具。"); return 4

    info = classify(args.pdf)
    if args.lang != "auto": info["lang"] = args.lang
    if info["kind"] == "scanned" and resolve("ocr") is None:
        print("[缺能力·熔断] 输入为扫描件但环境无可用 OCR 提供方（ocrmypdf/tesseract/easyocr 均不可用）。\n"
              "请安装/接入 OCR 工具，或改提供数字版 PDF / Word。"); return 4

    conv = convert(args.pdf, args.out, info)
    v = verify(args.pdf, args.out, info)
    manifest = {"input": str(args.pdf), "output": str(args.out), "classification": info,
                "conversion": conv, "verification": v}
    if args.manifest:
        args.manifest.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")

    print(f"分类: {info['kind']} | 语言: {info['lang']} | 栏数: {info['columns']} | 方法: {conv['method']}")
    print(f"提供方: " + ", ".join(f"{k}={v2}" for k, v2 in conv["providers"].items()))
    print(f"文字: {v['out_chars']} 字符 / {v['paragraphs']} 段 | 覆盖率: {v['coverage']} | 关键词: {len(v['key_tokens'])}")
    if v["ocr_confidence"]: print(f"OCR 置信: {v['ocr_confidence']}")
    print(f"准确性档位: {v['tier'].upper()} | 理由: {'; '.join(v['reasons']) or '高保真'}")
    for n in conv["notes"]: print("  注: " + n)
    if v["tier"] == "low":
        print("\n[熔断] 文字准确性不足，禁止直接审查。请提供更清晰的 PDF/原始 Word，或人工核对后再继续。"); return 3
    if v["tier"] == "medium":
        print("\n[可审查·待核对] 已生成标准 Word；报告须声明基于转换工作副本，并把待核对项交用户确认。")
    return 0

if __name__ == "__main__":
    sys.exit(main())
