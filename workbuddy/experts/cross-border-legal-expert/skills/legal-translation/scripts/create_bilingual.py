# -*- coding: utf-8 -*-
# © 深圳市法大大网络科技有限公司 版权所有
"""
生成中英双语对照法律文书 Word 文档。
用法:
  python create_bilingual.py <JSON路径> <输出路径>
  python create_bilingual.py --stdin <输出路径>     # 从 stdin 读取 JSON
依赖: pip install python-docx

输出格式：两列表格，左列中文，右列英文。
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')
import json
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("错误: 需要安装 python-docx (pip install python-docx)", file=sys.stderr)
    sys.exit(1)


# ── 工具函数 ──────────────────────────────────────────────────────────

def set_font(run, name_cn="宋体", name_en="Times New Roman", size=11, bold=False, color=None):
    run.font.size = Pt(size)
    run.bold = bold
    run.font.name = name_en
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), name_cn)
    if color:
        run.font.color.rgb = color


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_para(doc, text="", alignment=None, space_before=0, space_after=4,
             font_size=11, bold=False, font_cn="宋体", color=None):
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        set_font(run, name_cn=font_cn, size=font_size, bold=bold, color=color)
    return p


def write_cell_content(cell, text, font_cn="宋体", font_size=10.5, bold=False):
    """向单元格写入多行文本，保留段落格式。"""
    # 清空默认段落
    cell.text = ""
    lines = text.split("\n")
    for i, line in enumerate(lines):
        if i == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = Pt(16)

        # 检查是否为标题行（以"第"开头或"Article"开头）
        is_heading = (line.strip().startswith("第") and "条" in line[:8]) or \
                     line.strip().startswith("Article ")

        run = p.add_run(line)
        set_font(run, name_cn=font_cn, size=font_size,
                 bold=bold or is_heading)


# ── 主生成函数 ────────────────────────────────────────────────────────

def generate(data, output_path):
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    title_cn = data.get("title_cn", "")
    title_en = data.get("title_en", "")
    doc_type = data.get("doc_type", "")
    translation_date = data.get("translation_date", "")
    glossary = data.get("glossary_used", [])
    sections = data.get("sections", [])
    consistency = data.get("consistency_check", {})

    # ── 中文标题 ─────────────────────────────────────────────────────
    add_para(doc, title_cn,
             alignment=WD_ALIGN_PARAGRAPH.CENTER,
             font_size=16, bold=True, font_cn="黑体",
             space_after=4)

    # ── 英文标题 ─────────────────────────────────────────────────────
    add_para(doc, title_en,
             alignment=WD_ALIGN_PARAGRAPH.CENTER,
             font_size=14, bold=True, font_cn="Times New Roman",
             space_after=8)

    # ── 翻译信息 ─────────────────────────────────────────────────────
    if translation_date:
        info_text = "Translation Date: " + translation_date
        add_para(doc, info_text,
                 alignment=WD_ALIGN_PARAGRAPH.CENTER,
                 font_size=9, color=RGBColor(0x99, 0x99, 0x99),
                 space_after=12)

    # ── 双语对照表 ───────────────────────────────────────────────────
    if sections:
        # 表头
        table = doc.add_table(rows=1 + len(sections), cols=2)
        table.style = 'Table Grid'

        # 设置列宽
        for row in table.rows:
            row.cells[0].width = Cm(8.0)
            row.cells[1].width = Cm(8.5)

        # 表头行
        hdr_cn = table.rows[0].cells[0]
        hdr_en = table.rows[0].cells[1]
        hdr_cn.text = "中文原文"
        hdr_en.text = "English Translation"
        set_cell_bg(hdr_cn, "2F5496")
        set_cell_bg(hdr_en, "2F5496")
        for cell in [hdr_cn, hdr_en]:
            if cell.paragraphs[0].runs:
                set_font(cell.paragraphs[0].runs[0], size=11, bold=True,
                         color=RGBColor(0xFF, 0xFF, 0xFF), name_cn="黑体")
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 填充对照内容
        for i, sec in enumerate(sections):
            cn_text = sec.get("cn", "")
            en_text = sec.get("en", "")
            row = table.rows[i + 1]

            write_cell_content(row.cells[0], cn_text, font_cn="宋体", font_size=10.5)
            write_cell_content(row.cells[1], en_text, font_cn="Times New Roman", font_size=10.5)

            # 隔行浅色背景
            if i % 2 == 1:
                set_cell_bg(row.cells[0], "F5F5F5")
                set_cell_bg(row.cells[1], "F5F5F5")

    # ── 术语表附录 ───────────────────────────────────────────────────
    if glossary:
        doc.add_paragraph()
        add_para(doc, "附录：关键术语对照表 / Appendix: Key Terminology",
                 font_size=12, bold=True, font_cn="黑体",
                 space_before=16, space_after=8)

        g_table = doc.add_table(rows=1 + len(glossary), cols=2)
        g_table.style = 'Table Grid'

        g_hdr_cn = g_table.rows[0].cells[0]
        g_hdr_en = g_table.rows[0].cells[1]
        g_hdr_cn.text = "中文术语"
        g_hdr_en.text = "English Term"
        set_cell_bg(g_hdr_cn, "4472C4")
        set_cell_bg(g_hdr_en, "4472C4")
        for cell in [g_hdr_cn, g_hdr_en]:
            if cell.paragraphs[0].runs:
                set_font(cell.paragraphs[0].runs[0], size=10, bold=True,
                         color=RGBColor(0xFF, 0xFF, 0xFF), name_cn="黑体")
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        for j, term in enumerate(glossary):
            row = g_table.rows[j + 1]
            row.cells[0].text = term.get("cn", "")
            row.cells[1].text = term.get("en", "")
            if row.cells[0].paragraphs[0].runs:
                set_font(row.cells[0].paragraphs[0].runs[0], size=10, name_cn="宋体")
            if row.cells[1].paragraphs[0].runs:
                set_font(row.cells[1].paragraphs[0].runs[0], size=10)

    # ── 一致性检查声明 ───────────────────────────────────────────────
    doc.add_paragraph()
    check_lines = []
    if consistency.get("terminology_pass"):
        check_lines.append("Terminology Consistency: PASSED")
    if consistency.get("numbering_pass"):
        check_lines.append("Numbering Alignment: PASSED")
    if consistency.get("figures_pass"):
        check_lines.append("Figures Accuracy: PASSED")

    corrections = consistency.get("corrections", [])
    if corrections:
        check_lines.append("Corrections made: " + str(len(corrections)))
        for c in corrections:
            check_lines.append("  - " + c)

    if check_lines:
        add_para(doc, "Consistency Check Report",
                 font_size=10, bold=True, font_cn="黑体",
                 space_before=8, space_after=4)
        for line in check_lines:
            prefix = "pass" in line.lower()
            add_para(doc, line, font_size=9,
                     color=RGBColor(0x00, 0x80, 0x00) if prefix else RGBColor(0x66, 0x66, 0x66))

    # ── 免责声明 ─────────────────────────────────────────────────────
    add_para(doc, "", space_after=8)
    disclaimer = ("This bilingual document is prepared for reference purposes only. "
                   "In case of any discrepancy between the Chinese and English versions, "
                   "the Chinese version shall prevail.")
    disclaimer_cn = "本双语文本仅供参考。如中英文版本存在不一致之处，以中文版本为准。"
    add_para(doc, disclaimer_cn, font_size=9,
             color=RGBColor(0x99, 0x99, 0x99))
    add_para(doc, disclaimer, font_size=9,
             color=RGBColor(0x99, 0x99, 0x99))

    # ── 保存 ─────────────────────────────────────────────────────────
    doc.save(output_path)

    print("已生成：" + output_path)
    print("中文标题：" + title_cn)
    print("英文标题：" + title_en)
    print("对照段落数：" + str(len(sections)))
    print("术语表条目：" + str(len(glossary)))


def main():
    if len(sys.argv) >= 2 and sys.argv[1] == "--stdin":
        if len(sys.argv) < 3:
            print("用法: python create_bilingual.py --stdin <输出路径>", file=sys.stderr)
            sys.exit(1)
        output_path = str(Path(sys.argv[2]).resolve())
        text = sys.stdin.read()
    elif len(sys.argv) >= 3:
        json_path = Path(sys.argv[1]).resolve()
        output_path = str(Path(sys.argv[2]).resolve())
        if not json_path.exists():
            print("错误: JSON文件不存在 " + str(json_path), file=sys.stderr)
            sys.exit(1)
        with open(json_path, "r", encoding="utf-8") as f:
            text = f.read()
    else:
        print("用法: python create_bilingual.py <JSON路径> <输出路径>", file=sys.stderr)
        print("      python create_bilingual.py --stdin <输出路径>", file=sys.stderr)
        sys.exit(1)

    text = text.replace('\u201c', '\u2018').replace('\u201d', '\u2019')
    data = json.loads(text)
    generate(data, output_path)


if __name__ == "__main__":
    main()
