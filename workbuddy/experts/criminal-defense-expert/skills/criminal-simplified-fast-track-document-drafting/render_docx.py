#!/usr/bin/env python3
"""Render a selected Markdown template or prepared Markdown into a real DOCX."""

from __future__ import annotations

import argparse
import json
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


class RenderError(ValueError):
    """Raised when rendering cannot safely continue."""


def _resolved(path: str | Path) -> Path:
    return Path(path).expanduser().resolve()


def _inside(root: Path, path: Path) -> bool:
    return path == root or root in path.parents


def _load_data(data_json: str | None, data_file: str | None) -> dict[str, object]:
    if data_json and data_file:
        raise RenderError("data-json和data-file只能选择一个")
    if data_file:
        value = json.loads(_resolved(data_file).read_text(encoding="utf-8"))
    elif data_json:
        value = json.loads(data_json)
    else:
        value = {}
    if not isinstance(value, dict):
        raise RenderError("填充数据必须是JSON对象")
    return value


def _replace_fields(text: str, data: dict[str, object]) -> str:
    rendered = text
    for key, value in data.items():
        replacement = str(value)
        rendered = rendered.replace("{{" + key + "}}", replacement)
        rendered = rendered.replace("[" + key + "]", replacement)
    return rendered


def load_markdown(
    *,
    skill_root: str | Path,
    template: str | None,
    content_file: str | None,
    data: dict[str, object],
) -> tuple[str, str]:
    """Load exactly one source and return rendered Markdown and template id."""

    if bool(template) == bool(content_file):
        raise RenderError("template和content-file必须且只能提供一个")
    root = _resolved(skill_root)
    if template:
        templates_root = root / "templates"
        source = _resolved(template)
        if not _inside(templates_root, source) or not source.is_file():
            raise RenderError("模板必须是当前Skill templates目录中的现有文件")
        template_id = str(source.relative_to(templates_root))
    else:
        source = _resolved(content_file)
        if not source.is_file():
            raise RenderError("内容文件不存在")
        template_id = "prepared-markdown"
    return _replace_fields(source.read_text(encoding="utf-8"), data), template_id


def _set_run_font(run, *, size: int = 12, bold: bool = False, font_name: str = "宋体") -> None:
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def _set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def _add_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    width = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=width)
    table.style = "Table Grid"
    for row_index, row in enumerate(rows):
        for column_index in range(width):
            value = row[column_index] if column_index < len(row) else ""
            cell = table.cell(row_index, column_index)
            cell.text = value.strip()
            if row_index == 0:
                _set_cell_shading(cell, "E7E6E6")
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    _set_run_font(run, size=10, bold=row_index == 0)


def render_markdown(
    markdown: str,
    output_path: str | Path,
    *,
    title: str | None = None,
    format_profile: str = "I-Practical",
) -> Path:
    """Render a practical Markdown subset used by the criminal templates."""

    if not markdown.strip():
        raise RenderError("待渲染内容为空")
    output = _resolved(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    strict = format_profile in {"F-Strict", "II-Formal"}
    body_size = 16 if strict else 12
    heading_sizes = {1: 22 if strict else 18, 2: 16 if strict else 15, 3: 16 if strict else 13, 4: body_size}
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.6 if strict else 2.0)
    section.bottom_margin = Cm(2.6 if strict else 2.0)
    section.left_margin = Cm(2.8 if strict else 2.5)
    section.right_margin = Cm(2.6 if strict else 2.5)
    normal = document.styles["Normal"]
    body_font = "仿宋_GB2312" if strict else "宋体"
    normal.font.name = body_font
    normal.font.size = Pt(body_size)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), body_font)
    normal.paragraph_format.line_spacing = 1.5 if strict else 1.4
    first_title_written = False
    lines = markdown.splitlines()
    index = 0
    while index < len(lines):
        raw = lines[index].rstrip()
        if raw.startswith("|") and raw.endswith("|"):
            table_lines: list[str] = []
            while index < len(lines) and lines[index].strip().startswith("|") and lines[index].strip().endswith("|"):
                table_lines.append(lines[index].strip())
                index += 1
            rows = [[cell.strip() for cell in line.strip("|").split("|")] for line in table_lines]
            if len(rows) > 1 and all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in rows[1]):
                rows.pop(1)
            _add_table(document, rows)
            continue
        if not raw.strip():
            document.add_paragraph()
            index += 1
            continue
        heading = re.match(r"^(#{1,4})\s+(.+)$", raw)
        if heading:
            level = len(heading.group(1))
            paragraph = document.add_paragraph()
            run = paragraph.add_run(heading.group(2).strip())
            heading_font = "方正小标宋简体" if strict and level == 1 else ("黑体" if strict else "宋体")
            _set_run_font(run, size=heading_sizes[level], bold=True, font_name=heading_font)
            if level == 1 and not first_title_written:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                first_title_written = True
            index += 1
            continue
        list_match = re.match(r"^\s*(?:[-*]|\d+[.、])\s+(.+)$", raw)
        paragraph = document.add_paragraph(style="List Bullet" if list_match else None)
        text = list_match.group(1) if list_match else raw
        run = paragraph.add_run(text.strip())
        _set_run_font(run, size=body_size, font_name=body_font)
        index += 1

    if title:
        document.core_properties.title = title
    document.save(output)
    return output


def render(
    *,
    skill_root: str | Path,
    matter_root: str | Path,
    output_path: str | Path,
    template: str | None = None,
    content_file: str | None = None,
    data_json: str | None = None,
    data_file: str | None = None,
    doc_type: str = "",
) -> dict:
    matter = _resolved(matter_root)
    output = _resolved(output_path)
    if not matter.is_dir():
        raise RenderError("案件目录不存在")
    if not _inside(matter, output):
        raise RenderError("输出路径必须属于当前案件目录")
    data = _load_data(data_json, data_file)
    markdown, template_id = load_markdown(
        skill_root=skill_root,
        template=template,
        content_file=content_file,
        data=data,
    )
    manifest_path = _resolved(skill_root) / "meta" / "manifest.json"
    format_profile = "I-Practical"
    if manifest_path.is_file():
        manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
        format_profile = str(manifest.get("format_seriousness", format_profile))
    path = render_markdown(
        markdown,
        output,
        title=doc_type or template_id,
        format_profile=format_profile,
    )
    return {
        "outcome": "PASS",
        "document_path": str(path),
        "template_id": template_id,
        "doc_type": doc_type,
        "output_format": "docx",
        "format_profile": format_profile,
        "rendered": True,
    }


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description="将刑事Markdown模板渲染为DOCX")
    parser.add_argument("--skill-root", default=str(Path(__file__).parents[1]))
    parser.add_argument("--matter-root", required=True)
    parser.add_argument("--output", required=True)
    source = parser.add_mutually_exclusive_group(required=True)
    source.add_argument("--template")
    source.add_argument("--content-file")
    data = parser.add_mutually_exclusive_group()
    data.add_argument("--data-json")
    data.add_argument("--data-file")
    parser.add_argument("--doc-type", default="")
    args = parser.parse_args(argv)
    try:
        result = render(
            skill_root=args.skill_root,
            matter_root=args.matter_root,
            output_path=args.output,
            template=args.template,
            content_file=args.content_file,
            data_json=args.data_json,
            data_file=args.data_file,
            doc_type=args.doc_type,
        )
    except (OSError, ValueError, json.JSONDecodeError) as error:
        result = {"outcome": "BLOCKED", "rendered": False, "reason": str(error)}
    print(json.dumps(result, ensure_ascii=False))
    return 0 if result["outcome"] == "PASS" else 2


if __name__ == "__main__":
    raise SystemExit(main())
