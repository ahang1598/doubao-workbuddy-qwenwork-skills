#!/usr/bin/env python3
"""Deterministic, loop-safe criminal document delivery checks."""

from __future__ import annotations

import json
import hashlib
import re
import sys
from html.parser import HTMLParser
from pathlib import Path
from typing import Any, Dict, Iterable, List, Mapping

from docx import Document


VALID_SCENES = {
    "agency_submission",
    "family_communication",
    "lawyer_working",
}
VALID_MODES = {"draft", "formal"}
AUTHORITY_MARKERS = {
    "police": ("公安", "侦查机关"),
    "procuratorate": ("检察",),
    "court": ("法院", "人民法院", "中院", "高院", "基层法院"),
}
DOC_TYPE_AUTHORITIES = {
    "non_arrest_opinion": {"procuratorate"},
    "juvenile_non_arrest_opinion": {"procuratorate"},
    "custody_review_application": {"procuratorate"},
    "custody_assessment_police": {"police"},
    "sentencing_negotiation_opinion": {"procuratorate"},
    "sentencing_recommendation_negotiation_opinion": {"procuratorate"},
    "conditional_non_prosecution_application": {"procuratorate"},
    "pretrial_conference_application": {"court"},
    "recusal_application": {"court"},
    "jurisdiction_objection": {"court"},
    "closed_hearing_application": {"court"},
    "defendant_pretrial_participation_application": {"court"},
    "ordinary_procedure_recommendation": {"court"},
    "criminal_appeal": {"court"},
    "second_instance_defense_opinion": {"court"},
    "second_instance_open_hearing_application": {"court"},
    "second_instance_supplemental_opinion": {"court"},
    "first_instance_defense_speech": {"court"},
    "post_trial_written_opinion": {"court"},
    "criminal_retrial_application": {"court"},
    "compulsory_medical_representation_opinion": {"court"},
    "compulsory_medical_release_application": {"court"},
    "illicit_proceeds_participation_opinion": {"court"},
    "property_ownership_objection": {"court"},
}
PLACEHOLDER_RE = re.compile(r"\[(?:待补|待核实|待确认)[^\]]*\]|【(?:待补|待核实|待确认)[^】]*】")
SIGNATURE_RE = re.compile(r"(申请人|辩护人|上诉人|代理人)[:：]\s*(?:\[待|【待|$)")
DATE_RE = re.compile(r"(日期|年\s*月\s*日)[:：]?\s*(?:\[待|【待|$)")


class _HtmlTextParser(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.values: List[str] = []

    def handle_data(self, data: str) -> None:
        if data.strip():
            self.values.append(data.strip())


def _resolved(path: str | Path) -> Path:
    return Path(path).expanduser().resolve()


def _inside(root: Path, path: Path) -> bool:
    return path == root or root in path.parents


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as source:
        for block in iter(lambda: source.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def _read_artifact(path: Path) -> tuple[str, str]:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        document = Document(path)
        values = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                values.extend(cell.text for cell in row.cells)
        return "\n".join(value for value in values if value.strip()), "docx"
    if suffix in {".html", ".htm"}:
        parser = _HtmlTextParser()
        parser.feed(path.read_text(encoding="utf-8"))
        return "\n".join(parser.values), "html"
    if suffix in {".md", ".txt"}:
        return path.read_text(encoding="utf-8"), suffix.lstrip(".")
    raise ValueError(f"不支持的实际制品格式: {suffix or '无扩展名'}")


def _artifact_from_payload(payload: Mapping[str, Any]) -> tuple[Dict[str, Any] | None, Dict[str, str] | None]:
    raw_path = str(payload.get("document_path", "")).strip()
    if not raw_path:
        return None, None
    raw_root = str(payload.get("matter_root", "")).strip()
    if not raw_root:
        return None, {"code": "MATTER_ROOT_MISSING", "message": "提供实际制品时必须提供当前案件根目录"}
    root = _resolved(raw_root)
    path = _resolved(raw_path)
    if not root.is_dir():
        return None, {"code": "MATTER_ROOT_INVALID", "message": "当前案件根目录不存在"}
    if not _inside(root, path):
        return None, {"code": "CROSS_MATTER_PATH", "message": "实际制品路径不属于当前案件"}
    if not path.is_file():
        return None, {"code": "DOCUMENT_MISSING", "message": "实际制品不存在"}
    try:
        text, artifact_format = _read_artifact(path)
    except Exception as error:
        return None, {"code": "DOCUMENT_DAMAGED", "message": f"实际制品无法打开或解析: {error}"}
    return {
        "path": str(path),
        "text": text.strip(),
        "format": artifact_format,
        "sha256": _sha256(path),
    }, None


def _add_delivery_contract(
    result: Dict[str, Any],
    payload: Mapping[str, Any],
    artifact: Mapping[str, Any] | None,
) -> Dict[str, Any]:
    scene = str(payload.get("output_scene", "")).strip()
    mode = str(payload.get("submission_mode", "draft")).strip() or "draft"
    receipt = payload.get("render_receipt")
    rendered = isinstance(receipt, Mapping) and receipt.get("rendered") is True
    template_id = str(receipt.get("template_id", "")).strip() if isinstance(receipt, Mapping) else ""
    submission_ready = bool(
        scene == "agency_submission"
        and mode == "formal"
        and result["status"] == "PASS"
        and artifact
        and rendered
        and template_id
    )
    if scene == "agency_submission":
        deliverable_as = "agency_submission" if submission_ready else "draft"
    elif scene in VALID_SCENES and result["deliverable"]:
        deliverable_as = scene
    else:
        deliverable_as = "not_deliverable"
    result.update(
        {
            "matter_id": str(payload.get("matter_id", "")).strip(),
            "document_path": artifact.get("path", "") if artifact else str(payload.get("document_path", "")).strip(),
            "artifact_checked": bool(artifact),
            "artifact_format": artifact.get("format", "") if artifact else "text",
            "artifact_sha256": artifact.get("sha256", "") if artifact else "",
            "template_id": template_id,
            "render_receipt_verified": bool(artifact and rendered and template_id),
            "deliverable_as": deliverable_as,
            "submission_ready": submission_ready,
        }
    )
    return result


def _result(
    status: str,
    *,
    findings: Iterable[Mapping[str, str]] = (),
    warnings: Iterable[str] = (),
    missing_fields: Iterable[str] = (),
    previous_codes: Iterable[str] = (),
) -> Dict[str, Any]:
    finding_list = [dict(item) for item in findings]
    codes = [item["code"] for item in finding_list]
    previous = set(previous_codes)
    fused = status == "BLOCKED" and bool(previous.intersection(codes))
    retry_allowed = status == "BLOCKED" and not fused
    return {
        "status": status,
        "deliverable": status in {"PASS", "PASS_WITH_WARNINGS"},
        "finding_codes": codes,
        "findings": finding_list,
        "warnings": list(warnings),
        "missing_fields": list(missing_fields),
        "retry_allowed": retry_allowed,
        "max_correction_attempts": 1,
        "fused": fused,
    }


def _authority_mismatch(doc_type: str, authority: str) -> str | None:
    if not authority:
        return None
    actual = next(
        (kind for kind, markers in AUTHORITY_MARKERS.items() if any(marker in authority for marker in markers)),
        None,
    )
    allowed = DOC_TYPE_AUTHORITIES.get(doc_type)
    if actual and allowed and actual not in allowed:
        names = {"police": "公安机关", "procuratorate": "人民检察院", "court": "人民法院"}
        expected = "或".join(names[item] for item in sorted(allowed))
        return f"{doc_type}应致送{expected}，当前机关为{authority}"
    return None


def _has_material_conflict(conflicts: Any) -> bool:
    if not isinstance(conflicts, list):
        return False
    protected = {"subject", "party", "case_no", "amount", "主体", "案号", "金额"}
    for item in conflicts:
        if not isinstance(item, Mapping) or item.get("resolved") is True:
            continue
        category = str(item.get("category", ""))
        values = item.get("values", [])
        distinct = {str(value).strip() for value in values if str(value).strip()}
        if category in protected and len(distinct) > 1:
            return True
    return False


def check_document(payload: Mapping[str, Any]) -> Dict[str, Any]:
    """Return a deterministic delivery status without automatic retries."""
    missing = [key for key in ("doc_type", "output_scene") if key not in payload]
    if "document" not in payload and "document_path" not in payload:
        missing.append("document_path_or_document")
    previous = payload.get("previous_finding_codes") or []
    if missing:
        return _add_delivery_contract(_result(
            "NEEDS_INPUT",
            missing_fields=missing,
            warnings=["缺少最小校验输入；保留已有草稿并一次性补充。"],
            previous_codes=previous,
        ), payload, None)

    doc_type = str(payload.get("doc_type", "")).strip()
    scene = str(payload.get("output_scene", "")).strip()
    document_value = payload.get("document")
    mode = str(payload.get("submission_mode", "draft")).strip() or "draft"

    input_missing: List[str] = []
    if not doc_type:
        input_missing.append("doc_type")
    if not scene:
        input_missing.append("output_scene")
    if input_missing:
        return _add_delivery_contract(_result(
            "NEEDS_INPUT",
            missing_fields=input_missing,
            warnings=["最小字段为空；保留草稿，不自动重试。"],
            previous_codes=previous,
        ), payload, None)

    if scene not in VALID_SCENES:
        return _add_delivery_contract(_result(
            "NEEDS_INPUT",
            missing_fields=["output_scene"],
            warnings=[f"output_scene 必须为三类场景之一，当前为 {scene!r}。"],
            previous_codes=previous,
        ), payload, None)
    if mode not in VALID_MODES:
        return _add_delivery_contract(_result(
            "NEEDS_INPUT",
            missing_fields=["submission_mode"],
            warnings=[f"submission_mode 仅支持 draft/formal，当前为 {mode!r}。"],
            previous_codes=previous,
        ), payload, None)

    findings: List[Dict[str, str]] = []
    artifact, artifact_finding = _artifact_from_payload(payload)
    if artifact_finding:
        findings.append(artifact_finding)
        document = ""
    elif artifact:
        document = str(artifact["text"])
    elif payload.get("damaged") is True:
        findings.append({"code": "DOCUMENT_DAMAGED", "message": "文件被标记为损坏或无法读取"})
        document = ""
    elif document_value is None:
        document = ""
    else:
        document = str(document_value).strip()

    if not document:
        findings.append({"code": "DOCUMENT_EMPTY", "message": "文书内容为空"})

    if mode == "formal" and document:
        nonempty_lines = [line.strip() for line in document.splitlines() if line.strip()]
        body_text = "".join(nonempty_lines[1:]) if len(nonempty_lines) > 1 else ""
        if not body_text or len(body_text) < 12:
            findings.append({"code": "FORMAL_TITLE_ONLY", "message": "正式稿只有标题或没有实质正文"})

    authority = str(payload.get("handling_authority", "")).strip()
    mismatch = _authority_mismatch(doc_type, authority)
    if mismatch:
        findings.append({"code": "AUTHORITY_DOCUMENT_MISMATCH", "message": mismatch})

    if scene == "family_communication" and document:
        sensitive_hits = [
            marker for marker in ("同案犯供述", "卷宗摘录", "供述称", "讯问笔录原文")
            if marker in document
        ]
        if len(sensitive_hits) >= 2 or any(marker in document for marker in ("同案犯供述", "卷宗摘录")):
            findings.append({
                "code": "FAMILY_SENSITIVE_DISCLOSURE",
                "message": "家属稿包含高置信度供述或卷宗敏感披露",
            })

    if _has_material_conflict(payload.get("conflicts")):
        findings.append({
            "code": "MATERIAL_VALUE_CONFLICT",
            "message": "主体、案号或金额存在无法判断真伪的确定值冲突",
        })

    if findings:
        unique = {item["code"]: item for item in findings}
        return _add_delivery_contract(_result(
            "BLOCKED",
            findings=unique.values(),
            previous_codes=previous,
        ), payload, artifact)

    warnings: List[str] = []
    if (
        not str(payload.get("case_stage", "")).strip()
        and doc_type not in DOC_TYPE_AUTHORITIES
        and not authority
    ):
        warnings.append("case_stage 未显式提供；可由文种或机关推导时无需补填。")
    if PLACEHOLDER_RE.search(document):
        warnings.append("文书含待补/待核实占位；草稿可交付，正式提交前补齐。")
    if mode == "formal" and SIGNATURE_RE.search(document):
        warnings.append("签署信息待补。")
    if mode == "formal" and DATE_RE.search(document):
        warnings.append("日期待补。")
    if payload.get("legal_authority_verified") is False:
        warnings.append("一般法律依据尚待核验。")
    if payload.get("format_warnings"):
        warnings.append("存在非根本排版问题，不阻断内容草稿交付。")
    receipt = payload.get("render_receipt")
    if artifact is None:
        warnings.append("尚未读取实际制品；当前结果只能作为文本草稿检查，不能标记为可正式提交。")
    elif not isinstance(receipt, Mapping) or receipt.get("rendered") is not True:
        warnings.append("缺少专项Skill真实渲染回执；制品可供复核，但不能标记为可正式提交。")
    elif not str(receipt.get("template_id", "")).strip():
        warnings.append("渲染回执缺少template_id；不能确认实体模板命中。")

    status = "PASS_WITH_WARNINGS" if warnings else "PASS"
    return _add_delivery_contract(
        _result(status, warnings=warnings, previous_codes=previous),
        payload,
        artifact,
    )


def main(argv: List[str] | None = None) -> int:
    args = list(sys.argv[1:] if argv is None else argv)
    if len(args) > 1:
        raise SystemExit("usage: check_delivery.py [input.json]")
    if args:
        payload = json.loads(Path(args[0]).read_text(encoding="utf-8"))
    else:
        payload = json.load(sys.stdin)
    json.dump(check_document(payload), sys.stdout, ensure_ascii=False, indent=2)
    sys.stdout.write("\n")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
