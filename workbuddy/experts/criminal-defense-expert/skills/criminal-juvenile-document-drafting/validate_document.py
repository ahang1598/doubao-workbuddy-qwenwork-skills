#!/usr/bin/env python3
"""Validate an actual DOCX produced by a criminal document Skill."""

from __future__ import annotations

import argparse
import json
import re
from pathlib import Path

from docx import Document


class DocumentValidationError(ValueError):
    """Raised when a document cannot be inspected."""


def _resolved(path: str | Path) -> Path:
    return Path(path).expanduser().resolve()


def _inside(root: Path, path: Path) -> bool:
    return path == root or root in path.parents


def extract_docx_text(path: str | Path) -> str:
    document = Document(_resolved(path))
    values = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            values.extend(cell.text for cell in row.cells)
    return "\n".join(value.strip() for value in values if value.strip())


def validate_document(
    *,
    document_path: str | Path,
    matter_root: str | Path,
    doc_type: str = "",
    template_id: str = "",
    output_scene: str = "lawyer_working",
) -> dict:
    matter = _resolved(matter_root)
    path = _resolved(document_path)
    if not matter.is_dir() or not _inside(matter, path):
        return _blocked(path, "cross_matter_path", "文档路径不属于当前案件")
    if not path.is_file():
        return _blocked(path, "file_missing", "文档不存在")
    if path.suffix.lower() != ".docx":
        return _blocked(path, "wrong_file_type", "文档不是DOCX")
    try:
        text = extract_docx_text(path)
    except Exception as error:
        return _blocked(path, "file_unreadable", f"DOCX无法打开: {error}")
    if not text:
        return _blocked(path, "empty_document", "DOCX正文为空")
    compact = re.sub(r"\s+", "", text)
    if len(compact) < 20:
        return _blocked(path, "title_only", "DOCX缺少实质正文")

    warnings: list[dict] = []
    if re.search(r"\[待(?:补|核|确认)[^\]]*\]", text):
        warnings.append({"code": "unresolved_placeholder", "message": "文档仍含待补或待核占位"})
    internal_markers = ["resumeSubSessionId", "preloadedContexts", "submission_ready", "[INTERNAL", "{{INTERNAL"]
    found = [marker for marker in internal_markers if marker in text]
    if found:
        return _blocked(path, "internal_marker", f"文档含内部控制标记: {', '.join(found)}")
    if not template_id:
        warnings.append({"code": "missing_template_receipt", "message": "缺少模板ID回执"})

    status = "PASS_WITH_WARNINGS" if warnings else "PASS"
    return {
        "outcome": status,
        "document_path": str(path),
        "readable": True,
        "doc_type": doc_type,
        "template_id": template_id,
        "output_scene": output_scene,
        "text_length": len(compact),
        "warnings": warnings,
        "retryable": False,
    }


def _blocked(path: Path, code: str, message: str) -> dict:
    return {
        "outcome": "BLOCKED",
        "document_path": str(path),
        "readable": False,
        "issues": [{"code": code, "message": message}],
        "warnings": [],
        "retryable": True,
    }


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description="检查刑事Skill生成的实际DOCX")
    parser.add_argument("--document", required=True)
    parser.add_argument("--matter-root", required=True)
    parser.add_argument("--doc-type", default="")
    parser.add_argument("--template-id", default="")
    parser.add_argument("--output-scene", default="lawyer_working")
    args = parser.parse_args(argv)
    result = validate_document(
        document_path=args.document,
        matter_root=args.matter_root,
        doc_type=args.doc_type,
        template_id=args.template_id,
        output_scene=args.output_scene,
    )
    print(json.dumps(result, ensure_ascii=False))
    return 0 if result["outcome"] in {"PASS", "PASS_WITH_WARNINGS"} else 2


if __name__ == "__main__":
    raise SystemExit(main())

