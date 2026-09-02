#!/usr/bin/env python3
"""Build a Chinese contract review report DOCX."""

from __future__ import annotations

# 同目录模块（emoji_text / skill_paths）在部分宿主环境下不会自动进入 sys.path，
# 显式注入脚本所在目录，避免 import 失败连锁触发调用方逐条试错。
import os
import sys

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

import argparse
import json
import re
import unicodedata
from pathlib import Path

from emoji_text import sanitize_data
from skill_paths import generated_path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


DISCLAIMER = (
    "本文档由 AI 辅助生成，仅供参考，不构成正式法律意见，"
    "不能替代具有执业资格的律师。"
)

# Richee output standard colors
RICHEE_BLACK = "0a0d12"
RICHEE_WHITE = "ffffff"
RICHEE_LINE = "e2e5ea"
RICHEE_BG = "f6f7f9"
REPORT_FONT_SIZES = {
    "body": 12,
    "title": 22,
    "heading1": 18,
    "heading2": 16,
    "heading3": 14,
    "disclaimer": 10.5,
    "table": 10.5,
}
INDEX_HEADER_RE = re.compile(r"^(?:序号|编号|no\.?|id)$", re.IGNORECASE)
COMPACT_HEADER_MARKERS = (
    "等级",
    "标识",
    "重要性",
    "优先级",
    "状态",
    "是否",
    "结果",
)


def set_font(
    run,
    size: float = REPORT_FONT_SIZES["body"],
    bold: bool = False,
    color_hex: str | None = None,
) -> None:
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "PingFang SC"
    if color_hex:
        run.font.color.rgb = RGBColor(
            int(color_hex[0:2], 16),
            int(color_hex[2:4], 16),
            int(color_hex[4:6], 16),
        )
    r_pr = run._element.get_or_add_rPr()
    fonts = r_pr.find(qn("w:rFonts"))
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, fonts)
    fonts.set(qn("w:ascii"), "Arial")
    fonts.set(qn("w:hAnsi"), "Arial")
    fonts.set(qn("w:eastAsia"), "PingFang SC")


def set_cell_shading(cell, color_hex: str) -> None:
    """Set background shading on a table cell (Richee compliance)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), color_hex)


def style_header_row(row, widths: list[int]) -> None:
    """Apply Richee black-header-white-text style to a header row."""
    for idx, cell in enumerate(row.cells):
        set_cell_shading(cell, RICHEE_BLACK)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.bold = True


def configure(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.3)
    section.bottom_margin = Cm(2.3)
    section.left_margin = Cm(2.6)
    section.right_margin = Cm(2.6)
    normal = doc.styles["Normal"]
    normal.font.name = "PingFang SC"
    normal.font.size = Pt(REPORT_FONT_SIZES["body"])
    # 中文 PingFang SC、西文 Arial——ascii/hAnsi 必须显式设 Arial，
    # 否则英文落入 PingFang SC，造成"字体不统一"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(5)
    # Heading 4-9 也必须显式设黑：validate_review_outputs.py 的 check_black_heading_styles
    # 检查全部 heading 样式，python-docx 默认给 Heading 4-9 蓝色，不显式覆盖则交付闸门必失败。
    for name, size in (
        ("Title", REPORT_FONT_SIZES["title"]),
        ("Heading 1", REPORT_FONT_SIZES["heading1"]),
        ("Heading 2", REPORT_FONT_SIZES["heading2"]),
        ("Heading 3", REPORT_FONT_SIZES["heading3"]),
        ("Heading 4", REPORT_FONT_SIZES["heading3"]),
        ("Heading 5", REPORT_FONT_SIZES["heading3"]),
        ("Heading 6", REPORT_FONT_SIZES["heading3"]),
        ("Heading 7", REPORT_FONT_SIZES["heading3"]),
        ("Heading 8", REPORT_FONT_SIZES["heading3"]),
        ("Heading 9", REPORT_FONT_SIZES["heading3"]),
    ):
        style = doc.styles[name]
        style.font.name = "PingFang SC"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0x0A, 0x0D, 0x12)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.paragraph_format.keep_with_next = True


def set_table_width(table, widths: list[int]) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _display_units(text: str) -> float:
    """Estimate rendered text demand for mixed Chinese and Latin content."""
    units = 0.0
    for character in re.sub(r"\s+", " ", text.strip()):
        if character.isspace():
            units += 0.35
        elif unicodedata.east_asian_width(character) in {"W", "F"}:
            units += 1.0
        elif character.isupper():
            units += 0.65
        else:
            units += 0.55
    return units


def _column_kind(header: str) -> str:
    normalized = re.sub(r"[\s：:()\uff08\uff09【】\[\]]+", "", header).lower()
    if INDEX_HEADER_RE.fullmatch(normalized):
        return "index"
    if any(marker in normalized for marker in COMPACT_HEADER_MARKERS):
        return "compact"
    return "narrative"


def allocate_content_widths(table, total_width: int) -> list[int]:
    """Allocate exact DXA widths by semantics and real cell text demand."""
    headers = [cell.text.strip() for cell in table.rows[0].cells]
    columns = [
        [row.cells[index].text.strip() for row in table.rows]
        for index in range(len(headers))
    ]
    count = len(headers)
    base_min_ratio = min(0.12, 0.68 / count)
    demands: list[float] = []
    minimums: list[int] = []
    maximums: list[int] = []
    for header, values in zip(headers, columns):
        kind = _column_kind(header)
        measured = [min(max(_display_units(value), 1.0), 48.0) for value in values]
        peak = max(measured, default=1.0)
        average = sum(measured) / len(measured) if measured else 1.0
        demands.append(
            max(_display_units(header) + 1.5, peak * 0.68 + average * 0.32, 1.0)
        )
        if kind == "index":
            minimums.append(round(total_width * 0.065))
            maximums.append(round(total_width * 0.09))
        elif kind == "compact":
            minimums.append(round(total_width * min(base_min_ratio, 0.11)))
            maximums.append(round(total_width * 0.20))
        else:
            minimums.append(round(total_width * base_min_ratio))
            maximums.append(round(total_width * (0.80 if count <= 3 else 0.65)))

    while sum(minimums) > total_width:
        minimums[max(range(count), key=lambda item: minimums[item])] -= 1
    widths = minimums[:]
    remaining = total_width - sum(widths)
    active = {index for index in range(count) if widths[index] < maximums[index]}
    while remaining > 0 and active:
        total_demand = sum(demands[index] for index in active)
        planned = []
        for index in active:
            raw = remaining * demands[index] / total_demand
            planned.append((index, int(raw), raw - int(raw)))
        assigned = 0
        for index, amount, _ in planned:
            amount = min(amount, maximums[index] - widths[index])
            widths[index] += amount
            assigned += amount
        leftover = remaining - assigned
        for index, _, _ in sorted(planned, key=lambda item: item[2], reverse=True):
            if leftover <= 0:
                break
            if widths[index] < maximums[index]:
                widths[index] += 1
                assigned += 1
                leftover -= 1
        if assigned == 0:
            break
        remaining -= assigned
        active = {index for index in active if widths[index] < maximums[index]}

    # All caps are soft when their sum is smaller than the physical table.
    order = sorted(range(count), key=lambda item: demands[item], reverse=True)
    while remaining > 0:
        for index in order:
            if remaining <= 0:
                break
            widths[index] += 1
            remaining -= 1
    return widths


def rebalance_table_widths(table, total_width: int) -> list[int]:
    widths = allocate_content_widths(table, total_width)
    set_table_width(table, widths)
    return widths


def configure_table_pagination(table) -> None:
    for index, row in enumerate(table.rows):
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))
        if index == 0 and tr_pr.find(qn("w:tblHeader")) is None:
            header = OxmlElement("w:tblHeader")
            header.set(qn("w:val"), "true")
            tr_pr.append(header)


def add_bilingual_block(
    doc: Document,
    en: str,
    zh: str,
    en_size: float = REPORT_FONT_SIZES["body"],
    zh_size: float = REPORT_FONT_SIZES["body"],
    bullet: bool = False,
) -> None:
    """zh-primary：中文报告默认只渲染中文段，仅当中文缺失时回退英文。

    （修复人工评测 B2：en+zh 双段全渲染导致整段英文嵌入中文报告、严重冗余。
    数据层 bilingual dict 保持不变，双语版交付物另行复用 en。）
    """
    text = zh or en
    if not text:
        return
    add_plain_para(doc, text, zh_size, bullet=bullet)


def _is_bilingual_dict(content) -> bool:
    """True when content is a dict with at least one of 'en'/'zh' keys (and not a simple item dict)."""
    if not isinstance(content, dict):
        return False
    return ("en" in content or "zh" in content) and not any(
        k in content for k in ("title", "issue", "item", "target")
    )


# 模型在 report.json 中夹带的 Markdown 标记（## 标题、**加粗**、`code`、- 列表）
# 不会被 docx 渲染，会以字面符号泄漏到报告里。脚本层兜底转换：
# **…** → 真实加粗 run；行首 #… → 加粗小节标题段；其余记号剥离。
MD_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
MD_HEADING_RE = re.compile(r"^\s{0,3}#{1,6}\s+")
MD_BULLET_RE = re.compile(r"^\s{0,3}[-*•]\s+")


def strip_inline_md(text: str) -> str:
    """表格单元格等无富文本场景：剥离 Markdown 记号只留文字。"""
    if not text:
        return text
    text = MD_HEADING_RE.sub("", text)
    text = MD_BOLD_RE.sub(r"\1", text)
    return text.replace("`", "")


def add_md_runs(paragraph, text: str, size: float, bold: bool = False) -> None:
    """把含 **…** 的文本拆成普通/加粗 run，全部经 set_font。"""
    pos = 0
    for match in MD_BOLD_RE.finditer(text):
        if match.start() > pos:
            set_font(paragraph.add_run(text[pos:match.start()].replace("`", "")), size, bold)
        set_font(paragraph.add_run(match.group(1).replace("`", "")), size, True)
        pos = match.end()
    if pos < len(text):
        set_font(paragraph.add_run(text[pos:].replace("`", "")), size, bold)


def add_plain_para(
    doc: Document,
    text: str,
    size: float = REPORT_FONT_SIZES["body"],
    bullet: bool = False,
) -> None:
    """所有正文段落统一经 set_font（修复 B11 字体不统一），并做 Markdown 兜底转换。"""
    text = text or ""
    if "\n" in text:
        # 多行字符串（模型整段 Markdown 直出）逐行渲染，行首 # 各自成小节标题
        for line in text.split("\n"):
            line = line.strip()
            if line:
                add_plain_para(doc, line, size, bullet=bullet)
        return
    heading = MD_HEADING_RE.match(text)
    if heading:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        add_md_runs(p, text[heading.end():].strip(), size, bold=True)
        return
    if bullet:
        text = MD_BULLET_RE.sub("", text)
    p = doc.add_paragraph(style="List Bullet" if bullet else None)
    add_md_runs(p, text, size)


def add_section(doc: Document, title: str, content) -> None:
    """Enhanced section renderer supporting:
    - bilingual dict: {"en": "...", "zh": "...", "items": [...]} — zh-primary 渲染
    - explained empty: {"empty_reason": "..."} shows reason instead of 无。
    - backward-compatible: plain string, list of dicts, list of strings

    章节永不整体跳过（修复 B1 编号跳跃）：旧版 {"omit_if_empty": true}
    兼容为渲染"本节不适用。"一行，编号保持连续。
    """
    if isinstance(content, dict) and content.get("omit_if_empty"):
        content = {"empty_reason": content.get("empty_reason", "本节不适用。")}

    if isinstance(content, dict) and "empty_reason" in content and not content.get("en") and not content.get("zh"):
        doc.add_heading(title, level=1)
        add_plain_para(doc, content["empty_reason"])
        return

    doc.add_heading(title, level=1)

    if not content:
        add_plain_para(doc, "无。")
        return

    # --- Bilingual dict format（zh-primary）---
    if _is_bilingual_dict(content):
        add_bilingual_block(doc, content.get("en", ""), content.get("zh", ""))
        items = content.get("items", []) or content.get("checklist", [])
        for item in items:
            if isinstance(item, dict):
                add_bilingual_block(doc, item.get("en", ""), item.get("zh", ""), bullet=True)
            else:
                add_plain_para(doc, str(item), REPORT_FONT_SIZES["body"], bullet=True)
        return

    # --- Plain string ---
    if isinstance(content, str):
        add_plain_para(doc, content)
        return

    # --- List of items ---
    for item in content:
        if isinstance(item, dict):
            if _is_bilingual_dict(item):
                add_bilingual_block(doc, item.get("en", ""), item.get("zh", ""), bullet=True)
                continue

            heading = item.get("title") or item.get("issue") or item.get("item") or "事项"
            detail = item.get("detail") or item.get("text") or item.get("status") or ""
            basis = item.get("basis_tag", "")
            add_plain_para(
                doc,
                f"{basis} {heading}：{detail}".strip(),
                REPORT_FONT_SIZES["body"],
                bullet=True,
            )
        else:
            add_plain_para(doc, str(item), REPORT_FONT_SIZES["body"], bullet=True)


def _fill_basis_column(rows: list[dict]) -> list[dict]:
    """结构参数硬检查表的「依据」列需要分类标签 + 详细依据说明，而非裸标签。

    真机缺陷：依据列只渲染 `basis_tag`（如 `[惯例]`），看不到具体依据。此函数把
    每行的 `basis_tag` 与详细说明（`basis_detail` / `basis`）合并到 `basis` 列；
    详细说明缺失时降级为标签 + 提示，促使补全而非留裸标签。
    """
    filled = []
    for row in rows:
        if not isinstance(row, dict):
            filled.append(row)
            continue
        row = dict(row)
        tag = str(row.get("basis_tag", "")).strip()
        detail = str(row.get("basis_detail") or row.get("basis") or "").strip()
        # 若 detail 本身已含标签则不重复叠加
        if detail and tag and not detail.startswith(tag):
            row["basis"] = f"{tag} {detail}"
        elif detail:
            row["basis"] = detail
        elif tag:
            row["basis"] = f"{tag} 详见对应条款审查意见"
        else:
            row["basis"] = "基于合理判断，建议专业律师确认"
        filled.append(row)
    return filled


def add_matrix(
    doc: Document,
    title: str,
    rows: list[dict],
    headers: list[str],
    keys: list[str],
    widths: list[int],
    empty_reason: str = "无。",
) -> None:
    """Matrix renderer。rows 为空时不再跳节（修复 B1 编号跳跃），
    而是渲染 empty_reason 一行向读者交代不适用原因。"""
    # 入参容错（真机 17e2c418）：模型常把 ip_analysis 等写成 {"items":[...]}
    # （该形状在 add_section 里是合法的），矩阵渲染需要 list，此处先解包，
    # 避免 build 崩溃、逼主 Agent 手改 report.json 兜底。
    if isinstance(rows, dict):
        rows = rows.get("items") or rows.get("checklist") or rows.get("list") or []
    if not rows:
        doc.add_heading(title, level=1)
        add_plain_para(doc, empty_reason)
        return
    doc.add_heading(title, level=1)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths)
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
        for run in cell.paragraphs[0].runs:
            set_font(run, REPORT_FONT_SIZES["table"], True)
    style_header_row(table.rows[0], widths)
    for item in rows:
        cells = table.add_row().cells
        for cell, key in zip(cells, keys):
            cell.text = strip_inline_md(str(item.get(key, ""))).strip()
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.line_spacing = 1.5
                for run in paragraph.runs:
                    set_font(run, REPORT_FONT_SIZES["table"])
    rebalance_table_widths(table, sum(widths))
    configure_table_pagination(table)


def add_risk_table(doc: Document, title: str, risks: list[dict]) -> None:
    doc.add_heading(title, level=1)
    if not risks:
        add_plain_para(doc, "未识别到需要单列的风险，但仍需完成专业复核。")
        return
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    widths = [680, 1350, 1700, 1700, 2150, 1220]
    set_table_width(table, widths)
    headers = ["编号/等级", "位置", "问题", "影响与可能性", "建议及备选", "依据"]
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
        for run in cell.paragraphs[0].runs:
            set_font(run, REPORT_FONT_SIZES["table"], True)
    style_header_row(table.rows[0], widths)
    for item in risks:
        values = [
            "\n".join(
                value
                for value in (
                    str(item.get("issue_id", "")).strip(),
                    str(item.get("level", "待定")).strip(),
                )
                if value
            ),
            item.get("location", ""),
            item.get("issue", ""),
            item.get("impact_likelihood", ""),
            item.get("recommendation", ""),
            item.get("basis_tag", "[要点]"),
        ]
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            cell.text = strip_inline_md(str(value)).strip()
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.line_spacing = 1.5
                for run in paragraph.runs:
                    set_font(run, REPORT_FONT_SIZES["table"])
    rebalance_table_widths(table, sum(widths))
    configure_table_pagination(table)


def validate_risk_ids(risks: list[dict]) -> None:
    issue_ids = [str(item.get("issue_id", "")).strip() for item in risks]
    if any(not issue_id for issue_id in issue_ids):
        raise ValueError("every report risk requires issue_id")
    if len(issue_ids) != len(set(issue_ids)):
        raise ValueError("report risk issue_id values must be unique")


def renumber_risks(risks: list[dict]) -> None:
    """按报告最终排序重编号 R-001…R-N（修复 B7：模型按发现顺序赋 ID、
    按优先级排序输出导致 R-017 插在 R-007/R-008 之间）。原 ID 仅内部使用。"""
    for index, item in enumerate(risks, 1):
        original = str(item.get("issue_id", "")).strip()
        new_id = f"R-{index:03d}"
        if original and original != new_id:
            print(f"risk renumbered: {original} -> {new_id}")
        item["issue_id"] = new_id


def aggregate_pending(data: dict) -> None:
    """交叉校验"待核查"一致性（修复 B5：十二节列待核查项、
    十三节却称"无"的前后矛盾）。pending 为空时自动聚合全文待核查条目。"""
    found: list[str] = []

    def scan(obj) -> None:
        if isinstance(obj, str):
            if "待核查" in obj or "待核实" in obj:
                found.append(obj if len(obj) <= 120 else obj[:117] + "…")
        elif isinstance(obj, list):
            for entry in obj:
                scan(entry)
        elif isinstance(obj, dict):
            for value in obj.values():
                scan(value)

    for key in ("verification", "risks", "structural_parameters", "missing_terms", "executive_summary"):
        scan(data.get(key))

    pending = data.get("pending")
    is_empty = (
        not pending
        or (isinstance(pending, str) and pending.strip() in {"无", "无。", "暂无"})
        or (isinstance(pending, list) and not pending)
    )
    if found and is_empty:
        deduped = list(dict.fromkeys(found))
        data["pending"] = deduped
        print(
            f"warning: pending was empty but {len(deduped)} 待核查 item(s) found "
            "elsewhere in the report — auto-aggregated into 待核查事项"
        )


IDENTIFIER_RE_KEYS = ("item", "direction", "status")


def validate_coverage(coverage) -> None:
    """coverage 结构校验（修复 B4：模型把 playbook 内部英文键名
    如 universal_rules 直接塞入第七节、渲染为裸标签无展开）。"""
    import re as _re

    bare_identifier = _re.compile(r"^[A-Za-z0-9_\-]+$")
    has_cjk = _re.compile(r"[一-鿿]")

    def fail(detail: str) -> None:
        raise ValueError(
            f"coverage 数据不合格：{detail}。"
            "请使用结构化条目 {item(中文), direction, status, location}，"
            "或无审查原则时使用 {\"empty_reason\": \"…\"}（见 input-schema.md）"
        )

    if coverage is None or isinstance(coverage, str):
        return
    if isinstance(coverage, dict):
        if "empty_reason" in coverage or _is_bilingual_dict(coverage):
            return
        bare = [k for k in coverage if bare_identifier.match(str(k))]
        if bare:
            fail(f"字典键为裸英文标识符 {bare}，疑似 playbook 内部键名直出")
        return
    if isinstance(coverage, list):
        for entry in coverage:
            if isinstance(entry, str) and bare_identifier.match(entry):
                fail(f"列表项为裸英文标识符 '{entry}'")
            if isinstance(entry, dict):
                if not entry.get("item") or not entry.get("status"):
                    fail(f"条目缺少 item/status 字段：{entry}")
                if not has_cjk.search(str(entry.get("item", ""))):
                    fail(f"item 须为中文描述，得到：{entry.get('item')}")


class ChapterNumberer:
    """动态章节编号（修复 B1）：序号按实际渲染顺序分配，
    任何章节的有无都不会造成编号跳跃。"""

    CN = ["一", "二", "三", "四", "五", "六", "七", "八", "九", "十",
          "十一", "十二", "十三", "十四", "十五", "十六"]

    def __init__(self) -> None:
        self.index = 0

    def title(self, name: str) -> str:
        self.index += 1
        return f"{self.CN[self.index - 1]}、{name}"


def warn_missing_basis_quote(risks: list[dict]) -> None:
    """高/中风险的 [法规]/[用规]/[要点] 依据须带「」原文摘录；缺失则警告（非致命）。"""
    need = ("[法规]", "[用规]", "[要点]")
    quote = re.compile(r"[「『\"“][^」』\"”]{2,}")
    hits = [f"{r.get('issue_id','?')}({r.get('location','')})"
            for r in risks
            if str(r.get("level", "")).strip().lower() in ("高", "中", "high", "medium")
            and any(t in str(r.get("basis_tag", "")) for t in need)
            and not quote.search(str(r.get("basis_tag", "")))]
    if hits:
        print("[警告] 以下高/中风险依据缺「」原文摘录，请补全原文或降级为待核查: " + "; ".join(hits))


def build(data: dict, output: Path) -> None:
    output = generated_path(output, "review report")
    validate_risk_ids(data.get("risks", []))
    renumber_risks(data.get("risks", []))
    validate_coverage(data.get("coverage"))
    warn_missing_basis_quote(data.get("risks", []))
    aggregate_pending(data)
    doc = Document()
    configure(doc)
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(
        title.add_run(data.get("report_title", "合同审查报告")),
        REPORT_FONT_SIZES["title"],
        True,
    )
    first = doc.add_paragraph()
    first.paragraph_format.space_before = Pt(8)
    set_font(first.add_run(DISCLAIMER), REPORT_FONT_SIZES["disclaimer"], True)

    ch = ChapterNumberer()
    add_section(doc, ch.title("审查范围与已知事实"), data.get("scope"))
    add_section(doc, ch.title("事实、假设与用户立场"), data.get("facts"))
    add_section(doc, ch.title("客户规则与审查立场校准"), data.get("playbook_status"))
    add_section(doc, ch.title("执行摘要"), data.get("executive_summary"))
    add_matrix(
        doc,
        ch.title("结构参数硬检查"),
        _fill_basis_column(data.get("structural_parameters", [])),
        ["参数", "首选", "底线", "合同实际", "状态", "依据"],
        ["parameter", "preferred", "bottom_line", "actual", "status", "basis"],
        [1150, 1450, 1350, 1750, 1400, 1700],
        empty_reason="经核实不涉及：本次审查未启用结构参数硬检查规则。",
    )
    add_risk_table(doc, ch.title("风险清单"), data.get("risks", []))
    coverage = data.get("coverage")
    if isinstance(coverage, list) and coverage and all(isinstance(e, dict) for e in coverage):
        add_matrix(
            doc,
            ch.title("双向审查覆盖状态"),
            coverage,
            ["审查项", "方向", "状态", "位置"],
            ["item", "direction", "status", "location"],
            [3200, 1200, 1600, 2800],
        )
    else:
        add_section(doc, ch.title("双向审查覆盖状态"), coverage)
    add_section(doc, ch.title("缺失保护与系统性不一致"), data.get("missing_terms"))
    add_matrix(
        doc,
        ch.title("权利对称性检查"),
        data.get("symmetry", []),
        ["权利", "客户立场", "相对方立场", "评估"],
        ["right", "client_position", "counterparty_position", "assessment"],
        [1600, 2100, 2100, 3000],
        empty_reason="本节不适用：合同未涉及需对称性比较的实质性权利分配条款。",
    )
    add_matrix(
        doc,
        ch.title("知识产权三层检查"),
        data.get("ip_analysis", []),
        ["对象", "归属", "申请或控制", "维权"],
        ["category", "ownership", "control", "enforcement"],
        [1500, 2350, 2400, 2550],
        empty_reason="本节不适用：合同未涉及知识产权归属、许可或维权条款。",
    )
    add_section(doc, ch.title("建议文本与谈判备选"), data.get("recommendations"))
    add_section(doc, ch.title("准据法核验状态"), data.get("verification"))
    add_section(doc, ch.title("待核查事项"), data.get("pending"))
    add_section(doc, ch.title("交付文件与验证结果"), data.get("deliverables"))

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    print(f"created {output}")


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("input", type=Path)
    parser.add_argument("output", type=Path)
    args = parser.parse_args()
    data = json.loads(args.input.read_text(encoding="utf-8"))
    # DOCX hard rule (D3-S2): no emoji. Model-authored status markers map to
    # bracket labels; anything else in emoji ranges is stripped.
    build(sanitize_data(data), args.output)


if __name__ == "__main__":
    main()
