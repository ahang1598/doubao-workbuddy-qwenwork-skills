#!/usr/bin/env python3
"""
可复用 Word 正式文档样式规范化模块。

遵循 Richee 输出规范 1.2.0（word-report / word-revision profile）。
适用于法律报告、意见书、备忘录、函件等所有正式法律文档。

使用方式：
    import sys
    from docx import Document
    from scripts.doc_styler import apply_doc_style

    doc = Document("pandoc_output.docx")
    violations = apply_doc_style(doc, profile="word-report", doc_name="文档名称")
    if violations:
        print("校验失败:", violations)
        sys.exit(1)
    doc.save("final.docx")
"""

import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---- 配置常量 ----

# 字体
CN_FONT = "宋体"
EN_FONT = "宋体"
FONT_EAST_ASIA = qn("w:eastAsia")
FONT_ASCII = qn("w:ascii")
FONT_H_ANSI = qn("w:hAnsi")
FONT_CS = qn("w:cs")

# 字号阶梯（单位 Pt）
SIZE_TITLE = 22
SIZE_H1 = 16
SIZE_H2 = 15
SIZE_H3 = 14
SIZE_BODY = 12
SIZE_SMALL = 10.5

# 配色（灰度 + 克制绿色点缀）
COLOR_BLACK = RGBColor(0x1A, 0x1A, 0x1A)
COLOR_DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
COLOR_MID_GRAY = RGBColor(0x99, 0x99, 0x99)
COLOR_LIGHT_GRAY = RGBColor(0xE8, 0xE8, 0xE8)
COLOR_TABLE_HEADER_BG = RGBColor(0x2D, 0x2D, 0x2D)
COLOR_TABLE_HEADER_FG = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_TABLE_BORDER = RGBColor(0xBF, 0xBF, 0xBF)
COLOR_TABLE_ZEBRA = RGBColor(0xF5, 0xF5, 0xF5)
COLOR_GREEN_ACCENT = RGBColor(0x00, 0x6D, 0x4A)
COLOR_HYPERLINK = RGBColor(0x00, 0x70, 0xC0)

# 页面
A4_WIDTH = Cm(21)
A4_HEIGHT = Cm(29.7)
MARGIN_DEFAULT = Cm(2.54)

# 行距
LINE_SPACING_BODY = 1.5
LINE_SPACING_TABLE = 1.15

# 表格安全余量（约 2mm）
TABLE_WIDTH_PERCENT = 97  # 页面内容区的 97%

# 列宽比例参考
SEQ_COL_MAX = 9    # 序号列 ≤9%
SHORT_COL_MAX = 20  # 短值列（等级/状态/标识）≤20%

# 表格内容用于列类型推断的关键词扩展
_SEQ_KEYWORDS = ("序号", "no", "no.", "#", "编号", "项次")
_SHORT_KEYWORDS = ("等级", "级别", "风险", "状态", "标识", "程度", "level", "status", "type", "类型", "优先级", "是否", "有无")
_NARRATIVE_KEYWORDS = ("说明", "描述", "内容", "建议", "影响", "备注", "分析", "description", "remark", "note", "suggestion", "措施", "方案", "意见", "结论")

# ---- 内部辅助 ----

def _set_font(run, size_pt: float, bold: bool = False, color: RGBColor = COLOR_BLACK,
              cn: str = CN_FONT, en: str = EN_FONT, preserve_color: bool = False,
              preserve_user_color: bool = True):
    """在单个 run 上统一设置字体、字号、颜色、粗细。

    参数说明：
    - preserve_color=True：保留 run 原有颜色（用于超链接等场景）
    - preserve_user_color=True（默认）：检测 run 是否已有显式 <w:color> 子元素
      （由 Lua 过滤器或用户在 Markdown 中指定的颜色），有则保留，不覆盖
      用于保护 <span style="color:red">XXX</span> 这类用户显式颜色
      标题/表头等结构性样式应传 preserve_user_color=False 强制覆盖为 black
    """
    run.font.size = Pt(size_pt)
    run.font.bold = bold

    # 颜色处理优先级：
    # 1. preserve_color=True → 保留原色（超链接）
    # 2. preserve_user_color=True 且 run 已有 <w:color> → 保留用户指定色
    # 3. 其他 → 强制应用 color 参数（默认 COLOR_BLACK）
    if preserve_color:
        pass  # 不动颜色
    elif preserve_user_color:
        # 检查 run 是否已有显式 <w:color>（直接子元素，非样式继承）
        rPr_elem = run._element.find(qn("w:rPr"))
        if rPr_elem is not None and rPr_elem.find(qn("w:color")) is not None:
            pass  # 保留用户/lua 指定的颜色
        else:
            run.font.color.rgb = color
    else:
        # 强制覆盖（标题、表头等结构性样式）
        run.font.color.rgb = color

    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(FONT_EAST_ASIA, cn)
    rFonts.set(FONT_ASCII, en)
    rFonts.set(FONT_H_ANSI, en)
    rFonts.set(FONT_CS, en)


def _get_all_runs(para):
    """获取段落中所有 run，包括超链接 <w:hyperlink> 内的 run。

    python-docx 的 para.runs 仅返回段落直接子级的 <w:r>，
    超链接内的 run 需要通过 XML 遍历获取。
    """
    from docx.text.run import Run
    return [Run(r, para) for r in para._element.iter(qn("w:r"))]


def _is_in_hyperlink(run) -> bool:
    """判断 run 是否位于 <w:hyperlink> 元素内。"""
    parent = run._element.getparent()
    while parent is not None:
        if parent.tag == qn("w:hyperlink"):
            return True
        parent = parent.getparent()
    return False


def _set_paragraph_spacing(para, line_spacing: float = LINE_SPACING_BODY,
                           space_before: float = 0, space_after: float = 6):
    """设置段落行距与段前段后间距。"""
    pf = para.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)


def _set_cell_border(cell, borders: dict):
    """设置单元格边框（一次设置所有边）。"""
    tcPr = cell._tc.get_or_add_tcPr()
    # 移除旧边框
    existing = tcPr.find(qn("w:tcBorders"))
    if existing is not None:
        tcPr.remove(existing)
    tcBorders = OxmlElement("w:tcBorders")
    for edge, val in borders.items():
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), val.get("val", "single"))
        border.set(qn("w:sz"), val.get("sz", "4"))
        border.set(qn("w:color"), val.get("color", "BFBFBF"))
        tcBorders.append(border)
    tcPr.append(tcBorders)


def _set_cell_margins(cell, top: int = 40, bottom: int = 40, left: int = 80, right: int = 80):
    """设置单元格内边距（单位：dxa，1pt=20dxa）。"""
    tcPr = cell._tc.get_or_add_tcPr()
    existing = tcPr.find(qn("w:tcMar"))
    if existing is not None:
        tcPr.remove(existing)
    tcMar = OxmlElement("w:tcMar")
    for edge, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)


def _set_cell_shading(cell, color_hex: str):
    """设置单元格底纹。"""
    tcPr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), color_hex)
    existing = tcPr.find(qn("w:shd"))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(shading)


def _estimate_heading_level(para) -> int:
    """推断标题级别（1-3），0 表示非标题。

    仅依赖 Pandoc 生成的段落样式名（Heading 1/2/3/Title）。
    Pandoc 已将 Markdown 的 #/##/### 转为正确的 Heading 样式，无需文本回退。
    """
    style_name = (para.style.name if para.style else "").lower()

    if "heading 1" in style_name:
        return 1
    if "heading 2" in style_name:
        return 2
    if "heading 3" in style_name:
        return 3
    if "title" in style_name:
        return 0
    return 0


def _add_horizontal_line(doc, color_hex: str = "006D4A", width_pt: float = 1.0):
    """在文档末尾追加一条水平分隔线。"""
    para = doc.add_paragraph()
    pPr = para._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(int(width_pt * 8)))  # 1pt ≈ 8 EMU
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return para


# ---- 公开 API ----

def setup_page(doc: Document, profile: str = "word-report"):
    """设置页面尺寸与页边距。"""
    for section in doc.sections:
        section.page_width = A4_WIDTH
        section.page_height = A4_HEIGHT
        section.top_margin = MARGIN_DEFAULT
        section.bottom_margin = MARGIN_DEFAULT
        section.left_margin = MARGIN_DEFAULT
        section.right_margin = MARGIN_DEFAULT


def remove_horizontal_rules(doc: Document):
    """清除文档中 Pandoc 生成的水平线（<v:rect o:hr="t"/>）。

    Pandoc 把 Markdown 的 --- 渲染为 VML 矩形（带 o:hr="t" 标记），
    在 Word 中显示为可见横线。法律文档中通常只用于内容分隔，
    横线视觉突兀，统一移除。

    Lua 过滤器 horizontal-rule-to-spacer.lua 已把 HorizontalRule 转为
    空段落，但兜底处理仍有必要：覆盖直接调用 pandoc（不走完整过滤器
    链）或残留 VML 元素的场景。
    """
    body = doc.element.body
    # 命名空间前缀 v: 在 lxml 中需要显式声明
    nsmap = {'v': 'urn:schemas-microsoft-com:vml',
             'o': 'urn:schemas-microsoft-com:office:office'}
    removed = 0
    # 找所有 v:rect 带 o:hr 属性的元素
    for rect in body.findall('.//v:rect', nsmap):
        hr_attr = rect.get('{urn:schemas-microsoft-com:office:office}hr')
        if hr_attr == 't':
            # 找到包含该 v:rect 的 <w:r>，再找包含 <w:r> 的 <w:p>
            # 整个 <w:p> 就是水平线段落（通常只含 v:rect），移除整段
            parent = rect.getparent()
            while parent is not None and parent.tag != qn('w:p'):
                parent = parent.getparent()
            if parent is not None:
                grandparent = parent.getparent()
                if grandparent is not None:
                    grandparent.remove(parent)
                    removed += 1
    if removed > 0:
        print(f"[doc_styler] 移除水平线段落 {removed} 处")


def normalize_fonts(doc: Document, profile: str = "word-report"):
    """
    规范化全文字体与字号：
    - 显式声明中文宋体 + 西文宋体（全文统一宋体）
    - 按字号阶梯覆盖每个 run（含超链接内 run，保留超链接颜色）
    - 覆盖正文段落 + 表格内段落 + 页眉页脚段落
    """
    def _process_paragraph(para):
        level = _estimate_heading_level(para)
        style_name = (para.style.name if para.style else "").lower()

        if "title" in style_name:
            size = SIZE_TITLE
        elif level == 1:
            size = SIZE_H1
        elif level == 2:
            size = SIZE_H2
        elif level == 3:
            size = SIZE_H3
        else:
            text = para.text[:50] if para.text else ""
            if any(k in text for k in ("免责声明", "AI 辅助", "仅供参考", "不构成")):
                size = SIZE_SMALL
            else:
                size = SIZE_BODY

        for run in _get_all_runs(para):
            # 超链接内的 run 保留原有颜色（蓝色超链接样式）
            _set_font(run, size, preserve_color=_is_in_hyperlink(run))

    # 正文段落
    for para in doc.paragraphs:
        _process_paragraph(para)

    # 表格内段落
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _process_paragraph(para)  # 先统一设置为正文字体，style_tables 会覆盖

    # 页眉页脚段落
    for section in doc.sections:
        for para in section.header.paragraphs:
            _process_paragraph(para)
        for para in section.footer.paragraphs:
            _process_paragraph(para)


def style_headings(doc: Document, profile: str = "word-report"):
    """标题样式：统一黑色加粗，不添加装饰色。"""
    if profile != "word-report":
        return

    # 样式级颜色：统一黑色，覆盖 Word 默认主题蓝
    for style in doc.styles:
        name = (style.name or "").lower()
        if name.startswith("heading") or name == "title":
            style.font.color.rgb = COLOR_BLACK

    for para in doc.paragraphs:
        level = _estimate_heading_level(para)
        if level == 0:
            style_name = (para.style.name if para.style else "").lower()
            if "title" in style_name:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in _get_all_runs(para):
                    _set_font(run, SIZE_TITLE, bold=True, color=COLOR_BLACK,
                              preserve_user_color=False)
            continue

        # 标题统一黑色加粗（不保留用户颜色，标题样式应统一）
        for run in _get_all_runs(para):
            _set_font(run, {1: SIZE_H1, 2: SIZE_H2, 3: SIZE_H3}[level],
                      bold=True, color=COLOR_BLACK,
                      preserve_user_color=False)

        # 行距
        _set_paragraph_spacing(para, LINE_SPACING_BODY, space_before=12, space_after=6)


def style_hyperlinks(doc: Document):
    """确保外部超链接为蓝色 + 下划线，内部锚点链接保持普通文本样式。

    Pandoc 模板的 Hyperlink 样式可能缺少下划线。
    外部链接（r:id）保留蓝色+下划线；内部锚点（w:anchor）移除超链接样式，
    使其视觉上与普通文本一致但仍可点击跳转。
    """
    # 1. 确保样式定义存在（供外部链接使用）
    for style in doc.styles:
        name = (style.name or "").lower()
        if name in ("hyperlink", "followedhyperlink"):
            if name == "hyperlink":
                style.font.color.rgb = COLOR_HYPERLINK
            style.font.underline = True

    # 2. 内部锚点链接（w:anchor）移除超链接样式，显式设为普通文本外观
    #    Word/WPS 对 <w:hyperlink> 元素有内置默认渲染（蓝+下划线），
    #    仅删除 rStyle 不够，必须显式覆盖颜色和下划线属性
    body = doc.element.body
    for hl in body.findall('.//' + qn('w:hyperlink')):
        if hl.get(qn('w:anchor')) is not None:
            for run in hl.findall(qn('w:r')):
                rPr = run.get_or_add_rPr()
                # 移除 rStyle（不再引用 Hyperlink 样式）
                rStyle = rPr.find(qn('w:rStyle'))
                if rStyle is not None:
                    rPr.remove(rStyle)
                # 显式覆盖：黑色文字 + 无下划线（抵消 W/WS 内置渲染）
                color = OxmlElement('w:color')
                color.set(qn('w:val'), '000000')  # 黑色
                existing_color = rPr.find(qn('w:color'))
                if existing_color is not None:
                    rPr.remove(existing_color)
                rPr.append(color)

                underline = OxmlElement('w:u')
                underline.set(qn('w:val'), 'none')  # 无下划线
                existing_u = rPr.find(qn('w:u'))
                if existing_u is not None:
                    rPr.remove(existing_u)
                rPr.append(underline)


def fix_internal_hyperlinks(doc: Document):
    """将 w:anchor 方式的内部超链接转换为 relationship-based 方式（WPS 兼容）。

    Pandoc 生成的内部锚点跳转使用 w:anchor 属性（OOXML 标准），
    但 WPS 对此支持不完善，点击会报"无法打开指定的文件"。
    本函数将其转换为通过 r:id + 内部关系引用书签的方式，
    WPS / Word 均可正常跳转。

    转换前：<w:hyperlink w:anchor="risk-compliance">...</w:hyperlink>
    转换后：<w:hyperlink r:id="rIdXxx">...</w:hyperlink>
           关系：rIdXxx -> Target="#risk-compliance" TargetMode="Internal"
    """
    from lxml import etree

    HYPERLINK_REL_TYPE = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink"
    )
    RELS_NS = "http://schemas.openxmlformats.org/package/2006/relationships"

    body = doc.element.body
    hyperlinks = body.findall('.//' + qn('w:hyperlink'))

    # 收集需要转换的 anchor
    anchor_links = []
    for hl in hyperlinks:
        anchor = hl.get(qn('w:anchor'))
        if anchor:
            anchor_links.append((hl, anchor))

    if not anchor_links:
        return

    # 通过 doc.part.rels 的内部 _rId_to_rel 映射找到最大 rId
    rels = doc.part.rels
    max_rid = 0
    for r_id in rels._rId_to_rel.keys():
        m = re.match(r'rId(\d+)', r_id)
        if m:
            max_rid = max(max_rid, int(m.group(1)))

    # 获取关系 XML 元素（Relationships 对象的底层 XML）
    # 遍历 Relationships 对象的属性找 XML 根元素
    rels_xml_root = None
    for attr_name in ('_element', 'xml', '_root'):
        if hasattr(rels, attr_name):
            val = getattr(rels, attr_name)
            if hasattr(val, 'tag') or hasattr(val, 'append'):
                rels_xml_root = val
                break

    # 如果找不到，直接从 part 的 blob 解析
    if rels_xml_root is None:
        rels_part_name = '_rels/' + doc.part.basename + '.rels'
        rels_blob = doc.part.package.related_parts.get(rels_part_name)
        if rels_blob is not None:
            if hasattr(rels_blob, 'blob'):
                rels_xml_root = etree.fromstring(rels_blob.blob)
            elif hasattr(rels_blob, '_blob'):
                rels_xml_root = etree.fromstring(rels_blob._blob)

    if rels_xml_root is None:
        print("[doc_styler] 警告：无法获取关系 XML，跳过内部超链接转换")
        return

    converted = 0
    new_rels = []
    for hl, anchor in anchor_links:
        # 移除 w:anchor 属性
        del hl.attrib[qn('w:anchor')]

        # 在关系 XML 中创建新关系
        max_rid += 1
        new_rid = f'rId{max_rid}'
        new_rel = etree.Element(
            f'{{{RELS_NS}}}Relationship',
            Id=new_rid,
            Type=HYPERLINK_REL_TYPE,
            Target=f'#{anchor}',
            TargetMode='Internal',
        )
        rels_xml_root.append(new_rel)
        new_rels.append((new_rid, anchor))

        # 设置 r:id 到超链接元素
        hl.set(qn('r:id'), new_rid)
        converted += 1

    print(f"[doc_styler] 已转换 {converted} 个内部超链接为 WPS 兼容格式")


def _get_content_width_dxa(doc: Document) -> int:
    """计算页面内容区宽度（dxa/twips，1 dxa = 635 EMU）。"""
    section = doc.sections[0]
    content_width_emu = section.page_width - section.left_margin - section.right_margin
    return int(content_width_emu / 635)


def style_tables(doc: Document, profile: str = "word-report"):
    """
    表格样式规范化：
    - 黑底白字表头
    - 细灰边框
    - 单元格内边距
    - 首行跨页重复
    - 数据行禁跨页拆分
    - 单元格内边距
    - 列宽按语义 + 内容分析分配
    - 表头居中、数据行左对齐（短列居中）
    """
    for table in doc.tables:
        num_rows = len(table.rows)
        # 使用 tblGrid 列数作为权威列数（兼容合并单元格表格）
        # table.columns 对 gridSpan 表格返回的是首行物理单元格数，可能小于实际列数
        tbl_grid = table._tbl.find(qn('w:tblGrid'))
        if tbl_grid is not None:
            num_cols = len(tbl_grid.findall(qn('w:gridCol')))
        else:
            num_cols = len(table.columns)
        if num_rows == 0:
            continue

        # ---- 表格属性 ----
        tblPr = table._tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            table._tbl.insert(0, tblPr)

        # 首行跨页重复
        for old in tblPr.findall(qn("w:tblHeader")):
            tblPr.remove(old)
        tblHeader = OxmlElement("w:tblHeader")
        tblHeader.set(qn("w:val"), "true")
        tblPr.append(tblHeader)

        # 表格总宽度（绝对宽度 dxa）
        content_width_dxa = _get_content_width_dxa(doc)
        table_width_dxa = int(content_width_dxa * TABLE_WIDTH_PERCENT / 100)
        tblW = OxmlElement("w:tblW")
        tblW.set(qn("w:w"), str(table_width_dxa))
        tblW.set(qn("w:type"), "dxa")
        existing_w = tblPr.find(qn("w:tblW"))
        if existing_w is not None:
            tblPr.remove(existing_w)
        tblPr.append(tblW)

        # 固定列宽布局（防止 Word 自动均分列宽）
        for old_layout in tblPr.findall(qn("w:tblLayout")):
            tblPr.remove(old_layout)
        tblLayout = OxmlElement("w:tblLayout")
        tblLayout.set(qn("w:type"), "fixed")
        tblPr.append(tblLayout)

        # ---- 列宽分配 ----
        # 检测是否为"用户显式指定列宽"的表格（markdown 中 <th style="width:NN%">）
        # 此类表格由 Lua 过滤器加 user-widths bookmark 标记，跳过动态计算
        if _has_user_widths_marker(table):
            _apply_user_column_widths(table, num_cols, table_width_dxa)
        else:
            _apply_column_widths(table, num_cols, table_width_dxa)

        # ---- 确定每列的文本对齐方式 ----
        col_alignments = _infer_column_alignments(table, num_cols)

        # ---- 行处理 ----
        for i, row in enumerate(table.rows):
            is_header = (i == 0)

            for j, cell in enumerate(row.cells):
                # 边框
                _set_cell_border(cell, {
                    "top":    {"val": "single", "sz": "4", "color": "BFBFBF"},
                    "bottom": {"val": "single", "sz": "4", "color": "BFBFBF"},
                    "left":   {"val": "single", "sz": "4", "color": "BFBFBF"},
                    "right":  {"val": "single", "sz": "4", "color": "BFBFBF"},
                })

                # 单元格内边距
                _set_cell_margins(cell, top=40, bottom=40, left=80, right=80)

                # 底纹：统一白色背景（无斑马纹）
                _set_cell_shading(cell, "FFFFFF")

                # 禁止数据行跨页拆分
                if not is_header:
                    tcPr = cell._tc.get_or_add_tcPr()
                    cant_split = OxmlElement("w:cantSplit")
                    cant_split.set(qn("w:val"), "true")
                    existing_cs = tcPr.find(qn("w:cantSplit"))
                    if existing_cs is not None:
                        tcPr.remove(existing_cs)
                    tcPr.append(cant_split)

                # 垂直居中
                tcPr = cell._tc.get_or_add_tcPr()
                vAlign = OxmlElement("w:vAlign")
                vAlign.set(qn("w:val"), "center")
                existing_va = tcPr.find(qn("w:vAlign"))
                if existing_va is not None:
                    tcPr.remove(existing_va)
                tcPr.append(vAlign)

                # 单元格文本样式
                cell_align = col_alignments[j] if not is_header else WD_ALIGN_PARAGRAPH.CENTER
                for para in cell.paragraphs:
                    # 尊重 HTML <td align>/<th align> 显式对齐（pandoc 已写入 w:jc），
                    # 不覆盖用户指定；仅对未指定对齐的单元格应用关键词推断
                    if not _has_explicit_alignment(para):
                        para.alignment = cell_align
                    para.paragraph_format.first_line_indent = Pt(0)  # 表格内禁止首行缩进
                    _set_paragraph_spacing(para, LINE_SPACING_TABLE, space_before=2, space_after=2)
                    for run in _get_all_runs(para):
                        if is_header:
                            _set_font(run, SIZE_SMALL, bold=True, color=COLOR_BLACK,
                                      preserve_user_color=False)
                        else:
                            # 数据行保留用户显式颜色（如 <span style="color:red">已签署</span>）
                            _set_font(run, SIZE_SMALL, color=COLOR_BLACK,
                                      preserve_color=_is_in_hyperlink(run),
                                      preserve_user_color=True)


def _has_explicit_alignment(para) -> bool:
    """检测段落是否已有显式对齐（w:jc）。

    来自 HTML <td align>/<th align> 属性的对齐由 pandoc 写成 w:jc，
    应尊重用户指定，避免被列对齐推断覆盖。
    """
    ppr = para._p.pPr
    if ppr is None:
        return False
    return ppr.find(qn("w:jc")) is not None


def _infer_column_alignments(table, num_cols: int):
    """根据表头关键词推断每列的文本对齐方式。

    仅对未显式指定 align 的单元格生效（显式指定的由 _has_explicit_alignment 保护）。
    """
    alignments = []
    header_row = table.rows[0]
    for j in range(num_cols):
        txt = header_row.cells[j].text.strip().lower()
        if any(k in txt for k in _SEQ_KEYWORDS + _SHORT_KEYWORDS):
            alignments.append(WD_ALIGN_PARAGRAPH.CENTER)
        else:
            alignments.append(WD_ALIGN_PARAGRAPH.LEFT)
    return alignments


def _estimate_text_width(text: str) -> float:
    """估算文本显示宽度（CJK 字符算 2，ASCII 算 1）。"""
    width = 0.0
    for ch in text:
        if ord(ch) > 0x2E80:
            width += 2
        else:
            width += 1
    return width


def _has_user_widths_marker(table) -> bool:
    """检测表格是否被 Lua 过滤器标记为"用户显式指定列宽"。

    Lua 过滤器 html-table-to-ast.lua 在解析到 <th style="width:NN%"> 时，
    会给 Pandoc Table 加 id="user-widths"，pandoc writer 把 id 序列化为
    <w:bookmarkStart w:name="user-widths"/>。检测此 bookmark 即可识别。
    """
    # bookmark 通常出现在 <w:tbl> 之前或 <w:tblPr> 之后
    # 在 table._tbl 前后查找
    parent = table._tbl.getparent()
    if parent is None:
        return False
    # bookmark 在 tbl 的兄弟节点中
    idx = parent.index(table._tbl)
    # 检查 tbl 之前最多 5 个兄弟节点（避免误检远处的 bookmark）
    for i in range(max(0, idx - 5), idx):
        sibling = parent[i]
        if sibling.tag == qn('w:bookmarkStart'):
            name = sibling.get(qn('w:name'))
            if name == 'user-widths':
                return True
    return False


def _apply_user_column_widths(table, num_cols: int, table_width_dxa: int):
    """按用户在 markdown 中显式指定的列宽比例，分配绝对宽度（dxa）。

    与 _apply_column_widths 的内容驱动策略不同：
    - 不分析单元格文本内容
    - 读取 pandoc 已写入的 tblGrid 比例
    - 按比例缩放到 table_width_dxa（保证总和精确等于表格总宽）
    - 写回 tblGrid（覆盖 pandoc 默认的 5000pct 计算结果，统一对齐 A4 内容区）
    - 写入每个单元格的 tcW
    """
    if num_cols == 0 or table_width_dxa <= 0:
        return

    tbl = table._tbl
    tbl_grid = tbl.find(qn('w:tblGrid'))
    if tbl_grid is None:
        # 没有 tblGrid：退回到均分
        per_col = table_width_dxa // num_cols
        col_widths_dxa = [per_col] * num_cols
        # 末列补差
        diff = table_width_dxa - sum(col_widths_dxa)
        col_widths_dxa[-1] += diff
    else:
        grid_cols = tbl_grid.findall(qn('w:gridCol'))
        # 读取 pandoc 写入的相对宽度
        user_widths = [int(gc.get(qn('w:w'))) for gc in grid_cols]
        total = sum(user_widths)
        if total <= 0 or len(user_widths) != num_cols:
            # 异常：退回均分
            per_col = table_width_dxa // num_cols
            col_widths_dxa = [per_col] * num_cols
            col_widths_dxa[-1] += table_width_dxa - sum(col_widths_dxa)
        else:
            # 按用户比例分配绝对宽度
            col_widths_dxa = []
            for i, uw in enumerate(user_widths):
                if i == num_cols - 1:
                    # 末列补差，避免舍入误差
                    remaining = table_width_dxa - sum(col_widths_dxa)
                    col_widths_dxa.append(max(800, remaining))
                else:
                    w = int(table_width_dxa * uw / total)
                    col_widths_dxa.append(max(800, w))

    # 写回 tblGrid
    if tbl_grid is not None:
        tbl.remove(tbl_grid)
    tbl_grid = OxmlElement("w:tblGrid")
    for w in col_widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(w))
        tbl_grid.append(grid_col)
    tblPr = tbl.tblPr
    if tblPr is not None:
        tblPr.addnext(tbl_grid)
    else:
        tbl.insert(0, tbl_grid)

    # 写入每个单元格的 tcW
    for row in table.rows:
        for j, cell in enumerate(row.cells):
            if j >= num_cols:
                continue
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = OxmlElement("w:tcW")
            tcW.set(qn("w:w"), str(col_widths_dxa[j]))
            tcW.set(qn("w:type"), "dxa")
            existing = tcPr.find(qn("w:tcW"))
            if existing is not None:
                tcPr.remove(existing)
            tcPr.append(tcW)

    pct_list = [round(w / table_width_dxa * 100, 1) for w in col_widths_dxa]
    print(f"[doc_styler] 表格 {num_cols} 列用户指定宽度: {pct_list}% (总宽 {table_width_dxa} dxa)")


def _apply_column_widths(table, num_cols: int, table_width_dxa: int):
    """按实际内容宽度动态分配列宽（绝对宽度 dxa）。

    纯内容驱动策略：
    1. 统计每列所有单元格的内容显示宽度（CJK×2）
    2. 以「最大宽度」为主指标（确保最宽内容不被截断）
    3. 以「平均宽度 × 1.5」为辅指标（避免极端长单元格独占过多空间）
    4. 需求宽度 = max(最大宽度, 平均宽度 × 1.5)
    5. 按需求宽度比例分配绝对列宽，设最小列宽保底
    6. 写入 tblGrid + tcW（dxa），配合 tblLayout=fixed 生效
    """
    if num_cols == 0 or table_width_dxa <= 0:
        return

    MIN_COL_DXA = 800   # 最小列宽 ~1.4cm，防止序号等窄列被压扁到不可读
    MAX_COL_PCT = 40    # 单列最大占比40%，防止长内容列独占过多空间
    HEADER_BONUS = 0.3  # 表头文本额外加权系数

    num_rows = len(table.rows)
    if num_rows == 0:
        return

    # ---- 收集每列所有单元格的内容宽度（合并单元格行跳过）----
    col_all_widths = [[] for _ in range(num_cols)]

    def _row_has_merge(row) -> bool:
        """检测行是否包含合并单元格（gridSpan > 1）。

        合并单元格行的内容跨多列，参与列宽计算会导致分摊不均。
        此类行直接跳过，只用非合并行计算各列需求宽度。
        """
        seen_tcs = set()
        for cell in row.cells:
            tc_id = id(cell._tc)
            if tc_id in seen_tcs:
                continue  # python-docx 对 gridSpan 返回的重复引用
            seen_tcs.add(tc_id)
            tcPr = cell._tc.find(qn('w:tcPr'))
            if tcPr is not None:
                gs = tcPr.find(qn('w:gridSpan'))
                if gs is not None:
                    try:
                        if int(gs.get(qn('w:val'), '1')) > 1:
                            return True
                    except (ValueError, TypeError):
                        pass
        return False

    def _collect_row_cells(row, *, is_header: bool = False):
        """收集行内各列内容宽度。

        对 gridSpan > 1 的单元格，内容宽度均摊到所跨各列。
        用 set 去重 python-docx 返回的重复 cell 引用。
        """
        col_idx = 0
        seen_tcs = set()
        for cell in row.cells:
            tc_id = id(cell._tc)
            if tc_id in seen_tcs:
                continue  # 跳过 python-docx 对 gridSpan 的重复引用
            seen_tcs.add(tc_id)

            span = 1
            tcPr = cell._tc.find(qn('w:tcPr'))
            if tcPr is not None:
                gs = tcPr.find(qn('w:gridSpan'))
                if gs is not None:
                    try:
                        span = int(gs.get(qn('w:val'), '1'))
                    except (ValueError, TypeError):
                        pass

            w = _estimate_text_width(cell.text)
            if is_header:
                w *= (1 + HEADER_BONUS)
            per_col_w = w / max(span, 1)
            for s in range(span):
                target_col = col_idx + s
                if target_col < num_cols:
                    col_all_widths[target_col].append(per_col_w)
            col_idx += span

    # 表头行参与计算（即使有合并单元格也参与，按实际位置映射到各列）
    _collect_row_cells(table.rows[0], is_header=True)

    # 数据行（跳过含合并单元格的行，避免分摊不均）
    for i in range(1, num_rows):
        if _row_has_merge(table.rows[i]):
            continue
        _collect_row_cells(table.rows[i])

    # ---- 计算每列需求宽度 ----
    # 字符宽度单位 (CJK×2) 与 dxa 的近似映射：1 字符宽度单位 ≈ 180 dxa
    # 用于确保短内容列（如"条文编号"）有足够宽度避免换行
    CHAR_TO_DXA = 180

    col_demands = []
    col_min_dxa = []  # 每列基于内容的动态最小宽度
    for j in range(num_cols):
        widths = col_all_widths[j]
        if not widths:
            col_demands.append(MIN_COL_DXA)
            col_min_dxa.append(MIN_COL_DXA)
            continue
        max_w = max(widths)
        avg_w = sum(widths) / len(widths)
        # 需求以最大宽度为主（×1.25 留余量），平均值作保底下限
        demand = max(max_w * 1.25, avg_w)
        col_demands.append(demand)
        # 动态最小宽度：至少容纳最宽的单行内容不换线
        # 上限为均分的 150%，避免长内容列独占过多空间
        max_allowed = table_width_dxa / num_cols * 1.5
        col_min_dxa.append(max(MIN_COL_DXA, min(int(max_w * CHAR_TO_DXA), int(max_allowed))))

    total_demand = sum(col_demands)
    if total_demand == 0:
        return

    # ---- 按比例分配绝对宽度 dxa ----
    max_col_dxa = int(table_width_dxa * MAX_COL_PCT / 100)
    col_widths_dxa = []
    for j in range(num_cols):
        w = int(table_width_dxa * col_demands[j] / total_demand)
        # 使用动态最小宽度保底（确保短内容列不换行）
        w = max(w, col_min_dxa[j])
        # 单列不超过上限（防止长内容列独占过多空间）
        w = min(w, max_col_dxa)
        col_widths_dxa.append(w)

    # 若保底总和超出表格宽度，按等比缩放（保底 MIN_COL_DXA）
    actual_sum = sum(col_widths_dxa)
    if actual_sum > table_width_dxa:
        scale = table_width_dxa / actual_sum
        col_widths_dxa = [max(MIN_COL_DXA, int(w * scale)) for w in col_widths_dxa]

    # 兜底：确保每列不低于 MIN_COL_DXA，从最宽列削减补偿
    for j in range(num_cols):
        if col_widths_dxa[j] < MIN_COL_DXA:
            deficit = MIN_COL_DXA - col_widths_dxa[j]
            col_widths_dxa[j] = MIN_COL_DXA
            # 从最宽列削减
            max_idx = max(range(num_cols), key=lambda i: col_widths_dxa[i] if i != j else 0)
            col_widths_dxa[max_idx] = max(MIN_COL_DXA, col_widths_dxa[max_idx] - deficit)

    # 修正舍入误差，使总和精确等于 table_width_dxa
    diff = table_width_dxa - sum(col_widths_dxa)
    if diff != 0 and col_widths_dxa:
        # 将差值分配给最宽的列（视觉影响最小），但不低于 MIN_COL_DXA
        max_idx = max(range(num_cols), key=lambda i: col_widths_dxa[i])
        col_widths_dxa[max_idx] = max(MIN_COL_DXA, col_widths_dxa[max_idx] + diff)
    # ---- 写入 tblGrid（固定布局下列宽的决定性来源）----
    tbl = table._tbl
    existing_grid = tbl.find(qn("w:tblGrid"))
    if existing_grid is not None:
        tbl.remove(existing_grid)
    tblGrid = OxmlElement("w:tblGrid")
    for w in col_widths_dxa:
        gridCol = OxmlElement("w:gridCol")
        gridCol.set(qn("w:w"), str(w))
        tblGrid.append(gridCol)
    tblPr = tbl.tblPr
    if tblPr is not None:
        tblPr.addnext(tblGrid)
    else:
        tbl.insert(0, tblGrid)

    # ---- 写入每个单元格的 tcW ----
    for row in table.rows:
        for j, cell in enumerate(row.cells):
            if j >= num_cols:
                continue
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = OxmlElement("w:tcW")
            tcW.set(qn("w:w"), str(col_widths_dxa[j]))
            tcW.set(qn("w:type"), "dxa")
            existing = tcPr.find(qn("w:tcW"))
            if existing is not None:
                tcPr.remove(existing)
            tcPr.append(tcW)

    # 调试输出
    pct_list = [round(w / table_width_dxa * 100, 1) for w in col_widths_dxa]
    print(f"[doc_styler] 表格 {num_cols} 列宽度分配: {pct_list}% (总宽 {table_width_dxa} dxa)")


def style_body(doc: Document, profile: str = "word-report"):
    """正文段落：1.5 倍行距、设置首行缩进、列表项间距等。"""
    for para in doc.paragraphs:
        level = _estimate_heading_level(para)
        if level > 0:
            continue
        style_name = (para.style.name if para.style else "").lower()
        if "title" in style_name or "heading" in style_name:
            continue

        # 判断是否为列表项
        is_list = _is_list_paragraph(para)

        if is_list:
            _set_paragraph_spacing(para, LINE_SPACING_BODY, space_before=2, space_after=2)
        else:
            _set_paragraph_spacing(para, LINE_SPACING_BODY, space_after=6)


def _is_list_paragraph(para) -> bool:
    """判断段落是否为列表项（有序/无序）。"""
    style_name = (para.style.name if para.style else "").lower()
    if "list" in style_name:
        return True
    text = para.text.strip()
    if len(text) < 2:
        return False
    # 有序列表模式：1. / (1) / ① / a) / A.
    if re.match(r'^[\d]+[.)]\s*', text):
        return True
    if re.match(r'^[a-zA-Z][.)]\s*', text):
        return True
    if re.match(r'^\([0-9a-zA-Z]+\)\s*', text):
        return True
    # 无序列表模式：- /* / •
    if text.startswith(('- ', '* ', '• ')):
        return True
    # 中文序号：一、 / 1）
    if re.match(r'^[一二三四五六七八九十][、）)]\s*', text):
        return True
    return False


def style_footnotes(doc: Document):
    """对脚注文本应用字体样式（宋体 + 小字号 + 紧凑行距）。

    Pandoc 生成的脚注文本默认无显式字体声明，
    需统一设置为宋体 + SIZE_SMALL，与正文风格一致。
    """
    # 查找脚注 part（通过文档关系定位 word/footnotes.xml）
    footnotes_part = None
    for rel in doc.part.rels.values():
        if "footnotes" in rel.reltype:
            footnotes_part = rel.target_part
            break

    if footnotes_part is None:
        return

    # 获取脚注 XML 根元素（兼容 Part / XmlPart 两种加载方式）
    # 通用 Part 需手动 parse blob 并在修改后写回
    from lxml import etree
    is_generic_part = not hasattr(footnotes_part, 'element') and not hasattr(footnotes_part, '_element')
    if hasattr(footnotes_part, 'element'):
        footnotes_root = footnotes_part.element
    elif hasattr(footnotes_part, '_element'):
        footnotes_root = footnotes_part._element
    else:
        footnotes_root = etree.fromstring(footnotes_part.blob)

    sz_val = str(int(SIZE_SMALL * 2))  # 半磅值

    for footnote in footnotes_root.findall(qn("w:footnote")):
        # 跳过分隔符脚注（id=-1 连续分隔符, id=0 普通分隔符）
        fn_id = footnote.get(qn("w:id"))
        if fn_id in ("-1", "0"):
            continue

        for p in footnote.findall(qn("w:p")):
            # 段落行距（紧凑，与表格一致）
            pPr = p.find(qn("w:pPr"))
            if pPr is None:
                pPr = OxmlElement("w:pPr")
                p.insert(0, pPr)
            spacing = pPr.find(qn("w:spacing"))
            if spacing is None:
                spacing = OxmlElement("w:spacing")
                pPr.append(spacing)
            spacing.set(qn("w:line"), str(int(LINE_SPACING_TABLE * 240)))
            spacing.set(qn("w:lineRule"), "auto")

            # 每个 run 设置字体 + 字号（iter 递归包含超链接内 run）
            for r in p.iter(qn("w:r")):
                rPr = r.find(qn("w:rPr"))
                if rPr is None:
                    rPr = OxmlElement("w:rPr")
                    r.insert(0, rPr)

                rFonts = rPr.find(qn("w:rFonts"))
                if rFonts is None:
                    rFonts = OxmlElement("w:rFonts")
                    rPr.insert(0, rFonts)
                rFonts.set(FONT_EAST_ASIA, CN_FONT)
                rFonts.set(FONT_ASCII, EN_FONT)
                rFonts.set(FONT_H_ANSI, EN_FONT)
                rFonts.set(FONT_CS, EN_FONT)

                for tag in ("w:sz", "w:szCs"):
                    elem = rPr.find(qn(tag))
                    if elem is None:
                        elem = OxmlElement(tag)
                        rPr.append(elem)
                    elem.set(qn("w:val"), sz_val)

    # 通用 Part 需将修改后的 XML 写回 blob，否则保存时丢失
    if is_generic_part:
        footnotes_part._blob = etree.tostring(
            footnotes_root, xml_declaration=True, encoding='UTF-8', standalone=True
        )


def add_disclaimer(doc: Document, profile: str = "word-report"):
    """在文末插入 AI 免责声明段落。"""
    if profile != "word-report":
        return

    # 检查是否已存在
    for para in doc.paragraphs:
        if "AI 辅助" in para.text or "不构成正式法律意见" in para.text:
            return

    disclaimer_text = (
        "免责声明：本文档由 AI 辅助生成，仅供参考，不构成正式法律意见。"
        "如涉及重大决策，建议咨询执业律师并依据现行有效法律法规作出独立判断。"
    )
    _add_horizontal_line(doc, color_hex="006D4A", width_pt=0.5)
    para = doc.add_paragraph()
    run = para.add_run(disclaimer_text)
    _set_font(run, SIZE_SMALL, color=COLOR_MID_GRAY)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(para, LINE_SPACING_BODY, space_before=6, space_after=0)


def _enable_field_update(doc: Document) -> None:
    """在 settings.xml 写入 updateFields，让 Word 打开时自动计算域（PAGE/NUMPAGES）。

    不设置此开关时，Word 打开新生成的文档不会自动刷新域，页码显示为空白。
    """
    settings = doc.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def add_footer(doc: Document, doc_name: str = "", footer_format: str = "page-of-total"):
    """添加页码页脚。

    Args:
        doc: Document 对象
        doc_name: 文档名称（用于页脚）
        footer_format: 页脚格式
            - "page-only": 仅页码
            - "name-page": 文档名 + 页码
            - "page-of-total": 页码 / 总页数（默认）
            - "name-page-of-total": 文档名 + 页码 / 总页数
            - "none": 无页脚
    """
    if footer_format == "none":
        return

    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False

        # 清除页脚所有现有内容（含 SDT 内容控件包裹的模板页码）
        ftr = footer._element
        for child in list(ftr):
            ftr.remove(child)

        para = footer.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 构造页脚内容片段
        # 正式文件页脚（页码/文档名）统一黑色，不用灰色
        def _add_text(text):
            run = para.add_run(text)
            _set_font(run, SIZE_SMALL, color=COLOR_BLACK)
            return run

        def _add_page_field(instr=" PAGE "):
            """添加一个域代码（PAGE 或 NUMPAGES）"""
            run = para.add_run()
            fldChar_begin = OxmlElement("w:fldChar")
            fldChar_begin.set(qn("w:fldCharType"), "begin")
            # dirty=true：Word 打开文档时强制重新计算此域，避免页码显示空白
            fldChar_begin.set(qn("w:dirty"), "true")
            run._element.append(fldChar_begin)
            instrText = OxmlElement("w:instrText")
            instrText.set(qn("xml:space"), "preserve")
            instrText.text = instr
            run._element.append(instrText)
            fldChar_end = OxmlElement("w:fldChar")
            fldChar_end.set(qn("w:fldCharType"), "end")
            run._element.append(fldChar_end)
            _set_font(run, SIZE_SMALL, color=COLOR_BLACK)
            return run

        if footer_format == "page-only":
            _add_page_field(" PAGE ")

        elif footer_format == "name-page":
            if doc_name:
                _add_text(f"{doc_name}  |  ")
            _add_page_field(" PAGE ")

        elif footer_format == "page-of-total":
            _add_page_field(" PAGE ")
            _add_text(" / ")
            _add_page_field(" NUMPAGES ")

        elif footer_format == "name-page-of-total":
            if doc_name:
                _add_text(f"{doc_name}  |  ")
            _add_page_field(" PAGE ")
            _add_text(" / ")
            _add_page_field(" NUMPAGES ")

        else:
            # 未知格式，降级为 name-page
            if doc_name:
                _add_text(f"{doc_name}  |  ")
            _add_page_field(" PAGE ")


def validate_output(doc: Document, profile: str = "word-report") -> list[str]:
    """
    交付前自检，返回违规列表（空列表 = 通过）。

    检查项：
    - 无绝对化法律结论
    - 免责声明存在（word-report）
    """
    violations = []

    # 1. 绝对化用语检查
    absolute_terms = ["保证胜诉", "绝无风险", "完全合规", "一定合法", "100%胜诉", "绝对安全"]
    for i, para in enumerate(doc.paragraphs):
        for term in absolute_terms:
            if term in para.text:
                violations.append(f"绝对化用语: 段落#{i} 含\"{term}\"")

    # 3. 免责声明检查
    if profile == "word-report":
        has_disclaimer = any(
            "不构成正式法律意见" in p.text or "AI 辅助" in p.text
            for p in doc.paragraphs
        )
        if not has_disclaimer:
            violations.append("缺少 AI 免责声明")

    return violations


def apply_doc_style(doc: Document, profile: str = "word-report",
                    doc_name: str = "", add_page_numbers: bool = True,
                    footer_format: str = "page-of-total"):
    """
    主入口：对 Document 对象应用完整的正式文档样式规范化。

    Args:
        doc: python-docx Document 对象
        profile: "word-report" 或 "word-revision"
        doc_name: 文档名称（用于页脚）
        add_page_numbers: 是否添加页码（已弃用，保留向后兼容）
        footer_format: 页脚格式
            - "page-only": 仅页码
            - "name-page": 文档名 + 页码
            - "page-of-total": 页码 / 总页数（默认）
            - "name-page-of-total": 文档名 + 页码 / 总页数
            - "none": 无页脚

    Returns:
        list[str]: 校验违规列表（空 = 全部通过）

    Usage:
        doc = Document("draft.docx")
        violations = apply_doc_style(doc, "word-report", "租赁合同法律研究报告")
        if violations:
            print("违规:", violations)
        doc.save("final.docx")
    """
    # 1. 页面设置
    setup_page(doc, profile)

    # 1.1 启用 Word 打开时自动更新域（PAGE/NUMPAGES 页码）
    # 不设置时 Word 默认不刷新域，页码显示为空白
    _enable_field_update(doc)

    # 1.5 移除水平线（---）残留的 VML 矩形，避免横线影响后续样式判断
    remove_horizontal_rules(doc)

    # 2. 字体规范化（必须最先执行，确保所有 run 有显式字体声明）
    normalize_fonts(doc, profile)

    # 3. 标题样式
    style_headings(doc, profile)

    # 3.5 超链接样式（蓝色 + 下划线，所有 profile 生效）
    style_hyperlinks(doc)

    # 注：Pandoc 生成的内部锚点跳转使用 w:anchor 属性（OOXML 标准），
    # Word 可正常跳转，WPS 部分版本可能不支持。如需 WPS 完全兼容，
    # 后续可在 doc.save() 后通过 ZIP 后处理将 w:anchor 转为 r:id 关系引用。

    # 4. 正文样式
    style_body(doc, profile)

    # 5. 表格样式
    style_tables(doc, profile)

    # 5.5 脚注样式（宋体 + 小字号）
    style_footnotes(doc)

    # 6. 免责声明
    add_disclaimer(doc, profile)

    # 7. 页脚（页码 + 文档名）
    # add_page_numbers 为 False 时退化为 "none"（向后兼容）
    fmt = "none" if not add_page_numbers else footer_format
    add_footer(doc, doc_name, fmt)

    # 8. 交付前校验
    return validate_output(doc, profile)
