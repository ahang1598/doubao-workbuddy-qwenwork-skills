#!/usr/bin/env python3
"""Validate a generated cross-examination opinion DOCX against its source JSON."""

from __future__ import annotations

import argparse
import json
import re
import sys
import zipfile
from pathlib import Path
from typing import Any

try:
    from docx import Document
except ImportError as exc:  # pragma: no cover - environment-specific
    raise SystemExit("缺少 python-docx。请先安装 python-docx，再校验 DOCX。") from exc


DIMENSIONS = {
    "authenticity": "真实性",
    "legality": "合法性",
    "relevance": "关联性",
    "probative_force": "证明力",
}
PLACEHOLDER_RE = re.compile(
    r"(?:\[(?:待填写|待补充|填写[^\]]*|请输入[^\]]*)\]"
    r"|【(?:待填写|待补充|填写[^】]*|请输入[^】]*)】"
    r"|\b(?:TODO|TBD|XXX)\b|_{4,})",
    re.IGNORECASE,
)
MIN_DOCX_BYTES = 5_000


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--input", required=True, type=Path, help="Source UTF-8 JSON")
    parser.add_argument("--docx", required=True, type=Path, help="Generated .docx")
    return parser.parse_args()


def load_input(path: Path) -> dict[str, Any]:
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError) as exc:
        raise ValueError(f"无法读取输入 JSON：{exc}") from exc
    if not isinstance(data, dict):
        raise ValueError("JSON 顶层必须是对象")
    items = data.get("evidence_items")
    if not isinstance(items, list) or not items:
        raise ValueError("evidence_items 必须是非空数组")
    return data


def validate_package(path: Path, errors: list[str]) -> None:
    if not path.is_file():
        errors.append(f"文件不存在：{path}")
        return
    if path.suffix.lower() != ".docx":
        errors.append("文件扩展名不是 .docx")
    if path.stat().st_size < MIN_DOCX_BYTES:
        errors.append(f"文件过小：{path.stat().st_size} bytes")
    if not zipfile.is_zipfile(path):
        errors.append("文件不是有效 ZIP/OOXML 包")
        return
    try:
        with zipfile.ZipFile(path) as package:
            broken = package.testzip()
            if broken:
                errors.append(f"ZIP 成员损坏：{broken}")
            names = set(package.namelist())
            for required in ("[Content_Types].xml", "word/document.xml", "word/styles.xml"):
                if required not in names:
                    errors.append(f"缺少 OOXML 必需成员：{required}")
    except (OSError, zipfile.BadZipFile) as exc:
        errors.append(f"无法读取 OOXML 包：{exc}")


def extract_document(path: Path, errors: list[str]) -> Any | None:
    try:
        return Document(path)
    except Exception as exc:  # python-docx may raise several parser exceptions
        errors.append(f"python-docx 无法打开文档：{exc}")
        return None


def all_text(document: Any) -> str:
    blocks = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            blocks.extend(cell.text.strip() for cell in row.cells if cell.text.strip())
    return "\n".join(blocks)


def find_overview_table(document: Any) -> Any | None:
    expected = ["编号", "名称", "真实性", "合法性", "关联性", "证明力", "单项结论"]
    for table in document.tables:
        if not table.rows:
            continue
        header = [cell.text.strip() for cell in table.rows[0].cells]
        if header == expected:
            return table
    return None


def validate_table(table: Any, items: list[dict[str, Any]], errors: list[str]) -> None:
    rows = table.rows[1:]
    if len(rows) != len(items):
        errors.append(f"速览表证据数量不一致：表格 {len(rows)}，输入 {len(items)}")
        return
    actual_ids: list[str] = []
    by_id: dict[str, list[str]] = {}
    for row in rows:
        values = [cell.text.strip() for cell in row.cells]
        if len(values) != 7:
            errors.append("速览表存在列数不为 7 的数据行")
            continue
        actual_ids.append(values[0])
        by_id[values[0]] = values
    expected_ids = [str(item.get("id", "")).strip() for item in items]
    if actual_ids != expected_ids:
        errors.append(f"速览表证据编号或顺序不一致：{actual_ids} != {expected_ids}")

    dimension_keys = list(DIMENSIONS)
    for item in items:
        evidence_id = item["id"].strip()
        row = by_id.get(evidence_id)
        if row is None:
            errors.append(f"速览表缺少证据：{evidence_id}")
            continue
        expected = [
            evidence_id,
            item["name"].strip(),
            *[item["dimensions"][key]["status"].strip() for key in dimension_keys],
            item["conclusion"].strip(),
        ]
        if row != expected:
            errors.append(f"速览表与输入不一致：{evidence_id}")


def validate_body_and_summary(document: Any, items: list[dict[str, Any]], errors: list[str]) -> None:
    paragraph_texts = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    paragraph_set = set(paragraph_texts)
    for item in items:
        evidence_id = item["id"].strip()
        name = item["name"].strip()
        conclusion = item["conclusion"].strip()
        required_lines = [
            f"证据编号：{evidence_id}",
            f"证据名称：{name}",
            f"单项结论：{conclusion}",
            f"{evidence_id}（{name}）：{conclusion}",
        ]
        for key, label in DIMENSIONS.items():
            detail = item["dimensions"][key]
            required_lines.append(f"{label}（{detail['status'].strip()}）：{detail['reason'].strip()}")
        for line in required_lines:
            if line not in paragraph_set:
                errors.append(f"正文或总结缺少/不一致：{line[:80]}")


def validate_core_fields(document: Any, data: dict[str, Any], errors: list[str]) -> None:
    paragraph_set = {p.text.strip() for p in document.paragraphs if p.text.strip()}
    checks = [
        f"致送法院：{data.get('court_name', '').strip()}",
        f"案号：{data.get('case_no', '').strip()}",
        f"总体请求：{data.get('overall_conclusion', '').strip()}",
        f"日期：{data.get('submission_date', '').strip()}",
    ]
    for line in checks:
        if line not in paragraph_set:
            errors.append(f"核心字段缺少/不一致：{line[:80]}")


def validate_page_size(document: Any, errors: list[str]) -> None:
    if not document.sections:
        errors.append("文档没有页面节")
        return
    section = document.sections[0]
    width_cm = section.page_width.cm
    height_cm = section.page_height.cm
    if abs(width_cm - 21.0) > 0.1 or abs(height_cm - 29.7) > 0.1:
        errors.append(f"页面不是 A4：{width_cm:.2f}cm × {height_cm:.2f}cm")


def validate_heading_colors(document: Any, errors: list[str]) -> None:
    for style_name in ("Title", "Heading 1", "Heading 2"):
        color = document.styles[style_name].font.color.rgb
        if color is None or str(color).upper() != "000000":
            errors.append(f"标题样式不是固定黑色：{style_name}")


def main() -> int:
    args = parse_args()
    errors: list[str] = []
    try:
        data = load_input(args.input)
    except ValueError as exc:
        print(f"校验失败：{exc}", file=sys.stderr)
        return 1

    validate_package(args.docx, errors)
    if errors and (not args.docx.is_file() or not zipfile.is_zipfile(args.docx)):
        for error in errors:
            print(f"- {error}", file=sys.stderr)
        return 1

    document = extract_document(args.docx, errors)
    if document is not None:
        text = all_text(document)
        if not text.strip():
            errors.append("文档没有可见文本")
        placeholder = PLACEHOLDER_RE.search(text)
        if placeholder:
            errors.append(f"文档包含占位符：{placeholder.group(0)}")

        table = find_overview_table(document)
        if table is None:
            errors.append("未找到质证意见速览表")
        else:
            validate_table(table, data["evidence_items"], errors)
        validate_body_and_summary(document, data["evidence_items"], errors)
        validate_core_fields(document, data, errors)
        validate_page_size(document, errors)
        validate_heading_colors(document, errors)

    if errors:
        print("校验失败：", file=sys.stderr)
        for error in errors:
            print(f"- {error}", file=sys.stderr)
        return 1

    result = {
        "status": "passed",
        "docx": str(args.docx.resolve()),
        "bytes": args.docx.stat().st_size,
        "evidence_count": len(data["evidence_items"]),
        "checks": [
            "valid_ooxml_package",
            "openable_by_python_docx",
            "nonempty",
            "a4_page",
            "black_heading_styles",
            "evidence_coverage",
            "table_body_summary_consistency",
            "four_dimensions",
            "no_placeholders",
        ],
    }
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
