#!/usr/bin/env python3
"""Generate a formal cross-examination opinion DOCX from structured JSON."""

from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path
from typing import Any

try:
    from docx import Document
    from docx.enum.section import WD_SECTION_START
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor
except ImportError as exc:  # pragma: no cover - environment-specific
    raise SystemExit("缺少 python-docx。请先安装 python-docx，再生成 DOCX。") from exc


DIMENSIONS = {
    "authenticity": "真实性",
    "legality": "合法性",
    "relevance": "关联性",
    "probative_force": "证明力",
}
ALLOWED_STATUSES = {"无异议", "有异议", "待核验", "不适用"}
PLACEHOLDER_RE = re.compile(
    r"(?:\[(?:待填写|待补充|填写[^\]]*|请输入[^\]]*)\]"
    r"|【(?:待填写|待补充|填写[^】]*|请输入[^】]*)】"
    r"|\b(?:TODO|TBD|XXX)\b|_{4,})",
    re.IGNORECASE,
)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--input", required=True, type=Path, help="UTF-8 JSON input")
    parser.add_argument("--output", required=True, type=Path, help="Output .docx path")
    return parser.parse_args()


def load_json(path: Path) -> dict[str, Any]:
    if not path.is_file():
        raise ValueError(f"输入文件不存在：{path}")
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError) as exc:
        raise ValueError(f"无法读取 JSON：{exc}") from exc
    if not isinstance(data, dict):
        raise ValueError("JSON 顶层必须是对象")
    return data


def required_text(obj: dict[str, Any], key: str, context: str) -> str:
    value = obj.get(key)
    if not isinstance(value, str) or not value.strip():
        raise ValueError(f"{context}.{key} 必须是非空字符串")
    return value.strip()


def text_list(obj: dict[str, Any], key: str, context: str) -> list[str]:
    value = obj.get(key, [])
    if value is None:
        return []
    if not isinstance(value, list) or any(not isinstance(x, str) or not x.strip() for x in value):
        raise ValueError(f"{context}.{key} 必须是非空字符串数组")
    return [x.strip() for x in value]


def find_placeholders(value: Any, path: str = "$") -> list[str]:
    findings: list[str] = []
    if isinstance(value, str) and PLACEHOLDER_RE.search(value):
        findings.append(path)
    elif isinstance(value, dict):
        for key, child in value.items():
            findings.extend(find_placeholders(child, f"{path}.{key}"))
    elif isinstance(value, list):
        for index, child in enumerate(value):
            findings.extend(find_placeholders(child, f"{path}[{index}]"))
    return findings


def validate_data(data: dict[str, Any]) -> list[dict[str, Any]]:
    for key in ("court_name", "case_no", "overall_conclusion", "submission_date"):
        required_text(data, key, "root")

    submitter = data.get("submitter")
    if not isinstance(submitter, dict):
        raise ValueError("root.submitter 必须是对象")
    required_text(submitter, "role", "submitter")
    required_text(submitter, "name", "submitter")

    items = data.get("evidence_items")
    if not isinstance(items, list) or not items:
        raise ValueError("root.evidence_items 必须是非空数组")

    seen: set[str] = set()
    normalized: list[dict[str, Any]] = []
    for index, item in enumerate(items):
        context = f"evidence_items[{index}]"
        if not isinstance(item, dict):
            raise ValueError(f"{context} 必须是对象")
        for key in ("id", "name", "submitter", "purpose", "locator", "conclusion"):
            required_text(item, key, context)
        evidence_id = item["id"].strip()
        if evidence_id in seen:
            raise ValueError(f"证据编号重复：{evidence_id}")
        seen.add(evidence_id)

        dimensions = item.get("dimensions")
        if not isinstance(dimensions, dict):
            raise ValueError(f"{context}.dimensions 必须是对象")
        if set(dimensions) != set(DIMENSIONS):
            missing = set(DIMENSIONS) - set(dimensions)
            extra = set(dimensions) - set(DIMENSIONS)
            raise ValueError(f"{context}.dimensions 键不正确；缺少 {sorted(missing)}，多出 {sorted(extra)}")
        for key in DIMENSIONS:
            detail = dimensions[key]
            if not isinstance(detail, dict):
                raise ValueError(f"{context}.dimensions.{key} 必须是对象")
            status = required_text(detail, "status", f"{context}.dimensions.{key}")
            required_text(detail, "reason", f"{context}.dimensions.{key}")
            if status not in ALLOWED_STATUSES:
                raise ValueError(f"{context}.dimensions.{key}.status 不在允许值中：{status}")

        text_list(item, "procedural_requests", context)
        text_list(item, "counter_evidence_suggestions", context)
        recognized = item.get("recognized_parts", "")
        if not isinstance(recognized, str):
            raise ValueError(f"{context}.recognized_parts 必须是字符串")
        normalized.append(item)

    text_list(data, "procedural_applications", "root")
    placeholders = find_placeholders(data)
    if placeholders:
        raise ValueError("输入包含占位符：" + ", ".join(placeholders))
    return normalized


def set_run_font(run: Any, name: str = "宋体", size: float = 12) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        fonts.set(qn(key), name)


def configure_styles(document: Any) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "宋体"
    normal.font.size = Pt(12)
    normal_fonts = normal._element.get_or_add_rPr().get_or_add_rFonts()
    for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        normal_fonts.set(qn(key), "宋体")
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(0)

    for style_name, font_name, size in (
        ("Title", "黑体", 22),
        ("Heading 1", "黑体", 15),
        ("Heading 2", "黑体", 14),
    ):
        style = document.styles[style_name]
        style.font.name = font_name
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(0, 0, 0)
        style_fonts = style._element.get_or_add_rPr().get_or_add_rFonts()
        for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
            style_fonts.set(qn(key), font_name)


def configure_page(document: Any) -> None:
    section = document.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.6)
    section.bottom_margin = Cm(2.6)
    section.left_margin = Cm(2.7)
    section.right_margin = Cm(2.7)


def add_text(document: Any, text: str, *, bold_label: str | None = None, align: Any = None) -> Any:
    paragraph = document.add_paragraph()
    if align is not None:
        paragraph.alignment = align
    if bold_label and text.startswith(bold_label):
        label_run = paragraph.add_run(bold_label)
        label_run.bold = True
        set_run_font(label_run)
        body_run = paragraph.add_run(text[len(bold_label):])
        set_run_font(body_run)
    else:
        run = paragraph.add_run(text)
        set_run_font(run)
    return paragraph


def keep_table_row_together(row: Any) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def repeat_table_header(row: Any) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def format_cell(cell: Any, text: str, *, bold: bool = False, size: float = 8) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = bold
    set_run_font(run, size=size)


def add_overview_table(document: Any, items: list[dict[str, Any]]) -> None:
    headers = ["编号", "名称", "真实性", "合法性", "关联性", "证明力", "单项结论"]
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    repeat_table_header(table.rows[0])
    for cell, header in zip(table.rows[0].cells, headers):
        format_cell(cell, header, bold=True)

    for item in items:
        row = table.add_row()
        keep_table_row_together(row)
        values = [
            item["id"].strip(),
            item["name"].strip(),
            item["dimensions"]["authenticity"]["status"].strip(),
            item["dimensions"]["legality"]["status"].strip(),
            item["dimensions"]["relevance"]["status"].strip(),
            item["dimensions"]["probative_force"]["status"].strip(),
            item["conclusion"].strip(),
        ]
        for cell, value in zip(row.cells, values):
            format_cell(cell, value)


def add_evidence_body(document: Any, items: list[dict[str, Any]]) -> None:
    for index, item in enumerate(items, start=1):
        document.add_heading(f"{index}. {item['id']}：{item['name']}", level=2)
        add_text(document, f"证据编号：{item['id']}", bold_label="证据编号：")
        add_text(document, f"证据名称：{item['name']}", bold_label="证据名称：")
        add_text(document, f"提交方：{item['submitter']}", bold_label="提交方：")
        add_text(document, f"证明目的：{item['purpose']}", bold_label="证明目的：")
        add_text(document, f"材料定位：{item['locator']}", bold_label="材料定位：")
        if item.get("evidence_type"):
            add_text(document, f"证据类型：{item['evidence_type']}", bold_label="证据类型：")
        recognized = item.get("recognized_parts", "").strip()
        add_text(document, f"认可部分：{recognized or '无单独认可部分。'}", bold_label="认可部分：")

        for key, label in DIMENSIONS.items():
            detail = item["dimensions"][key]
            add_text(
                document,
                f"{label}（{detail['status']}）：{detail['reason']}",
                bold_label=f"{label}（{detail['status']}）：",
            )
        add_text(document, f"单项结论：{item['conclusion']}", bold_label="单项结论：")

        requests = text_list(item, "procedural_requests", f"evidence {item['id']}")
        if requests:
            add_text(document, "本项程序性请求：", bold_label="本项程序性请求：")
            for request in requests:
                add_text(document, f"• {request}")


def add_procedural_section(document: Any, data: dict[str, Any], items: list[dict[str, Any]]) -> None:
    requests = text_list(data, "procedural_applications", "root")
    for item in items:
        for request in text_list(item, "procedural_requests", f"evidence {item['id']}"):
            tagged = f"{item['id']}：{request}"
            if tagged not in requests:
                requests.append(tagged)
    document.add_heading("三、程序性申请", level=1)
    if not requests:
        add_text(document, "根据现有材料，暂无单独程序性申请。")
    else:
        for index, request in enumerate(requests, start=1):
            add_text(document, f"{index}. {request}")


def add_counter_evidence_section(document: Any, items: list[dict[str, Any]]) -> None:
    suggestions: list[tuple[str, str]] = []
    for item in items:
        for suggestion in text_list(item, "counter_evidence_suggestions", f"evidence {item['id']}"):
            suggestions.append((item["id"], suggestion))
    if not suggestions:
        return
    document.add_heading("五、反证与核验建议", level=1)
    for index, (evidence_id, suggestion) in enumerate(suggestions, start=1):
        add_text(document, f"{index}. 针对{evidence_id}：{suggestion}")


def build_document(data: dict[str, Any], items: list[dict[str, Any]]) -> Any:
    document = Document()
    configure_page(document)
    configure_styles(document)
    document.core_properties.title = data.get("title", "质证意见")
    document.core_properties.subject = "民事诉讼质证意见"

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(data.get("title", "质证意见"))
    run.bold = True
    set_run_font(run, name="黑体", size=22)

    submitter = data["submitter"]
    add_text(document, f"致送法院：{data['court_name']}", bold_label="致送法院：")
    add_text(document, f"案号：{data['case_no']}", bold_label="案号：")
    if data.get("case_type"):
        add_text(document, f"案由：{data['case_type']}", bold_label="案由：")
    add_text(document, f"提交人：{submitter['role']} {submitter['name']}", bold_label="提交人：")
    if submitter.get("agent"):
        add_text(document, f"代理人：{submitter['agent']}", bold_label="代理人：")
    if submitter.get("contact"):
        add_text(document, f"联系方式：{submitter['contact']}", bold_label="联系方式：")
    if data.get("opposing_party"):
        add_text(document, f"证据提交方：{data['opposing_party']}", bold_label="证据提交方：")
    if data.get("scope_note"):
        add_text(document, f"材料范围及局限：{data['scope_note']}", bold_label="材料范围及局限：")

    document.add_heading("一、质证意见速览", level=1)
    add_overview_table(document, items)

    document.add_heading("二、逐项质证意见", level=1)
    add_evidence_body(document, items)

    add_procedural_section(document, data, items)

    document.add_heading("四、质证总结论", level=1)
    for item in items:
        add_text(document, f"{item['id']}（{item['name']}）：{item['conclusion']}")
    add_text(document, f"总体请求：{data['overall_conclusion']}", bold_label="总体请求：")

    add_counter_evidence_section(document, items)

    add_text(document, "")
    add_text(document, f"提交人：{submitter['name']}", align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_text(document, f"日期：{data['submission_date']}", align=WD_ALIGN_PARAGRAPH.RIGHT)
    return document


def main() -> int:
    args = parse_args()
    try:
        data = load_json(args.input)
        items = validate_data(data)
        if args.output.suffix.lower() != ".docx":
            raise ValueError("输出文件扩展名必须是 .docx")
        args.output.parent.mkdir(parents=True, exist_ok=True)
        document = build_document(data, items)
        document.save(args.output)
        if not args.output.is_file() or args.output.stat().st_size == 0:
            raise ValueError("DOCX 未成功写入")
    except (OSError, ValueError) as exc:
        print(f"生成失败：{exc}", file=sys.stderr)
        return 1
    print(f"生成成功：{args.output.resolve()}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
