#!/usr/bin/env python3
"""Generate a DOCX evidence list from structured data.

Usage:
  python3 scripts/generate_evidence_docx.py \
    --title "证据清单（原告提供）" \
    --evidence "2021年3月20日签署的协议书、46万元打款凭证|1-3|复印件|1. 证明原被告之间存在居间关系...|2. 证明原告向被告支付了46万元..." \
    --evidence "2022年9月4日签署的协议书、13.4万元收款凭证|4-8|复印件|1. 证明双方约定被告应按约返还...|2. 证明被告仅返还原告13.4万元..." \
    --submitter "张三律师" \
    --submit-date "2025年4月7日" \
    --output "证据清单.docx"

Evidence format (pipe-separated):
  证据名称|页码|原/复印件|证明对象1|证明对象2|...

Or use JSON input:
  echo '{"title": "证据清单", "evidence": [...], "submitter": "...", "submit_date": "..."}' | \
    python3 scripts/generate_evidence_docx.py --output "证据清单.docx"
"""

from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path
from typing import List, Dict, Optional

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Pt, Cm
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ImportError as exc:
    raise SystemExit(
        "Missing dependency: python-docx. Install it with: "
        "python3 -m pip install python-docx"
    ) from exc


def set_run_font(run, size: int = 12, bold: bool = False, font_name: str = "宋体") -> None:
    """设置字体样式"""
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_paragraph_spacing(paragraph, before: float = 0, after: float = 6, line: float = 1.5) -> None:
    """设置段落间距"""
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def add_title(document: Document, title: str) -> None:
    """添加标题（居中，黑体，二号）"""
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(paragraph, before=0, after=20, line=1.2)
    run = paragraph.add_run(title)
    set_run_font(run, size=22, bold=True, font_name="黑体")


def add_table_header(document: Document) -> None:
    """添加表格表头"""
    headers = ["序号", "证据名称", "页码", "原/复印件", "证明对象"]
    
    table = document.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    
    # 设置列宽（序号 8%、证据名称 30%、页码 10%、原/复印件 12%、证明对象 40%）
    widths = [Cm(1.8), Cm(6.8), Cm(2.3), Cm(2.7), Cm(9.0)]
    for idx, width in enumerate(widths):
        table.columns[idx].width = width
    
    # 填充表头
    header_row = table.rows[0]
    for idx, header in enumerate(headers):
        cell = header_row.cells[idx]
        cell.text = header
        
        # 设置表头样式
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(paragraph, before=6, after=6, line=1.2)
        
        for run in paragraph.runs:
            set_run_font(run, size=12, bold=True, font_name="宋体")
    
    return table


def add_evidence_row(table, seq: int, name: str, page: str, copy_type: str, proof_objects: List[str]) -> None:
    """添加证据行"""
    row = table.add_row()
    
    # 序号
    cell = row.cells[0]
    cell.text = str(seq)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(paragraph, before=4, after=4, line=1.5)
    for run in paragraph.runs:
        set_run_font(run, size=12, font_name="宋体")
    
    # 证据名称
    cell = row.cells[1]
    cell.text = name
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(paragraph, before=4, after=4, line=1.5)
    for run in paragraph.runs:
        set_run_font(run, size=12, font_name="宋体")
    
    # 页码
    cell = row.cells[2]
    cell.text = page
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(paragraph, before=4, after=4, line=1.5)
    for run in paragraph.runs:
        set_run_font(run, size=12, font_name="宋体")
    
    # 原/复印件
    cell = row.cells[3]
    cell.text = copy_type
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(paragraph, before=4, after=4, line=1.5)
    for run in paragraph.runs:
        set_run_font(run, size=12, font_name="宋体")
    
    # 证明对象（可能多项）
    cell = row.cells[4]
    cell.text = ""  # 清空默认文本
    
    if proof_objects:
        # 合并证明对象
        proof_text = "\n".join([f"{i+1}. {obj}" for i, obj in enumerate(proof_objects)])
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_paragraph_spacing(paragraph, before=4, after=4, line=1.5)
        
        # 分条添加
        run = paragraph.add_run(proof_text)
        set_run_font(run, size=12, font_name="宋体")
    else:
        paragraph = cell.paragraphs[0]
        set_paragraph_spacing(paragraph, before=4, after=4, line=1.5)


def add_footer(document: Document, submitter: str, submit_date: str) -> None:
    """添加底部提交人和日期"""
    # 空行
    document.add_paragraph()
    
    # 提交人
    paragraph = document.add_paragraph()
    set_paragraph_spacing(paragraph, before=0, after=6, line=1.5)
    run = paragraph.add_run(f"提交人（代理律师）：{submitter}")
    set_run_font(run, size=12, font_name="宋体")
    
    # 提交日期
    paragraph = document.add_paragraph()
    set_paragraph_spacing(paragraph, before=0, after=0, line=1.5)
    run = paragraph.add_run(f"提交日期：{submit_date}")
    set_run_font(run, size=12, font_name="宋体")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Generate DOCX evidence list.")
    parser.add_argument("--title", type=str, default="证据清单（原告提供）", help="标题")
    parser.add_argument("--evidence", action="append", type=str, help="证据条目（管道分隔格式）")
    parser.add_argument("--submitter", type=str, default="", help="提交人")
    parser.add_argument("--submit-date", type=str, default="", help="提交日期")
    parser.add_argument("--output", "-o", type=Path, required=True, help="输出 .docx 路径")
    parser.add_argument("--json-input", type=Path, help="JSON 输入文件")
    return parser.parse_args()


def parse_evidence_pipe(pipe_str: str) -> Dict:
    """解析管道分隔的证据字符串"""
    parts = pipe_str.split("|")
    if len(parts) < 4:
        raise ValueError(f"证据格式错误，至少需要4个字段（证据名称|页码|原/复印件|证明对象）：{pipe_str}")
    
    name = parts[0].strip()
    page = parts[1].strip()
    copy_type = parts[2].strip()
    proof_objects = [p.strip() for p in parts[3:] if p.strip()]
    
    return {
        "name": name,
        "page": page,
        "copy_type": copy_type,
        "proof_objects": proof_objects
    }


def main() -> None:
    args = parse_args()
    
    # 读取输入数据
    evidence_list = []
    title = args.title
    submitter = args.submitter
    submit_date = args.submit_date
    
    # 优先从 JSON 文件读取
    if args.json_input:
        with open(args.json_input, 'r', encoding='utf-8') as f:
            data = json.load(f)
        title = data.get("title", title)
        submitter = data.get("submitter", submitter)
        submit_date = data.get("submit_date", submit_date)
        
        for item in data.get("evidence", []):
            if isinstance(item, str):
                evidence_list.append(parse_evidence_pipe(item))
            elif isinstance(item, dict):
                evidence_list.append({
                    "name": item.get("name", ""),
                    "page": item.get("page", ""),
                    "copy_type": item.get("copy_type", "复印件"),
                    "proof_objects": item.get("proof_objects", [])
                })
    else:
        # 从命令行参数读取
        if args.evidence:
            for pipe_str in args.evidence:
                evidence_list.append(parse_evidence_pipe(pipe_str))
    
    if not evidence_list:
        # 尝试从 stdin 读取 JSON
        try:
            stdin_data = sys.stdin.read()
            if stdin_data.strip():
                data = json.loads(stdin_data)
                title = data.get("title", title)
                submitter = data.get("submitter", submitter)
                submit_date = data.get("submit_date", submit_date)
                
                for item in data.get("evidence", []):
                    if isinstance(item, str):
                        evidence_list.append(parse_evidence_pipe(item))
                    elif isinstance(item, dict):
                        evidence_list.append({
                            "name": item.get("name", ""),
                            "page": item.get("page", ""),
                            "copy_type": item.get("copy_type", "复印件"),
                            "proof_objects": item.get("proof_objects", [])
                        })
        except (json.JSONDecodeError, Exception):
            pass
    
    if not evidence_list:
        raise SystemExit("No evidence data provided. Use --evidence, --json-input, or pipe JSON through stdin.")
    
    # 生成 Word 文档
    document = Document()
    
    # 设置页面边距
    section = document.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)
    
    # 添加标题
    add_title(document, title)
    
    # 添加表格
    table = add_table_header(document)
    
    # 添加证据行
    for idx, evidence in enumerate(evidence_list, 1):
        add_evidence_row(
            table,
            seq=idx,
            name=evidence["name"],
            page=evidence["page"],
            copy_type=evidence["copy_type"],
            proof_objects=evidence["proof_objects"]
        )
    
    # 添加底部信息
    add_footer(document, submitter, submit_date)
    
    # 保存文档
    args.output.parent.mkdir(parents=True, exist_ok=True)
    document.save(args.output)
    print(f"Evidence list DOCX generated: {args.output}")


if __name__ == "__main__":
    main()
