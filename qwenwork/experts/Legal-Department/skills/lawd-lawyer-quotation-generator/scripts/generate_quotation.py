#!/usr/bin/env python3
"""
律师法律服务报价方案 DOCX 生成脚本

用法：
    python3 scripts/generate_quotation.py --json '<JSON数据或文件路径>' [--output <输出路径>]

JSON 数据结构说明见脚本底部的 EXAMPLE_JSON。
"""

import argparse
import json
import sys
import os
from datetime import datetime
from pathlib import Path
from typing import List, Tuple

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("缺少依赖 python-docx：请先执行 pip3 install python-docx", file=sys.stderr)
    sys.exit(1)


class InputError(Exception):
    """入参错误，退出码 2。"""


# ── 路径解析 ──────────────────────────────────────────────
# 本脚本位于 <技能根目录>/scripts/ 下，数据/配置文件（如 quotation_data.json）
# 仍存放在技能根目录。这里基于 __file__ 向上定位技能根目录，
# 保证无论从哪个工作目录调用脚本，都能解析到技能自带的数据文件。

SKILL_ROOT = Path(__file__).resolve().parent.parent


# ── 样式常量 ──────────────────────────────────────────────

FONT_TITLE = "方正小标宋简体"
FONT_HEADING = "黑体"
FONT_BODY = "仿宋"
FONT_BODY_FALLBACK = "宋体"

COLOR_DARK = RGBColor(0x00, 0x00, 0x00)
COLOR_ACCENT = RGBColor(0x1F, 0x49, 0x7D)  # 深蓝，用于标题


# ── 工具函数 ──────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    """设置单元格底色"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val.get("val", "single")}" '
            f'w:sz="{val.get("sz", "4")}" w:space="0" '
            f'w:color="{val.get("color", "000000")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)


def set_run_font(run, font_cn, size_pt, bold=False, color=None):
    """统一设置 run 的中英文字体"""
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.name = font_cn
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_cn)
    if color:
        run.font.color.rgb = color


def add_paragraph_styled(doc, text, font_cn=FONT_BODY, size_pt=12,
                         bold=False, alignment=None, color=None,
                         space_before=0, space_after=0, first_line_indent=None):
    """添加一个带样式的段落"""
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(22)
    if first_line_indent is not None:
        p.paragraph_format.first_line_indent = Cm(first_line_indent)
    run = p.add_run(text)
    set_run_font(run, font_cn, size_pt, bold=bold, color=color)
    return p


def add_table_with_data(doc, headers, rows, col_widths=None):
    """添加一个格式化的表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # 表头
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(header))
        set_run_font(run, FONT_HEADING, 11, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        set_cell_shading(cell, "1F497D")

    # 数据行
    for r_idx, row_data in enumerate(rows):
        row_data = row_data if isinstance(row_data, (list, tuple)) else [row_data]
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[1 + r_idx].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(cell_text))
            first_cell = str(row_data[0]) if row_data else ""
            is_bold = first_cell.startswith("**") or first_cell.startswith("风险") or first_cell.startswith("成果")
            clean_text = str(cell_text).replace("**", "")
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(clean_text)
            set_run_font(run, FONT_BODY, 11, bold=is_bold)
            # 斑马纹
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F2F2F2")

    # 列宽
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    return table


# ── 主生成逻辑 ────────────────────────────────────────────

def filter_services_by_case_type(services, case_type):
    """根据案件类型过滤服务内容（已禁用自动过滤，忠实呈现用户输入）
    
    Args:
        services: 服务内容列表
        case_type: 案件类型（目前不再用于自动过滤，仅作为元数据保留）
    
    Returns:
        原始服务内容列表（仅过滤掉 items 为空的组）
    """
    if not services:
        return []
    
    # 【修复】移除基于关键词的自动过滤逻辑，避免误删用户明确要求的服务组
    # 现在脚本将忠实呈现用户在 JSON 中提供的所有服务组
    filtered = [s for s in services if s.get("items")]
    
    if not filtered and services:
        print(f"警告：提供的服务内容中所有组的 items 均为空。")
    
    return filtered


def add_additional_services(services, additional_items):
    """将用户补充的服务内容添加到服务列表中
    
    Args:
        services: 现有服务内容列表
        additional_items: 用户补充的服务项目列表
    
    Returns:
        更新后的服务内容列表
    """
    if not additional_items:
        return services
    
    # 如果已有服务内容，追加到最后一组
    if services:
        last_group = services[-1]
        last_group["items"].extend(additional_items)
    else:
        # 如果没有现有服务，创建新的服务组
        services.append({
            "title": "其他法律服务",
            "items": additional_items
        })
    
    return services


def generate_quotation(data, output_path):
    """根据 JSON 数据生成报价方案 DOCX"""
    doc = Document()

    # ── 页面设置 ──
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    law_firm = data.get("law_firm", "XX律师事务所")
    doc_title = data.get("doc_title", "法律服务报价方案")
    client_name = data.get("client_name", "XXXXXXX")
    date_str = data.get("date", "")
    opening = data.get("opening", "")
    template_type = data.get("template_type", "报价+服务内容")
    case_type = data.get("case_type", "")  # litigation 或 non_litigation

    # ── 封面/标题区 ──
    add_paragraph_styled(doc, law_firm, font_cn=FONT_HEADING, size_pt=22,
                         bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         color=COLOR_ACCENT, space_before=60, space_after=8)

    add_paragraph_styled(doc, doc_title, font_cn=FONT_HEADING, size_pt=18,
                         bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         color=COLOR_DARK, space_before=8, space_after=30)

    # ── 称谓 ──
    honorific = data.get("honorific", "尊敬的")
    add_paragraph_styled(doc, f"{honorific} {client_name} ：",
                         size_pt=12, space_before=6, space_after=6)

    # ── 开场白 ──
    if opening:
        add_paragraph_styled(doc, opening, size_pt=12,
                             first_line_indent=0.75,
                             space_before=3, space_after=6)

    # ── 法规标准引用（标准+费用 模板） ──
    if template_type == "标准+费用" and data.get("fee_regulations"):
        add_paragraph_styled(doc, "一、律师法律服务收费标准性规定",
                             font_cn=FONT_HEADING, size_pt=14, bold=True,
                             space_before=12, space_after=6)
        for reg in data["fee_regulations"]:
            add_paragraph_styled(doc, reg, size_pt=11,
                                 first_line_indent=0.75,
                                 space_before=2, space_after=2)

    # ── 报价表 ──
    fee_section_num = "二" if (template_type == "标准+费用" and data.get("fee_regulations")) else "一"
    if template_type == "服务内容+团队+价格":
        fee_section_title = "三、法律服务费用"
    elif template_type == "团队+费用":
        fee_section_title = "二、法律服务费用报价"
    else:
        fee_section_title = f"{fee_section_num}、法律服务报价"

    # 服务内容（仅"服务内容+团队+价格"模板在报价前显示）
    if template_type == "服务内容+团队+价格" and data.get("services"):
        add_paragraph_styled(doc, "一、法律服务内容",
                             font_cn=FONT_HEADING, size_pt=14, bold=True,
                             space_before=12, space_after=6)
        
        # 根据案件类型过滤服务内容
        services = filter_services_by_case_type(data["services"], case_type)
        
        # 添加用户补充的服务内容
        additional_services = data.get("additional_services", [])
        if additional_services:
            services = add_additional_services(services, additional_services)
        
        for svc_group in services:
            group_title = svc_group.get("title", "")
            if group_title:
                add_paragraph_styled(doc, group_title,
                                     font_cn=FONT_HEADING, size_pt=12,
                                     bold=True, space_before=6, space_after=3)
            for idx, item in enumerate(svc_group.get("items", []), 1):
                add_paragraph_styled(doc, f"{idx}、{item}",
                                     size_pt=11, first_line_indent=0.75,
                                     space_before=1, space_after=1)

    # 团队介绍（服务内容+团队+价格 / 团队+费用）
    if template_type in ("服务内容+团队+价格", "团队+费用") and data.get("team"):
        team_section_num = "二" if template_type == "服务内容+团队+价格" else "一"
        add_paragraph_styled(doc, f"{team_section_num}、法律服务团队介绍",
                             font_cn=FONT_HEADING, size_pt=14, bold=True,
                             space_before=12, space_after=6)
        if data.get("team_intro"):
            add_paragraph_styled(doc, data["team_intro"], size_pt=11,
                                 first_line_indent=0.75,
                                 space_before=3, space_after=6)
        team_headers = ["姓名", "执业年限", "业务专长", "所内任职"]
        team_rows = []
        for member in data["team"]:
            team_rows.append([
                member.get("name", ""),
                member.get("years", ""),
                member.get("expertise", ""),
                member.get("position", "")
            ])
        add_table_with_data(doc, team_headers, team_rows,
                            col_widths=[3, 3, 5.5, 4])
        doc.add_paragraph()  # 间距

    # ── 报价表格 ──
    add_paragraph_styled(doc, fee_section_title,
                         font_cn=FONT_HEADING, size_pt=14, bold=True,
                         space_before=12, space_after=6)

    if data.get("fee_intro"):
        add_paragraph_styled(doc, data["fee_intro"], size_pt=11,
                             first_line_indent=0.75,
                             space_before=3, space_after=6)

    # 支持多个独立表格（fee_tables 数组）
    if data.get("fee_tables"):
        for table_data in data["fee_tables"]:
            # 添加表格标题
            table_title = table_data.get("title", "")
            if table_title:
                add_paragraph_styled(doc, table_title,
                                     font_cn=FONT_HEADING, size_pt=12, bold=True,
                                     space_before=8, space_after=4)
            # 添加表格
            headers = table_data.get("headers", ["项目", "付费时间/条件", "律师费金额（万元）"])
            rows = table_data.get("rows", [])
            col_widths = table_data.get("col_widths", [4, 6.5, 4])
            add_table_with_data(doc, headers, rows, col_widths=col_widths)
            doc.add_paragraph()  # 表格后间距
    elif data.get("fee_table"):
        # 兼容旧的单一表格格式
        ft = data["fee_table"]
        headers = ft.get("headers", ["项目", "付费时间/条件", "律师费金额（万元）"])
        rows = ft.get("rows", [])
        col_widths = ft.get("col_widths", [4, 6.5, 4])
        add_table_with_data(doc, headers, rows, col_widths=col_widths)
        doc.add_paragraph()  # 间距

    # ── 脚注 ──
    if data.get("footnotes"):
        for fn in data["footnotes"]:
            add_paragraph_styled(doc, fn, size_pt=10,
                                 first_line_indent=0.75,
                                 space_before=1, space_after=1)

    # ── 报价+服务内容 模板中，服务内容在报价后面 ──
    if template_type == "报价+服务内容" and data.get("services"):
        add_paragraph_styled(doc, "上述报价包括法律服务内容如下：",
                             size_pt=12, space_before=10, space_after=4,
                             first_line_indent=0.75)
        
        # 根据案件类型过滤服务内容
        services = filter_services_by_case_type(data["services"], case_type)
        
        # 添加用户补充的服务内容
        additional_services = data.get("additional_services", [])
        if additional_services:
            services = add_additional_services(services, additional_services)
        
        for svc_group in services:
            for idx, item in enumerate(svc_group.get("items", []), 1):
                add_paragraph_styled(doc, f"{idx}、{item}",
                                     size_pt=11, first_line_indent=0.75,
                                     space_before=1, space_after=1)

    # ── 结尾 ──
    closing = data.get("closing", "以上法律服务报价方案供贵方参酌。如有任何疑问，可随时与服务律师联系。")
    # 修复：避免 closing 中包含“顺祝商祺”导致重复
    if "顺祝商祺" in closing:
        # 截取“顺祝商祺”之前的内容
        closing = closing.split("顺祝商祺")[0].rstrip("。").rstrip("！").rstrip("，").strip()
    
    add_paragraph_styled(doc, closing, size_pt=12,
                         first_line_indent=0.75,
                         space_before=12, space_after=6)

    add_paragraph_styled(doc, "顺祝商祺！", size_pt=12,
                         space_before=6, space_after=20)

    add_paragraph_styled(doc, law_firm, size_pt=12,
                         alignment=WD_ALIGN_PARAGRAPH.RIGHT,
                         space_before=6, space_after=3)

    if date_str:
        add_paragraph_styled(doc, date_str, size_pt=12,
                             alignment=WD_ALIGN_PARAGRAPH.RIGHT,
                             space_before=0, space_after=6)

    # ── 保存 ──
    doc.save(output_path)
    return output_path


# ── 入参校验 ──────────────────────────────────────────────
# 所有「缺字段 / 类型错误」类入参问题统一走 InputError -> 退出码 2 的友好报错，
# 不再让脚本以 traceback 崩溃（P1：入参崩栈修复）。

SCALAR_FIELDS = ("law_firm", "doc_title", "client_name", "honorific", "date",
                 "opening", "fee_intro", "team_intro", "closing", "case_type",
                 "template_type")
LIST_FIELDS = ("services", "additional_services", "team", "footnotes",
               "fee_regulations", "fee_tables")
STR_LIST_FIELDS = ("footnotes", "fee_regulations")
# 各模板类型的必填字段（锚定 SKILL.md「各模板类型使用的字段」表）
TEMPLATE_REQUIRED = {
    "报价+服务内容": ("services",),
    "服务内容+团队+价格": ("services", "team"),
    "团队+费用": ("team",),
    "标准+费用": ("fee_regulations",),
}


def load_json_input(raw: str):
    """按文件路径或内联 JSON 解析报价数据；失败抛 InputError。"""
    json_file = None
    if os.path.isfile(raw):
        json_file = raw
    else:
        candidate = SKILL_ROOT / raw
        if candidate.is_file():
            json_file = str(candidate)
    try:
        if json_file:
            with open(json_file, "r", encoding="utf-8") as f:
                return json.load(f)
        return json.loads(raw)
    except (json.JSONDecodeError, UnicodeDecodeError, OSError) as exc:
        raise InputError(f"无法解析 JSON 数据：{exc}")


def validate_data(data) -> List[str]:
    """校验报价数据结构的类型与必填字段，返回问题列表（空列表=通过）。"""
    problems: List[str] = []
    if not isinstance(data, dict):
        problems.append(f"JSON 顶层必须是对象（{{...}}），实际为 {type(data).__name__}")
        return problems

    for f in SCALAR_FIELDS:
        v = data.get(f)
        if v is not None and not isinstance(v, str):
            problems.append(f"字段「{f}」应为字符串，实际为 {type(v).__name__}")
    for f in LIST_FIELDS:
        v = data.get(f)
        if v is not None and not isinstance(v, list):
            problems.append(f"字段「{f}」应为数组，实际为 {type(v).__name__}")
    for f in STR_LIST_FIELDS:
        for i, item in enumerate(data.get(f) or []):
            if not isinstance(item, str):
                problems.append(f"字段「{f}」第 {i + 1} 项应为字符串，实际为 {type(item).__name__}")

    if "fee_table" in data and data["fee_table"] is not None \
            and not isinstance(data["fee_table"], dict):
        problems.append(f"字段「fee_table」应为对象，实际为 {type(data['fee_table']).__name__}")

    for i, g in enumerate(data.get("services") or []):
        if not isinstance(g, dict):
            problems.append(f"services 第 {i + 1} 组应为对象，实际为 {type(g).__name__}")
            continue
        items = g.get("items")
        if items is not None and not isinstance(items, list):
            problems.append(f"services 第 {i + 1} 组的 items 应为数组，实际为 {type(items).__name__}")
    for i, m in enumerate(data.get("team") or []):
        if not isinstance(m, dict):
            problems.append(f"team 第 {i + 1} 名律师应为对象，实际为 {type(m).__name__}")

    # 报价表结构（fee_table / fee_tables 两种写法），并核对行宽不超过表头列数
    tables: List[Tuple[str, dict]] = []
    if isinstance(data.get("fee_table"), dict):
        tables.append(("fee_table", data["fee_table"]))
    for i, t in enumerate(data.get("fee_tables") or []):
        if isinstance(t, dict):
            tables.append((f"fee_tables 第 {i + 1} 个", t))
        else:
            problems.append(f"fee_tables 第 {i + 1} 个表格应为对象，实际为 {type(t).__name__}")
    for label, t in tables:
        headers = t.get("headers")
        if headers is not None and not isinstance(headers, list):
            problems.append(f"{label}表格的 headers 应为数组，实际为 {type(headers).__name__}")
        rows = t.get("rows")
        if rows is not None and not isinstance(rows, list):
            problems.append(f"{label}表格的 rows 应为数组，实际为 {type(rows).__name__}")
        ncols = len(headers) if isinstance(headers, list) and headers else 3
        for j, r in enumerate(rows or []):
            if not isinstance(r, list):
                problems.append(f"{label}表格 rows 第 {j + 1} 行应为数组，实际为 {type(r).__name__}")
            elif len(r) > ncols:
                problems.append(f"{label}表格 rows 第 {j + 1} 行有 {len(r)} 列，超过表头 {ncols} 列")

    # 必填字段（fee_table 所有模板均必填；其余按模板类型）
    has_fee = any(t.get("rows") for _, t in tables)
    if not has_fee:
        problems.append("缺少必填字段：fee_table（或 fee_tables）报价表须提供且至少含一行数据，"
                        "所有模板均须提供报价表")
    template_type = data.get("template_type")
    if template_type in TEMPLATE_REQUIRED:
        for field in TEMPLATE_REQUIRED[template_type]:
            if not data.get(field):
                problems.append(f"模板「{template_type}」缺少必填字段：{field}")
    return problems


# ── CLI 入口 ──────────────────────────────────────────────

EXAMPLE_JSON = r"""
{
  "law_firm": "XX律师事务所",
  "doc_title": "法律服务报价方案",
  "client_name": "某某公司",
  "honorific": "尊敬的",
  "date": "二〇二六年四月",
  "template_type": "报价+服务内容",
  "opening": "关于贵司与张某股东资格确认相关纠纷之诉讼代理（以下称\"委托事项\"），根据贵司的具体诉求及服务内容，现向贵司提供本法律服务报价方案。",
  "fee_intro": "本着就上述委托事项，本所提议采取基础法律服务费+风险代理费的方式收取法律服务费用。具体如下：",
  "fee_table": {
    "headers": ["项目", "付费时间/条件", "律师费金额（万元）"],
    "rows": [
      ["**基础法律服务费**", "", ""],
      ["一审阶段", "委托合同签订之日起5日内", "10"],
      ["", "立案进入诉前调解之日起5日内", "20"],
      ["二审阶段（如有）", "提起上诉之日起5日内", "20"],
      ["**风险代理费**", "", ""],
      ["调解/和解达成目的", "达成贵方认可的成果", "100"],
      ["判决胜诉", "判决生效之日起5日内", "150"]
    ],
    "col_widths": [4, 6.5, 4]
  },
  "footnotes": [
    "注：如案件涉及其他衍生诉讼，需另行协商收费及签订补充合同。"
  ],
  "services": [
    {
      "title": "",
      "items": [
        "对案件进行法律分析，进行法律风险研判，提出法律建议，并协助贵方制定诉讼策略；",
        "协助贵方搜集/调取必要证据；",
        "必要时介入案件相关谈判，并根据贵方需求出具相关协议，包括参与和解、调解等；",
        "参与诉讼过程，包括诉前调解、一审、二审，协助起草诉讼过程中涉及的各类法律文书。"
      ]
    }
  ],
  "closing": "以上法律服务方案，供贵司参考。如有任何疑问，可随时与本所律师联系。"
}
"""


def main() -> int:
    parser = argparse.ArgumentParser(description="生成律师法律服务报价方案 DOCX 文档")
    parser.add_argument("--json", default=None,
                        help="JSON 格式的报价数据（字符串或文件路径）")
    parser.add_argument("--output", default=None,
                        help="输出文件路径，默认为当前目录下的 法律服务报价方案.docx")
    parser.add_argument("--example", action="store_true",
                        help="打印 JSON 数据示例后退出（无需 --json）")
    args = parser.parse_args()

    if args.example:
        print(EXAMPLE_JSON)
        return 0

    if not args.json:
        print("输入错误：缺少 --json（报价数据 JSON 字符串或文件路径）。\n"
              "  示例：python3 scripts/generate_quotation.py --json '{\"client_name\": \"某某公司\", ...}'\n"
              "  查看数据结构示例：python3 scripts/generate_quotation.py --example",
              file=sys.stderr)
        return 2

    # 读取 JSON 数据：优先按文件路径解析（先相对当前工作目录，再相对技能根目录），
    # 都不存在时按内联 JSON 字符串解析。
    try:
        data = load_json_input(args.json)
    except InputError as exc:
        print(f"输入错误：{exc}", file=sys.stderr)
        return 2

    problems = validate_data(data)
    if problems:
        print("输入错误：报价数据校验未通过：", file=sys.stderr)
        for p in problems:
            print(f"- {p}", file=sys.stderr)
        return 2

    # 确定输出路径
    if args.output:
        out_dir = os.path.dirname(os.path.abspath(args.output))
        if not os.path.isdir(out_dir):
            print(f"输入错误：输出目录不存在：{out_dir}", file=sys.stderr)
            return 2
        output_path = args.output
    else:
        title = data.get("doc_title") or "法律服务报价方案"
        output_path = os.path.join(os.getcwd(), f"{title}.docx")

    try:
        result = generate_quotation(data, output_path)
    except Exception as exc:
        print(f"生成失败：{exc}", file=sys.stderr)
        return 1
    print(f"报价方案已生成: {result}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
