#!/usr/bin/env python3
"""lawd-defense-statement-draft 答辩状交付门禁（防幻觉 / 防缺漏 / 防模板空壳）。

用法
----
    python3 validate_answer.py --file 答辩状.md   --claims N   # 主路径：先出 Markdown 过门禁
    python3 validate_answer.py --file 答辩状.docx --claims N   # 复核路径：对最终 Word 成稿复核

支持格式：.md / .markdown / .txt / .docx。
（--file 亦可写作 --doc / --docx / --md，等价，便于旧命令兼容。）
docx 读取依赖 python-docx（pip3 install python-docx），会读取正文段落、表格、
页眉、页脚与文本框中的文字。

--claims（原告诉讼请求条数）为必填：缺省时脚本直接以退出码 2 报错，不做静默
放过——否则“防漏项”这道门禁等于自愿检查。取值口径见 --help。

检查项
------
1. 五段结构齐备：首部/标题（含“答辩状”）、答辩请求、事实与理由（“事实和理由”
   等写法均认）、法律依据（论证）、尾部（此致 + 致送法院 + 答辩人落款）。
2. 诉请逐项回应数量 >= --claims N。统计口径：在“事实与理由 / 答辩事项 / 逐项
   答辩”等回应型章节内，统计“有实质内容的回应条目”（编号小标题块、要素式表格
   行），并取其与该范围内“第X项诉请/诉讼请求/请求”不同序数个数的较大值。
   不再统计“答辩请求”章节的编号条目（该章节模板固定只有“驳回全部 + 诉讼费
   承担”，且零实质内容的模板空壳可凑数绕过门禁）。
3. 引用法条格式：法律条号须带《法律名》。允许“《法名》（以下简称《简称》）第X条”
   “《法名》第X条、第Y条第2款”等写法；“合同第三条”“公司章程第五条”等非法条
   条号不计入拦截。
4. 法院提交格式项：标题、致送法院（“此致 XX人民法院”）、答辩人落款。

退出码：0=通过；1=拦截；2=输入错误（含缺 --claims）。
"""

from __future__ import annotations

import argparse
import re
import sys
import zipfile
from pathlib import Path
from typing import Any, List, Optional, Sequence, Set, Tuple

# ------------------------------------------------------------------ 常量/正则

NUM_CHARS = "零〇○一二两三四五六七八九十百千0123456789"
CN_ONLY = "零〇○一二两三四五六七八九十百千"

_N = "[" + NUM_CHARS + "]+"
_SUFFIX = r"(?:\s*第" + _N + r"[款项])*"

# 《法律名》[（以下简称《简称》）] 第X条[第X款/项][、第Y条第Y款…]
CITATION_RE = re.compile(
    r"《([^《》\n]{2,80}?)》"                          # 法律名
    r"(?:\s*[（(][^）)\n]{0,60}[）)])?"                # 可选简称括注
    r"\s*(?:之|的)?\s*"
    r"第(" + _N + r")条" + _SUFFIX +                   # 首个条号（可带款/项）
    r"((?:\s*[、，,;；和与及]?\s*第" + _N + r"条" + _SUFFIX + r")*)"  # 并条列举
)

ARTICLE_RE = re.compile(r"第(" + _N + r")条")

# 诉请序数：第X项诉请 / 第X项诉讼请求 / 第X项请求 / 第X项主张
CLAIM_REF_RE = re.compile(
    r"第(" + _N + r")项\s*(?:之)?\s*(?:诉讼请求|诉讼诉请|诉请|请求|主张)"
)

# 事实与理由 / 事实和理由 / 事实理由 / 事实与答辩理由
FACT_REASON_RE = re.compile(r"事实\s*[与和及]?\s*(?:答辩)?\s*理由")

# 回应型章节（统计逐项回应的范围）
RESPONSE_SECTION_RE = re.compile(
    r"事实\s*[与和及]?\s*(?:答辩)?\s*理由"
    r"|答辩事项|逐项答辩|逐项回应|逐项意见|分项答辩|对诉讼请求的答辩|对诉请的答辩"
)
# 回应型章节的终止标记
SECTION_END_RE = re.compile(
    r"综上所述|综上，|综上,|此\s*致|证据清单|答辩辅助提示|辅助提示|风险提示|程序提醒"
)

# 编号小标题：（一） (1) 一、 1、 1. 第一项
SUBHEAD_RE = re.compile(
    r"^(?:#{1,6}\s*)?(?:\*{1,2})?\s*(?:"
    r"[（(]\s*[" + NUM_CHARS + r"]{1,4}\s*[)）]"
    r"|[" + CN_ONLY + r"]{1,4}\s*[、.．]"
    r"|\d{1,3}\s*[.、)）]"
    r"|第[" + NUM_CHARS + r"]{1,4}[项条点]"
    r")"
)

# 明确表态标记（要素式表格 / 逐项回应的态度）
STANCE_RE = re.compile(
    r"无异议|有异议|异议|不认可|不予认可|予以认可|认可|确认|驳回|☑|☒|■|√"
)

# 紧邻条号之前的“非法条载体”（合同/章程等），这类“第X条”不是法条引用
NON_STATUTE_PREFIX_RE = re.compile(
    r"(?:合同|协议|合约|契约|章程|规约|公约|条款|细则|规则|规程|议事规则|管理规约|"
    r"业主公约|标准|规范|附件|附则|备忘录|意向书|承诺书|授权书|决议|判决|裁定|裁决|"
    r"调解书|通知|保单|制度|手册|须知|说明|方案|纪要|清单|明细)"
    r"\s*(?:书|文本)?\s*(?:之|的|中|里)?\s*$"
)

# 可不带《》的“回指”表述（该法/本法第X条），需全文存在法律引用才放行
BACKREF_RE = re.compile(
    r"(?:该法|本法|上述法律|前述法律|同法|该司法解释|该解释|上述解释|该条例|"
    r"上述条例|该规定|上述规定)\s*(?:之|的)?\s*$"
)

# 法律/法规/司法解释名称特征（用于区分《XX买卖合同》与《民法典》）
STATUTE_NAME_RE = re.compile(
    r"(?:法|法典|宪法|条例|规定|解释|办法|细则|规则|通知|意见|纪要|批复|决定|准则|"
    r"标准|规范)$"
)

COURT_RE = r"[\u4e00-\u9fa5]{2,25}(?:人民法院|海事法院|知识产权法院|金融法院|法院)"

CN_DIGITS = {
    "零": 0, "〇": 0, "○": 0,
    "一": 1, "二": 2, "两": 2, "三": 3, "四": 4,
    "五": 5, "六": 6, "七": 7, "八": 8, "九": 9,
}
CN_UNITS = {"十": 10, "百": 100, "千": 1000}

TEXT_SUFFIXES = {".md", ".markdown", ".txt", ".text"}
MIN_SUBSTANCE = 30          # 一条“有实质内容的回应”的最少字数（去空白后）
LOOKBACK_FOR_LAW = 120      # 裸条号向前查找《法律名》的字符窗口


class InputError(Exception):
    """输入层错误，退出码 2。"""


# ------------------------------------------------------------------ 小工具


def cn_num_to_int(text: str) -> Optional[int]:
    s = text.strip()
    if not s:
        return None
    if s.isdigit():
        return int(s)
    total, current, valid = 0, 0, False
    for ch in s:
        if ch in CN_DIGITS:
            current = CN_DIGITS[ch]
            valid = True
        elif ch in CN_UNITS:
            unit = CN_UNITS[ch]
            if current == 0:
                current = 1
            total += current * unit
            current = 0
            valid = True
        else:
            return None
    return total + current if valid else None


def ordinal_key(raw: str) -> Tuple[str, Any]:
    num = cn_num_to_int(raw)
    return ("num", num) if num is not None else ("str", raw.strip())


def norm(text: str) -> str:
    return re.sub(r"\s+", "", text)


def strip_markup(line: str) -> str:
    """去掉 markdown 标记，便于判断是否像标题。"""
    s = line.strip()
    s = re.sub(r"^[#>\-*\s|]+", "", s)
    s = s.replace("**", "").replace("*", "").strip()
    return s.rstrip("|").strip()


def heading_like(line: str) -> bool:
    return len(strip_markup(line)) <= 40


def is_table_row(line: str) -> bool:
    s = line.strip()
    return s.startswith("|") and s.count("|") >= 3


def is_separator_row(line: str) -> bool:
    return bool(re.fullmatch(r"[|\s:：\-—–]+", line.strip()))


# ------------------------------------------------------------------ 文档读取


def _docx_block_items(document):
    """按文档顺序产出正文层的段落与表格。"""
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P

    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield Table(child, document)


def _table_lines(table) -> List[str]:
    """把表格渲染成便于统计的文本行。

    单列表格在中文法律文书里常被当作排版容器（整篇正文塞进一列），此时按普通
    段落输出，保证其中的“（一）……”仍能被识别为编号小标题；多列表格（要素式
    答辩状）保留 markdown 行形式，按“表格行=一条回应”统计。
    """
    lines: List[str] = []
    for row in table.rows:
        raw = [cell.text for cell in row.cells]
        if not any(c.strip() for c in raw):
            continue
        filled = {c.strip() for c in raw if c.strip()}
        if len(raw) == 1 or len(filled) == 1:
            lines.extend(part for part in raw[0].splitlines() if part.strip())
        else:
            cells = [re.sub(r"\s*\n\s*", " ", c).strip() for c in raw]
            lines.append("| " + " | ".join(cells) + " |")
    return lines


def _textbox_texts(element) -> List[str]:
    """抽取文本框（w:txbxContent，含 mc:Fallback 中的 v:textbox）内的文字。

    注意：w:txbxContent 不是 python-docx 注册的元素类，其 .xpath() 缺少命名空间
    映射会直接抛错，因此这里用 lxml iter() + 限定标签名遍历。
    """
    try:
        from docx.oxml.ns import qn
    except ImportError:
        return []
    tag_box, tag_p, tag_t = qn("w:txbxContent"), qn("w:p"), qn("w:t")
    out: List[str] = []
    try:
        boxes = list(element.iter(tag_box))
    except Exception:
        return out
    for box in boxes:
        for para in box.iter(tag_p):
            text = "".join(node.text or "" for node in para.iter(tag_t)).strip()
            if text:
                out.append(text)
    seen: Set[str] = set()
    uniq: List[str] = []
    for item in out:
        if item not in seen:
            seen.add(item)
            uniq.append(item)
    return uniq


def load_docx(path: Path) -> Tuple[List[str], List[str]]:
    """读取 docx，返回（正文块列表, 页眉/页脚/文本框附加块列表）。"""
    if not zipfile.is_zipfile(path):
        raise InputError("文件不是合法 DOCX ZIP 容器")
    try:
        from docx import Document
    except ImportError:
        raise InputError("当前环境未安装 python-docx，请先执行：pip3 install python-docx")
    try:
        document = Document(str(path))
    except Exception as exc:
        raise InputError(f"python-docx 无法打开文件：{exc}")

    from docx.table import Table

    blocks: List[str] = []
    for item in _docx_block_items(document):
        if isinstance(item, Table):
            blocks.extend(_table_lines(item))
        else:
            blocks.append(item.text)

    extras: List[str] = _textbox_texts(document.element.body)
    for section in document.sections:
        parts = []
        for attr in ("header", "first_page_header", "even_page_header",
                     "footer", "first_page_footer", "even_page_footer"):
            try:
                parts.append(getattr(section, attr))
            except Exception:
                continue
        for part in parts:
            if part is None:
                continue
            try:
                for para in part.paragraphs:
                    if para.text.strip():
                        extras.append(para.text)
                for table in part.tables:
                    extras.extend(_table_lines(table))
                extras.extend(_textbox_texts(part._element))
            except Exception:
                continue

    seen: Set[str] = set()
    uniq_extras: List[str] = []
    for item in extras:
        key = norm(item)
        if key and key not in seen:
            seen.add(key)
            uniq_extras.append(item)
    return blocks, uniq_extras


def load_document(path: Path) -> Tuple[List[str], List[str]]:
    """读取 md/txt/docx，返回（正文块列表, 附加块列表）。"""
    suffix = path.suffix.lower()
    if suffix not in TEXT_SUFFIXES and suffix != ".docx":
        raise InputError(
            f"不支持的文件类型：{path.name}"
            "（支持 .md/.markdown/.txt/.docx；.doc 请先转 .docx）"
        )
    if not path.exists():
        raise InputError(f"文件不存在：{path}")
    if path.stat().st_size == 0:
        raise InputError("文件为空")
    if suffix == ".docx":
        return load_docx(path)
    try:
        content = path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        try:
            content = path.read_text(encoding="gbk")
        except Exception as exc:
            raise InputError(f"无法以 UTF-8/GBK 解码文本文件：{exc}")
    if not content.strip():
        raise InputError("文件内容为空")
    return content.splitlines(), []


# ------------------------------------------------------------------ 章节切分


def response_sections(lines: Sequence[str]) -> List[Tuple[int, int]]:
    """定位所有回应型章节的行区间 [start, end)。"""
    starts = [idx for idx, line in enumerate(lines)
              if RESPONSE_SECTION_RE.search(line)
              and not is_table_row(line) and heading_like(line)]
    if not starts:
        starts = [idx for idx, line in enumerate(lines)
                  if RESPONSE_SECTION_RE.search(line) and not is_table_row(line)]
    spans: List[Tuple[int, int]] = []
    for pos, start in enumerate(starts):
        limit = starts[pos + 1] if pos + 1 < len(starts) else len(lines)
        end = limit
        for idx in range(start + 1, limit):
            if SECTION_END_RE.search(lines[idx]) and heading_like(lines[idx]):
                end = idx
                break
        if end > start:
            spans.append((start, end))
    return spans


def split_response_items(lines: Sequence[str]) -> List[Tuple[str, str]]:
    """把章节内容切成带类型的回应条目：table（要素式表格行）/ head（编号小标题
    块）/ prose（无编号的散段）。"""
    items: List[Tuple[str, str]] = []
    buffer: List[str] = []
    kind = "prose"
    for line in lines:
        if is_separator_row(line):
            continue
        if is_table_row(line):
            if buffer:
                items.append((kind, "\n".join(buffer)))
                buffer, kind = [], "prose"
            items.append(("table", line))
            continue
        if SUBHEAD_RE.match(line.strip()) and heading_like(line):
            if buffer:
                items.append((kind, "\n".join(buffer)))
            buffer, kind = [line], "head"
            continue
        if line.strip():
            buffer.append(line)
    if buffer:
        items.append((kind, "\n".join(buffer)))
    return items


def pick_response_items(items: Sequence[Tuple[str, str]]) -> List[str]:
    """挑出真正代表"逐项回应"的条目。

    要素式（表格逐项）与传统式（编号小标题逐项）各自成体系；一旦成体系，就只数
    该体系内的条目，避免"答辩依据""章节引言"等散段被当成一条回应虚增数量。
    """
    solid = [(k, t) for k, t in items if is_substantive(t)]
    tables = [t for k, t in solid if k == "table"]
    heads = [t for k, t in solid if k == "head"]
    if len(tables) >= 2:
        return tables
    if len(heads) >= 2:
        return heads
    return [t for _, t in solid]


def is_substantive(item: str) -> bool:
    """有实质内容 = 字数达标，或（要素式表格行）已明确表态。"""
    if len(norm(item)) >= MIN_SUBSTANCE:
        return True
    return bool(is_table_row(item) and STANCE_RE.search(item))


def count_defense_request_items(lines: Sequence[str]) -> int:
    """仅用于展示：答辩请求章节内的编号条目数（不计入回应数）。"""
    start = None
    for idx, line in enumerate(lines):
        if "答辩请求" in line:
            start = idx
            break
    if start is None:
        return 0
    end = len(lines)
    for idx in range(start + 1, len(lines)):
        line = lines[idx]
        if (RESPONSE_SECTION_RE.search(line) or SECTION_END_RE.search(line)) \
                and heading_like(line):
            end = idx
            break
    return sum(1 for line in lines[start + 1:end] if SUBHEAD_RE.match(line.strip()))


# ------------------------------------------------------------------ 法条引用


def sentence_head(text: str, pos: int) -> str:
    """返回 pos 所在句子中位于 pos 之前的部分。"""
    start = max((text.rfind(ch, 0, pos) for ch in "。；;！!？?\n"), default=-1)
    return text[start + 1:pos]


def statute_citations(text: str) -> List[re.Match]:
    return [m for m in CITATION_RE.finditer(text)
            if STATUTE_NAME_RE.search(m.group(1))]


def law_name_nearby(text: str, pos: int) -> bool:
    """裸条号之前 LOOKBACK_FOR_LAW 个字符内是否出现《法律名》。"""
    window = text[max(0, pos - LOOKBACK_FOR_LAW):pos]
    return any(STATUTE_NAME_RE.search(name)
               for name in re.findall(r"《([^《》\n]{2,80}?)》", window))


def classify_bare_articles(
    text: str, citations: Sequence[re.Match]
) -> Tuple[List[re.Match], List[re.Match], List[re.Match]]:
    """把未被《法名》直接包裹的第X条分成：非法条条号 / 可归属 / 真裸条号。"""
    spans = [m.span() for m in citations]
    has_statute = bool(statute_citations(text))
    non_statute: List[re.Match] = []
    attributed: List[re.Match] = []
    bare: List[re.Match] = []
    for m in ARTICLE_RE.finditer(text):
        if any(s <= m.start() < e for s, e in spans):
            continue
        prefix = text[max(0, m.start() - 20):m.start()]
        if NON_STATUTE_PREFIX_RE.search(prefix):
            non_statute.append(m)
            continue
        if BACKREF_RE.search(prefix) and has_statute:
            attributed.append(m)
            continue
        head = sentence_head(text, m.start())
        if any(STATUTE_NAME_RE.search(name)
               for name in re.findall(r"《([^《》\n]{2,80}?)》", head)):
            attributed.append(m)
            continue
        if law_name_nearby(text, m.start()):
            attributed.append(m)
            continue
        bare.append(m)
    return non_statute, attributed, bare


# ------------------------------------------------------------------ 主流程


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="答辩状交付门禁：五段结构 / 诉请逐项回应（防漏项、防模板空壳）"
                    " / 法条引用格式 / 法院提交格式项。支持 md/txt/docx。",
        epilog=(
            "--claims 取值口径：原告起诉状“诉讼请求”部分的实体诉请条数，即 Phase 2"
            "“逐项拆解原告诉讼请求”得到的诉请清单条数，不计“诉讼费用由被告负担”"
            "这类程序性请求。docx 读取依赖 python-docx（pip3 install python-docx）。"
        ),
    )
    parser.add_argument("--file", "--doc", "--docx", "--md", dest="path",
                        required=True, type=Path,
                        help="答辩状文件路径（.md/.markdown/.txt/.docx）")
    parser.add_argument("--claims", type=int, default=None,
                        help="【必填】原告诉讼请求条数，用于核对逐项回应数量 >= N")
    return parser


def main() -> int:
    args = build_parser().parse_args()

    if args.claims is None:
        print(
            "输入错误：缺少 --claims。\n"
            "  防漏项检查依赖原告诉请条数，缺省时无法校验“逐项回应不遗漏”，"
            "本门禁不做静默放过。\n"
            "  请按 SKILL.md 流程传入：--claims <原告起诉状“诉讼请求”部分的实体诉请条数>\n"
            "  例：python3 validate_answer.py --file 答辩状.md --claims 4",
            file=sys.stderr,
        )
        return 2
    if args.claims < 1:
        print("输入错误：--claims 必须 >= 1", file=sys.stderr)
        return 2

    try:
        blocks, extras = load_document(args.path)
    except InputError as exc:
        print(f"输入错误：{exc}", file=sys.stderr)
        return 2

    main_text = "\n".join(blocks)
    body = "\n".join(blocks + extras) if extras else main_text
    extras_offset = len(main_text) if extras else len(body)
    lines = body.splitlines()

    checks: List[str] = []
    warns: List[str] = []
    errors: List[str] = []

    if extras:
        warns.append(f"已额外读取页眉/页脚/文本框中的 {len(extras)} 段文字参与校验")

    # --- ① 五段结构 ---
    head_blocks = [b for b in blocks if b.strip()][:10]
    title_in_head = any("答辩状" in norm(b) for b in head_blocks)
    title_anywhere = any("答辩状" in norm(b) for b in blocks + extras)
    if title_in_head:
        checks.append("五段结构-首部/标题：文首找到“答辩状”标题")
    elif title_anywhere:
        errors.append("五段结构-首部/标题：“答辩状”未置于文首（仅出现在正文中）")
    else:
        errors.append("五段结构-首部/标题：全文未找到“答辩状”标题")

    if re.search(r"答辩请求|答辩之请求|请求事项", body):
        checks.append("五段结构-答辩请求：找到“答辩请求”")
    else:
        errors.append("五段结构-答辩请求：缺少“答辩请求”部分")

    m_fact = FACT_REASON_RE.search(body)
    if m_fact:
        checks.append("五段结构-事实与理由：找到“%s”" % norm(m_fact.group(0)))
    else:
        errors.append("五段结构-事实与理由：缺少“事实与理由/事实和理由”部分")

    citations = list(CITATION_RE.finditer(body))
    statutes = statute_citations(body)
    if re.search(r"法律依据|法律论证|法律适用|法律规定", body):
        checks.append("五段结构-法律依据（论证）：找到 法律依据/法律论证 表述")
    elif statutes:
        checks.append("五段结构-法律依据（论证）：无独立章节，但含 %d 处格式正确的"
                      "《法律名》第X条论证引用" % len(statutes))
    else:
        errors.append("五段结构-法律依据（论证）：缺少 法律依据/法律论证 章节，"
                      "且无《法律名》第X条引用")

    m_cizhi = re.search(r"此\s*致", body)
    court_anywhere = re.search(COURT_RE, body)
    sign_anywhere = re.search(r"答辩人|具状人", body)
    cizhi_in_extras = bool(m_cizhi and m_cizhi.start() >= extras_offset)
    m_court = None
    sign_after = None
    court_name = ""
    if m_cizhi:
        m_court = re.search(r"此\s*致[\s\S]{0,100}?(" + COURT_RE + r")", body)
        sign_after = re.search(r"答辩人|具状人", body[m_cizhi.end():])
        if m_court:
            court_name = m_court.group(1)
        elif court_anywhere:
            m_court = court_anywhere
            court_name = court_anywhere.group(0)
            warns.append("致送法院“%s”与“此致”不相邻（可能分布在页眉/页脚/文本框或跨"
                         "表格单元格），已按齐备计" % court_name)
        if not sign_after and sign_anywhere and cizhi_in_extras:
            sign_after = sign_anywhere
            warns.append("“此致”来自页眉/页脚/文本框，落款先后无法判断，已按齐备计")

    if m_cizhi and m_court and sign_after:
        checks.append("五段结构-尾部：此致 + 致送法院“%s” + 答辩人落款齐备" % court_name)
    else:
        if not m_cizhi:
            errors.append("五段结构-尾部：缺少“此致”")
        elif not m_court:
            errors.append("五段结构-尾部：“此致”后未有致送法院（格式：此致 XX人民法院）")
        if not sign_after:
            errors.append("五段结构-尾部：致送法院后缺少答辩人落款（“答辩人：”/“具状人：”）")

    # --- ② 诉请逐项回应数量（防漏项 + 防模板空壳）---
    spans = response_sections(lines)
    best_items = 0
    best_ordinals: Set[Tuple[str, Any]] = set()
    section_details: List[str] = []
    for start, end in spans:
        items = split_response_items(lines[start + 1:end])
        solid = pick_response_items(items)
        ordinals: Set[Tuple[str, Any]] = set()
        if solid:
            for it in solid + [lines[start]]:
                ordinals |= {ordinal_key(m.group(1))
                             for m in CLAIM_REF_RE.finditer(it)}
        section_details.append(
            "“%s”：实质回应条目 %d/%d、诉请序数 %d"
            % (strip_markup(lines[start])[:24], len(solid), len(items), len(ordinals))
        )
        best_items = max(best_items, len(solid))
        if len(ordinals) > len(best_ordinals):
            best_ordinals = ordinals

    responses = max(best_items, len(best_ordinals))
    req_items = count_defense_request_items(lines)
    detail = "；".join(section_details) if section_details else "未定位到回应型章节"
    detail += "；参考：“答辩请求”章节编号条目 %d 条（不计入回应数）" % req_items

    if not spans:
        errors.append("诉请逐项回应：未定位到“事实与理由/答辩事项/逐项答辩”等回应型章节，"
                      "无法核验逐项回应，视为漏答风险")
    elif responses >= args.claims:
        checks.append("诉请逐项回应：%d 项 >= 原告诉请 %d 项（%s）"
                      % (responses, args.claims, detail))
    else:
        errors.append("诉请逐项回应不足：统计 %d 项 < 原告诉请 %d 项（%s）；"
                      "注意仅统计回应型章节内“有实质内容”的条目，单纯“驳回第X项"
                      "诉讼请求”之类无论证的模板条目不计数"
                      % (responses, args.claims, detail))

    # --- ③ 法条引用格式：法条条号必须带《法律名》 ---
    non_statute, attributed, bare = classify_bare_articles(body, citations)
    if bare:
        for m in bare[:5]:
            ctx = body[max(0, m.start() - 15): m.end() + 5].replace("\n", " ")
            errors.append(f"法条引用格式：条号缺少《法律名》 → …{ctx}…")
        if len(bare) > 5:
            errors.append(f"法条引用格式：共 {len(bare)} 处裸条号，仅列出前 5 处")
    elif citations or attributed or non_statute:
        parts = [f"{len(citations)} 处《法律名》+条号引用"]
        if attributed:
            parts.append(f"{len(attributed)} 处同句/近邻已归属法律名的条号")
        if non_statute:
            parts.append(f"{len(non_statute)} 处非法条条号（合同/章程等，不作要求）")
        checks.append("法条引用格式：" + "、".join(parts) + "，未见裸条号")
    else:
        checks.append("法条引用格式：全文无“第X条”引用（建议法律依据论证写明具体条号）")

    # --- ④ 法院提交格式项 ---
    format_items = (
        ("标题（含“答辩状”）", title_in_head),
        ("致送法院（此致 XX人民法院）", bool(m_court)),
        ("答辩人落款", bool(sign_after)),
    )
    for name, ok in format_items:
        if ok:
            checks.append(f"法院提交格式项-{name}：齐备")
        else:
            errors.append(f"法院提交格式项-{name}：缺失")

    print(f"== 答辩状交付门禁：{args.path.name}（通过清单）==")
    for item in checks:
        print(f"- {item}")
    if warns:
        print("== 提示 ==")
        for item in warns:
            print(f"- {item}")
    if errors:
        print("== 拦截清单 ==", file=sys.stderr)
        for item in errors:
            print(f"- {item}", file=sys.stderr)
        print(f"校验不通过：共拦截 {len(errors)} 项", file=sys.stderr)
        return 1
    print("校验通过：%s 五段结构齐备；诉请逐项回应 %d 项 >= %d 项；"
          "法条引用与法院提交格式项均合格" % (args.path.name, responses, args.claims))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
