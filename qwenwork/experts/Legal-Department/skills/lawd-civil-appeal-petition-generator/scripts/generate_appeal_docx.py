#!/usr/bin/env python3
"""Generate a DOCX civil appeal petition with precise formatting.

Usage:
  python3 scripts/generate_appeal_docx.py \
    --appellant "张三，原审被告，男，1980年1月1日生，汉族，住北京市朝阳区，电话：13800000000" \
    --appellee "李四，原审原告，男，1985年5月5日生，汉族，住上海市浦东新区，电话：13900000000" \
    --cause "合同纠纷" \
    --first-court "北京市朝阳区人民法院" \
    --case-no "（2024）京0105民初12345号" \
    --judgment-date "2024年12月1日" \
    --appeal-court "北京市第三中级人民法院" \
    --requests "1. 撤销一审判决第二项\n2. 依法改判上诉人无需承担违约责任\n3. 本案一、二审诉讼费用由被上诉人承担" \
    --reasons "一、原审认定事实不清\n\n原审认为上诉人违约，但根据证据1可以证明...\n\n二、原审适用法律错误\n\n本案应适用合同法第X条..." \
    --output "民事上诉状_张三.docx"

Or use JSON input:
  echo '{"appellant": "...", "appellee": "...", "cause": "...", ...}' | \
    python3 scripts/generate_appeal_docx.py --output "民事上诉状.docx"
"""

from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path
from typing import List, Optional, Dict

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Pt, Cm, RGBColor
except ImportError as exc:
    raise SystemExit(
        "Missing dependency: python-docx. Install it with: "
        "python3 -m pip install python-docx"
    ) from exc


def set_run_font(run, size: int = 12, bold: bool = False, 
                 font_name: str = "宋体", color: Optional[str] = None) -> None:
    """设置字体样式"""
    # 设置字体大小
    run.font.size = Pt(size)
    run.bold = bold
    
    # 设置字体名称（包括中文和西文）
    run.font.name = font_name
    
    # 通过XML设置中文字体
    from lxml import etree
    rPr = run._element.get_or_add_rPr()
    
    # 查找或创建rFonts元素
    rFonts_elem = rPr.find(qn("w:rFonts"))
    if rFonts_elem is None:
        rFonts_elem = rPr.makeelement(qn("w:rFonts"), {})
        rPr.append(rFonts_elem)
    
    # 设置中文字体属性
    rFonts_elem.set(qn("w:eastAsia"), font_name)
    rFonts_elem.set(qn("w:ascii"), "Times New Roman")
    rFonts_elem.set(qn("w:hAnsi"), "Times New Roman")
    
    if color:
        run.font.color.rgb = RGBColor(*hex_to_rgb(color))


def hex_to_rgb(hex_color: str) -> tuple:
    """Convert hex color to RGB tuple"""
    hex_color = hex_color.lstrip('#')
    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))


def set_paragraph_spacing(paragraph, before: float = 0, after: float = 6, 
                          line_spacing: float = 1.5, first_line_indent: Optional[float] = None) -> None:
    """设置段落间距"""
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line_spacing
    if first_line_indent is not None:
        paragraph.paragraph_format.first_line_indent = Cm(first_line_indent)


def add_centered_title(document: Document, title: str) -> None:
    """添加标题（居中，宋体，二号=22pt）"""
    paragraph = document.add_paragraph()
    # 清除段落默认样式
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(12)
    paragraph.paragraph_format.line_spacing = 1.0
    # 设置居中对齐
    from docx.oxml.ns import qn as qn_ns
    pPr = paragraph._p.get_or_add_pPr()
    jc = pPr.makeelement(qn_ns('w:jc'), {qn_ns('w:val'): 'center'})
    pPr.append(jc)
    
    run = paragraph.add_run(title)
    set_run_font(run, size=22, bold=True, font_name="宋体")


def add_paragraph(document: Document, text: str, font_size: int = 16,
                  font_name: str = "仿宋", bold: bool = False,
                  alignment: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT,
                  first_line_indent: Optional[float] = 2.0,
                  before: float = 0, after: float = 6,
                  line_spacing: float = 1.5) -> None:
    """添加普通段落"""
    paragraph = document.add_paragraph()
    paragraph.alignment = alignment
    set_paragraph_spacing(paragraph, before=before, after=after, 
                         line_spacing=line_spacing, 
                         first_line_indent=first_line_indent if alignment == WD_ALIGN_PARAGRAPH.LEFT else None)
    
    run = paragraph.add_run(text)
    set_run_font(run, size=font_size, bold=bold, font_name=font_name)


def add_right_aligned_block(document: Document, lines: List[str], 
                            font_size: int = 16, font_name: str = "仿宋",
                            before: float = 12, after: float = 6) -> None:
    """添加右对齐文本块（用于尾部）"""
    from docx.oxml.ns import qn as qn_ns
    
    for line in lines:
        if line is None or line.strip() == "":
            # 空行
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(6)
            continue
            
        paragraph = document.add_paragraph()
        # 使用XML直接设置右对齐，确保生效
        paragraph.paragraph_format.space_before = Pt(before)
        paragraph.paragraph_format.space_after = Pt(after)
        
        pPr = paragraph._p.get_or_add_pPr()
        jc = pPr.makeelement(qn_ns('w:jc'), {qn_ns('w:val'): 'right'})
        pPr.append(jc)
        
        run = paragraph.add_run(line)
        set_run_font(run, size=font_size, font_name=font_name)
        before = 0  # 后续行不需要额外间距


def generate_appeal_docx(data: Dict, output_path: str) -> str:
    """生成民事上诉状Word文档"""
    document = Document()
    
    # 设置默认页面边距
    for section in document.sections:
        section.top_margin = Cm(3.7)
        section.bottom_margin = Cm(3.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.6)
    
    # 1. 标题（居中，宋体，二号）
    add_centered_title(document, "民事上诉状")
    
    # 2. 当事人信息
    appellant_text = data.get("appellant", "")
    appellee_text = data.get("appellee", "")
    
    if appellant_text:
        add_paragraph(document, appellant_text, first_line_indent=0, after=6)
    
    if appellee_text:
        add_paragraph(document, appellee_text, first_line_indent=0, after=12)
    
    # 3. 上诉起因
    cause = data.get("cause", "【案由待补充】")
    first_court = data.get("first-court", "【一审法院待补充】")
    judgment_date = data.get("judgment-date", "【裁判日期待补充】")
    case_no = data.get("case-no", "【案号待补充】")
    judgment_type = data.get("judgment-type", "判决")
    if judgment_type not in ("判决", "裁定"):
        judgment_type = "判决"

    appeal_intro = (
        f"上诉人因与被上诉人{cause}一案，不服{first_court}"
        f"于{judgment_date}作出的{case_no}民事{judgment_type}，现依法提起上诉。"
    )
    add_paragraph(document, appeal_intro, first_line_indent=2.0, after=12)
    
    # 4. 上诉请求
    add_paragraph(document, "上诉请求", font_size=14, bold=True, 
                 font_name="黑体", first_line_indent=0, after=6)
    
    requests = data.get("requests", "")
    if requests:
        for req_line in requests.strip().split("\n"):
            if req_line.strip():
                add_paragraph(document, req_line.strip(), first_line_indent=2.0, after=6)
    
    # 5. 事实与理由
    add_paragraph(document, "事实与理由", font_size=14, bold=True, 
                 font_name="黑体", first_line_indent=0, before=12, after=6)
    
    reasons = data.get("reasons", "")
    if reasons:
        # 按段落分割，保留原有结构
        reason_paragraphs = reasons.strip().split("\n\n")
        for idx, para in enumerate(reason_paragraphs):
            if para.strip():
                add_paragraph(document, para.strip(), first_line_indent=2.0, after=6)
    
    # 6. 尾部（右对齐）
    appeal_court = data.get("appeal-court", "【二审法院待补充】")
    signature = data.get("signature", "上诉人：【签名/盖章待补充】")
    date = data.get("date", "【日期待补充】")
    copies = data.get("copies", "附：本上诉状副本【】份")
    
    # 添加空行间隔
    add_paragraph(document, "", first_line_indent=0, before=24, after=0)
    
    # 右对齐尾部信息
    right_aligned_lines = [
        "此致",
        "",
        appeal_court,
        "",
        signature,
        "",
        date,
        "",
        copies
    ]
    add_right_aligned_block(document, right_aligned_lines, before=0, after=6)
    
    # 保存文档
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))
    
    return str(output_path)


def main():
    parser = argparse.ArgumentParser(description="生成民事上诉状Word文档")
    
    # 输入方式：JSON 或 命令行参数
    parser.add_argument("--json", type=str, help="JSON格式输入（文件或字符串）")
    parser.add_argument("--output", type=str, required=True, help="输出文件路径")
    
    # 命令行参数
    parser.add_argument("--appellant", type=str, help="上诉人信息")
    parser.add_argument("--appellee", type=str, help="被上诉人信息")
    parser.add_argument("--cause", type=str, help="案由")
    parser.add_argument("--first-court", type=str, help="一审法院")
    parser.add_argument("--case-no", type=str, help="一审案号")
    parser.add_argument("--judgment-date", type=str, help="裁判日期")
    parser.add_argument("--appeal-court", type=str, help="二审法院")
    parser.add_argument("--requests", type=str, help="上诉请求（多行用\\n分隔）")
    parser.add_argument("--reasons", type=str, help="事实与理由（多段落用\\n\\n分隔）")
    parser.add_argument("--signature", type=str, help="签名")
    parser.add_argument("--date", type=str, help="日期")
    parser.add_argument("--copies", type=str, help="副本份数")
    parser.add_argument("--judgment-type", type=str, choices=["判决", "裁定"], default="判决",
                        help="一审裁判类型：判决（默认）或裁定，用于上诉起因句渲染")
    
    args = parser.parse_args()
    
    # 解析输入数据
    if args.json:
        json_path = Path(args.json)
        if json_path.exists():
            with open(json_path, "r", encoding="utf-8") as f:
                data = json.load(f)
        else:
            data = json.loads(args.json)
    else:
        data = {
            "appellant": args.appellant,
            "appellee": args.appellee,
            "cause": args.cause,
            "first-court": args.first_court,
            "case-no": args.case_no,
            "judgment-date": args.judgment_date,
            "appeal-court": args.appeal_court,
            "requests": args.requests,
            "reasons": args.reasons,
            "signature": args.signature,
            "date": args.date,
            "copies": args.copies,
            "judgment-type": args.judgment_type,
        }
    
    # 生成文档
    output_file = generate_appeal_docx(data, args.output)
    print(f"✓ 民事上诉状已生成：{output_file}")


if __name__ == "__main__":
    main()
