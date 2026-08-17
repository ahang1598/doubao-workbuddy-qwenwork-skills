#!/usr/bin/env python3
"""Validate the generated DOCX and its case-set consistency."""

from __future__ import annotations

import argparse
import hashlib
import json
import re
import sys
import zipfile
from pathlib import Path
from typing import Any

from docx import Document


REQUIRED_PARTS = {"[Content_Types].xml", "word/document.xml"}
REQUIRED_HEADINGS = (
    "一、检索说明",
    "二、检索目标",
    "三、检索结果",
    "四、检索内容",
    "五、附件",
)
PLACEHOLDER_PATTERNS = (
    re.compile(r"\{\{.*?\}\}"),
    re.compile(r"\[(?:待补充|请填写|填写|律师|律所|案件名称)\]"),
    re.compile(r"\b(?:TODO|TBD)\b", re.IGNORECASE),
)


def case_fingerprint(cases: list[dict[str, Any]]) -> str:
    canonical = [
        {
            "case_id": str(case.get("case_id") or ""),
            "title": str(case.get("title") or ""),
            "case_no": str(case.get("case_no") or ""),
            "court": str(case.get("court") or ""),
            "decision_date": str(case.get("decision_date") or ""),
        }
        for case in cases
    ]
    payload = json.dumps(canonical, ensure_ascii=False, separators=(",", ":"))
    return hashlib.sha256(payload.encode("utf-8")).hexdigest()


def document_text(document: Document) -> str:
    blocks: list[str] = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            blocks.extend(cell.text for cell in row.cells)
    return "\n".join(blocks)


def validate_layout(document: Document, errors: list[str]) -> None:
    if not document.sections:
        errors.append("DOCX 没有页面节")
    else:
        section = document.sections[0]
        width_cm = section.page_width.cm
        height_cm = section.page_height.cm
        if abs(width_cm - 21.0) > 0.1 or abs(height_cm - 29.7) > 0.1:
            errors.append(f"页面不是 A4：{width_cm:.2f}cm × {height_cm:.2f}cm")

    for style_name in ("Title", "Heading 1", "Heading 2", "Heading 3"):
        color = document.styles[style_name].font.color.rgb
        if color is None or str(color).upper() != "000000":
            errors.append(f"标题样式不是固定黑色：{style_name}")


def validate(data: dict[str, Any], docx_path: Path) -> list[str]:
    errors: list[str] = []
    if not docx_path.exists():
        return [f"DOCX 不存在：{docx_path}"]
    if docx_path.stat().st_size == 0:
        return ["DOCX 文件为空"]
    if not zipfile.is_zipfile(docx_path):
        return ["文件不是合法 DOCX ZIP 容器"]

    try:
        with zipfile.ZipFile(docx_path) as archive:
            missing_parts = REQUIRED_PARTS - set(archive.namelist())
            if missing_parts:
                errors.append(f"DOCX 缺少必要部件：{', '.join(sorted(missing_parts))}")
    except (OSError, zipfile.BadZipFile) as exc:
        return [f"DOCX ZIP 无法打开：{exc}"]

    try:
        document = Document(docx_path)
    except Exception as exc:
        return errors + [f"python-docx 无法打开文件：{exc}"]

    body = document_text(document)
    validate_layout(document, errors)
    if len(body.strip()) < 100:
        errors.append("DOCX 正文为空或过短")
    for heading in REQUIRED_HEADINGS:
        if heading not in body:
            errors.append(f"缺少关键章节：{heading}")

    cases = data.get("cases", [])
    for index, case in enumerate(cases, start=1):
        title = str(case.get("title") or "").strip()
        case_no = str(case.get("case_no") or "").strip()
        court = str(case.get("court") or "").strip()
        if title and title not in body:
            errors.append(f"第 {index} 个案例名称未出现在 DOCX：{title}")
        if case_no and case_no not in body:
            errors.append(f"第 {index} 个案例案号未出现在 DOCX：{case_no}")
        if court and court not in body:
            errors.append(f"第 {index} 个案例法院未出现在 DOCX：{court}")

    expected_marker = f"lawd-case-set-sha256={case_fingerprint(cases)};count={len(cases)}"
    actual_marker = document.core_properties.comments or ""
    if actual_marker != expected_marker:
        errors.append("DOCX 内嵌案例集合摘要与输入 JSON 不一致")

    for pattern in PLACEHOLDER_PATTERNS:
        match = pattern.search(body)
        if match:
            errors.append(f"DOCX 含模板占位符：{match.group(0)}")

    return errors


def main() -> int:
    parser = argparse.ArgumentParser(description="校验类案检索报告 DOCX")
    parser.add_argument("input_json", type=Path)
    parser.add_argument("input_docx", type=Path)
    args = parser.parse_args()

    try:
        data = json.loads(args.input_json.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError) as exc:
        print(f"校验失败：无法读取输入 JSON：{exc}", file=sys.stderr)
        return 1
    if not isinstance(data, dict) or not isinstance(data.get("cases"), list):
        print("校验失败：输入 JSON 缺少 cases 数组", file=sys.stderr)
        return 1

    errors = validate(data, args.input_docx)
    if errors:
        print(f"DOCX 校验失败，共 {len(errors)} 项：", file=sys.stderr)
        for error in errors:
            print(f"- {error}", file=sys.stderr)
        return 1

    print(
        f"DOCX 校验通过：{args.input_docx}；{len(data['cases'])} 个案例；"
        "A4 页面、黑色标题、章节、案例集合和占位符检查均通过。"
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
