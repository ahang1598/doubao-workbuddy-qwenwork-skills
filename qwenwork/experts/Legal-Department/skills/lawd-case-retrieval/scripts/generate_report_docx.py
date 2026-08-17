#!/usr/bin/env python3
"""Generate a professionally formatted Chinese legal case retrieval report."""

from __future__ import annotations

import argparse
import hashlib
import json
import os
import subprocess
import sys
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


MISSING = "检索结果未提供"
BODY_FONT = os.environ.get("LAWDING_DOCX_BODY_FONT", "Songti SC")
HEADING_FONT = os.environ.get("LAWDING_DOCX_HEADING_FONT", "Heiti SC")


def text(value: Any) -> str:
    if value is None:
        return MISSING
    if isinstance(value, list):
        items = [str(item).strip() for item in value if str(item).strip()]
        return "、".join(items) if items else MISSING
    value = str(value).strip()
    return value or MISSING


def case_fingerprint(cases: list[dict[str, Any]]) -> str:
    canonical = [
        {
            "case_id": str(case.get("case_id") or ""),
            "title": str(case.get("title") or ""),
            "case_no": str(case.get("case_no") or ""),
            "court": str(case.get("court") or ""),
            "decision_date": str(case.get("decision_date") or ""),
        }
        for case in cases
    ]
    payload = json.dumps(canonical, ensure_ascii=False, separators=(",", ":"))
    return hashlib.sha256(payload.encode("utf-8")).hexdigest()


def set_run_font(run, chinese: str = BODY_FONT, size: float = 12):
    run.font.name = chinese
    run.font.size = Pt(size)
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        fonts.set(qn(key), chinese)


def configure_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(12)
    normal_fonts = normal._element.get_or_add_rPr().get_or_add_rFonts()
    for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        normal_fonts.set(qn(key), BODY_FONT)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_after = Pt(6)

    for name, font, size, bold in (
        ("Title", BODY_FONT, 22, True),
        ("Heading 1", HEADING_FONT, 16, True),
        ("Heading 2", HEADING_FONT, 14, True),
        ("Heading 3", HEADING_FONT, 12, True),
    ):
        style = styles[name]
        style.font.name = font
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = RGBColor(0, 0, 0)
        style_fonts = style._element.get_or_add_rPr().get_or_add_rFonts()
        for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
            style_fonts.set(qn(key), font)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)
        paragraph_properties = style._element.get_or_add_pPr()
        border = paragraph_properties.find(qn("w:pBdr"))
        if border is not None:
            paragraph_properties.remove(border)

    if "Case Quote" not in styles:
        quote = styles.add_style("Case Quote", WD_STYLE_TYPE.PARAGRAPH)
        quote.base_style = styles["Normal"]
        quote.paragraph_format.left_indent = Cm(0.74)
        quote.paragraph_format.right_indent = Cm(0.74)
        quote.paragraph_format.space_before = Pt(4)
        quote.paragraph_format.space_after = Pt(8)
        quote.font.size = Pt(11)
        quote.font.name = BODY_FONT
        quote_fonts = quote._element.get_or_add_rPr().get_or_add_rFonts()
        for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
            quote_fonts.set(qn(key), BODY_FONT)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instr, end))


def add_bullet(document: Document, value: Any) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.add_run(text(value))


def add_numbered_items(document: Document, values: Any, empty_message: str) -> None:
    if not isinstance(values, list) or not values:
        document.add_paragraph(empty_message)
        return
    for index, item in enumerate(values, start=1):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Cm(0.74)
        paragraph.paragraph_format.first_line_indent = Cm(-0.74)
        paragraph.add_run(f"{index}. {text(item)}")


def add_info_table(document: Document, rows: list[tuple[str, Any]]) -> None:
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].width = Cm(3.4)
        cells[1].width = Cm(12.2)
        cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cells[0].text = label
        cells[1].text = text(value)
        for run in cells[0].paragraphs[0].runs:
            run.bold = True
            set_run_font(run, chinese=HEADING_FONT, size=10.5)
        for run in cells[1].paragraphs[0].runs:
            set_run_font(run, size=10.5)


def add_constraints(document: Document, constraints: dict[str, Any]) -> None:
    date_from = str(constraints.get("date_from") or "").strip()
    date_to = str(constraints.get("date_to") or "").strip()
    regions = constraints.get("regions") or []
    levels = constraints.get("court_levels") or []
    other = constraints.get("other") or []

    add_bullet(document, f"时间限制：{date_from or '未设定'} 至 {date_to or '未设定'}")
    add_bullet(document, f"地域限制：{text(regions) if regions else '未设定特定地域限制'}")
    add_bullet(document, f"法院层级：{text(levels) if levels else '未设定特定法院层级'}")
    if other:
        add_bullet(document, f"其他明确约束：{text(other)}")


def generate(data: dict[str, Any], output: Path) -> None:
    document = Document()
    configure_styles(document)
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.7)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.2)
    section.footer_distance = Cm(1.2)

    report = data.get("report", {})
    cases = data.get("cases", [])
    title = text(report.get("title")) if report.get("title") else "类案检索报告"
    document.core_properties.title = title
    document.core_properties.subject = text(report.get("matter"))
    document.core_properties.author = text(report.get("prepared_by"))
    document.core_properties.comments = (
        f"lawd-case-set-sha256={case_fingerprint(cases)};count={len(cases)}"
    )

    header = section.header.paragraphs[0]
    header.text = title
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in header.runs:
        set_run_font(run, chinese=BODY_FONT, size=9)
    add_page_number(section.footer.paragraphs[0])

    paragraph = document.add_paragraph(style="Title")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run(title)
    if report.get("matter"):
        subtitle = document.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(text(report.get("matter")))
        set_run_font(run, chinese=HEADING_FONT, size=14)
    document.add_paragraph()
    add_info_table(
        document,
        [
            ("制作主体", report.get("prepared_by")),
            ("制作日期", report.get("prepared_at")),
            ("数据来源", report.get("data_source")),
            ("检索日期", report.get("retrieved_at")),
        ],
    )
    document.add_page_break()

    document.add_heading("一、检索说明", level=1)
    add_info_table(
        document,
        [
            ("检索平台/数据来源", report.get("data_source")),
            ("检索方法", report.get("methods")),
            ("检索查询", data.get("query")),
            ("纳入案例数量", str(len(cases))),
        ],
    )
    document.add_heading("用户明确约束", level=2)
    add_constraints(document, data.get("explicit_constraints", {}))

    document.add_heading("二、检索目标", level=1)
    add_numbered_items(document, data.get("retrieval_targets"), "未提供明确检索目标。")

    document.add_heading("三、检索结果", level=1)
    document.add_paragraph(f"以下结论仅基于本次纳入的 {len(cases)} 个案例：")
    add_numbered_items(document, data.get("conclusions"), "现有输入未提供可形成报告结论的分析内容。")

    document.add_heading("四、检索内容", level=1)
    for index, case in enumerate(cases, start=1):
        case_title = text(case.get("title"))
        case_no = text(case.get("case_no"))
        document.add_heading(f"{index}. {case_title}【案号：{case_no}】", level=2)
        add_info_table(
            document,
            [
                ("审理法院", case.get("court")),
                ("裁判日期", case.get("decision_date")),
                ("审判人员", case.get("judges")),
                ("数据来源", case.get("source_url")),
                ("原始记录定位", case.get("raw_record_locator")),
            ],
        )
        document.add_heading("基本案情", level=3)
        document.add_paragraph(text(case.get("facts")))
        document.add_heading("裁判要旨", level=3)
        document.add_paragraph(text(case.get("holding")))
        document.add_heading("裁判观点", level=3)
        reasoning = str(case.get("court_reasoning_quote") or "").strip()
        document.add_paragraph(
            reasoning or "该案例未提供‘本院认为’原文。",
            style="Case Quote",
        )
        document.add_heading("法律依据", level=3)
        legal_basis = case.get("legal_basis")
        if isinstance(legal_basis, list) and legal_basis:
            for basis in legal_basis:
                add_bullet(document, basis)
        else:
            document.add_paragraph(MISSING)

    document.add_heading("五、附件", level=1)
    any_full_text = any(str(case.get("full_text") or "").strip() for case in cases)
    if not any_full_text:
        document.add_paragraph("本次检索数据未包含案例全文。以下列明案例来源定位信息：")
        for index, case in enumerate(cases, start=1):
            add_bullet(
                document,
                f"{index}. {text(case.get('title'))}；案号：{text(case.get('case_no'))}；"
                f"来源：{text(case.get('source_url'))}；定位：{text(case.get('raw_record_locator'))}",
            )
    else:
        for index, case in enumerate(cases, start=1):
            full_text = str(case.get("full_text") or "").strip()
            if not full_text:
                add_bullet(
                    document,
                    f"{index}. {text(case.get('title'))}：本次检索数据未包含案例全文；"
                    f"来源定位：{text(case.get('raw_record_locator'))}",
                )
                continue
            document.add_page_break()
            document.add_heading(f"附件 {index}：{text(case.get('title'))}", level=2)
            document.add_paragraph(f"案号：{text(case.get('case_no'))}")
            for block in full_text.splitlines():
                if block.strip():
                    document.add_paragraph(block.strip())

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def main() -> int:
    parser = argparse.ArgumentParser(description="生成类案检索报告 DOCX")
    parser.add_argument("input_json", type=Path)
    parser.add_argument("output_docx", type=Path)
    args = parser.parse_args()

    if args.output_docx.suffix.lower() != ".docx":
        print("生成失败：输出文件必须使用 .docx 扩展名", file=sys.stderr)
        return 1

    validator = Path(__file__).with_name("validate_report_cases.py")
    validation = subprocess.run(
        [sys.executable, str(validator), str(args.input_json)],
        check=False,
        text=True,
        capture_output=True,
    )
    if validation.returncode != 0:
        print(validation.stderr or validation.stdout, file=sys.stderr, end="")
        return validation.returncode

    try:
        data = json.loads(args.input_json.read_text(encoding="utf-8"))
        generate(data, args.output_docx)
    except Exception as exc:
        print(f"生成失败：{exc}", file=sys.stderr)
        return 1

    print(f"DOCX 已生成：{args.output_docx}（{len(data['cases'])} 个案例）")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
