#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
要素式民事起诉状 Word 文档生成脚本

设计目标：
    针对最高人民法院《关于印发部分案件民事起诉状、答辩状示范文本（试行）的通知》
    要素式示范文本，提供两种生成路径：

    路径 A 「严格按模板填充」：当案由命中 references/templates/element-based/
        目录下的空白模板时，直接打开 docx 模板，**完整保留模板原有的字体、字号、
        表格、复选框、段落样式**，仅按调用方提供的 replacements / checkboxes /
        appendices 指令进行内容填充。

    路径 B 「通用要素式」：当案由未命中模板时，按通用 5 段骨架生成 Word，
        并在文档开头加入《通知》未覆盖的警示语。此路径的字体字号采用 python-docx
        默认值（不刻意模仿传统起诉状样式）。

使用方式：
    python3 generate_element_complaint.py --input data.json --output /path/out.docx
    python3 generate_element_complaint.py --input data.json --output-dir /tmp

JSON Schema 详见 references/script-usage.md「第九节 要素式脚本」。
"""

import argparse
import copy
import json
import os
import sys
from datetime import datetime

try:
    from docx import Document
except ImportError:
    sys.stderr.write(
        "[错误] 缺少依赖 python-docx。\n"
        "请执行：pip install python-docx\n"
    )
    sys.exit(2)


# ============== 路径常量 ==============
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATE_DIR = os.path.normpath(
    os.path.join(SCRIPT_DIR, "..", "references", "templates", "element-based")
)

# 《通知》明确覆盖的 11 类案由白名单（用于警示语判断）
OFFICIAL_11 = {
    "金融借款合同纠纷",
    "银行信用卡纠纷",
    "民间借贷纠纷",
    "保证保险合同纠纷",
    "买卖合同纠纷",
    "物业服务合同纠纷",
    "融资租赁合同纠纷",
    "证券虚假陈述责任纠纷",
    "机动车交通事故责任纠纷",
    "劳动争议纠纷",
    "离婚纠纷",
}


# ============== 各案由结构化字段映射表 ==============
# 格式：字段名 → 模板中对应的 find 字符串
# 脚本根据 template_fill.fields 中的键，查此表生成 replacements
# 如字段名在此表中无对应项，则 stderr 输出警告并跳过

FIELD_MAPPINGS = {
    # ---- 金融借款合同纠纷 ----
    "金融借款": {
        "截止日期_年":       "   年",
        "截止日期_月":       "  月",
        "截止日期_日":       "  日止",
        "尚欠本金":          "尚欠本金          元",
        "欠利息":            "欠利息     元",
        "欠复利":            "复利   元",
        "欠罚息":            "罚息(违约金)\n元",
        "利息计算方式":      "计算方式：\n",
        "是否支付至清偿日":  "是否请求支付至实际清偿之日止：是",
        "贷款合同签订日期":  "合同编号",
        "贷款金额":          "借款金额",
        "贷款期限":          "借款期限",
        "贷款用途":          "借款用途",
        "担保方式":          "担保方式",
        "担保人":            "担保人",
        "逾期起始日":        "逾期起始日",
        "管辖约定":          "管辖约定",
        "保全申请":          "保全申请",
        "送达地址":          "地址：\n",
        "电子送达方式":      "短信",
    },
    # ---- 民间借贷纠纷 ----
    "民间借贷": {
        "借款金额":          "借款金额",
        "借款日期":          "借款日期",
        "约定还款日":        "约定还款日",
        "尚欠本金":          "尚欠本金",
        "利率约定":          "利率约定",
        "截止日期":          "截止计算日",
        "欠付利息":          "欠付利息",
        "利息计算方式":      "利息计算方式",
        "是否支付至清偿日":  "是否支付至清偿日",
        "担保方式":          "担保方式",
        "担保人":            "担保人",
        "借款用途":          "借款用途",
        "出借方式":          "出借方式",
        "管辖约定":          "管辖约定",
        "保全申请":          "保全申请",
    },
    # ---- 买卖合同纠纷 ----
    "买卖合同": {
        "合同签订日期":      "合同签订日期",
        "货物名称规格":      "货物名称规格",
        "合同总价":          "合同总价",
        "已付款金额":        "已付款金额",
        "欠付货款":          "欠付货款",
        "交货日期":          "交货日期",
        "是否已交货":        "是否已交货",
        "逾期付款起算日":    "逾期付款起算日",
        "利率标准":          "利率标准",
        "违约金约定":        "违约金约定",
        "是否支付至清偿日":  "是否支付至清偿日",
        "管辖约定":          "管辖约定",
    },
    # ---- 建设工程施工合同纠纷 ----
    "建设工程": {
        "工程名称":          "工程名称",
        "工程地点":          "工程地点",
        "合同签订日期":      "合同签订日期",
        "合同总价":          "合同总价",
        "竣工验收日期":      "竣工验收日期",
        "已付款金额":        "已付款金额",
        "欠付工程款":        "欠付工程款",
        "逾期利息起算日":    "逾期利息起算日",
        "利率标准":          "利率标准",
        "优先受偿权金额":    "优先受偿权金额",
        "优先受偿权起算日":  "优先受偿权起算日",
        "实际施工人身份":    "实际施工人身份",
        "管辖约定":          "管辖约定",
    },
    # ---- 劳动争议纠纷 ----
    "劳动争议": {
        "用人单位名称":      "用人单位名称",
        "劳动合同签订日期":  "劳动合同签订日期",
        "劳动合同期限":      "劳动合同期限",
        "岗位":              "岗位",
        "月工资":            "月工资",
        "欠薪金额":          "欠薪金额",
        "欠薪期间":          "欠薪期间",
        "解除日期":          "解除日期",
        "解除原因":          "解除原因",
        "仲裁裁决书号":      "仲裁裁决书号",
        "仲裁裁决日期":      "仲裁裁决日期",
        "管辖约定":          "管辖约定",
    },
    # ---- 离婚纠纷 ----
    "离婚": {
        "结婚登记日期":      "结婚登记日期",
        "婚姻登记机关":      "婚姻登记机关",
        "子女信息":          "子女信息",
        "共同财产":          "共同财产",
        "共同债务":          "共同债务",
        "离婚原因":          "离婚原因",
        "管辖约定":          "管辖约定",
        # 扩展字段（第三轮）
        "marriage_date":     "结婚登记日期",
        "children_info":     "子女信息",
        "joint_property":    "共同财产",
        "divorce_fault":     "离婚原因",
        "monthly_support":   "每月抚养费",
        "damage_amount":     "损害赔偿金额",
    },
    # ---- 离婚后财产纠纷 ----
    "离婚后财产": {
        "divorce_date":          "离婚日期",
        "divorce_method":        "离婚方式",
        "omitted_property":      "遗漏财产",
        "property_value":        "财产价值",
        "compensation_amount":   "折价补偿款",
        "transfer_amount":       "转移金额",
        "transfer_date":         "转移日期",
        "agreement_clause":      "协议条款",
    },
    # ---- 抚养费纠纷 ----
    "抚养费": {
        "child_name":            "子女姓名",
        "child_birthdate":       "子女出生日期",
        "monthly_support":       "月抚养费",
        "support_start_date":    "抚养费起算日期",
        "arrears_months":        "欠付月数",
        "arrears_amount":        "欠付总额",
        "income_basis":          "月收入基数",
        "income_percentage":     "收入比例",
        "change_reason":         "增减事由",
    },
    # ---- 变更抚养关系纠纷 ----
    "变更抚养": {
        "child_name":            "子女姓名",
        "child_age":             "子女年龄",
        "current_custodian":     "现直接抚养方",
        "change_reason":         "变更事由",
        "actual_residence_start": "实际居住起始",
        "child_will":            "子女意愿",
        "new_monthly_support":   "变更后月抚养费",
        "visitation_frequency":  "探望安排",
    },
    # ---- 分家析产纠纷 ----
    "分家析产": {
        "property_address":      "房屋地址",
        "property_area":         "建筑面积",
        "family_members":        "家庭成员",
        "plaintiff_contribution": "原告出资/贡献",
        "property_value":        "财产总价值",
        "plaintiff_share_ratio": "原告份额比例",
        "compensation_amount":   "折价补偿款",
        "demolition_compensation": "拆迁补偿总额",
        "household_population":  "安置人口数",
    },
    # ---- 房屋买卖合同纠纷 ----
    "房屋买卖": {
        "house_address":         "房屋地址",
        "house_area":            "建筑面积",
        "property_cert_no":      "不动产权证号",
        "total_price":           "合同总价",
        "paid_amount":           "已付购房款",
        "contract_date":         "合同签订日期",
        "delivery_date":         "约定交付日期",
        "actual_delivery_date":  "实际交付日期",
        "overdue_days":          "逾期天数",
        "daily_penalty_rate":    "日违约金率",
        "penalty_amount":        "违约金金额",
        "breach_type":           "违约类型",
        "lawyer_fee":            "律师费",
    },
    # ---- 名誉权纠纷 ----
    "名誉权": {
        "infringement_platform":  "侵权平台",
        "infringement_content":   "侵权内容",
        "infringement_date":      "侵权日期",
        "spread_scope":           "传播范围",
        "apology_channel":        "道歉渠道",
        "apology_duration":       "道歉持续时间",
        "mental_damage_amount":   "精神损害抚慰金",
        "lawyer_fee":             "律师费",
        "platform_defendant":     "是否追加平台被告",
    },
    # ---- 生命权/身体权/健康权纠纷 ----
    "生命健康": {
        "injury_date":            "受伤日期",
        "injury_description":     "伤情描述",
        "disability_level":       "伤残等级",
        "medical_fee":            "医疗费",
        "lost_income":            "误工费",
        "nursing_fee":            "护理费",
        "nutrition_fee":          "营养费",
        "transport_fee":          "交通费",
        "hospitalization_allowance": "住院伙食补助费",
        "disability_compensation": "残疾赔偿金",
        "death_compensation":     "死亡赔偿金",
        "funeral_fee":            "丧葬费",
        "mental_damage_amount":   "精神损害抚慰金",
        "fault_ratio":            "被告过错比例",
        "admin_penalty_doc":      "行政处罚决定书",
        "total_claim":            "赔偿总额",
    },
}

# 案由 → 映射表 key 的模糊匹配
_CAUSE_TO_MAPPING_KEY = [
    ("金融借款", "金融借款"),
    ("银行信用卡", "金融借款"),
    ("民间借贷", "民间借贷"),
    ("买卖合同", "买卖合同"),
    ("房屋买卖", "房屋买卖"),
    ("建设工程", "建设工程"),
    ("劳动争议", "劳动争议"),
    ("离婚后财产", "离婚后财产"),
    ("离婚", "离婚"),
    ("抚养费", "抚养费"),
    ("变更抚养", "变更抚养"),
    ("分家析产", "分家析产"),
    ("名誉权", "名誉权"),
    ("生命权", "生命健康"),
    ("身体权", "生命健康"),
    ("健康权", "生命健康"),
    ("人身损害", "生命健康"),
]


def _get_mapping_for_cause(case_cause):
    """根据案由名称返回对应的字段映射表（dict），找不到返回 None。"""
    for keyword, key in _CAUSE_TO_MAPPING_KEY:
        if keyword in (case_cause or ""):
            return FIELD_MAPPINGS.get(key)
    return None


def _fields_to_replacements(fields, case_cause):
    """将 template_fill.fields 字典转换为 replacements 数组。
    返回 (converted_list, unrecognized_keys)。"""
    mapping = _get_mapping_for_cause(case_cause)
    converted = []
    unrecognized = []
    for field_name, field_value in (fields or {}).items():
        if not mapping or field_name not in mapping:
            unrecognized.append(field_name)
            continue
        find_str = mapping[field_name]
        if not field_value:
            field_value = "（未提供）"
            sys.stderr.write(f"[提示] 字段「{field_name}」值为空，已用「（未提供）」填充。\n")
        converted.append({"find": find_str, "replace": str(field_value), "once": False})
    return converted, unrecognized


# ============== 工具：模板查找 ==============
def find_template(case_cause):
    """根据案由查找空白模板文件路径，找不到返回 None。
    支持精确匹配与子串匹配（去除括号/空格后再比较）。"""
    if not os.path.isdir(TEMPLATE_DIR):
        return None
    target_clean = _clean(case_cause)
    candidates = []
    for fname in os.listdir(TEMPLATE_DIR):
        if not fname.endswith(".docx") or fname.startswith("~$"):
            continue
        # 文件名形如：要素式起诉状-XXX.docx
        stem = fname[:-5]  # remove .docx
        if stem.startswith("要素式起诉状-"):
            type_name = stem[len("要素式起诉状-"):]
        else:
            type_name = stem
        type_clean = _clean(type_name)
        if type_clean == target_clean:
            return os.path.join(TEMPLATE_DIR, fname)
        if target_clean and (target_clean in type_clean or type_clean in target_clean):
            candidates.append((len(type_clean), os.path.join(TEMPLATE_DIR, fname)))
    if candidates:
        candidates.sort()  # 最短的子串匹配优先（更可能精确）
        return candidates[0][1]
    return None


def _clean(s):
    if not s:
        return ""
    return s.replace(" ", "").replace("　", "").replace("（", "").replace("）", "") \
        .replace("(", "").replace(")", "").replace("、", "")


# ============== 工具：遍历 docx 段落 ==============
def _iter_paragraphs(doc):
    """遍历文档中所有 paragraph，包括表格单元格内的。"""
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p
                # 嵌套表格
                for nested in cell.tables:
                    for nrow in nested.rows:
                        for ncell in nrow.cells:
                            for np in ncell.paragraphs:
                                yield np


def _paragraph_text(p):
    return "".join(run.text for run in p.runs)


def _replace_in_paragraph(p, find_text, replace_text):
    """在段落内做字符串替换，**保留原 run 的格式**。
    实现策略：把 runs 合并到第一个 run 的 text 上，做 str.replace，
    再清空其余 run 的 text。这样字体属性以第一个 run 为准（其余 run
    若与之同格式则视觉无差异；模板中绝大多数 run 同段同格式）。"""
    if not p.runs:
        return 0
    full = _paragraph_text(p)
    if find_text not in full:
        return 0
    count = full.count(find_text)
    new_full = full.replace(find_text, replace_text)
    p.runs[0].text = new_full
    for r in p.runs[1:]:
        r.text = ""
    return count


# ============== 路径 A：模板填充 ==============
def fill_template(template_path, data, output_path):
    """打开模板，应用 replacements / checkboxes / appendices 后另存。"""
    doc = Document(template_path)

    instructions = data.get("template_fill") or {}
    total_replacements = 0
    unmatched = []

    # 0) 结构化字段映射：将 template_fill.fields 转换为 replacements 并前置合并
    fields = instructions.get("fields") or {}
    if fields:
        converted, unrecognized = _fields_to_replacements(fields, data.get("case_cause", ""))
        if unrecognized:
            sys.stderr.write(
                f"[提示] 以下字段名在案由「{data.get('case_cause','')}」的映射表中未找到，已跳过：\n"
            )
            for k in unrecognized:
                sys.stderr.write(f"  - {k}\n")
        # fields 转换结果优先，再拼接手动 replacements（手动可覆盖）
        manual_replacements = instructions.get("replacements") or []
        all_replacements = converted + manual_replacements
    else:
        all_replacements = instructions.get("replacements") or []

    # 1) 文本替换（按 find/replace 数组顺序执行；同一 find 出现多次按顺序逐个替换）
    replacements = all_replacements
    for item in replacements:
        find_text = item.get("find") or ""
        replace_text = item.get("replace") or ""
        if not find_text:
            continue
        # 支持 once: True/False —— 默认 False 即替换该 find 在文档中的所有出现
        once = bool(item.get("once", False))
        applied = 0
        for p in _iter_paragraphs(doc):
            n = _replace_in_paragraph(p, find_text, replace_text)
            if n > 0:
                applied += n
                total_replacements += n
                if once and applied > 0:
                    break
        if applied == 0:
            unmatched.append(find_text)

    # 2) 复选框勾选：anchor + select(是/否) → 把 anchor 之后最近的对应 ☐ 改为 ☑
    #    JSON 格式： {"checkboxes": [{"anchor": "是否约定管辖", "select": "是"}]}
    for cb in instructions.get("checkboxes") or []:
        anchor = cb.get("anchor") or ""
        select = cb.get("select") or ""
        if not anchor or select not in ("是", "否"):
            continue
        _apply_checkbox(doc, anchor, select)

    # 3) 在文档末尾追加段落（如附加证据清单、补充说明）
    for line in instructions.get("appendices") or []:
        doc.add_paragraph(line)

    if unmatched:
        sys.stderr.write(
            "[提示] 以下 find 字符串在模板中未找到（可能模板已更新或写法不一致）：\n"
        )
        for u in unmatched:
            sys.stderr.write(f"  - {u}\n")

    doc.save(output_path)
    return output_path, total_replacements


def _apply_checkbox(doc, anchor, select):
    """在 anchor 文本所在段落（或紧随其后的段落）中，根据 select=是/否，
    把对应的 ☐/□ 改为 ☑。同时把另一选项保持为 ☐。"""
    paragraphs = list(_iter_paragraphs(doc))
    for idx, p in enumerate(paragraphs):
        text = _paragraph_text(p)
        if anchor not in text:
            continue
        # 在当前段以及紧随其后的最多 2 段范围内查找复选框
        for target in paragraphs[idx: idx + 3]:
            t = _paragraph_text(target)
            if not any(box in t for box in ("☐", "□", "☑")):
                continue
            # 是/否结构：根据 select 找到对应字符旁的框
            # 例如 "是☐ 否☐" 或 "是□ 否□"
            new_t = _toggle_checkbox(t, select)
            if new_t != t and target.runs:
                target.runs[0].text = new_t
                for r in target.runs[1:]:
                    r.text = ""
                return True
        return False
    return False


def _toggle_checkbox(text, select):
    """把"是"前/后的框设为☑，"否"前/后的框设为☐（或反之）。
    支持 ☐/□ 两种空框写法。"""
    EMPTY = ("☐", "□")
    FILLED = "☑"

    def _set(label, on):
        """在 label 字符附近找到一个 EMPTY 框替换为 ☑ 或保持 ☐"""
        nonlocal text
        # 优先匹配 label + 框（如"是☐"），其次匹配 框 + label
        for empty in EMPTY:
            pattern = label + empty
            if pattern in text:
                text = text.replace(pattern, label + (FILLED if on else "☐"), 1)
                return True
        for empty in EMPTY:
            pattern = empty + label
            if pattern in text:
                text = text.replace(pattern, (FILLED if on else "☐") + label, 1)
                return True
        return False

    if select == "是":
        _set("是", True)
        _set("否", False)
    elif select == "否":
        _set("是", False)
        _set("否", True)
    return text


# ============== 路径 B：通用要素式 ==============
GENERAL_WARNING = (
    "⚠ 本案由不在最高人民法院《关于印发部分案件民事起诉状、答辩状示范文本"
    "（试行）的通知》要素式示范文本覆盖范围，本文按通用要素式结构输出，仅供"
    "参考。建议立案前向受理法院核实是否接受。"
)


def build_general_element(data, output_path):
    """无模板时生成通用要素式起诉状。
    字体字号使用 python-docx 默认值，不刻意模仿传统起诉状样式。"""
    doc = Document()

    doc.add_paragraph(GENERAL_WARNING)

    # 标题
    title = doc.add_paragraph()
    title.alignment = 1  # center
    run = title.add_run("民事起诉状")
    run.bold = True

    # ① 当事人信息
    doc.add_heading("一、当事人信息", level=2)
    for p_ in data.get("plaintiffs", []) or []:
        doc.add_paragraph(_format_party_inline("原告", p_))
    for d_ in data.get("defendants", []) or []:
        doc.add_paragraph(_format_party_inline("被告", d_))
    for tp_ in data.get("third_parties", []) or []:
        doc.add_paragraph(_format_party_inline("第三人", tp_))

    elements = data.get("elements") or {}
    if elements.get("service_address"):
        doc.add_paragraph(f"送达地址：{elements['service_address']}")
    if elements.get("electronic_service"):
        doc.add_paragraph(f"电子送达方式：{elements['electronic_service']}")

    # ② 诉讼请求和依据
    doc.add_heading("二、诉讼请求和依据", level=2)
    claim_elements = elements.get("claim_elements") or {}
    if isinstance(claim_elements, dict):
        for i, (k, v) in enumerate(claim_elements.items(), start=1):
            doc.add_paragraph(f"{i}. {k}：{v}")
    elif isinstance(claim_elements, list):
        for i, v in enumerate(claim_elements, start=1):
            doc.add_paragraph(f"{i}. {v}")
    if elements.get("claim_total"):
        doc.add_paragraph(f"标的总额：人民币 {elements['claim_total']} 元")
    if elements.get("claim_basis"):
        doc.add_paragraph(f"请求依据：{elements['claim_basis']}")

    # ③ 约定管辖和诉讼保全
    doc.add_heading("三、约定管辖和诉讼保全", level=2)
    doc.add_paragraph(f"约定管辖：{elements.get('jurisdiction_clause', '无')}")
    doc.add_paragraph(f"仲裁约定：{elements.get('arbitration_clause', '无')}")
    doc.add_paragraph(f"诉讼保全申请：{elements.get('preservation', '不申请')}")

    # ④ 事实和理由
    doc.add_heading("四、事实和理由", level=2)
    fact_elements = elements.get("fact_elements") or {}
    if isinstance(fact_elements, dict):
        for i, (k, v) in enumerate(fact_elements.items(), start=1):
            doc.add_paragraph(f"{i}. {k}：{v}")
    elif isinstance(fact_elements, list):
        for i, v in enumerate(fact_elements, start=1):
            doc.add_paragraph(f"{i}. {v}")

    # ⑤ 证据与附件
    doc.add_heading("五、证据与附件", level=2)
    if data.get("evidence_hint"):
        doc.add_paragraph(data["evidence_hint"])
    attachments = data.get("attachments") or []
    if attachments:
        doc.add_paragraph("附件：")
        for i, att in enumerate(attachments, start=1):
            doc.add_paragraph(f"{i}. {att}")

    # 收尾
    doc.add_paragraph("此致")
    doc.add_paragraph(data.get("court", "") or "")

    doc.add_paragraph("")
    p_signer = doc.add_paragraph()
    p_signer.alignment = 2  # right
    plaintiff_names = "、".join(
        [p.get("name", "") for p in (data.get("plaintiffs") or []) if p.get("name")]
    )
    p_signer.add_run(f"具状人（签字/盖章）：{plaintiff_names}")

    p_date = doc.add_paragraph()
    p_date.alignment = 2
    p_date.add_run(data.get("filing_date", "") or "")

    doc.save(output_path)
    return output_path


def _format_party_inline(role, party):
    name = party.get("name", "")
    ptype = party.get("type", "natural_person")
    bits = [f"{role}：{name}"]
    if ptype == "natural_person":
        for k, label in [
            ("gender", ""),
            ("ethnicity", "族"),
            ("birth_date", "出生"),
            ("id_number", "公民身份号码："),
            ("address", "住"),
            ("phone", "联系电话："),
        ]:
            v = party.get(k)
            if v:
                if label.endswith("："):
                    bits.append(f"{label}{v}")
                elif label == "出生":
                    bits.append(f"{v}{label}")
                elif label == "族":
                    bits.append(f"{v}{label}")
                elif label == "住":
                    bits.append(f"{label}{v}")
                else:
                    bits.append(v)
    elif ptype in ("legal_person", "unincorporated"):
        if party.get("address"):
            bits.append(f"住所地：{party['address']}")
        if party.get("uscc"):
            bits.append(f"统一社会信用代码：{party['uscc']}")
        rep_title = party.get("legal_rep_title") or (
            "法定代表人" if ptype == "legal_person" else "负责人"
        )
        if party.get("legal_rep_name"):
            bits.append(f"{rep_title}：{party['legal_rep_name']}")
        if party.get("phone"):
            bits.append(f"联系电话：{party['phone']}")
    return "，".join(bits) + "。"


# ============== 校验、命名、入口 ==============
def _validate(data):
    errors = []
    for k in ("court", "case_cause", "plaintiffs", "defendants", "filing_date"):
        if k not in data or not data[k]:
            errors.append(f"缺少必填字段：{k}")
    if isinstance(data.get("plaintiffs"), list) and len(data["plaintiffs"]) == 0:
        errors.append("至少需要 1 名原告")
    if isinstance(data.get("defendants"), list) and len(data["defendants"]) == 0:
        errors.append("至少需要 1 名被告")
    for i, p in enumerate(data.get("plaintiffs") or []):
        if not isinstance(p, dict) or not p.get("name"):
            errors.append(f"原告 #{i+1} 缺少 name")
    for i, p in enumerate(data.get("defendants") or []):
        if not isinstance(p, dict) or not p.get("name"):
            errors.append(f"被告 #{i+1} 缺少 name")
    if errors:
        sys.stderr.write("[校验失败]\n")
        for e in errors:
            sys.stderr.write(f"  - {e}\n")
        sys.exit(3)


def _short_name(party):
    name = (party or {}).get("name", "当事人") or "当事人"
    return name[:8]


def _default_filename(data):
    p_name = _short_name((data.get("plaintiffs") or [{}])[0])
    d_name = _short_name((data.get("defendants") or [{}])[0])
    cause = data.get("case_cause", "民事纠纷")
    today = datetime.now().strftime("%Y%m%d")
    return f"要素式民事起诉状-{p_name}诉{d_name}-{cause}-{today}.docx"


def main():
    parser = argparse.ArgumentParser(description="要素式民事起诉状 Word 生成器")
    parser.add_argument("--input", "-i", help="JSON 输入文件路径；不指定则从 stdin 读取")
    parser.add_argument("--output", "-o", help="输出 docx 文件路径")
    parser.add_argument("--output-dir", "-d", default=".", help="自动命名时的目录")
    parser.add_argument(
        "--force-general", action="store_true",
        help="即使匹配到模板也走通用要素式路径（调试用）"
    )
    args = parser.parse_args()

    # 读取 JSON
    try:
        if args.input:
            with open(args.input, "r", encoding="utf-8") as f:
                data = json.load(f)
        else:
            raw = sys.stdin.read()
            if not raw.strip():
                sys.stderr.write("[错误] 未提供 --input 且 stdin 为空。\n")
                sys.exit(1)
            data = json.loads(raw)
    except json.JSONDecodeError as e:
        sys.stderr.write(f"[错误] JSON 解析失败：{e}\n")
        sys.exit(1)
    except FileNotFoundError:
        sys.stderr.write(f"[错误] 输入文件不存在：{args.input}\n")
        sys.exit(1)

    _validate(data)

    # 输出路径
    if args.output:
        output_path = args.output
        out_dir = os.path.dirname(output_path)
        if out_dir and not os.path.isdir(out_dir):
            os.makedirs(out_dir, exist_ok=True)
    else:
        os.makedirs(args.output_dir, exist_ok=True)
        output_path = os.path.join(args.output_dir, _default_filename(data))

    case_cause = data.get("case_cause", "")
    template_path = None if args.force_general else find_template(case_cause)

    try:
        if template_path:
            sys.stderr.write(f"[信息] 命中模板：{os.path.basename(template_path)}\n")
            fill_template(template_path, data, output_path)
        else:
            in_official = case_cause in OFFICIAL_11
            if in_official:
                # 罕见：白名单案由却没找到模板 —— 报告错误
                sys.stderr.write(
                    f"[错误] 案由「{case_cause}」属于《通知》11 类白名单，"
                    f"但模板目录 {TEMPLATE_DIR} 中未找到对应文件。\n"
                )
                sys.exit(5)
            sys.stderr.write("[信息] 未匹配到模板，按通用要素式生成。\n")
            build_general_element(data, output_path)
    except SystemExit:
        raise
    except Exception as e:
        sys.stderr.write(f"[错误] 文档生成失败：{type(e).__name__}: {e}\n")
        sys.exit(4)

    sys.stdout.write(f"{output_path}\n")


if __name__ == "__main__":
    main()
