#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
民事起诉状 Word 文档生成脚本

设计目标：
    将民事起诉状的全部排版规范（字体、字号、段落、对齐、缩进、纸张）固化在脚本中，
    调用方仅需提供结构化 JSON 数据，避免大模型自行编排格式造成的格式漂移。

使用方式：
    1) 文件输入：
        python3 generate_complaint.py --input data.json --output /path/to/out.docx
    2) 标准输入：
        python3 generate_complaint.py --output /path/to/out.docx < data.json
    3) 自动命名（不指定 --output 时，按规则自动命名并输出至当前目录）：
        python3 generate_complaint.py --input data.json

JSON 输入字段说明详见 references/script-usage.md
"""

import argparse
import json
import os
import sys
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Cm, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    sys.stderr.write(
        "[错误] 缺少依赖 python-docx。\n"
        "请执行：pip install python-docx\n"
    )
    sys.exit(2)


# ============== 字体配置（带 fallback） ==============
# Word 中文字体设置需要同时设置 ASCII 字体和东亚字体
TITLE_FONT_CANDIDATES = ["微软雅黑", "Microsoft YaHei", "黑体", "SimHei", "宋体", "SimSun"]
HEADING_FONT_CANDIDATES = ["仿宋_GB2312", "仿宋", "FangSong", "FangSong_GB2312", "STFangsong", "宋体", "SimSun"]
BODY_FONT_CANDIDATES = ["仿宋_GB2312", "仿宋", "FangSong", "FangSong_GB2312", "STFangsong", "宋体", "SimSun"]

# ============== 字号配置（pt）==============
# 中文字号对照：三号=16pt（标题）；四号=14pt（段落标题、正文）
TITLE_SIZE_PT = 16    # 三号
HEADING_SIZE_PT = 14  # 四号
BODY_SIZE_PT = 14     # 四号

# 四号字（14pt）2 字符缩进 ≈ 28pt
INDENT_2_CHARS_PT = 28


def _set_run_font(run, east_asia_fonts, ascii_font="Times New Roman", size_pt=14):
    """设置 run 的字体（中文用 east_asia，西文用 ascii）。
    east_asia_fonts 是 fallback 列表，取首项作为主字体。"""
    primary = east_asia_fonts[0]
    run.font.name = ascii_font
    run.font.size = Pt(size_pt)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), ascii_font)
    rFonts.set(qn("w:hAnsi"), ascii_font)
    rFonts.set(qn("w:eastAsia"), primary)
    rFonts.set(qn("w:cs"), ascii_font)


def _set_paper_and_margins(doc):
    """设置 A4 纸张与页边距（上下 2.54cm，左右 3.17cm）。"""
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)


def _add_blank_line(doc, size_pt=BODY_SIZE_PT):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run("")
    _set_run_font(run, BODY_FONT_CANDIDATES, size_pt=size_pt)
    return p


def _add_title(doc, text="民事起诉状"):
    """标题：微软雅黑（fallback：黑体/宋体）三号 居中 加粗"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(text)
    run.bold = True
    _set_run_font(run, TITLE_FONT_CANDIDATES, size_pt=TITLE_SIZE_PT)
    return p


def _add_section_heading(doc, text):
    """段落标题：仿宋 四号 加粗 段首空两格"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Pt(INDENT_2_CHARS_PT)  # 段首空两格
    run = p.add_run(text)
    run.bold = True
    _set_run_font(run, HEADING_FONT_CANDIDATES, size_pt=HEADING_SIZE_PT)
    return p


def _add_cause_line(doc, cause_text):
    """案由行：段首空2字符，「案由：」加粗，案由内容正常字重。
    位置：当事人信息块之后、诉讼请求之前。
    格式示例：  案由：买卖合同纠纷
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.first_line_indent = Pt(INDENT_2_CHARS_PT)  # 段首空两格
    run_label = p.add_run("案由：")
    run_label.bold = True
    _set_run_font(run_label, BODY_FONT_CANDIDATES, size_pt=BODY_SIZE_PT)
    run_value = p.add_run(cause_text)
    _set_run_font(run_value, BODY_FONT_CANDIDATES, size_pt=BODY_SIZE_PT)
    return p


def _add_body_paragraph(doc, text, indent=True, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    """正文：仿宋 GB2312 四号 单倍行距 首行缩进 2 字符"""
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    if indent:
        p.paragraph_format.first_line_indent = Pt(INDENT_2_CHARS_PT)
    else:
        p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(text)
    _set_run_font(run, BODY_FONT_CANDIDATES, size_pt=BODY_SIZE_PT)
    return p


def _add_party_block(doc, role_label, party):
    """渲染当事人信息块。
    _format_party_text 返回字符串列表，每个元素对应一行段落。
    第一行：「原告/被告：」加粗，其余行（如法定代表人行）正常字重，段首空两格。
    """
    lines = _format_party_text(role_label, party)
    if isinstance(lines, str):
        lines = [lines]  # 兼容旧格式

    paragraphs = []
    for idx, line_text in enumerate(lines):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.first_line_indent = Pt(INDENT_2_CHARS_PT)  # 段首空两格
        p.paragraph_format.space_after = Pt(2)

        if idx == 0:
            # 第一行：角色标签加粗
            label_prefix = f"{role_label}："
            if line_text.startswith(label_prefix):
                run_label = p.add_run(label_prefix)
                run_label.bold = True
                _set_run_font(run_label, BODY_FONT_CANDIDATES, size_pt=BODY_SIZE_PT)
                run_rest = p.add_run(line_text[len(label_prefix):])
                _set_run_font(run_rest, BODY_FONT_CANDIDATES, size_pt=BODY_SIZE_PT)
            else:
                run = p.add_run(line_text)
                _set_run_font(run, BODY_FONT_CANDIDATES, size_pt=BODY_SIZE_PT)
        else:
            # 后续行（法定代表人/代表人/执行事务合伙人/经营者行）：正常字重
            run = p.add_run(line_text)
            _set_run_font(run, BODY_FONT_CANDIDATES, size_pt=BODY_SIZE_PT)
        paragraphs.append(p)

    return paragraphs[0] if paragraphs else None


def _format_party_text(role_label, party):
    """根据当事人类型生成规范化的描述文本列表。
    返回值为字符串列表，每个元素对应起诉状中的一行（一个段落）。
    第一行含「原告/被告：」前缀；后续行（法定代表人/代表人/经营者等）另起一行。
    """
    ptype = party.get("type", "natural_person")

    # ── 自然人 ──────────────────────────────────────────────────────────────
    if ptype == "natural_person":
        name = party.get("name", "")
        alias = party.get("alias", "")
        name_str = f"{name}（{alias}）" if alias else name

        line_parts = [f"{role_label}：{name_str}"]
        attrs = []
        if party.get("gender"):
            attrs.append(party["gender"])
        if party.get("ethnicity"):
            attrs.append(f"{party['ethnicity']}族")
        if party.get("birth_date"):
            attrs.append(f"{party['birth_date']}出生")
        if party.get("id_number"):
            attrs.append(f"公民身份号码：{party['id_number']}")
        # 住所（精确到门牌号）
        address = party.get("address", "")
        current_address = party.get("current_address", "")
        if address:
            attrs.append(f"住{address}")
        if current_address and current_address != address:
            attrs.append(f"现住{current_address}")
        if party.get("phone"):
            attrs.append(f"联系电话：{party['phone']}")
        if party.get("email"):
            attrs.append(f"电子邮箱：{party['email']}")
        # 国籍（港澳台 / 外国人）
        nationality = party.get("nationality", "")
        hkmo_tw = ("香港特别行政区居民", "澳门特别行政区居民", "台湾地区居民")
        if nationality in hkmo_tw:
            cert_no = party.get("cert_no", "")
            attrs.append(f"{nationality}，证件号码：{cert_no}" if cert_no else nationality)
        elif nationality and nationality not in ("中国", "中华人民共和国", ""):
            passport_no = party.get("passport_no", "")
            if passport_no:
                attrs.append(f"{nationality}国籍，护照号码：{passport_no}")
            else:
                attrs.append(f"{nationality}国籍")
        elif not nationality and party.get("passport_no"):
            # 无国籍人
            attrs.append(f"无国籍，护照号码：{party['passport_no']}")

        if attrs:
            line_parts.append("，" + "，".join(attrs))
        line_parts.append("。")
        return ["".join(line_parts)]

    # ── 法人 ────────────────────────────────────────────────────────────────
    elif ptype == "legal_person":
        attrs = []
        if party.get("address"):
            attrs.append(f"住所地：{party['address']}")
        if party.get("uscc"):
            attrs.append(f"统一社会信用代码：{party['uscc']}")
        if party.get("phone"):
            attrs.append(f"联系电话：{party['phone']}")
        line1 = f"{role_label}：{party.get('name', '')}"
        if attrs:
            line1 += "，" + "，".join(attrs)
        line1 += "。"

        # 法定代表人另起一行
        rep_name = party.get("legal_rep_name") or party.get("rep_name", "")
        rep_title = party.get("legal_rep_title") or "法定代表人"
        rep_position = party.get("rep_position", "")
        line2_parts = [f"{rep_title}：{rep_name}" if rep_name else f"{rep_title}："]
        if rep_position:
            line2_parts.append(f"，{rep_position}")
        line2_parts.append("。")
        return [line1, "".join(line2_parts)]

    # ── 非法人组织 ──────────────────────────────────────────────────────────
    elif ptype == "unincorporated":
        attrs = []
        if party.get("address"):
            attrs.append(f"住所地：{party['address']}")
        if party.get("uscc"):
            attrs.append(f"统一社会信用代码：{party['uscc']}")
        if party.get("phone"):
            attrs.append(f"联系电话：{party['phone']}")
        line1 = f"{role_label}：{party.get('name', '')}"
        if attrs:
            line1 += "，" + "，".join(attrs)
        line1 += "。"

        # 代表人另起一行（使用「代表人」称谓，严禁用「负责人」）
        rep_name = party.get("rep_name") or party.get("legal_rep_name", "")
        rep_title = "代表人"  # 非法人组织固定用「代表人」
        rep_position = party.get("rep_position", "")
        line2_parts = [f"{rep_title}：{rep_name}" if rep_name else f"{rep_title}："]
        if rep_position:
            line2_parts.append(f"，{rep_position}")
        line2_parts.append("。")
        return [line1, "".join(line2_parts)]

    # ── 合伙企业 ────────────────────────────────────────────────────────────
    elif ptype == "partnership_enterprise":
        attrs = []
        if party.get("address"):
            attrs.append(f"住所地：{party['address']}")
        if party.get("uscc"):
            attrs.append(f"统一社会信用代码：{party['uscc']}")
        if party.get("phone"):
            attrs.append(f"联系电话：{party['phone']}")
        line1 = f"{role_label}：{party.get('name', '')}"
        if attrs:
            line1 += "，" + "，".join(attrs)
        line1 += "。"

        # 执行事务合伙人另起一行
        gp_name = party.get("gp_name", "")
        gp_delegate = party.get("gp_delegate", "")
        if gp_name:
            if gp_delegate:
                line2 = f"执行事务合伙人：{gp_name}（委派代表：{gp_delegate}）。"
            else:
                line2 = f"执行事务合伙人：{gp_name}。"
        else:
            line2 = "执行事务合伙人：。"
        return [line1, line2]

    # ── 个人合伙 ────────────────────────────────────────────────────────────
    elif ptype == "partnership_individual":
        lines = []
        partners = party.get("partners", [])
        for i, partner in enumerate(partners):
            role = f"{role_label}{'一二三四五六七八九十'[i] if i < 10 else str(i+1)}"
            sub_lines = _format_party_text(role, {**partner, "type": "natural_person"})
            lines.extend(sub_lines)
        trade_name = party.get("trade_name", "")
        biz_addr = party.get("business_address", "")
        if trade_name:
            note = f"（经营字号：{trade_name}"
            if biz_addr:
                note += f"，经营场所：{biz_addr}"
            note += "）"
            lines.append(note)
        return lines if lines else [f"{role_label}：（待补充）。"]

    # ── 个体工商户 ──────────────────────────────────────────────────────────
    elif ptype == "individual_business":
        has_trade_name = party.get("has_trade_name", bool(party.get("operator_name")))
        if has_trade_name:
            # 有字号：以字号为当事人，另起一行写经营者信息
            attrs = []
            if party.get("address"):
                attrs.append(f"经营场所：{party['address']}")
            if party.get("phone"):
                attrs.append(f"联系电话：{party['phone']}")
            line1 = f"{role_label}：{party.get('name', '')}（个体工商户）"
            if attrs:
                line1 += "，" + "，".join(attrs)
            line1 += "。"

            # 经营者信息另起一行
            op_parts = [f"经营者：{party.get('operator_name', '')}"]
            op_attrs = []
            if party.get("operator_gender"):
                op_attrs.append(party["operator_gender"])
            if party.get("operator_ethnicity"):
                op_attrs.append(f"{party['operator_ethnicity']}族")
            if party.get("operator_birth_date"):
                op_attrs.append(f"{party['operator_birth_date']}出生")
            if party.get("id_number"):
                op_attrs.append(f"公民身份号码：{party['id_number']}")
            if party.get("operator_address"):
                op_attrs.append(f"住{party['operator_address']}")
            if op_attrs:
                op_parts.append("，" + "，".join(op_attrs))
            op_parts.append("。")
            return [line1, "".join(op_parts)]
        else:
            # 无字号：以经营者为当事人，同自然人格式，括注（个体工商户）
            name = party.get("name", "")
            alias = party.get("alias", "")
            name_str = f"{name}（{alias}）" if alias else name
            line_parts = [f"{role_label}：{name_str}（个体工商户）"]
            attrs = []
            if party.get("gender"):
                attrs.append(party["gender"])
            if party.get("ethnicity"):
                attrs.append(f"{party['ethnicity']}族")
            if party.get("birth_date"):
                attrs.append(f"{party['birth_date']}出生")
            if party.get("id_number"):
                attrs.append(f"公民身份号码：{party['id_number']}")
            if party.get("address"):
                attrs.append(f"住{party['address']}")
            if party.get("phone"):
                attrs.append(f"联系电话：{party['phone']}")
            if attrs:
                line_parts.append("，" + "，".join(attrs))
            line_parts.append("。")
            return ["".join(line_parts)]

    # ── 未成年人 ────────────────────────────────────────────────────────────
    elif ptype == "minor":
        line_parts = [f"{role_label}：{party.get('name', '')}"]
        attrs = []
        if party.get("gender"):
            attrs.append(party["gender"])
        if party.get("birth_date"):
            attrs.append(f"{party['birth_date']}出生（未成年人）")
        if party.get("id_number"):
            attrs.append(f"公民身份号码：{party['id_number']}")
        if party.get("address"):
            attrs.append(f"住{party['address']}")
        if attrs:
            line_parts.append("，" + "，".join(attrs))
        line_parts.append("。")
        line1 = "".join(line_parts)

        # 法定代理人另起一行
        guardian = party.get("guardian") or {}
        if guardian:
            g_parts = [f"法定代理人：{guardian.get('name', '')}"]
            g_attrs = []
            if guardian.get("relation"):
                g_attrs.append(f"系{role_label}{guardian['relation']}")
            if guardian.get("id_number"):
                g_attrs.append(f"公民身份号码：{guardian['id_number']}")
            if guardian.get("address"):
                g_attrs.append(f"住{guardian['address']}")
            if guardian.get("phone"):
                g_attrs.append(f"联系电话：{guardian['phone']}")
            if g_attrs:
                g_parts.append("，" + "，".join(g_attrs))
            g_parts.append("。")
            return [line1, "".join(g_parts)]
        return [line1]

    # ── 兜底 ────────────────────────────────────────────────────────────────
    else:
        return [f"{role_label}：{party.get('name', '')}。"]


def _short_name(party):
    """提取当事人简称，用于文件命名。"""
    name = party.get("name", "当事人") if party else "当事人"
    # 截断过长名称
    if len(name) > 8:
        return name[:8]
    return name


def _validate(data):
    """对必填字段进行校验，缺失即报错退出。"""
    errors = []
    required_top = ["court", "plaintiffs", "defendants", "case_cause", "claims", "facts_and_reasons", "filing_date"]
    for k in required_top:
        if k not in data or not data[k]:
            errors.append(f"缺少必填字段：{k}")

    if "plaintiffs" in data and not isinstance(data["plaintiffs"], list):
        errors.append("plaintiffs 必须为列表")
    if "defendants" in data and not isinstance(data["defendants"], list):
        errors.append("defendants 必须为列表")
    if "claims" in data and not isinstance(data["claims"], list):
        errors.append("claims 必须为列表")

    # 至少一个原告 / 一个被告
    if isinstance(data.get("plaintiffs"), list) and len(data["plaintiffs"]) == 0:
        errors.append("至少需要 1 名原告")
    if isinstance(data.get("defendants"), list) and len(data["defendants"]) == 0:
        errors.append("至少需要 1 名被告")
    if isinstance(data.get("claims"), list) and len(data["claims"]) == 0:
        errors.append("诉讼请求 claims 不得为空")

    # 当事人 name 必填
    for i, p in enumerate(data.get("plaintiffs", []) or []):
        if not isinstance(p, dict) or not p.get("name"):
            errors.append(f"原告 #{i+1} 缺少 name")
    for i, p in enumerate(data.get("defendants", []) or []):
        if not isinstance(p, dict) or not p.get("name"):
            errors.append(f"被告 #{i+1} 缺少 name")

    if errors:
        sys.stderr.write("[校验失败]\n")
        for e in errors:
            sys.stderr.write(f"  - {e}\n")
        sys.exit(3)


def _default_filename(data, mode="court"):
    p_name = _short_name((data.get("plaintiffs") or [{}])[0])
    d_name = _short_name((data.get("defendants") or [{}])[0])
    cause = data.get("case_cause", "民事纠纷")
    today = datetime.now().strftime("%Y%m%d")
    if mode == "arbitration":
        return f"仲裁申请书-{p_name}与{d_name}-{cause}-{today}.docx"
    return f"民事起诉状-{p_name}诉{d_name}-{cause}-{today}.docx"


def _split_paragraphs(text):
    """按 \n\n 分段，并清理空白。"""
    if not text:
        return []
    chunks = [s.strip() for s in str(text).replace("\r\n", "\n").split("\n\n")]
    return [c for c in chunks if c]


def build_document(data, output_path, mode="court"):
    doc = Document()
    _set_paper_and_margins(doc)

    # 模式文案（诉讼 vs 仲裁）
    is_arb = mode == "arbitration"
    doc_title = "仲裁申请书" if is_arb else "民事起诉状"
    claimant_label = "申请人" if is_arb else "原告"
    respondent_label = "被申请人" if is_arb else "被告"
    claims_heading = "仲裁请求" if is_arb else "诉讼请求"

    # 标题
    _add_title(doc, doc_title)

    # 当事人信息（顶格）
    for plaintiff in data.get("plaintiffs", []):
        _add_party_block(doc, claimant_label, plaintiff)
    for defendant in data.get("defendants", []):
        _add_party_block(doc, respondent_label, defendant)
    for tp in data.get("third_parties", []) or []:
        _add_party_block(doc, "第三人", tp)

    _add_blank_line(doc)

    # 案由行（当事人信息下方、诉讼请求上方）
    cause = data.get("case_cause", "")
    if cause:
        _add_cause_line(doc, cause)
        _add_blank_line(doc)

    # 诉讼请求 / 仲裁请求
    _add_section_heading(doc, claims_heading)
    for claim in data.get("claims", []):
        # 若 claim 已经包含 "1." 前缀就保留，否则不强行加序号（由调用方控制）
        _add_body_paragraph(doc, claim, indent=True)

    _add_blank_line(doc)

    # 事实和理由
    _add_section_heading(doc, "事实和理由")
    paragraphs = _split_paragraphs(data.get("facts_and_reasons", ""))
    for para in paragraphs:
        _add_body_paragraph(doc, para, indent=True)

    # 证据指引（可选）
    if data.get("evidence_hint"):
        _add_body_paragraph(doc, data["evidence_hint"], indent=True)

    # 法律依据 / 综上（可选）
    if data.get("legal_basis"):
        _add_body_paragraph(doc, data["legal_basis"], indent=True)

    _add_blank_line(doc)

    # 此致（段首空两格）+ 法院/仲裁机构（顶格；仲裁模式下 data["court"] 填仲裁委员会）
    p_to = doc.add_paragraph()
    p_to.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_to.paragraph_format.line_spacing = 1.5
    p_to.paragraph_format.first_line_indent = Pt(INDENT_2_CHARS_PT)  # 段首空两格
    run_to = p_to.add_run("此致")
    _set_run_font(run_to, BODY_FONT_CANDIDATES, size_pt=BODY_SIZE_PT)

    p_court = doc.add_paragraph()
    p_court.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_court.paragraph_format.line_spacing = 1.5
    p_court.paragraph_format.first_line_indent = Cm(0)  # 法院名顶格
    run_court = p_court.add_run(data.get("court", "") or "")
    _set_run_font(run_court, BODY_FONT_CANDIDATES, size_pt=BODY_SIZE_PT)

    _add_blank_line(doc)
    _add_blank_line(doc)

    # 落款区（右对齐）
    # 根据原告主体类型生成不同落款行
    for plaintiff in data.get("plaintiffs", []):
        ptype = plaintiff.get("type", "natural_person")
        name = plaintiff.get("name", "")
        if not name:
            continue

        if ptype == "natural_person":
            # 自然人：原告（签字）：姓名
            p_signer = doc.add_paragraph()
            p_signer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p_signer.paragraph_format.line_spacing = 1.5
            p_signer.paragraph_format.first_line_indent = Cm(0)
            run_signer = p_signer.add_run(f"{claimant_label}（签字）：{name}")
            _set_run_font(run_signer, BODY_FONT_CANDIDATES, size_pt=BODY_SIZE_PT)

        elif ptype == "legal_person":
            # 法人：原告（盖章）：公司名 + 法定代表人（签字）：___
            p_signer = doc.add_paragraph()
            p_signer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p_signer.paragraph_format.line_spacing = 1.5
            p_signer.paragraph_format.first_line_indent = Cm(0)
            run_signer = p_signer.add_run(f"{claimant_label}（盖章）：{name}")
            _set_run_font(run_signer, BODY_FONT_CANDIDATES, size_pt=BODY_SIZE_PT)

            rep_name = plaintiff.get("legal_rep_name", "")
            rep_title = plaintiff.get("legal_rep_title", "法定代表人")
            p_rep = doc.add_paragraph()
            p_rep.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p_rep.paragraph_format.line_spacing = 1.5
            p_rep.paragraph_format.first_line_indent = Cm(0)
            run_rep = p_rep.add_run(f"{rep_title}（签字）：          ")  # 留白10字符供手写签名
            _set_run_font(run_rep, BODY_FONT_CANDIDATES, size_pt=BODY_SIZE_PT)

        elif ptype == "unincorporated":
            # 非法人组织（合伙企业等）：原告（盖章）：名称 + 执行事务合伙人/代表人（签字）：留白
            p_signer = doc.add_paragraph()
            p_signer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p_signer.paragraph_format.line_spacing = 1.5
            p_signer.paragraph_format.first_line_indent = Cm(0)
            run_signer = p_signer.add_run(f"{claimant_label}（盖章）：{name}")
            _set_run_font(run_signer, BODY_FONT_CANDIDATES, size_pt=BODY_SIZE_PT)

            rep_title = plaintiff.get("legal_rep_title", "执行事务合伙人")
            p_rep = doc.add_paragraph()
            p_rep.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p_rep.paragraph_format.line_spacing = 1.5
            p_rep.paragraph_format.first_line_indent = Cm(0)
            run_rep = p_rep.add_run(f"{rep_title}（签字）：          ")  # 留白10字符供手写签名
            _set_run_font(run_rep, BODY_FONT_CANDIDATES, size_pt=BODY_SIZE_PT)

        elif ptype == "individual_business":
            # 个体工商户：原告（签字）：经营者姓名
            operator = plaintiff.get("operator_name", "") or name
            p_signer = doc.add_paragraph()
            p_signer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p_signer.paragraph_format.line_spacing = 1.5
            p_signer.paragraph_format.first_line_indent = Cm(0)
            run_signer = p_signer.add_run(f"{claimant_label}（签字）：{operator}")
            _set_run_font(run_signer, BODY_FONT_CANDIDATES, size_pt=BODY_SIZE_PT)

        else:
            # 其他类型（minor/foreign等）：原告（签字）：姓名
            p_signer = doc.add_paragraph()
            p_signer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p_signer.paragraph_format.line_spacing = 1.5
            p_signer.paragraph_format.first_line_indent = Cm(0)
            run_signer = p_signer.add_run(f"{claimant_label}（签字）：{name}")
            _set_run_font(run_signer, BODY_FONT_CANDIDATES, size_pt=BODY_SIZE_PT)

    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_date.paragraph_format.line_spacing = 1.5
    p_date.paragraph_format.first_line_indent = Cm(0)
    run_date = p_date.add_run(data.get("filing_date", "") or "")
    _set_run_font(run_date, BODY_FONT_CANDIDATES, size_pt=BODY_SIZE_PT)

    _add_blank_line(doc)

    # 附件清单
    attachments = data.get("attachments") or []
    if attachments:
        _add_section_heading(doc, "附件")
        for idx, att in enumerate(attachments, start=1):
            _add_body_paragraph(doc, f"{idx}. {att}", indent=True)

    doc.save(output_path)
    return output_path


def main():
    parser = argparse.ArgumentParser(description="民事起诉状 / 仲裁申请书 Word 生成器")
    parser.add_argument("--input", "-i", help="JSON 输入文件路径；不指定则从 stdin 读取")
    parser.add_argument("--output", "-o", help="输出 docx 文件路径；不指定则按规则自动命名")
    parser.add_argument("--output-dir", "-d", default=".", help="自动命名时使用的目录（默认当前目录）")
    parser.add_argument("--mode", "-m", choices=["court", "arbitration"], default="court",
                        help="文书模式：court=民事起诉状（默认）；arbitration=仲裁申请书（标题/当事人称谓/请求标题/落款按仲裁文书渲染；JSON 键仍用 plaintiffs/defendants，court 字段填仲裁委员会）")
    args = parser.parse_args()

    # 读取 JSON
    try:
        if args.input:
            with open(args.input, "r", encoding="utf-8") as f:
                data = json.load(f)
        else:
            raw = sys.stdin.read()
            if not raw.strip():
                sys.stderr.write("[错误] 未提供 --input 且 stdin 为空。\n")
                sys.exit(1)
            data = json.loads(raw)
    except json.JSONDecodeError as e:
        sys.stderr.write(f"[错误] JSON 解析失败：{e}\n")
        sys.exit(1)
    except FileNotFoundError:
        sys.stderr.write(f"[错误] 输入文件不存在：{args.input}\n")
        sys.exit(1)

    _validate(data)

    # 确定输出路径
    if args.output:
        output_path = args.output
        out_dir = os.path.dirname(output_path)
        if out_dir and not os.path.isdir(out_dir):
            os.makedirs(out_dir, exist_ok=True)
    else:
        os.makedirs(args.output_dir, exist_ok=True)
        output_path = os.path.join(args.output_dir, _default_filename(data, args.mode))

    try:
        build_document(data, output_path, mode=args.mode)
    except Exception as e:
        sys.stderr.write(f"[错误] 文档生成失败：{type(e).__name__}: {e}\n")
        sys.exit(4)

    sys.stdout.write(f"{output_path}\n")


if __name__ == "__main__":
    main()
