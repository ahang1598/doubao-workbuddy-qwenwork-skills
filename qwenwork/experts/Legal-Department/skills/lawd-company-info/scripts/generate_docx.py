#!/usr/bin/env python3
"""
企业尽调报告 Markdown → Word 转换脚本（跨平台）

将 Markdown 格式的企业尽调报告转换为排版规范的 .docx 文件，
自动将 Markdown 语法（**加粗**、*斜体*、# 标题、> 引用、表格、列表等）
转换为 Word 原生格式，确保最终文档中不残留任何 Markdown 符号。

排版规范：
  - 纸张：A4，上下边距 2.54cm，左右边距 3.17cm
  - 正文字体：宋体 12pt（中文）+ Times New Roman 12pt（英文/数字）
  - 一级标题：黑体 22pt，居中加粗
  - 二级标题：黑体 16pt，加粗
  - 三级标题：黑体 14pt，加粗
  - 四级标题：黑体 12pt，加粗
  - 引用块：宋体 11pt，左缩进 1cm，灰色左边框效果
  - 行间距：1.5 倍
  - 表格：统一边框样式，表头加粗

用法:
  python3 scripts/generate_docx.py --workspace output --file /tmp/report_draft.md
  python3 scripts/generate_docx.py --workspace output --file draft.md --filename "小米科技_企业尽调报告.docx"
  python3 scripts/generate_docx.py --workspace output --content "Markdown 内容..."

依赖:
  pip install python-docx
"""

import argparse
import os
import re
import sys
from datetime import datetime
from pathlib import Path
from typing import Optional


# ======================================================================
# Inline 格式解析 — 将 **bold** 和 *italic* 转换为 Word runs
# ======================================================================

def parse_inline(text: str) -> list:
    """
    将单行文本中的 Markdown inline 格式解析为 Word runs。
    支持: **加粗** / __加粗__、*斜体* / _斜体_
    返回: [{'text': str, 'bold': bool, 'italic': bool}, ...]
    """
    segments = []
    position = 0

    while position < len(text):
        match = re.search(r'(\*\*|__|\*(?!\*))', text[position:])
        if not match:
            segments.append(("plain", text[position:]))
            break

        if match.start() > 0:
            segments.append(("plain", text[position:position + match.start()]))

        marker = match.group(1)
        search_start = position + match.end()

        if marker in ("**", "__"):
            close_pattern = r'\*\*' if marker == "**" else r'__'
            close_match = re.search(close_pattern, text[search_start:])
            if close_match:
                inner_text = text[search_start:search_start + close_match.start()]
                segments.append(("bold", inner_text))
                position = search_start + close_match.end()
            else:
                segments.append(("plain", text[position:]))
                break
        elif marker == "*":
            close_match = re.search(r'\*', text[search_start:])
            if close_match:
                inner_text = text[search_start:search_start + close_match.start()]
                segments.append(("italic", inner_text))
                position = search_start + close_match.end()
            else:
                segments.append(("plain", text[position:]))
                break
        else:
            segments.append(("plain", text[position:]))
            break

    merged = []
    for kind, content in segments:
        if content:
            if merged and merged[-1][0] == kind:
                merged[-1] = (kind, merged[-1][1] + content)
            else:
                merged.append((kind, content))

    return [
        {"text": content, "bold": kind == "bold", "italic": kind == "italic"}
        for kind, content in merged
    ]


# ======================================================================
# Markdown 结构解析
# ======================================================================

def detect_heading_level(line: str) -> tuple:
    match = re.match(r'^(#{1,4})\s+(.+)$', line.strip())
    if match:
        return len(match.group(1)), match.group(2).strip()
    return 0, None


def detect_list_item(line: str) -> tuple:
    stripped = line.strip()
    unordered_match = re.match(r'^[-*]\s+(.*)$', stripped)
    if unordered_match:
        indent = len(line) - len(line.lstrip())
        return True, "bullet", indent // 2, unordered_match.group(1)
    ordered_match = re.match(r'^(\d+)[.、]\s+(.*)$', stripped)
    if ordered_match:
        indent = len(line) - len(line.lstrip())
        return True, "number", indent // 2, ordered_match.group(2)
    return False, "", 0, ""


def detect_blockquote(line: str) -> tuple:
    match = re.match(r'^>\s*(.*)$', line.strip())
    if match:
        return True, match.group(1)
    return False, ""


def is_table_row(line: str) -> bool:
    stripped = line.strip()
    return (
        stripped.startswith("|")
        and stripped.endswith("|")
        and stripped.count("|") >= 3
    )


def is_table_separator(line: str) -> bool:
    stripped = line.strip()
    if not stripped.startswith("|") or not stripped.endswith("|"):
        return False
    content = stripped[1:-1]
    cells = content.split("|")
    for cell in cells:
        cell_stripped = cell.strip()
        if cell_stripped and not re.match(r'^:?-+:?$', cell_stripped):
            return False
    return True


def parse_table_cells(line: str) -> list:
    stripped = line.strip()
    if stripped.startswith("|"):
        stripped = stripped[1:]
    if stripped.endswith("|"):
        stripped = stripped[:-1]
    return [cell.strip() for cell in stripped.split("|")]


def sanitize_html_tags(content: str) -> str:
    """清洗混入的 HTML 标签，转为 Markdown 等价写法或删除。"""
    content = re.sub(r'<h1[^>]*>(.*?)</h1>', r'# \1', content, flags=re.IGNORECASE | re.DOTALL)
    content = re.sub(r'<h2[^>]*>(.*?)</h2>', r'## \1', content, flags=re.IGNORECASE | re.DOTALL)
    content = re.sub(r'<h3[^>]*>(.*?)</h3>', r'### \1', content, flags=re.IGNORECASE | re.DOTALL)
    content = re.sub(r'<h4[^>]*>(.*?)</h4>', r'#### \1', content, flags=re.IGNORECASE | re.DOTALL)
    content = re.sub(r'<strong[^>]*>(.*?)</strong>', r'**\1**', content, flags=re.IGNORECASE | re.DOTALL)
    content = re.sub(r'<b[^>]*>(.*?)</b>', r'**\1**', content, flags=re.IGNORECASE | re.DOTALL)
    content = re.sub(r'<em[^>]*>(.*?)</em>', r'*\1*', content, flags=re.IGNORECASE | re.DOTALL)
    content = re.sub(r'<i[^>]*>(.*?)</i>', r'*\1*', content, flags=re.IGNORECASE | re.DOTALL)
    content = re.sub(r'<br\s*/?\s*>', '\n', content, flags=re.IGNORECASE)
    content = re.sub(r'<p[^>]*>(.*?)</p>', r'\n\1\n', content, flags=re.IGNORECASE | re.DOTALL)
    content = re.sub(r'</?center[^>]*>', '', content, flags=re.IGNORECASE)
    content = re.sub(r'</?div[^>]*>', '', content, flags=re.IGNORECASE)
    content = re.sub(r'</?span[^>]*>', '', content, flags=re.IGNORECASE)
    content = re.sub(r'<hr\s*/?\s*>', '', content, flags=re.IGNORECASE)
    content = re.sub(r'<[^>]+>', '', content)
    content = re.sub(r'\n{4,}', '\n\n\n', content)
    return content


def parse_markdown_to_blocks(content: str) -> list:
    """将 Markdown 内容解析为结构化块列表。"""
    content = sanitize_html_tags(content)
    lines = content.strip().split("\n")
    blocks = []
    index = 0

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        if not stripped:
            blocks.append({"type": "blank"})
            index += 1
            continue

        heading_level, heading_text = detect_heading_level(stripped)
        if heading_level > 0:
            blocks.append({
                "type": "heading",
                "level": heading_level,
                "text": heading_text,
                "runs": parse_inline(heading_text),
            })
            index += 1
            continue

        if is_table_row(stripped):
            table_rows = []
            while index < len(lines):
                current_line = lines[index].strip()
                if is_table_separator(current_line):
                    index += 1
                    continue
                if is_table_row(current_line):
                    cells = parse_table_cells(current_line)
                    is_header = (
                        len(table_rows) == 0
                        and index + 1 < len(lines)
                        and is_table_separator(lines[index + 1].strip())
                    )
                    table_rows.append({"cells": cells, "is_header": is_header})
                    index += 1
                else:
                    break
            blocks.append({"type": "table", "rows": table_rows})
            continue

        is_quote, quote_text = detect_blockquote(stripped)
        if is_quote:
            quote_lines = [quote_text]
            index += 1
            while index < len(lines):
                next_is_quote, next_text = detect_blockquote(lines[index].strip())
                if next_is_quote:
                    quote_lines.append(next_text)
                    index += 1
                else:
                    break
            full_quote = " ".join(line for line in quote_lines if line)
            blocks.append({
                "type": "blockquote",
                "text": full_quote,
                "runs": parse_inline(full_quote),
            })
            continue

        is_list, list_type, indent_level, list_text = detect_list_item(line)
        if is_list:
            blocks.append({
                "type": "list_item",
                "list_type": list_type,
                "indent": indent_level,
                "text": list_text,
                "runs": parse_inline(list_text),
            })
            index += 1
            continue

        blocks.append({
            "type": "paragraph",
            "text": stripped,
            "runs": parse_inline(stripped),
        })
        index += 1

    return blocks


# ======================================================================
# Word 文档生成
# ======================================================================

def add_runs_to_paragraph(paragraph, runs, font_name="宋体", font_size_pt=12, base_bold=False):
    """将解析后的 runs 添加到段落中，自动设置字体。"""
    from docx.shared import Pt
    from docx.oxml.ns import qn

    for run_data in runs:
        run = paragraph.add_run(run_data["text"])
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        run.font.size = Pt(font_size_pt)
        run.font.bold = base_bold or run_data.get("bold", False)
        run.font.italic = run_data.get("italic", False)


def set_paragraph_spacing(paragraph, before_pt=0, after_pt=6, line_spacing=1.5):
    """设置段落间距。"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    paragraph_properties = paragraph._element.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(int(before_pt * 20)))
    spacing.set(qn("w:after"), str(int(after_pt * 20)))
    spacing.set(qn("w:line"), str(int(line_spacing * 240)))
    spacing.set(qn("w:lineRule"), "auto")
    paragraph_properties.append(spacing)


def set_paragraph_indent(paragraph, left_pt=0):
    """设置段落左缩进。"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    paragraph_properties = paragraph._element.get_or_add_pPr()
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), str(int(left_pt * 20)))
    paragraph_properties.append(indent)


def generate_docx(workspace: str, content: str, filename: Optional[str] = None) -> Path:
    """将 Markdown 内容生成为排版规范的 .docx 企业尽调报告。"""
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        print("错误：请先安装 python-docx。运行: pip install python-docx", file=sys.stderr)
        sys.exit(1)

    workspace_path = Path(workspace)
    if not workspace_path.is_dir():
        print(f"错误：工作空间目录不存在: {workspace}", file=sys.stderr)
        sys.exit(1)

    if not content or not content.strip():
        print("错误：文档内容不能为空", file=sys.stderr)
        sys.exit(1)

    if not filename:
        filename = f"企业尽调报告_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    if not filename.lower().endswith(".docx"):
        filename = f"{filename}.docx"

    output_path = workspace_path / filename
    if output_path.exists():
        base = output_path.stem
        suffix = output_path.suffix
        counter = 1
        while output_path.exists():
            output_path = workspace_path / f"{base}_{counter}{suffix}"
            counter += 1

    document = Document()

    # 页面设置：A4，标准边距
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    # 默认字体
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal_style.font.size = Pt(12)
    normal_style.font.color.rgb = RGBColor(0, 0, 0)

    heading_config = {
        1: {"font": "黑体", "size": 22, "center": True},
        2: {"font": "黑体", "size": 16, "center": False},
        3: {"font": "黑体", "size": 14, "center": False},
        4: {"font": "黑体", "size": 12, "center": False},
    }

    blocks = parse_markdown_to_blocks(content)

    for block in blocks:
        block_type = block.get("type")

        if block_type == "blank":
            blank_paragraph = document.add_paragraph()
            set_paragraph_spacing(blank_paragraph, before_pt=0, after_pt=0)
            continue

        if block_type == "heading":
            level = block.get("level", 1)
            config = heading_config.get(level, heading_config[4])
            heading_paragraph = document.add_paragraph()
            if config["center"]:
                heading_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_runs_to_paragraph(
                heading_paragraph,
                block["runs"],
                font_name=config["font"],
                font_size_pt=config["size"],
                base_bold=True,
            )
            set_paragraph_spacing(heading_paragraph, before_pt=12, after_pt=6)
            continue

        if block_type == "blockquote":
            quote_paragraph = document.add_paragraph()
            add_runs_to_paragraph(
                quote_paragraph,
                block["runs"],
                font_name="宋体",
                font_size_pt=11,
            )
            set_paragraph_spacing(quote_paragraph, before_pt=3, after_pt=3)
            set_paragraph_indent(quote_paragraph, left_pt=28)

            paragraph_properties = quote_paragraph._element.get_or_add_pPr()
            borders = OxmlElement("w:pBdr")
            left_border = OxmlElement("w:left")
            left_border.set(qn("w:val"), "single")
            left_border.set(qn("w:sz"), "12")
            left_border.set(qn("w:space"), "4")
            left_border.set(qn("w:color"), "999999")
            borders.append(left_border)
            paragraph_properties.append(borders)
            continue

        if block_type == "table":
            table_rows = block.get("rows", [])
            if not table_rows:
                continue

            num_cols = max(len(row["cells"]) for row in table_rows)
            num_rows = len(table_rows)

            table = document.add_table(rows=num_rows, cols=num_cols)
            table.style = "Table Grid"

            for row_index, row_data in enumerate(table_rows):
                row = table.rows[row_index]
                cells = row_data.get("cells", [])
                is_header = row_data.get("is_header", False)

                for col_index, cell_text in enumerate(cells):
                    if col_index >= num_cols:
                        break
                    cell = row.cells[col_index]
                    cell.text = ""
                    cell_paragraph = cell.paragraphs[0]

                    cell_runs = parse_inline(cell_text)
                    for run_data in cell_runs:
                        cell_run = cell_paragraph.add_run(run_data["text"])
                        cell_run.font.name = "Times New Roman"
                        cell_run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                        cell_run.font.size = Pt(10)
                        cell_run.font.bold = is_header or run_data.get("bold", False)
                        cell_run.font.italic = run_data.get("italic", False)

            after_table = document.add_paragraph()
            set_paragraph_spacing(after_table, before_pt=3, after_pt=3)
            continue

        if block_type == "list_item":
            list_type = block.get("list_type", "bullet")
            indent_level = block.get("indent", 0)
            list_paragraph = document.add_paragraph()

            prefix = "• " if list_type == "bullet" else ""
            if prefix:
                prefix_run = list_paragraph.add_run(prefix)
                prefix_run.font.name = "Times New Roman"
                prefix_run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                prefix_run.font.size = Pt(12)

            add_runs_to_paragraph(list_paragraph, block["runs"])
            set_paragraph_spacing(list_paragraph, before_pt=0, after_pt=3)
            left_indent = 21 + indent_level * 21
            set_paragraph_indent(list_paragraph, left_pt=left_indent)
            continue

        if block_type == "paragraph":
            normal_paragraph = document.add_paragraph()
            add_runs_to_paragraph(normal_paragraph, block["runs"])
            set_paragraph_spacing(normal_paragraph, before_pt=0, after_pt=6)
            continue

    document.save(str(output_path))
    return output_path


# ======================================================================
# 命令行入口
# ======================================================================

def main():
    parser = argparse.ArgumentParser(
        description="企业尽调报告 Markdown → Word 转换脚本",
    )
    parser.add_argument(
        "--workspace",
        required=True,
        help="工作空间目录（Word 文件保存位置）",
    )
    content_group = parser.add_mutually_exclusive_group(required=True)
    content_group.add_argument(
        "--content",
        help="Markdown 内容字符串",
    )
    content_group.add_argument(
        "--file",
        help="Markdown 文件路径（从文件读取内容）",
    )
    parser.add_argument(
        "--filename",
        default=None,
        help="输出文件名（默认：企业尽调报告_时间戳.docx）",
    )
    args = parser.parse_args()

    if args.file:
        file_path = Path(args.file)
        if not file_path.exists():
            print(f"错误：文件不存在: {args.file}", file=sys.stderr)
            sys.exit(1)
        markdown_content = file_path.read_text(encoding="utf-8")
    else:
        markdown_content = args.content

    workspace_path = Path(args.workspace)
    workspace_path.mkdir(parents=True, exist_ok=True)

    output_path = generate_docx(
        workspace=args.workspace,
        content=markdown_content,
        filename=args.filename,
    )
    print(f"FILE_PATH:{output_path}")


if __name__ == "__main__":
    main()
