#!/usr/bin/env python3
"""lawd-regulation-retrieval 法条引用核验门禁（防幻觉 / 防缺漏）。

用法
----
    python3 verify_laws.py --json <本次检索落盘的归一化 JSON> --doc <交付物 md/txt>
    python3 verify_laws.py --json result.json --doc report.md --require-citations

门禁逻辑
--------
1. 防幻觉（严格字段比对）：从交付物中提取形如《法律名》第X条的被引法条，逐条要求在
   落盘 JSON 的**专用字段**中同时命中「法规名字段」与「条号字段」：
   - 法规名只取 lawName / law_name / 法规名称 / 法律名称 等专用字段；
   - 条号只取 lawOrder / law_order / articleNumber / articleNo / 条号 / 条款编号 等专用字段。
   **绝不扫描条文正文（lawSourceContent 等），绝不扫描 JSON 原文。**
   任何一条无法按上述判据溯源即拦截。
2. 契约校验：若 JSON 中识别不出任何一条含法规名字段的法条记录，直接报错阻断
   （提示“检索结果未按归一化契约落盘”），**不做任何降级/宽松匹配**。
3. 防缺漏：交付物开头必须存在「检索说明」块，且 query、keywords、size（条数）、
   执行时间四个要素**均须有非空实值**（只写字样、值留空同样拦截）。

比对口径
--------
- 引用格式统一为 `《法律名》第X条`，兼容：
  `《民法典》第五百七十七条`、`《中华人民共和国民法典》（以下简称《民法典》）第五百七十七条`、
  `《民法典》第五百七十七条第一款（第X项）`、`《民法典》第五百七十七条、第五百七十八条`（并列列举）、
  `《刑法》第一百三十三条之一`。
- 不带书名号法律名的条号（如“合同第三条”“公司章程第五条”）不视为法条引用，不参与溯源。
- 法规名比对宽容但不模糊：允许有无“中华人民共和国”前缀（《民法典》↔《中华人民共和国民法典》），
  不允许部分/模糊匹配到不同法律（《民法典》≠《民法典合同编通则司法解释》）。
- 条号支持中文数字与阿拉伯数字互转（第五百七十七条 ↔ 第577条）；
  “第X条第X款第X项”中的款、项不参与溯源比对，只要条号对上即可。

依赖说明
--------
- 交付物**以 .md / .txt 为主要支持格式**（平台口径：先出 Markdown 过门禁、再转 Word）；
- .docx 为可选支持，需要 python-docx（pip3 install python-docx）。

退出码：0=通过（可能含警告）；1=拦截；2=输入错误。
"""

from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path
from typing import Any, Dict, List, Optional, Set, Tuple

# ---------------------------------------------------------------------------
# 字段契约（以 references/response-schema.md 为准）
# ---------------------------------------------------------------------------
# 法规名称专用字段：归一化契约为 lawDomain.lawName；兼容常见等价写法。
# 注意：**不含** name/title 等通用键，避免把 JSON 中无关对象误认成法条记录。
NAME_KEYS = {
    "lawname", "law_name",
    "法规名称", "法律名称",
}
# 条号专用字段：归一化契约为 lawDomain.lawOrder；兼容常见等价写法。
# 注意：**不含** order/article/序号 等通用键，也**不含** lawTitle/lawSourceContent
#（标题与正文可能夹带交叉引用条号，不得作为溯源依据）。
ORDER_KEYS = {
    "laworder", "law_order", "articlenumber", "article_number",
    "articleno", "article_no",
    "条号", "条款编号", "条款号",
}

# 书名号内属于“非法律法规文书”的名称后缀：这类条号不作为需溯源的法条引用。
NON_STATUTE_SUFFIXES = (
    "合同", "协议", "章程", "承诺书", "确认书", "备忘录", "纪要",
    "说明书", "投标书", "标书", "报告书", "招股书",
)

CN_DIGITS = {
    "零": 0, "〇": 0, "○": 0, "OO": 0,
    "一": 1, "壹": 1, "二": 2, "贰": 2, "两": 2, "三": 3, "叁": 3,
    "四": 4, "肆": 4, "五": 5, "伍": 5, "六": 6, "陆": 6,
    "七": 7, "柒": 7, "八": 8, "捌": 8, "九": 9, "玖": 9,
}
CN_UNITS = {"十": 10, "拾": 10, "百": 100, "佰": 100, "千": 1000, "仟": 1000}
NUM_CHARS = "零〇○一壹二贰两三叁四肆五伍六陆七柒八捌九玖十拾百佰千仟0123456789０-９"

_NUM = "[" + NUM_CHARS + "]+"
# 款/项后缀（不参与溯源比对，仅用于把引用整体吃掉，避免打断并列列举）
_SUB = r"(?:\s*第" + _NUM + r"[款项目]\s*)*"
# 条号本体：第X条（可带“之一/之二”）
_ART = r"第" + _NUM + r"条(?:之" + _NUM + r")?"
# 并列列举分隔符
_SEP = r"(?:\s*[、，,;；]\s*(?:和|与|及|以及)?\s*|\s*(?:和|与|及|以及)\s*)"

ARTICLE_FULL_RE = re.compile(r"第(" + _NUM + r")条(?:之(" + _NUM + r"))?")

CITATION_RE = re.compile(
    r"《(?P<law>[^《》\n]{2,80}?)》"
    # 可选括注，如“（以下简称《民法典》）”“(2020年修正)”
    r"(?:\s*(?:（[^（）\n]{0,60}）|\([^()\n]{0,60}\)))?"
    r"\s*(?P<first>" + _ART + r")" + _SUB +
    r"(?P<more>(?:" + _SEP + _ART + _SUB + r")*)"
)

VALUE_PLACEHOLDERS = {
    "", "-", "--", "—", "–", "/", "n/a", "na", "none", "null", "nil",
    "未提供", "未填写", "未知", "无", "空", "略", "待补充", "待填写", "tbd",
    "...", "…", "xxx", "xx", "待定", "省略",
}

SEARCH_NOTE_FIELDS: Tuple[Tuple[str, Tuple[str, ...]], ...] = (
    ("query", (r"query", r"检\s*索\s*词", r"查\s*询\s*词", r"查\s*询")),
    ("keywords", (r"keywords", r"关\s*键\s*词", r"关\s*键\s*字")),
    ("size", (r"--\s*size", r"\bsize\b", r"返\s*回\s*条\s*数", r"条\s*数", r"每\s*页")),
    ("执行时间", (r"执\s*行\s*时\s*间", r"检\s*索\s*时\s*间", r"时\s*间")),
)


class InputError(Exception):
    """输入层错误（文件不存在、无法解析等），退出码 2。"""


# ---------------------------------------------------------------------------
# 数字与条号归一化
# ---------------------------------------------------------------------------
def cn_num_to_int(text: str) -> Optional[int]:
    """中文数字/阿拉伯数字 → int，无法解析返回 None。"""
    s = re.sub(r"\s+", "", str(text)).strip()
    if not s:
        return None
    # 全角数字归一
    s = s.translate({ord(c): ord("0") + i for i, c in enumerate("０１２３４５６７８９")})
    if s.isdigit():
        return int(s)
    total, current, valid = 0, 0, False
    for ch in s:
        if ch in CN_DIGITS:
            current = CN_DIGITS[ch]
            valid = True
        elif ch in CN_UNITS:
            unit = CN_UNITS[ch]
            if current == 0:
                current = 1  # “十”→10，“百”→100
            total += current * unit
            current = 0
            valid = True
        else:
            return None
    return total + current if valid else None


ArticleKey = Tuple[str, Any, int]


def article_key(raw: str) -> ArticleKey:
    """条号归一化键。

    ("num", 577, 0)  ← 第五百七十七条 / 第577条 / 577
    ("num", 133, 1)  ← 第一百三十三条之一（“之一”单独成条，不与第133条混同）
    ("str", 原文, 0) ← 无法解析为数字时退回字符串精确比较
    """
    s = re.sub(r"\s+", "", str(raw))
    m = ARTICLE_FULL_RE.search(s)
    if m:
        num = cn_num_to_int(m.group(1))
        sub = cn_num_to_int(m.group(2)) if m.group(2) else 0
        if num is not None and sub is not None:
            return ("num", num, sub)
    num = cn_num_to_int(s)
    if num is not None:
        return ("num", num, 0)
    return ("str", s, 0)


def article_keys_of_field(value: str) -> Set[ArticleKey]:
    """从**条号字段值**中解析出全部条号键（字段可能写作“第X条”或“577”或少数并列写法）。"""
    keys: Set[ArticleKey] = set()
    for m in ARTICLE_FULL_RE.finditer(str(value)):
        keys.add(article_key(m.group(0)))
    if not keys:
        keys.add(article_key(value))
    return keys


# ---------------------------------------------------------------------------
# 法规名归一化与比对
# ---------------------------------------------------------------------------
PRC_PREFIX = "中华人民共和国"


def normalize_law_name(name: str) -> str:
    s = re.sub(r"\s+", "", str(name))
    s = s.strip("《》 \u3000")
    return s


def names_match(cited: str, record_name: str) -> bool:
    """法规名比对：允许有无“中华人民共和国”前缀（覆盖《民法典》↔《中华人民共和国民法典》），
    其余一律要求完全一致——不做模糊/部分匹配，避免匹配到不同的法律。"""
    a, b = normalize_law_name(cited), normalize_law_name(record_name)
    if not a or not b:
        return False
    if a == b:
        return True
    return a.removeprefix(PRC_PREFIX) == b.removeprefix(PRC_PREFIX) and bool(
        a.removeprefix(PRC_PREFIX)
    )


# ---------------------------------------------------------------------------
# JSON 记录抽取（严格字段，不扫正文、不扫原文）
# ---------------------------------------------------------------------------
class LawRecord:
    __slots__ = ("path", "names", "orders")

    def __init__(self, path: str, names: List[Tuple[str, str]],
                 orders: List[Tuple[str, str, Set[ArticleKey]]]) -> None:
        self.path = path
        self.names = names          # [(字段名, 法规名值)]
        self.orders = orders        # [(字段名, 条号值, 条号键集合)]


def collect_records(node: Any) -> List[LawRecord]:
    """递归收集法条记录：**含法规名专用字段**的 dict 才算一条记录。"""
    records: List[LawRecord] = []

    def walk(n: Any, path: str) -> None:
        if isinstance(n, dict):
            names: List[Tuple[str, str]] = []
            orders: List[Tuple[str, str, Set[ArticleKey]]] = []
            for key, value in n.items():
                k = str(key).strip().lower()
                text = value.strip() if isinstance(value, str) else (
                    str(value) if isinstance(value, (int, float)) else ""
                )
                if not text:
                    continue
                if k in NAME_KEYS:
                    names.append((str(key), text))
                elif k in ORDER_KEYS:
                    orders.append((str(key), text, article_keys_of_field(text)))
            if names:
                records.append(LawRecord(path or "$", names, orders))
            for key, value in n.items():
                walk(value, f"{path}.{key}" if path else str(key))
        elif isinstance(n, list):
            for i, item in enumerate(n):
                walk(item, f"{path}[{i}]")

    walk(node, "")
    return records


def match_citation(law: str, article: str,
                   records: List[LawRecord]) -> Tuple[Optional[str], Optional[str]]:
    """严格字段比对：返回 (命中说明, 未命中原因)。"""
    target = article_key(article)
    name_hits: List[Tuple[LawRecord, str, str]] = []
    for rec in records:
        for name_key, name_value in rec.names:
            if names_match(law, name_value):
                name_hits.append((rec, name_key, name_value))
                break
    if not name_hits:
        return None, "法规名未出现在检索 JSON 的法规名字段（lawName 等）中"
    for rec, name_key, name_value in name_hits:
        for field_name, field_value, keys in rec.orders:
            if target in keys:
                return (f"{rec.path}（{name_key}=\"{name_value}\"，"
                        f"{field_name}=\"{field_value}\"）"), None
    known: List[str] = []
    for rec, _k, _v in name_hits:
        for _f, field_value, _keys in rec.orders:
            if field_value not in known:
                known.append(field_value)
    if known:
        detail = "该法在 JSON 中仅有条号：" + "、".join(known)
    else:
        detail = "该法在 JSON 中无条号字段（lawOrder 等），仅命中整部法律，未取到具体条文"
    return None, f"法规名命中 {len(name_hits)} 条记录，但条号字段不匹配（{detail}）"


def validate_contract(data: Any, records: List[LawRecord]) -> Tuple[List[str], List[str], List[str]]:
    """校验落盘 JSON 是否符合归一化契约。返回 (通过项, 警告, 拦截项)。"""
    checks: List[str] = []
    warns: List[str] = []
    errors: List[str] = []

    law_result: Any = None
    if isinstance(data, dict) and isinstance(data.get("data"), dict):
        law_result = data["data"].get("lawResult")
    explicit_empty = isinstance(law_result, list) and not law_result

    if records:
        with_order = sum(1 for r in records if r.orders)
        checks.append(
            f"JSON 识别出 {len(records)} 条法条记录（含法规名专用字段），"
            f"其中 {with_order} 条带条号字段"
        )
        if isinstance(law_result, list):
            checks.append(f"JSON 符合归一化契约：data.lawResult 为数组（{len(law_result)} 项）")
        else:
            warns.append(
                "JSON 未见 data.lawResult 数组（结构偏离 references/response-schema.md），"
                f"已按专用字段名识别出 {len(records)} 条记录；建议按契约落盘"
            )
        if with_order == 0:
            warns.append("JSON 中所有记录均无条号字段（lawOrder 等），任何“第X条”引用都无法溯源")
    elif explicit_empty:
        checks.append("JSON 符合归一化契约，但 data.lawResult 为 0 条（本次检索无结果）")
        warns.append("本次检索无任何法条结果，交付物不得引用具体法条")
    else:
        errors.append(
            "检索结果未按归一化契约落盘：JSON 中识别不出任何含法规名专用字段"
            "（lawName / law_name / 法规名称 / 法律名称）的法条记录，无法进行引用溯源。"
            "请按 references/response-schema.md 归一化后重新落盘再跑门禁"
            "（本脚本不会降级为宽松匹配放行）"
        )

    if isinstance(data, dict) and data.get("success") is False:
        warns.append("JSON 中 success=false（本次检索标记为失败），请确认交付物是否应引用其结果")
    return checks, warns, errors


# ---------------------------------------------------------------------------
# 交付物引用提取
# ---------------------------------------------------------------------------
def extract_citations(text: str) -> Tuple[List[Tuple[str, str]], List[Tuple[str, str]]]:
    """提取被引法条，返回 (需溯源的法条引用, 已忽略的非法条条号)。

    - 只认《法律名》第X条形态；“合同第三条”等不带书名号法律名的条号不予提取；
    - 书名号内为合同/协议/章程等非法律法规文书时，条号不作为法条引用。
    """
    cited: List[Tuple[str, str]] = []
    skipped: List[Tuple[str, str]] = []
    for m in CITATION_RE.finditer(text):
        law = normalize_law_name(m.group("law"))
        if not law:
            continue
        articles = [m.group("first")]
        articles.extend(a.group(0) for a in ARTICLE_FULL_RE.finditer(m.group("more") or ""))
        bucket = skipped if law.endswith(NON_STATUTE_SUFFIXES) else cited
        for art in articles:
            item = (law, re.sub(r"\s+", "", art))
            if item not in bucket:
                bucket.append(item)
    return cited, skipped


# ---------------------------------------------------------------------------
# 「检索说明」块四要素（校验值非空）
# ---------------------------------------------------------------------------
def _clean_line(line: str) -> str:
    return re.sub(r"[*`_>#]", "", line).replace("\u3000", " ")


def _clean_value(value: str) -> str:
    v = value.strip()
    v = re.sub(r"^[\s：:=\-·•]+", "", v)
    v = v.strip(" \t|")
    v = v.strip("`\"'“”‘’")
    v = re.sub(r"^[\[（(【]+|[\]）)】]+$", "", v).strip()
    return v


def _is_empty_value(value: str) -> bool:
    v = _clean_value(value)
    return v.lower() in VALUE_PLACEHOLDERS


def _value_from_line(line: str, pattern: str) -> Optional[str]:
    """从单行中取出 label 对应的值；支持“label：值”与 Markdown 表格行。"""
    clean = _clean_line(line)
    if clean.count("|") >= 2:
        cells = [c.strip() for c in clean.strip().strip("|").split("|")]
        for idx, cell in enumerate(cells):
            m = re.search(pattern, cell, re.IGNORECASE)
            if not m:
                continue
            tail = cell[m.end():]
            m2 = re.match(r"\s*[：:=]\s*(.*)$", tail)
            if m2 and m2.group(1).strip():
                return m2.group(1)
            for nxt in cells[idx + 1:]:
                if nxt:
                    return nxt
            return ""
        return None
    m = re.search(pattern, clean, re.IGNORECASE)
    if not m:
        return None
    tail = clean[m.end():]
    m2 = re.match(r"\s*[：:=]\s*(.*)$", tail)
    if m2:
        return m2.group(1)
    # 没有分隔符：可能是“query 网络购物 七日无理由退货”这类写法
    rest = tail.strip()
    return rest if rest else ""


def _all_label_patterns() -> List[str]:
    pats: List[str] = []
    for _name, patterns in SEARCH_NOTE_FIELDS:
        pats.extend(patterns)
    return pats


def find_note_field(block: str, patterns: Tuple[str, ...]) -> Tuple[bool, Optional[str]]:
    """在「检索说明」块中查找要素，返回 (是否出现字样, 值)。"""
    lines = block.splitlines()
    found_label = False
    for i, line in enumerate(lines):
        for pattern in patterns:
            value = _value_from_line(line, pattern)
            if value is None:
                continue
            found_label = True
            if not _is_empty_value(value):
                return True, _clean_value(value)
            # 本行值为空：尝试取下一行（排除下一行本身是别的要素标签）
            for nxt in lines[i + 1:]:
                if not nxt.strip():
                    continue
                nxt_clean = _clean_line(nxt)
                if any(re.search(p, nxt_clean, re.IGNORECASE) for p in _all_label_patterns()):
                    break
                if re.match(r"\s*(#{1,6}\s|[-*]\s*$)", nxt_clean):
                    break
                if not _is_empty_value(nxt_clean):
                    return True, _clean_value(nxt_clean)
                break
    return found_label, None


def _note_block_at(text: str, match: re.Match) -> str:
    window = text[match.start(): match.start() + 2500]
    rel = match.end() - match.start()
    cut = re.search(r"\n\s*(?:#{1,6}\s|(?:[一二三四五六七八九十]+)、)", window[rel:])
    return window[: rel + cut.start()] if cut else window


def locate_search_note_block(text: str) -> Optional[str]:
    """定位「检索说明」块。

    交付物标题也可能含“检索说明”字样，故遍历前置窗口内的所有出现位置，
    取“四要素有实值数量最多”的候选块，避免锚到标题上得出空块。
    """
    head = text[:2000]
    candidates: List[Tuple[int, int, str]] = []
    for order, m in enumerate(re.finditer(r"检\s*索\s*说\s*明", head)):
        block = _note_block_at(text, m)
        score = 0
        for _name, patterns in SEARCH_NOTE_FIELDS:
            found_label, value = find_note_field(block, patterns)
            score += 2 if value is not None else (1 if found_label else 0)
        candidates.append((score, -order, block))
    if not candidates:
        return None
    candidates.sort(key=lambda item: (item[0], item[1]), reverse=True)
    return candidates[0][2]


def check_search_note(text: str) -> Tuple[List[str], List[str]]:
    """检查交付物开头的「检索说明」块，且 query/keywords/size/执行时间四要素值非空。"""
    checks: List[str] = []
    errors: List[str] = []
    block = locate_search_note_block(text)
    if block is None:
        errors.append(
            "交付物开头缺少「检索说明」块（应含 query、keywords、size/条数、执行时间四要素及其实值）"
        )
        return checks, errors
    checks.append("交付物开头存在「检索说明」块")

    for name, patterns in SEARCH_NOTE_FIELDS:
        found_label, value = find_note_field(block, patterns)
        if not found_label:
            errors.append(f"「检索说明」块缺少要素：{name}")
            continue
        if value is None:
            errors.append(f"「检索说明」块要素 {name} 值为空（只写了字样，未填实值）")
            continue
        if name == "size" and not re.search(r"[0-9０-９" + NUM_CHARS + r"]", value):
            errors.append(f"「检索说明」块要素 size 值不含条数数字：{value!r}")
            continue
        if name == "执行时间" and not re.search(
            r"\d{4}\s*[-/.年]\s*\d{1,2}|\d{4}\s*年", value
        ):
            errors.append(f"「检索说明」块要素 执行时间 值不含具体日期：{value!r}")
            continue
        shown = value if len(value) <= 60 else value[:57] + "..."
        checks.append(f"「检索说明」要素 {name} = {shown}")
    return checks, errors


# ---------------------------------------------------------------------------
# 交付物读取
# ---------------------------------------------------------------------------
def load_doc_text(path: Path) -> str:
    if not path.exists():
        raise InputError(f"交付物不存在：{path}")
    suffix = path.suffix.lower()
    if suffix in {".md", ".markdown", ".txt", ".text", ""}:
        try:
            return path.read_text(encoding="utf-8", errors="replace")
        except OSError as exc:
            raise InputError(f"无法读取交付物：{exc}")
    if suffix == ".docx":
        try:
            from docx import Document
        except ImportError:
            raise InputError(
                ".docx 交付物需要可选依赖 python-docx，当前环境未安装；"
                "请执行 pip3 install python-docx，或按平台口径先用 .md/.txt 过门禁再转 Word"
            )
        try:
            document = Document(str(path))
        except Exception as exc:
            raise InputError(f"python-docx 无法打开 docx：{exc}")
        blocks = [p.text for p in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                blocks.extend(cell.text for cell in row.cells)
        return "\n".join(blocks)
    raise InputError(
        f"不支持的交付物类型：{suffix}（主要支持 .md/.txt，可选支持 .docx）"
    )


# ---------------------------------------------------------------------------
# 主流程
# ---------------------------------------------------------------------------
def main() -> int:
    parser = argparse.ArgumentParser(
        description="法条引用核验门禁：从交付物提取《法律名》第X条，"
                    "与落盘检索 JSON 的法规名/条号专用字段严格比对；"
                    "并校验交付物开头「检索说明」块四要素有实值",
        epilog="交付物以 .md/.txt 为主要支持格式（平台口径：先出 Markdown 过门禁再转 Word）；"
               ".docx 为可选支持，依赖 python-docx。退出码：0=通过，1=拦截，2=输入错误。",
    )
    parser.add_argument("--json", required=True, type=Path,
                        help="本次检索返回落盘的归一化 JSON 文件路径")
    parser.add_argument("--doc", required=True, type=Path,
                        help="交付物文本（.md/.txt 为主，.docx 可选）")
    parser.add_argument("--require-citations", action="store_true",
                        help="要求交付物必须至少含一条《法律名》第X条引用；"
                             "默认零引用只告警（命中整部法律、未取具体条文的交付属合规情形）")
    args = parser.parse_args()

    try:
        if not args.json.exists():
            raise InputError(f"JSON 文件不存在：{args.json}")
        data = json.loads(args.json.read_text(encoding="utf-8"))
    except InputError as exc:
        print(f"输入错误：{exc}", file=sys.stderr)
        return 2
    except (OSError, json.JSONDecodeError) as exc:
        print(f"输入错误：无法解析检索 JSON：{exc}", file=sys.stderr)
        return 2

    try:
        doc_text = load_doc_text(args.doc)
    except InputError as exc:
        print(f"输入错误：{exc}", file=sys.stderr)
        return 2

    checks: List[str] = []
    warns: List[str] = []
    errors: List[str] = []

    records = collect_records(data)
    c_checks, c_warns, c_errors = validate_contract(data, records)
    checks.extend(c_checks)
    warns.extend(c_warns)
    errors.extend(c_errors)
    contract_ok = not c_errors

    citations, skipped = extract_citations(doc_text)
    for law, article in skipped:
        checks.append(f"[非法条条号已忽略] 《{law}》{article}（非法律法规名称，不参与溯源）")

    if not citations:
        message = ("交付物未提取到任何《法律名》第X条引用"
                   "（如本次检索仅命中整部法律、尚未取具体条文，属合规情形）")
        if args.require_citations:
            errors.append(f"{message}；--require-citations 已开启，按缺漏拦截")
        else:
            warns.append(f"{message}；已按告警处理，如需强制要求请加 --require-citations")

    verified = 0
    for law, article in citations:
        label = f"《{law}》{article}"
        if not contract_ok:
            errors.append(f"[引用无法核验] {label} —— JSON 未按归一化契约落盘，拒绝降级放行")
            continue
        hit, reason = match_citation(law, article, records)
        if hit:
            verified += 1
            checks.append(f"[引用可溯源] {label} → {hit}")
        else:
            errors.append(
                f"[引用不可溯源] {label} 未见于检索 JSON 的法规名+条号字段"
                f"（疑似幻觉或超出本次检索范围）：{reason}"
            )

    note_checks, note_errors = check_search_note(doc_text)
    checks.extend(note_checks)
    errors.extend(note_errors)

    print("== 法条引用核验：通过清单 ==")
    for item in checks:
        print(f"- {item}")
    if warns:
        print("== 警告清单（不阻断） ==")
        for item in warns:
            print(f"- {item}")
    if errors:
        print("== 拦截清单 ==", file=sys.stderr)
        for item in errors:
            print(f"- {item}", file=sys.stderr)
        print(f"核验不通过：共拦截 {len(errors)} 项（禁止交付）", file=sys.stderr)
        return 1
    print(
        f"核验通过：{verified}/{len(citations)} 条法条引用均可溯源至检索 JSON 的"
        f"法规名+条号字段，「检索说明」块四要素齐备且有实值"
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
