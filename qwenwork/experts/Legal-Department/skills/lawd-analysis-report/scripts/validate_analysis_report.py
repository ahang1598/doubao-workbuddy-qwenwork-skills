#!/usr/bin/env python3
"""
法律分析建议书交付前门禁（连接器口径）

校验四件事，任一不通过即以非零退出码阻断交付：
  1. 报告含「数据来源与局限」章节，且含 `工商数据获取状态：已取得/未取得` 声明行。
  2. 工商数据缺失硬门禁：声明「未取得」时，主体信息表不得出现被当作确定结论的
     企业登记值（USCC 18 位码 / 具体注册资本 / 具体成立日期），此类字段必须为
     "未取得 / 待核验 / 未提供" 等占位；否则视为无据编造，拦截。
  3. 供应商泄漏：交付物不得出现写死的供应商/工具名（天眼查、tianyancha、
     WebSearch、dws law、mcpServerConfig、mcp__tianyancha 等）。
  4. 引用溯源（可选，传入检索结果 JSON 时启用）：报告引用的法条（法规名+条号）
     与案号必须能在对应检索结果 JSON 中找到；对不上即拦截。

用法：
    python3 validate_analysis_report.py <报告.md 或 .docx>
    python3 validate_analysis_report.py <报告> --cases cases.json --laws laws.json

说明：
  - 报告为 .docx 时需要 python-docx；未安装则请传入 Markdown 草稿。
  - 检索结果 JSON 接受本技能归一化契约（caseResult / lawResult），也兼容
    形如 {"cases":[...]} / {"laws":[...]} 的简化结构。
"""

import argparse
import json
import re
import sys
from pathlib import Path


# 写死供应商/工具名黑名单（作为语义关键词举例出现在说明性文字中不算，
# 但交付物正文出现这些字面量即视为泄漏）
FORBIDDEN_LITERALS = [
    "天眼查",
    "tianyancha",
    "websearch",
    "dws law",
    "mcpserverconfig",
    "mcp__tianyancha",
    "mcp.tianyancha.com",
]

STATUS_OBTAINED = "已取得"
STATUS_MISSING = "未取得"

# 缺数据时禁止出现的“确定值”模式（覆盖同义表述，防"实缴1000万元""成立于2020年"绕过）
USCC_PATTERN = re.compile(r"[0-9A-HJ-NP-RT-UW-Y]{18}")
REGISTERED_CAPITAL_PATTERN = re.compile(
    r"(?:注册资本|实缴[^，。；|]{0,4}|认缴[^，。；|]{0,4}|出资额|出资[^，。；|]{0,6})"
    r"[^\n|]{0,8}[\d,]+\.?\d*\s*(?:万元|万|亿元|亿|元)"
)
FOUND_DATE_PATTERN = re.compile(
    r"(?:成立日期|成立于|设立日期|设立于|注册日期)"
    r"[^\n|]{0,8}\d{4}(?:[-年./]\d{1,2})?(?:年)?"
)

PLACEHOLDER_TOKENS = ("未取得", "待核验", "待核实", "未提供", "未查询", "N/A", "—", "-")


def read_report(path: Path) -> str:
    if not path.exists():
        raise FileNotFoundError(f"报告文件不存在：{path}")
    if path.suffix.lower() == ".docx":
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError(
                "读取 .docx 需要 python-docx（pip install python-docx），"
                "或改为传入 Markdown 草稿"
            ) from exc
        document = Document(str(path))
        blocks = [p.text for p in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                blocks.extend(cell.text for cell in row.cells)
        return "\n".join(blocks)
    return path.read_text(encoding="utf-8")


def check_data_source_section(text: str, errors: list) -> str:
    if "数据来源与局限" not in text:
        errors.append("缺少「数据来源与局限」章节")
    match = re.search(r"工商数据获取状态[：:]\s*([^\n|]+)", text)
    if not match:
        errors.append("缺少 `工商数据获取状态：已取得/未取得` 声明行")
        return ""
    value = match.group(1)
    if STATUS_OBTAINED in value and STATUS_MISSING not in value:
        return STATUS_OBTAINED
    if STATUS_MISSING in value:
        return STATUS_MISSING
    # 声明行仍是模板二选一原文（同时含两者且未删减）
    if STATUS_OBTAINED in value and STATUS_MISSING in value:
        errors.append("`工商数据获取状态` 声明行仍是模板二选一，未据实填写")
    else:
        errors.append(f"`工商数据获取状态` 声明值无法识别：{value.strip()}")
    return ""


def check_missing_company_hardline(text: str, status: str, errors: list) -> None:
    if status != STATUS_MISSING:
        return
    hits = []
    for label, pattern in (
        ("统一社会信用代码(USCC)", USCC_PATTERN),
        ("注册资本具体金额", REGISTERED_CAPITAL_PATTERN),
        ("成立日期具体日期", FOUND_DATE_PATTERN),
    ):
        for m in pattern.finditer(text):
            # 若命中值所在行含占位标记，则视为合规
            line_start = text.rfind("\n", 0, m.start()) + 1
            line_end = text.find("\n", m.end())
            line = text[line_start: line_end if line_end != -1 else len(text)]
            if any(tok in line for tok in PLACEHOLDER_TOKENS):
                continue
            hits.append(f"{label}: “{m.group(0)}”")
    if hits:
        errors.append(
            "工商数据声明为「未取得」，但报告出现确定性企业登记值（疑似编造）："
            + "；".join(hits[:5])
        )


def check_forbidden_literals(text: str, errors: list) -> None:
    lowered = text.lower()
    found = [lit for lit in FORBIDDEN_LITERALS if lit in lowered]
    if found:
        errors.append("交付物泄漏写死供应商/工具名：" + "、".join(found))


def load_json(path: Path):
    return json.loads(path.read_text(encoding="utf-8"))


def extract_case_nos(data) -> set:
    result = set()

    def walk(obj):
        if isinstance(obj, dict):
            for key, value in obj.items():
                if key in ("caseNo", "case_no") and isinstance(value, str) and value.strip():
                    result.add(value.strip())
                else:
                    walk(value)
        elif isinstance(obj, list):
            for item in obj:
                walk(item)

    walk(data)
    return result


def extract_law_refs(data) -> set:
    """返回 (法规名, 条号) 集合。"""
    result = set()

    def walk(obj):
        if isinstance(obj, dict):
            name = obj.get("lawName") or obj.get("law_name")
            order = obj.get("lawOrder") or obj.get("law_order")
            if isinstance(name, str) and isinstance(order, str) and name.strip() and order.strip():
                result.add((name.strip(), order.strip()))
            for value in obj.values():
                walk(value)
        elif isinstance(obj, list):
            for item in obj:
                walk(item)

    walk(data)
    return result


# 报告中案号的宽松识别：形如「（2023）京0105民初12345号」
DOC_CASE_NO_PATTERN = re.compile(r"[（(]\d{4}[)）][^\s，,。；;、]{2,}?号")
# 报告中法条引用：《法规名》第X条
DOC_LAW_REF_PATTERN = re.compile(r"《([^》]{2,40})》第([一二三四五六七八九十百零〇\d]+条(?:之[一二三四五六七八九十]+)?)")


def check_case_traceability(text: str, cases_json: Path, errors: list) -> None:
    data = load_json(cases_json)
    retrieved = extract_case_nos(data)
    cited = {m.group(0) for m in DOC_CASE_NO_PATTERN.finditer(text)}
    orphan = sorted(c for c in cited if c not in retrieved)
    if orphan:
        errors.append(
            "报告引用的案号在案例检索结果中找不到（疑似编造）：" + "；".join(orphan[:5])
        )


def normalize_order(order: str) -> str:
    return order.replace("第", "").replace("条", "").strip()


def check_law_traceability(text: str, laws_json: Path, errors: list) -> None:
    data = load_json(laws_json)
    retrieved = extract_law_refs(data)
    retrieved_norm = {(name, normalize_order(order)) for name, order in retrieved}
    retrieved_names = {name for name, _ in retrieved}
    orphan = []
    for m in DOC_LAW_REF_PATTERN.finditer(text):
        name = m.group(1).strip()
        order = normalize_order(m.group(2))
        # 名称需在检索结果出现，且该名称下条号需匹配
        name_hit = any(name in rn or rn in name for rn in retrieved_names)
        pair_hit = any(
            (name in rn or rn in name) and order == ro
            for rn, ro in retrieved_norm
        )
        if not name_hit or not pair_hit:
            orphan.append(f"《{name}》第{m.group(2)}")
    if orphan:
        errors.append(
            "报告引用的法条在法规检索结果中找不到对应条号（疑似凭记忆罗列）："
            + "；".join(orphan[:5])
        )


def main() -> int:
    parser = argparse.ArgumentParser(
        description="法律分析建议书交付前门禁（连接器口径）",
    )
    parser.add_argument("report", type=Path, help="报告文件（.md 或 .docx）")
    parser.add_argument("--cases", type=Path, default=None, help="案例检索结果 JSON（可选，启用案号溯源）")
    parser.add_argument("--laws", type=Path, default=None, help="法规检索结果 JSON（可选，启用法条溯源）")
    args = parser.parse_args()

    try:
        text = read_report(args.report)
    except (FileNotFoundError, RuntimeError) as exc:
        print(f"门禁失败：{exc}", file=sys.stderr)
        return 1

    errors: list = []
    status = check_data_source_section(text, errors)
    check_missing_company_hardline(text, status, errors)
    check_forbidden_literals(text, errors)

    if args.cases:
        try:
            check_case_traceability(text, args.cases, errors)
        except (OSError, json.JSONDecodeError) as exc:
            errors.append(f"无法读取案例检索 JSON：{exc}")
    if args.laws:
        try:
            check_law_traceability(text, args.laws, errors)
        except (OSError, json.JSONDecodeError) as exc:
            errors.append(f"无法读取法规检索 JSON：{exc}")

    if errors:
        print(f"交付门禁未通过，共 {len(errors)} 项：", file=sys.stderr)
        for error in errors:
            print(f"- {error}", file=sys.stderr)
        return 1

    print("交付门禁通过：数据来源声明、工商数据缺失硬门禁、供应商泄漏、引用溯源检查均通过。")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
