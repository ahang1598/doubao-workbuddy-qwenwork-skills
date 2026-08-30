#!/usr/bin/env python3
"""Build the six static quick-template DOCX assets from their Markdown sources."""

from __future__ import annotations

import argparse
import re
import time
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


SUITE_ROOT = Path(__file__).resolve().parents[1]
SKILLS = (
    "cn-prenuptial-property-agreement",
    "cn-marital-property-agreement",
    "cn-divorce-agreement",
    "cn-cohabitation-agreement",
    "cn-family-property-partition",
    "cn-adult-voluntary-guardianship-agreement",
)
EAST_ASIA_FONT = "Arial Unicode MS"
LATIN_FONT = "Times New Roman"
INK = RGBColor(31, 41, 55)
NAVY = RGBColor(31, 58, 95)
MUTED = RGBColor(89, 96, 105)
PLACEHOLDER = RGBColor(0, 82, 136)


def set_run_font(run, size: float | None = None, bold: bool | None = None, color=None):
    run.font.name = LATIN_FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), LATIN_FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), LATIN_FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def add_paragraph_border(paragraph, color="AAB4C3", size="12"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), size)
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), color)
    p_bdr.append(left)


def add_shading(paragraph, fill="F3F5F8"):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    label = paragraph.add_run("第 ")
    set_run_font(label, 9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    result = OxmlElement("w:t")
    result.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    set_run_font(run, 9, color=MUTED)
    run._r.extend([begin, instr, separate, result, end])
    tail = paragraph.add_run(" 页")
    set_run_font(tail, 9, color=MUTED)


INLINE_TOKEN = re.compile(r"(\*\*.+?\*\*|`.+?`|\[[^\]]+\])")


def add_inline(paragraph, text: str, default_size=10.5, default_bold=False, default_color=INK):
    for part in filter(None, INLINE_TOKEN.split(text)):
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, default_size, True, default_color)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, default_size, True, PLACEHOLDER)
            run.font.highlight_color = 15
        elif part.startswith("[") and part.endswith("]"):
            run = paragraph.add_run(part)
            set_run_font(run, default_size, True, PLACEHOLDER)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, default_size, default_bold, default_color)


def configure_styles(doc: Document):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = LATIN_FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), LATIN_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), LATIN_FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.35
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for name, size, before, after, color in (
        ("Heading 1", 15, 14, 7, NAVY),
        ("Heading 2", 12.5, 11, 5, NAVY),
        ("Heading 3", 11, 8, 4, INK),
    ):
        style = styles[name]
        style.font.name = LATIN_FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), LATIN_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), LATIN_FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Template Note" not in styles:
        note = styles.add_style("Template Note", WD_STYLE_TYPE.PARAGRAPH)
    else:
        note = styles["Template Note"]
    note.font.name = LATIN_FONT
    note._element.rPr.rFonts.set(qn("w:ascii"), LATIN_FONT)
    note._element.rPr.rFonts.set(qn("w:hAnsi"), LATIN_FONT)
    note._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    note.font.size = Pt(9.5)
    note.font.color.rgb = MUTED
    note.paragraph_format.left_indent = Cm(0.35)
    note.paragraph_format.right_indent = Cm(0.2)
    note.paragraph_format.space_before = Pt(4)
    note.paragraph_format.space_after = Pt(7)
    note.paragraph_format.line_spacing = 1.25


def configure_document(doc: Document, running_title: str):
    section = doc.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.0)
    section.footer_distance = Cm(1.0)

    header = section.header.paragraphs[0]
    header.text = ""
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run(f"{running_title}｜空白通用工作稿")
    set_run_font(run, 8.5, color=MUTED)
    add_page_field(section.footer.paragraphs[0])


def add_title(doc: Document, text: str, first=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4 if first else 10)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    add_inline(p, text, 18 if first else 16, True, NAVY if first else INK)


def add_body_line(doc: Document, line: str):
    stripped = line.strip()
    if not stripped:
        return
    if stripped.startswith("> "):
        p = doc.add_paragraph(style="Template Note")
        add_shading(p)
        add_paragraph_border(p)
        add_inline(p, stripped[2:], 9.5, False, MUTED)
        return
    if stripped.startswith("### "):
        p = doc.add_paragraph(style="Heading 3")
        add_inline(p, stripped[4:], 11, True, INK)
        return
    if stripped.startswith("## "):
        p = doc.add_paragraph(style="Heading 2")
        add_inline(p, stripped[3:], 12.5, True, NAVY)
        return
    if stripped.startswith("# "):
        add_title(doc, stripped[2:])
        return
    if stripped.startswith("- "):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.74)
        p.paragraph_format.first_line_indent = Cm(-0.37)
        add_inline(p, "• " + stripped[2:])
        return
    p = doc.add_paragraph()
    if re.match(r"^\d+\.\s", stripped):
        p.paragraph_format.left_indent = Cm(0.74)
        p.paragraph_format.first_line_indent = Cm(-0.74)
    if stripped.startswith(("甲方签字", "乙方签字", "委托人签字", "拟任监护人签字", "备选监护人", "监督人", "律师/见证", "见证/律师")):
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.keep_together = True
    add_inline(p, stripped)


def build_one(skill_name: str) -> tuple[Path, float]:
    source = SUITE_ROOT / "skills" / skill_name / "references" / "template.md"
    target = SUITE_ROOT / "skills" / skill_name / "assets" / "quick-template.docx"
    lines = source.read_text(encoding="utf-8").splitlines()
    first_title = next(line[2:] for line in lines if line.startswith("# "))

    started = time.perf_counter()
    doc = Document()
    configure_styles(doc)
    configure_document(doc, first_title.split("（", 1)[0])
    doc.core_properties.title = first_title
    doc.core_properties.subject = "中国大陆婚姻家事空白通用工作模板"
    doc.core_properties.author = "婚姻家事法律专家"
    doc.core_properties.keywords = "draft, quick template, 中国大陆, 婚姻家事"

    add_title(doc, first_title, first=True)
    status = doc.add_paragraph()
    status.alignment = WD_ALIGN_PARAGRAPH.CENTER
    status.paragraph_format.space_after = Pt(10)
    add_inline(status, "状态：DRAFT｜空白通用工作稿｜签署前须完成个案核验与律师实质复核", 9.5, True, MUTED)

    consumed_first_heading = False
    for line in lines:
        if line.startswith("# ") and not consumed_first_heading:
            consumed_first_heading = True
            continue
        add_body_line(doc, line)

    target.parent.mkdir(parents=True, exist_ok=True)
    doc.save(target)
    return target, time.perf_counter() - started


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--skill", choices=SKILLS, action="append", help="Build only the selected skill; repeatable")
    args = parser.parse_args()
    selected = tuple(args.skill) if args.skill else SKILLS
    total_started = time.perf_counter()
    for skill_name in selected:
        target, elapsed = build_one(skill_name)
        print(f"built {target.relative_to(SUITE_ROOT)} in {elapsed:.3f}s")
    print(f"built {len(selected)} templates in {time.perf_counter() - total_started:.3f}s")


if __name__ == "__main__":
    main()
