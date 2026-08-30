#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
将 Markdown 文章转换为 Word 文档，支持插入配图。
经过验证的脚本模板，Agent 执行时基于此模板修改路径和图片映射即可。

用法：
    python3 md_to_docx.py --md <md文件路径> --docx <docx输出路径> [--images <key1:path1> <key2:path2> ...]

示例（无配图）：
    python3 md_to_docx.py --md article.md --docx article.docx

示例（有配图）：
    python3 md_to_docx.py --md article.md --docx article.docx \
        --images cover:封面图.png diagram:示意图.png
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re
import os
import argparse

# ============================================================
# 安全路径校验
# ============================================================

# 允许的输入文件扩展名
ALLOWED_INPUT_EXTENSIONS = {'.md', '.markdown', '.txt'}
# 允许的输出文件扩展名
ALLOWED_OUTPUT_EXTENSIONS = {'.docx'}
# 允许的图片文件扩展名
ALLOWED_IMAGE_EXTENSIONS = {'.png', '.jpg', '.jpeg', '.gif', '.webp', '.bmp', '.svg'}


def _get_allowed_dirs():
    """获取允许访问的目录列表（仅当前工作目录）"""
    return [os.path.realpath(os.getcwd())]


def validate_read_path(file_path, allowed_extensions=None, description="文件"):
    """校验读取文件路径的安全性

    Args:
        file_path: 待校验的文件路径
        allowed_extensions: 允许的文件扩展名集合
        description: 文件描述（用于错误提示）

    Returns:
        解析后的绝对路径

    Raises:
        ValueError: 路径不合法时抛出
    """
    if not file_path or not file_path.strip():
        raise ValueError(f"{description}路径不能为空")

    # 解析为绝对路径（消除 ../ 等相对路径）
    resolved = os.path.realpath(os.path.abspath(file_path))

    # 检查路径遍历：只允许在当前工作目录下
    allowed_dirs = _get_allowed_dirs()
    path_allowed = any(
        resolved.startswith(d + os.sep) or resolved == d
        for d in allowed_dirs
    )
    if not path_allowed:
        raise ValueError(
            f"{description}路径不安全：只允许访问当前工作目录下的文件。"
            f"路径：{file_path} -> {resolved}"
        )

    # 检查文件扩展名
    if allowed_extensions:
        ext = os.path.splitext(resolved)[1].lower()
        if ext not in allowed_extensions:
            raise ValueError(
                f"{description}扩展名不合法：'{ext}'，允许的扩展名：{allowed_extensions}"
            )

    # 检查文件是否存在
    if not os.path.exists(resolved):
        raise ValueError(f"{description}不存在：{resolved}")

    # 检查是否为符号链接指向外部（防止 symlink 攻击）
    if os.path.islink(file_path):
        link_target = os.path.realpath(file_path)
        if not any(link_target.startswith(d + os.sep) for d in allowed_dirs):
            raise ValueError(f"{description}路径不安全：符号链接指向不允许的位置")

    return resolved


def validate_write_path(file_path, allowed_extensions, input_file_path, description="输出文件"):
    """校验写入文件路径的安全性（严格限制写入范围）

    输出文件只允许写入到：输入文件所在目录 或 当前工作目录（二者取其一）。

    Args:
        file_path: 待校验的输出文件路径
        allowed_extensions: 允许的文件扩展名集合
        input_file_path: 已校验的输入文件绝对路径（用于确定允许写入的目录）
        description: 文件描述（用于错误提示）

    Returns:
        解析后的绝对路径

    Raises:
        ValueError: 路径不合法时抛出
    """
    if not file_path or not file_path.strip():
        raise ValueError(f"{description}路径不能为空")

    # 解析为绝对路径
    resolved = os.path.normpath(os.path.abspath(file_path))

    # 写入范围严格限制：只允许写入输入文件所在目录 或 当前工作目录
    cwd = os.path.realpath(os.getcwd())
    input_dir = os.path.dirname(os.path.realpath(input_file_path))
    allowed_write_dirs = [cwd, input_dir]

    output_dir = os.path.dirname(resolved)
    output_dir_real = os.path.realpath(output_dir)

    write_allowed = any(
        output_dir_real == d or output_dir_real.startswith(d + os.sep)
        for d in allowed_write_dirs
    )
    if not write_allowed:
        raise ValueError(
            f"{description}路径不安全：输出文件只允许写入当前工作目录或输入文件所在目录。"
            f"路径：{file_path}，允许的目录：{allowed_write_dirs}"
        )

    # 确保输出目录已存在（不允许创建任意目录）
    if not os.path.isdir(output_dir):
        raise ValueError(
            f"{description}的父目录不存在：{output_dir}。请先创建目录或指定已存在的目录。"
        )

    # 检查文件扩展名
    ext = os.path.splitext(resolved)[1].lower()
    if ext not in allowed_extensions:
        raise ValueError(
            f"{description}扩展名不合法：'{ext}'，允许的扩展名：{allowed_extensions}"
        )

    # 防止覆盖输入文件
    if os.path.realpath(resolved) == os.path.realpath(input_file_path):
        raise ValueError(f"{description}不能与输入文件相同")

    return resolved


def add_image_with_caption(doc, image_path, caption_text):
    """插入图片并添加图注（经过验证的正确方式）

    ⚠️ 关键：通过 paragraph > run > add_picture 插入，
    不要用 doc.add_picture()，否则某些 Word 版本不显示图片。
    """
    # 安全校验图片路径
    try:
        validated_path = validate_read_path(
            image_path,
            allowed_extensions=ALLOWED_IMAGE_EXTENSIONS,
            description="图片"
        )
    except ValueError as e:
        print(f"警告：图片路径校验失败 - {e}")
        return False

    # 图片段落（居中）
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(validated_path, width=Inches(5.5))

    # 图注：居中、灰色、小字号
    caption = doc.add_paragraph(caption_text)
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.runs[0].font.size = Pt(9)
    caption.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    # 图片后空一行
    doc.add_paragraph()
    return True


def set_paragraph_spacing(paragraph, before_pt=0, after_pt=6, line_spacing=1.5):
    """设置段落间距和行距"""
    from docx.shared import Pt as PtShared
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = PtShared(before_pt)
    paragraph_format.space_after = PtShared(after_pt)
    paragraph_format.line_spacing = line_spacing


def add_styled_paragraph(doc, text, bold=False, font_size=12, alignment=None, color=None):
    """添加带样式的段落，正确处理内联加粗"""
    paragraph = doc.add_paragraph()
    if alignment:
        paragraph.alignment = alignment

    # 处理内联 **加粗** 格式
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif bold:
            run = paragraph.add_run(part)
            run.bold = True
        else:
            run = paragraph.add_run(part)
        run.font.size = Pt(font_size)
        if color:
            run.font.color.rgb = color

    set_paragraph_spacing(paragraph)
    return paragraph


def is_separator_line(text):
    """判断是否为分隔线（支持多种格式）"""
    stripped = text.strip()
    if stripped.startswith('---'):
        return True
    # 匹配 ···、◆◆◆、🔥🔥🔥、———、*** 等装饰性分隔符
    if len(stripped) <= 10 and stripped and all(c == stripped[0] for c in stripped):
        return True
    # 匹配纯装饰字符组成的行（emoji、符号等）
    if len(stripped) >= 2 and not any(c.isalnum() for c in stripped):
        return True
    return False


def convert_md_to_docx(md_file, docx_file, images_dict=None):
    """将 Markdown 文件转换为 Word 文档

    Args:
        md_file: Markdown 文件路径
        docx_file: 输出的 Word 文件路径
        images_dict: 图片映射字典 {关键词: 图片路径}，
                     其中 "cover" 为封面图（插在主标题下方），
                     其他 key 按小标题关键词匹配位置插入
    """
    if images_dict is None:
        images_dict = {}

    # 安全校验输入文件路径（必须存在且在工作目录下）
    validated_md = validate_read_path(
        md_file,
        allowed_extensions=ALLOWED_INPUT_EXTENSIONS,
        description="Markdown 输入文件"
    )

    # 安全校验输出文件路径（仅允许写入输入文件同目录或工作目录）
    validated_docx = validate_write_path(
        docx_file,
        allowed_extensions=ALLOWED_OUTPUT_EXTENSIONS,
        input_file_path=validated_md,
        description="Word 输出文件"
    )

    # 安全校验所有图片路径（必须存在且在工作目录下）
    validated_images = {}
    for key, img_path in images_dict.items():
        try:
            validated_img = validate_read_path(
                img_path,
                allowed_extensions=ALLOWED_IMAGE_EXTENSIONS,
                description=f"图片({key})"
            )
            validated_images[key] = validated_img
        except ValueError as e:
            print(f"警告：跳过图片 '{key}' - {e}")

    with open(validated_md, 'r', encoding='utf-8') as f:
        lines = f.read().split('\n')

    doc = Document()

    # 设置默认中文字体和段落样式
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft YaHei'
    font._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    font.size = Pt(12)
    # 设置默认段落间距
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.5

    inserted_images = 0
    inserted_keys = set()  # 跟踪已插入的图片 key，确保每张只插入一次
    skip_image_suggestions = False
    prev_was_empty = False  # 防止连续空行

    for line in lines:
        line_stripped = line.rstrip()

        # 跳过"配图建议"和"元信息"部分（md 末尾的辅助信息不需要写入 Word）
        if (line_stripped.startswith("## 配图建议") or line_stripped.startswith("## 配图需求")
                or "【配图建议】" in line_stripped or "【配图需求】" in line_stripped
                or line_stripped.startswith("**元信息**") or line_stripped.startswith("## 元信息")):
            skip_image_suggestions = True
            continue
        if skip_image_suggestions:
            continue

        # 空行处理：最多保留一个空行，用段前间距代替
        if not line_stripped.strip():
            if not prev_was_empty:
                prev_was_empty = True
            continue
        prev_was_empty = False

        # 分隔线（支持 ---、···、◆◆◆、———等）
        if is_separator_line(line_stripped):
            separator = doc.add_paragraph()
            separator.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = separator.add_run('— — —')
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(180, 180, 180)
            set_paragraph_spacing(separator, before_pt=12, after_pt=12)
            continue

        # 标题处理
        if line_stripped.startswith('### '):
            text = line_stripped[4:].strip()
            heading = doc.add_heading(text, level=2)
            set_paragraph_spacing(heading, before_pt=18, after_pt=8)
        elif line_stripped.startswith('## '):
            text = line_stripped[3:].strip()
            heading = doc.add_heading(text, level=1)
            set_paragraph_spacing(heading, before_pt=24, after_pt=10)
        elif line_stripped.startswith('# '):
            text = line_stripped[2:].strip()
            heading = doc.add_heading(text, level=0)
            for run in heading.runs:
                run.font.size = Pt(22)
            set_paragraph_spacing(heading, before_pt=0, after_pt=12)
            # 封面图插在主标题下方
            if 'cover' in validated_images:
                if add_image_with_caption(doc, validated_images['cover'], "封面图 | AI 生成"):
                    inserted_images += 1
        # 引用块
        elif line_stripped.startswith('> '):
            text = line_stripped[2:].strip()
            # 用带左边框效果的段落代替 Quote 样式（兼容性更好）
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Pt(24)
            set_paragraph_spacing(paragraph, before_pt=8, after_pt=8)
            # 处理内联加粗
            parts = re.split(r'(\*\*[^*]+\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = paragraph.add_run(part[2:-2])
                    run.bold = True
                else:
                    run = paragraph.add_run(part)
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(100, 100, 100)
                run.italic = True
        # 无序列表
        elif line_stripped.startswith('- ') or line_stripped.startswith('* '):
            text = line_stripped[2:].strip()
            paragraph = doc.add_paragraph(style='List Bullet')
            # 处理内联加粗
            clean_text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
            parts = re.split(r'(\*\*[^*]+\*\*)', clean_text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = paragraph.add_run(part[2:-2])
                    run.bold = True
                else:
                    paragraph.add_run(part)
            set_paragraph_spacing(paragraph, before_pt=2, after_pt=2)
        # 有序列表
        elif re.match(r'^\d+\.\s', line_stripped):
            text = re.sub(r'^\d+\.\s', '', line_stripped).strip()
            paragraph = doc.add_paragraph(style='List Number')
            clean_text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
            parts = re.split(r'(\*\*[^*]+\*\*)', clean_text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = paragraph.add_run(part[2:-2])
                    run.bold = True
                else:
                    paragraph.add_run(part)
            set_paragraph_spacing(paragraph, before_pt=2, after_pt=2)
        # 普通段落（处理加粗、链接等内联格式）
        else:
            text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', line_stripped)  # 去链接
            text = text.replace('`', '')
            paragraph = doc.add_paragraph()
            # 处理 **加粗** 片段
            parts = re.split(r'(\*\*[^*]+\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = paragraph.add_run(part[2:-2])
                    run.bold = True
                    run.font.size = Pt(12)
                else:
                    run = paragraph.add_run(part)
                    run.font.size = Pt(12)
            set_paragraph_spacing(paragraph)

        # 在标题行后插入对应配图（仅匹配 #/##/### 标题行，每个 key 最多插入一次）
        if line_stripped.startswith('#'):
            for key, img_path in validated_images.items():
                if key != 'cover' and key not in inserted_keys and key.lower() in line_stripped.lower():
                    if add_image_with_caption(doc, img_path, f"{key} | AI 生成"):
                        inserted_images += 1
                        inserted_keys.add(key)

    doc.save(validated_docx)
    print(f"✓ Word 文档已生成：{validated_docx}")
    print(f"✓ 已成功插入 {inserted_images} 张图片")


def main():
    parser = argparse.ArgumentParser(description='将 Markdown 文章转换为 Word 文档')
    parser.add_argument('--md', required=True, help='Markdown 文件路径')
    parser.add_argument('--docx', required=True, help='Word 输出文件路径')
    parser.add_argument('--images', nargs='*', default=[],
                        help='图片映射，格式为 key:path，例如 cover:封面图.png')

    args = parser.parse_args()

    # 解析图片映射
    images_dict = {}
    for img_arg in args.images:
        if ':' in img_arg:
            key, path = img_arg.split(':', 1)
            images_dict[key.strip()] = path.strip()

    try:
        convert_md_to_docx(args.md, args.docx, images_dict)
    except ValueError as e:
        print(f"错误：{e}")
        exit(1)


if __name__ == "__main__":
    main()
