#!/usr/bin/env python3
"""
Word文档格式化工具

功能：将合同文本转换为标准格式的Word文档
- 全文字体：宋体
- 合同标题：小三（15pt）加粗居中
- 一级标题（如"第一条"）：小四（12pt）加粗，Heading 1样式
- 二级标题（如"1.1"）：小四（12pt）不加粗，Heading 2样式
- 段后：0.5行
- 支持Markdown表格转Word真实表格
- 支持分页符（附件前、签署页前）
- 支持页码（第X页 共Y页）
- 支持Markdown内联格式：**加粗**、*斜体*

使用方式：
python docx-formatter.py --input <输入文件路径> --output <输出文件路径>
python docx-formatter.py --content "<文本内容>" --output <输出文件路径>
"""

import argparse
import os
import sys
import re
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("错误：缺少 python-docx 库，请先安装：pip install python-docx")
    sys.exit(1)


def set_run_font(run, font_name='宋体', font_size=12, bold=False, italic=False):
    """设置文本运行格式"""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold
    run.italic = italic
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)


def parse_inline_formatting(text):
    """解析Markdown内联格式，返回(text, bold, italic)列表"""
    result = []
    pattern = r'(\*\*\*[^*]+\*\*\*)|(\*\*[^*]+\*\*)|(\*[^*]+\*)|(_[^_]+_)'
    last_end = 0
    for match in re.finditer(pattern, text):
        if match.start() > last_end:
            plain_text = text[last_end:match.start()]
            if plain_text:
                result.append((plain_text, False, False))
        matched = match.group(0)
        if matched.startswith('***') and matched.endswith('***'):
            result.append((matched[3:-3], True, True))
        elif matched.startswith('**') and matched.endswith('**'):
            result.append((matched[2:-2], True, False))
        elif matched.startswith('*') and matched.endswith('*'):
            result.append((matched[1:-1], False, True))
        elif matched.startswith('_') and matched.endswith('_'):
            result.append((matched[1:-1], False, True))
        last_end = match.end()
    if last_end < len(text):
        remaining = text[last_end:]
        if remaining:
            result.append((remaining, False, False))
    if not result:
        result.append((text, False, False))
    return result


def add_formatted_paragraph(doc, text, font_size=12, bold=False, alignment=None, style=None):
    """添加一个段落，支持内联格式和样式"""
    if style:
        p = doc.add_paragraph(style=style)
    else:
        p = doc.add_paragraph()
    parts = parse_inline_formatting(text)
    for run_text, is_bold, is_italic in parts:
        run = p.add_run(run_text)
        set_run_font(run, font_size=font_size, bold=is_bold or bold, italic=is_italic)
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_after = Pt(6)
    return p


def is_contract_title(line, is_first):
    """判断是否为合同标题"""
    if is_first and ('协议' in line or '合同' in line or '附录' in line):
        return True
    if line.startswith('# ') and not line.startswith('## '):
        return True
    return False


def strip_markdown_heading(text):
    """移除Markdown标题标记"""
    return re.sub(r'^#{1,6}\s+', '', text)


def clean_markdown_artifacts(lines):
    """预处理：清除正文中的markdown残留标记"""
    cleaned = []
    for line in lines:
        if re.match(r'^-{3,}$|^\*{3,}$|^_{3,}$', line.strip()):
            continue
        line = re.sub(r'^(\s*)[-*]\s+', r'\1', line)
        line = re.sub(r'`([^`]+)`', r'\1', line)
        line = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', line)
        cleaned.append(line)
    return cleaned


def is_level_one_title(line):
    """判断是否为一级标题"""
    if re.match(r'^第[一二三四五六七八九十百千\d]+条', line):
        return True
    if re.match(r'^[一二三四五六七八九十]+、', line):
        return True
    if re.match(r'^#\s+', line) and not line.startswith('##'):
        return True
    return False


def is_level_two_title(line):
    """判断是否为二级标题"""
    if re.match(r'^\d+\.\d+', line):
        return True
    if re.match(r'^（[一二三四五六七八九十]+）', line):
        return True
    if re.match(r'^#{2,3}\s+', line):
        return True
    if re.match(r'^\d+\.\s', line) and not re.match(r'^\d+\.\d+', line):
        return True
    return False


def is_markdown_table_line(line):
    """判断是否为Markdown表格行"""
    stripped = line.strip()
    return stripped.startswith('|') and stripped.endswith('|') and stripped.count('|') >= 3


def is_table_separator(line):
    """判断是否为表格分隔行"""
    stripped = line.strip()
    return bool(re.match(r'^\|[\s\-:|]+\|$', stripped)) and '-' in stripped


def parse_table_block(lines, start_idx):
    """解析Markdown表格块，返回(rows, end_idx)"""
    rows = []
    i = start_idx
    while i < len(lines) and is_markdown_table_line(lines[i]):
        line = lines[i].strip()
        if not is_table_separator(line):
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            rows.append(cells)
        i += 1
    return rows, i


def add_table(doc, rows):
    """将解析的表格数据添加为Word真实表格"""
    if not rows:
        return
    num_cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            if j < num_cols:
                cell = table.cell(i, j)
                cell.text = ''
                p = cell.paragraphs[0]
                parts = parse_inline_formatting(cell_text)
                for run_text, is_bold, is_italic in parts:
                    run = p.add_run(run_text)
                    is_header = (i == 0)
                    set_run_font(run, font_size=10, bold=is_bold or is_header, italic=is_italic)
                p.paragraph_format.space_after = Pt(2)
    doc.add_paragraph()


def add_page_break(doc):
    """添加分页符"""
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(7)


def add_page_numbers(doc):
    """在页脚添加页码（第X页 共Y页）"""
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run1 = p.add_run('第 ')
    set_run_font(run1, font_size=9)

    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run2 = p.add_run()
    run2._element.append(fldChar1)
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    run3 = p.add_run()
    run3._element.append(instrText)
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run4 = p.add_run()
    run4._element.append(fldChar2)

    run5 = p.add_run(' 页 共 ')
    set_run_font(run5, font_size=9)

    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'begin')
    run6 = p.add_run()
    run6._element.append(fldChar3)
    instrText2 = OxmlElement('w:instrText')
    instrText2.set(qn('xml:space'), 'preserve')
    instrText2.text = ' NUMPAGES '
    run7 = p.add_run()
    run7._element.append(instrText2)
    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')
    run8 = p.add_run()
    run8._element.append(fldChar4)

    run9 = p.add_run(' 页')
    set_run_font(run9, font_size=9)


def is_signature_block_start(line):
    """判断是否为签署栏开始"""
    return '以下无正文' in line or '（以下无正文）' in line


def is_appendix_title(line):
    """判断是否为附件标题"""
    return bool(re.match(r'^附件[一二三四五六七八九十\d]+[：:]', line))


def create_formatted_document(content_lines):
    """根据内容行创建格式化的Word文档"""
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    add_page_numbers(doc)

    i = 0
    is_first_title = True

    while i < len(content_lines):
        line = content_lines[i].strip()

        if not line:
            i += 1
            continue

        # 签署栏前分页
        if is_signature_block_start(line):
            add_page_break(doc)
            add_formatted_paragraph(doc, line, font_size=12, bold=False,
                                    alignment=WD_ALIGN_PARAGRAPH.CENTER)
            i += 1
            continue

        # 附件标题前分页
        if is_appendix_title(line):
            add_page_break(doc)
            add_formatted_paragraph(doc, line, font_size=12, bold=True, style='Heading 1')
            i += 1
            continue

        # Markdown表格
        if is_markdown_table_line(line):
            rows, end_idx = parse_table_block(content_lines, i)
            add_table(doc, rows)
            i = end_idx
            continue

        # 合同标题
        if is_contract_title(line, is_first_title):
            title_text = strip_markdown_heading(line)
            add_formatted_paragraph(doc, title_text, font_size=15, bold=True,
                                    alignment=WD_ALIGN_PARAGRAPH.CENTER)
            is_first_title = False
            i += 1
            continue

        # 一级标题
        if is_level_one_title(line):
            title_text = strip_markdown_heading(line)
            add_formatted_paragraph(doc, title_text, font_size=12, bold=True, style='Heading 1')
            is_first_title = False
            i += 1
            continue

        # 二级标题
        if is_level_two_title(line):
            title_text = strip_markdown_heading(line)
            add_formatted_paragraph(doc, title_text, font_size=12, bold=False, style='Heading 2')
            i += 1
            continue

        # 普通正文
        add_formatted_paragraph(doc, line, font_size=12, bold=False)
        is_first_title = False
        i += 1

    return doc


def format_contract_text(input_path, output_path):
    """格式化合同文本为Word文档"""
    with open(input_path, 'r', encoding='utf-8') as f:
        content = f.read()
    lines = content.split('\n')
    lines = clean_markdown_artifacts(lines)
    doc = create_formatted_document(lines)
    doc.save(output_path)
    print(f"成功生成格式化Word文档：{output_path}")


def format_contract_content(content, output_path):
    """将合同内容直接格式化为Word文档"""
    lines = content.split('\n')
    lines = clean_markdown_artifacts(lines)
    doc = create_formatted_document(lines)
    doc.save(output_path)
    return output_path


def main():
    parser = argparse.ArgumentParser(description='合同文本格式化工具')
    parser.add_argument('--input', '-i', required=False, help='输入文本文件路径')
    parser.add_argument('--output', '-o', required=True, help='输出Word文档路径')
    parser.add_argument('--content', '-c', required=False, help='直接传入文本内容（可选）')
    args = parser.parse_args()

    if args.content:
        format_contract_content(args.content, args.output)
    elif args.input:
        format_contract_text(args.input, args.output)
    else:
        print("错误：必须提供 --input 或 --content 参数")
        sys.exit(1)


if __name__ == '__main__':
    main()
