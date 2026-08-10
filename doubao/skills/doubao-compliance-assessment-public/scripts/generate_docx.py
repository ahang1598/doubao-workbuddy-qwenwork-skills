#!/usr/bin/env python3
"""Convert a compliance assessment Markdown report into a Word document."""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


BLACK = RGBColor(0x00, 0x00, 0x00)


HEADING_RE = re.compile(r"^(#{1,6})\s+(.+?)\s*$")
ORDERED_RE = re.compile(r"^\s*(\d+)[.)]\s+(.+?)\s*$")
UNORDERED_RE = re.compile(r"^\s*[-*+]\s+(.+?)\s*$")
HR_RE = re.compile(r"^\s*([-*_])(?:\s*\1){2,}\s*$")

# 正文里模型常把联网检索来源写成句尾角标，如 ["https://..."] 或
# ["https://a","https://b"]。按模板正文不留 URL（尾注仍保留裸链接），
# 这里统一剥除这种角标。要求方括号内是引号包裹的 http(s) 链接，逗号可分隔多条，
# 从而只命中来源角标，不误伤 [1] 脚注标记、[文字](链接) 或尾注的 "URL：https://"。
_QUOTED_URL = r'["“”]https?://[^"“”]*["“”]'
INLINE_URL_CITATION_RE = re.compile(
    r'[\[［]\s*' + _QUOTED_URL + r'(?:\s*[,，]\s*' + _QUOTED_URL + r')*\s*[\]］]'
)


def clean_inline(text: str) -> str:
    """Remove light Markdown marks while preserving readable text."""
    text = INLINE_URL_CITATION_RE.sub("", text)
    # 剥除角标后，若两侧留下多余空格（英文语境），压回单空格；中文无空格不受影响。
    text = re.sub(r"[ \t]{2,}", " ", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1（\2）", text)
    return text.strip()


def set_east_asia_font(run, font_name: str = "宋体") -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_cell_text(cell, text: str) -> None:
    cell.text = clean_inline(text)
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            set_east_asia_font(run)
            run.font.size = Pt(10)
            run.font.color.rgb = BLACK


def set_table_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "auto")
        borders.append(border)
    tbl_pr.append(borders)


def is_table_separator(line: str) -> bool:
    cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell or "") for cell in cells)


def is_table_start(lines: list[str], index: int) -> bool:
    return (
        index + 1 < len(lines)
        and "|" in lines[index]
        and "|" in lines[index + 1]
        and is_table_separator(lines[index + 1])
    )


def parse_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def add_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    width = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=width)
    table.style = "Table Grid"
    set_table_borders(table)
    for row_index, row in enumerate(rows):
        for col_index in range(width):
            cell_text = row[col_index] if col_index < len(row) else ""
            set_cell_text(table.cell(row_index, col_index), cell_text)
        if row_index == 0:
            for cell in table.rows[row_index].cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True


def add_paragraph(document: Document, text: str, style: str | None = None) -> None:
    paragraph = document.add_paragraph(style=style)
    run = paragraph.add_run(clean_inline(text))
    set_east_asia_font(run)
    run.font.size = Pt(11)
    run.font.color.rgb = BLACK


def configure_document(document: Document) -> None:
    styles = document.styles
    styles["Normal"].font.name = "宋体"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].font.color.rgb = BLACK
    for style_name in ("Title", "Heading 1", "Heading 2", "Heading 3", "Intense Quote"):
        style = styles[style_name]
        style.font.name = "宋体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        style.font.color.rgb = BLACK


def convert_markdown_to_docx(markdown_path: Path, output_path: Path) -> None:
    lines = markdown_path.read_text(encoding="utf-8").splitlines()
    document = Document()
    configure_document(document)

    in_code_block = False
    code_buffer: list[str] = []
    index = 0

    while index < len(lines):
        line = lines[index].rstrip()

        if line.strip().startswith("```"):
            if in_code_block:
                if code_buffer:
                    add_paragraph(document, "\n".join(code_buffer))
                code_buffer = []
                in_code_block = False
            else:
                in_code_block = True
            index += 1
            continue

        if in_code_block:
            code_buffer.append(line)
            index += 1
            continue

        if not line.strip():
            index += 1
            continue

        if HR_RE.match(line):
            index += 1
            continue

        if is_table_start(lines, index):
            table_rows = [parse_table_row(lines[index])]
            index += 2
            while index < len(lines) and "|" in lines[index] and lines[index].strip():
                table_rows.append(parse_table_row(lines[index]))
                index += 1
            add_table(document, table_rows)
            continue

        heading = HEADING_RE.match(line)
        if heading:
            level = min(len(heading.group(1)), 3)
            title = clean_inline(heading.group(2))
            if level == 1:
                paragraph = document.add_paragraph(style="Title")
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run(title)
                set_east_asia_font(run)
                run.bold = True
                run.font.size = Pt(16)
                run.font.color.rgb = BLACK
            else:
                heading_paragraph = document.add_heading(title, level=level - 1)
                for run in heading_paragraph.runs:
                    set_east_asia_font(run)
                    run.font.color.rgb = BLACK
            index += 1
            continue

        unordered = UNORDERED_RE.match(line)
        if unordered:
            add_paragraph(document, unordered.group(1), style="List Bullet")
            index += 1
            continue

        ordered = ORDERED_RE.match(line)
        if ordered:
            # 方案 A：序号直接写进段落文本，用普通段落输出，
            # 不用 Word 的 List Number 自动编号，避免跨段/跨章节连续累加。
            add_paragraph(document, f"{ordered.group(1)}. {ordered.group(2)}")
            index += 1
            continue

        if line.startswith(">"):
            add_paragraph(document, line.lstrip("> "), style="Intense Quote")
            index += 1
            continue

        add_paragraph(document, line)
        index += 1

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("markdown", type=Path, help="Path to the Markdown report.")
    parser.add_argument(
        "-o",
        "--output",
        type=Path,
        help="Path to the generated .docx file. Defaults to the Markdown path with .docx suffix.",
    )
    args = parser.parse_args()

    markdown_path = args.markdown.expanduser().resolve()
    if not markdown_path.exists():
        raise FileNotFoundError(f"Markdown file not found: {markdown_path}")

    output_path = args.output.expanduser().resolve() if args.output else markdown_path.with_suffix(".docx")
    convert_markdown_to_docx(markdown_path, output_path)
    print(output_path)


if __name__ == "__main__":
    main()
