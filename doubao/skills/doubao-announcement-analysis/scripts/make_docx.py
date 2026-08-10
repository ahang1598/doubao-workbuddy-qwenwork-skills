#!/usr/bin/env python3
"""
可选工具，仅在用户明确要求 Word/DOCX 导出时才使用，不是默认交付物。

默认交付路径是：display markdown 直接展示 + 创建飞书文档。只有当用户明确说
"要 Word 文件" / "导出 docx" 之类的诉求时，才对 finalize_report.py 产出的
display markdown 跑这个脚本，额外生成一份 .docx。

只支持 display markdown 会用到的一小部分 Markdown 语法：
  - # / ## / ### 标题
  - | a | b | 形式的表格（含表头分隔行 |---|---|）
  - - 开头的无序列表
  - > 开头的引用块
  - **粗体**（行内，简单替换）
  - 普通段落
不支持嵌套列表、行内图片语法解析（图片请用 --image 参数额外插入）。

用法：
  python3 make_docx.py <display_markdown_path> -o <output.docx> \\
      [--title "文档标题"]
"""
import argparse
import re
import sys

_BOLD_PATTERN = re.compile(r"\*\*(.+?)\*\*")
_IMAGE_PATTERN = re.compile(r"^!\[([^\]]*)\]\(([^)]+)\)$")


def _add_paragraph_with_bold(doc, text, style=None):
    p = doc.add_paragraph(style=style)
    pos = 0
    for m in _BOLD_PATTERN.finditer(text):
        if m.start() > pos:
            p.add_run(text[pos:m.start()])
        run = p.add_run(m.group(1))
        run.bold = True
        pos = m.end()
    if pos < len(text):
        p.add_run(text[pos:])
    return p


def _is_table_line(line):
    return line.strip().startswith("|") and line.strip().endswith("|")


def _is_table_separator(line):
    return bool(re.match(r"^\s*\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?\s*$", line))


def _parse_table_row(line):
    cells = line.strip().strip("|").split("|")
    return [c.strip() for c in cells]


def convert(markdown_text: str, title: str = None):
    from docx import Document
    from docx.shared import Inches

    doc = Document()
    if title:
        doc.add_heading(title, level=0)

    lines = markdown_text.split("\n")
    i = 0
    n = len(lines)

    while i < n:
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped == "---":
            doc.add_paragraph("―" * 20)
            i += 1
            continue

        image_match = _IMAGE_PATTERN.match(stripped)
        if image_match:
            image_path = image_match.group(2)
            try:
                doc.add_picture(image_path, width=Inches(6))
            except Exception as exc:  # noqa: BLE001
                doc.add_paragraph(f"[图片未能嵌入: {image_path}（{exc}）]")
            caption = image_match.group(1).strip()
            if caption:
                cap_p = doc.add_paragraph(caption)
                cap_p.alignment = 1  # center
            i += 1
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if heading_match:
            level = min(len(heading_match.group(1)), 4)
            doc.add_heading(heading_match.group(2), level=level)
            i += 1
            continue

        if _is_table_line(stripped):
            table_lines = []
            while i < n and _is_table_line(lines[i].strip()):
                table_lines.append(lines[i].strip())
                i += 1
            rows = [_parse_table_row(l) for l in table_lines if not _is_table_separator(l)]
            if rows:
                n_cols = max(len(r) for r in rows)
                table = doc.add_table(rows=len(rows), cols=n_cols)
                table.style = "Light Grid Accent 1"
                for r_idx, row in enumerate(rows):
                    for c_idx in range(n_cols):
                        table.cell(r_idx, c_idx).text = row[c_idx] if c_idx < len(row) else ""
            continue

        if stripped.startswith("- ") or stripped.startswith("* "):
            _add_paragraph_with_bold(doc, stripped[2:], style="List Bullet")
            i += 1
            continue

        if stripped.startswith(">"):
            _add_paragraph_with_bold(doc, stripped.lstrip("> ").strip(), style="Intense Quote")
            i += 1
            continue

        _add_paragraph_with_bold(doc, stripped)
        i += 1

    return doc


def main():
    parser = argparse.ArgumentParser(description="把 display markdown 转换成 .docx（可选、非默认交付物）")
    parser.add_argument("markdown_path", help="display markdown 文件路径")
    parser.add_argument("-o", "--output", required=True, help="输出 .docx 路径")
    parser.add_argument("--title", default=None, help="文档标题（可选）")
    args = parser.parse_args()

    if args.markdown_path.endswith(".json"):
        raise SystemExit("[错误] 输入应为 display markdown 文件，不是 facts.json。")

    with open(args.markdown_path, "r", encoding="utf-8") as f:
        text = f.read()

    doc = convert(text, title=args.title)
    doc.save(args.output)
    print(f"docx 已生成: {args.output}", file=sys.stderr)


if __name__ == "__main__":
    main()
