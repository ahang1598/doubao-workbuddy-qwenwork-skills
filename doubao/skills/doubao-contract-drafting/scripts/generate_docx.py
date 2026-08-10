#!/usr/bin/env python3
"""Generate a black-text, no-comment DOCX from structured V5.7 draft JSON."""
from __future__ import annotations

import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from contract_model import validate_schema


BODY_FONT = "宋体"
HEADING_FONT = "黑体"


def set_run_font(run, font_name: str, font_size: float, bold: bool = False) -> None:
    """Set every Word font slot, including East Asian glyphs."""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for slot in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{slot}"), font_name)
    language = r_pr.find(qn("w:lang"))
    if language is None:
        language = OxmlElement("w:lang")
        r_pr.append(language)
    language.set(qn("w:val"), "zh-CN")
    language.set(qn("w:eastAsia"), "zh-CN")


def configure_document_defaults(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = Pt(12)
    r_pr = style._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for slot in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{slot}"), BODY_FONT)
    language = r_pr.find(qn("w:lang"))
    if language is None:
        language = OxmlElement("w:lang")
        r_pr.append(language)
    language.set(qn("w:val"), "zh-CN")
    language.set(qn("w:eastAsia"), "zh-CN")


def add_text(doc: Document, text: str, level: int = 0) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    is_heading = level in (0, 1)
    set_run_font(run, HEADING_FONT if is_heading else BODY_FONT, 16 if level == 0 else 12, is_heading)
    if level == 0:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 1:
        p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.line_spacing = 1.5


def set_cell_text(cell, value: object) -> None:
    cell.text = str(value)
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.line_spacing = 1.5
        for run in paragraph.runs:
            set_run_font(run, BODY_FONT, 10)


def add_table(doc: Document, table_data: list[list[object]]) -> None:
    if not table_data:
        return
    if not all(isinstance(row, list) and row for row in table_data):
        raise ValueError("表格必须由非空行数组组成")
    column_count = max(len(row) for row in table_data)
    table = doc.add_table(rows=0, cols=column_count)
    for row_data in table_data:
        cells = table.add_row().cells
        for column_index, cell in enumerate(cells):
            value = row_data[column_index] if column_index < len(row_data) else ""
            set_cell_text(cell, value)


def main() -> None:
    if len(sys.argv) != 3:
        raise SystemExit("Usage: generate_docx.py <contract.json> <output.docx>")
    data = json.loads(Path(sys.argv[1]).read_text(encoding="utf-8"))
    schema_errors = validate_schema(data)
    if schema_errors:
        raise SystemExit("Invalid contract JSON:\n" + "\n".join("- " + error for error in schema_errors))
    doc = Document()
    configure_document_defaults(doc)
    add_text(doc, data["title"], 0)
    for item in data.get("sections", []):
        add_text(doc, item["text"], item.get("level", 2))
    if data.get("allow_tables", True):
        for table_data in data.get("tables", []):
            add_table(doc, table_data)
    if data.get("appendices"):
        add_text(doc, "附件清单", 1)
        for item in data["appendices"]:
            add_text(doc, f"{item['number']}：{item['title']}", 2)
            if not item.get("list_only"):
                for paragraph in item.get("content", []):
                    add_text(doc, str(paragraph), 2)
                if data.get("allow_tables", True):
                    for table_data in item.get("tables", []):
                        add_table(doc, table_data)
    signature = data.get("signature")
    if signature:
        add_text(doc, "签署页", 1)
        for party in signature:
            add_text(doc, f"{party}：________________", 2)
            add_text(doc, "签署日期：______年____月____日", 2)
    doc.save(sys.argv[2])


if __name__ == "__main__":
    main()
