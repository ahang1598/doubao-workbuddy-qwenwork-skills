"""DOCX paragraph + table text extraction via python-docx.

No image rendering — see ``trace_llm_judge`` history for the rationale
(would require LibreOffice round-trip, marginal benefit for text-heavy
docs).
"""

from __future__ import annotations

import io
import logging

logger = logging.getLogger(__name__)


def extract_docx_text(docx_bytes: bytes) -> tuple[str, dict]:
    """Extract paragraphs + tables in authored order.

    Returns
    -------
    (text, summary)
        ``text`` interleaves paragraph lines and ``[Table N (M rows)]``
        blocks in the same order they appear in the document body.
        ``summary`` is ``{"paragraph_count", "table_count"}`` for the
        caller's manifest / debug output.

    On missing python-docx or unparseable bytes returns ``("", {...0})``.
    """
    try:
        import docx as python_docx
    except ImportError:
        logger.warning("python-docx not installed — docx extract skipped")
        return "", {"paragraph_count": 0, "table_count": 0}

    try:
        doc = python_docx.Document(io.BytesIO(docx_bytes))
    except Exception as exc:
        logger.warning("[docx] open failed: %s", exc)
        return "", {"paragraph_count": 0, "table_count": 0}

    body = doc.element.body
    para_by_el = {p._element: p for p in doc.paragraphs}
    table_by_el = {t._element: t for t in doc.tables}

    text_parts: list[str] = []
    n_para = 0
    n_table = 0
    for child in body.iterchildren():
        if child in para_by_el:
            line = (para_by_el[child].text or "").strip()
            if line:
                text_parts.append(line)
                n_para += 1
        elif child in table_by_el:
            tbl = table_by_el[child]
            rows = tbl.rows
            n_table += 1
            lines = [f"[Table {n_table} ({len(rows)} rows)]"]
            for r_idx, row in enumerate(rows, start=1):
                cells = [
                    (c.text or "").replace("\n", " ").replace("|", "\\|").strip()
                    for c in row.cells
                ]
                lines.append(f"R{r_idx}: " + " | ".join(cells))
            text_parts.append("\n".join(lines))

    return (
        "\n\n".join(text_parts),
        {"paragraph_count": n_para, "table_count": n_table},
    )
