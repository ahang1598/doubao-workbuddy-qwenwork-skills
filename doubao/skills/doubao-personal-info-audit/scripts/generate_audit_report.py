#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
审计报告 Markdown 转 Word 工具

功能：
1. 读取 Markdown 格式的审计报告
2. 生成格式化的 Word 文档（符合法律文书风格）
3. 自动处理标题、表格、列表、标注说明、页眉和页码

格式要求（依据审计报告格式规范）：
- 中文：宋体
- 英文/数字：Times New Roman
- 正文：小四号（12磅）
- 主标题：16磅，26磅固定行距，段前0磅/段后18磅，加粗居中
- 一级标题：14磅，24磅固定行距，段前12磅/段后6磅，加粗左对齐
- 二级标题：12磅，23磅固定行距，段前9磅/段后4磅，加粗左对齐
- 三级、四级标题：12磅，22磅固定行距，段前6磅/段后3磅，加粗左对齐
- 正文行距：20磅；段前、段后各6磅
- 正文首行缩进：2字符
- 页边距：普通（上下2.54cm，左右3.17cm）

使用方式：
    python scripts/generate_audit_report.py ./审计报告.md
    python scripts/generate_audit_report.py ./审计报告.md --output ./输出/审计报告.docx
"""

import sys
import os
import re
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, Twips, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("错误：缺少 python-docx 库，请安装：pip install python-docx")
    sys.exit(1)


# 字号对照表（磅值）
FONT_SIZE = {
    '小四': 12,
    '四号': 14,
    '主标题': 16,
    '五号': 10.5,
}

# 正文与标题的确定版式参数，禁止依赖Word模板默认值。
BODY_LINE_SPACING_PT = 20
BODY_SPACE_BEFORE_PT = 6
BODY_SPACE_AFTER_PT = 6
LEGAL_BASIS_COLUMN_RATIOS = (0.16, 0.48, 0.36)
LEGAL_BASIS_HEADERS = ('层级', '规范名称', '关键条款')
TITLE_FORMATS = {
    0: {'size': 16, 'line': 26, 'before': 0, 'after': 18, 'alignment': WD_ALIGN_PARAGRAPH.CENTER},
    1: {'size': 14, 'line': 24, 'before': 12, 'after': 6, 'alignment': WD_ALIGN_PARAGRAPH.LEFT},
    2: {'size': 12, 'line': 23, 'before': 9, 'after': 4, 'alignment': WD_ALIGN_PARAGRAPH.LEFT},
    3: {'size': 12, 'line': 22, 'before': 6, 'after': 3, 'alignment': WD_ALIGN_PARAGRAPH.LEFT},
    4: {'size': 12, 'line': 22, 'before': 6, 'after': 3, 'alignment': WD_ALIGN_PARAGRAPH.LEFT},
}


def set_run_font(run, chinese_font='宋体', western_font='Times New Roman', size_pt=12, bold=False):
    """设置文本字体（中英文分别设置）"""
    run.font.name = western_font
    run._element.rPr.rFonts.set(qn('w:eastAsia'), chinese_font)
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)


def set_paragraph_format(
    para,
    first_line_indent=False,
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    left_indent_cm=0,
):
    """设置段落格式"""
    para.paragraph_format.line_spacing = Pt(BODY_LINE_SPACING_PT)
    para.paragraph_format.space_before = Pt(BODY_SPACE_BEFORE_PT)
    para.paragraph_format.space_after = Pt(BODY_SPACE_AFTER_PT)
    para.alignment = alignment
    para.paragraph_format.left_indent = Cm(left_indent_cm)
    
    if first_line_indent:
        # 使用Word原生“2字符”，避免以磅值近似或用空格模拟。
        indentation = para._p.get_or_add_pPr().get_or_add_ind()
        indentation.set(qn('w:firstLineChars'), '200')
        indentation.attrib.pop(qn('w:firstLine'), None)


def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)


def add_page_number(paragraph):
    """在段落中插入Word PAGE域。"""
    run = paragraph.add_run()
    begin = OxmlElement('w:fldChar')
    begin.set(qn('w:fldCharType'), 'begin')
    instruction = OxmlElement('w:instrText')
    instruction.set(qn('xml:space'), 'preserve')
    instruction.text = ' PAGE '
    separate = OxmlElement('w:fldChar')
    separate.set(qn('w:fldCharType'), 'separate')
    end = OxmlElement('w:fldChar')
    end.set(qn('w:fldCharType'), 'end')
    run._r.extend([begin, instruction, separate, end])
    set_run_font(run, '宋体', 'Times New Roman', FONT_SIZE['五号'])


def create_document():
    """创建 Word 文档并设置页面格式"""
    doc = Document()
    
    # 设置页边距（普通页边距）
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)
        # 封面不显示页眉页脚；正文仍保留报告简称和页码。
        section.different_first_page_header_footer = True
        header = section.header.paragraphs[0]
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_run = header.add_run('个人信息保护合规审计报告')
        set_run_font(header_run, '宋体', 'Times New Roman', FONT_SIZE['五号'])
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_page_number(footer)
    
    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(FONT_SIZE['小四'])
    style.font.color.rgb = RGBColor(0, 0, 0)
    style.paragraph_format.line_spacing = Pt(BODY_LINE_SPACING_PT)
    style.paragraph_format.space_before = Pt(BODY_SPACE_BEFORE_PT)
    style.paragraph_format.space_after = Pt(BODY_SPACE_AFTER_PT)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    style.paragraph_format.first_line_indent = Pt(24)

    # 使用真实标题样式，便于Word导航和自动目录。
    title_style = doc.styles['Title']
    title_config = TITLE_FORMATS[0]
    title_style.font.name = 'Times New Roman'
    title_style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    title_style.font.size = Pt(title_config['size'])
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(0, 0, 0)
    title_style.paragraph_format.line_spacing = Pt(title_config['line'])
    title_style.paragraph_format.space_before = Pt(title_config['before'])
    title_style.paragraph_format.space_after = Pt(title_config['after'])
    title_style.paragraph_format.alignment = title_config['alignment']
    title_style.paragraph_format.first_line_indent = Pt(0)
    title_style.paragraph_format.keep_with_next = True

    heading_config = {
        'Heading 1': TITLE_FORMATS[1],
        'Heading 2': TITLE_FORMATS[2],
        'Heading 3': TITLE_FORMATS[3],
        'Heading 4': TITLE_FORMATS[4],
    }
    for style_name, config in heading_config.items():
        heading_style = doc.styles[style_name]
        heading_style.font.name = 'Times New Roman'
        heading_style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        heading_style.font.size = Pt(config['size'])
        heading_style.font.bold = True
        heading_style.font.color.rgb = RGBColor(0, 0, 0)
        heading_style.paragraph_format.line_spacing = Pt(config['line'])
        heading_style.paragraph_format.space_before = Pt(config['before'])
        heading_style.paragraph_format.space_after = Pt(config['after'])
        heading_style.paragraph_format.alignment = config['alignment']
        heading_style.paragraph_format.first_line_indent = Pt(0)
        heading_style.paragraph_format.keep_with_next = True
    
    return doc


def add_title(doc, text, level=1):
    """添加标题"""
    # 移除标题文本中的 Markdown 加粗符号
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
    
    para = doc.add_paragraph(style='Title') if level == 0 else doc.add_paragraph(style=f'Heading {min(level, 4)}')
    run = para.add_run(text)
    
    config = TITLE_FORMATS[min(level, 4)]
    set_run_font(run, '宋体', 'Times New Roman', config['size'], bold=True)
    para.alignment = config['alignment']
    
    # 标题不缩进
    para.paragraph_format.first_line_indent = Cm(0)
    para.paragraph_format.line_spacing = Pt(config['line'])
    para.paragraph_format.space_before = Pt(config['before'])
    para.paragraph_format.space_after = Pt(config['after'])
    para.paragraph_format.keep_with_next = True


def add_paragraph(
    doc,
    text,
    *,
    first_line_indent=True,
    left_indent_cm=0,
    shading=None,
):
    """添加段落，支持行内加粗"""
    if not text.strip():
        return
    
    para = doc.add_paragraph()
    
    # 解析行内加粗 **文字**
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # 加粗文本
            run = para.add_run(part[2:-2])
            set_run_font(run, '宋体', 'Times New Roman', FONT_SIZE['小四'], bold=True)
        elif part:
            # 普通文本
            run = para.add_run(part)
            set_run_font(run, '宋体', 'Times New Roman', FONT_SIZE['小四'], bold=False)
    
    set_paragraph_format(
        para,
        first_line_indent=first_line_indent,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        left_indent_cm=left_indent_cm,
    )
    if shading:
        paragraph_properties = para._p.get_or_add_pPr()
        shade = OxmlElement('w:shd')
        shade.set(qn('w:fill'), shading)
        paragraph_properties.append(shade)


def add_table(doc, rows):
    """添加表格"""
    if not rows or len(rows) < 1:
        return
    
    # 创建表格
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # 以字段语义分配列宽，避免多列表格被机械等分后逐字换行。
    usable_width_cm = 14.66
    narrative_headers = {
        '已确认事实', '尚不能确认的事项', '需要补充的证据', '业务环境与关联能力',
        '判断依据', '待补证事项', '实际处理活动', '事实与证据', '尚待核实',
        '问题', '材料名称', '用途', '未取得影响',
        '具体规则内容', '本案对应要件', '效力与适用说明', '证据及精确定位',
        '证据状态与证明力', '冲突/相反证据', '尚不能确认事项及影响',
        '待补数据与核验程序', '事实证据及定位', '冲突/不确定性',
        '不确定性及待补数据', '分析内容',
    }
    medium_headers = {
        '事实来源', '是否个人信息', '是否敏感个人信息', '是否匿名化/去标识化',
        '初步角色或行为', '直接法源', '证据', '审计判断', '对应问题/事实',
        '真实文件名', '版本/日期', '具体定位', '支持事实', '完整性说明',
        '层级/强制属性', '规范名称及条款', '效力/适用说明', '官方来源及核验',
        '适用规则表行',
    }
    compact_headers = {
        '事实编号', '发现编号', '序号', '风险等级', '整改期限', '证据编号',
        '证据等级', '数据项目', '主体/链路',
        '法源编号', '法源编号及层级', '分析项目', '讨论项目', '链条编号',
        '数据编号', '对象编号', '顺序及环节', '编号与直接上游',
    }
    headers = [re.sub(r'\*\*([^*]+)\*\*', r'\1', str(value).strip()) for value in rows[0]]
    if tuple(headers) == LEGAL_BASIS_HEADERS:
        weights = list(LEGAL_BASIS_COLUMN_RATIOS)
    elif headers in (['分析项目', '分析内容'], ['事实编号', '讨论项目', '内容'], ['链条编号', '讨论项目', '内容'], ['数据项目', '讨论项目', '内容'], ['数据编号', '讨论项目', '内容'], ['主体/链路', '讨论项目', '内容'], ['对象编号', '讨论项目', '内容']):
        weights = [1.0, 1.15, 3.85] if len(headers) == 3 else [1.0, 3.0]
    elif headers == ['顺序及环节', '编号与直接上游', '分析内容']:
        weights = [1.15, 1.65, 4.2]
    elif headers == ['法源编号及层级', '规范名称及条款', '具体规则内容', '本案对应要件及效力说明']:
        weights = [1.15, 1.65, 3.4, 2.3]
    elif headers == ['法源编号及层级', '规范名称及条款', '具体规则内容', '效力、适用及官方核验']:
        weights = [1.15, 1.65, 3.4, 2.3]
    else:
        weights = []
        for header in headers:
            if header in narrative_headers:
                weights.append(2.6)
            elif header in medium_headers:
                weights.append(2.0)
            elif header in compact_headers:
                weights.append(1.25)
            else:
                weights.append(1.6)
    total_weight = sum(weights)
    column_widths = [usable_width_cm * weight / total_weight for weight in weights]
    for index, width in enumerate(column_widths):
        table.columns[index].width = Cm(width)
    
    # 填充数据
    for i, row_data in enumerate(rows):
        row = table.rows[i]
        cant_split = OxmlElement('w:cantSplit')
        row._tr.get_or_add_trPr().append(cant_split)
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.width = Cm(column_widths[j])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_mar = tc_pr.first_child_found_in('w:tcMar')
            if tc_mar is None:
                tc_mar = OxmlElement('w:tcMar')
                tc_pr.append(tc_mar)
            for side in ('top', 'left', 'bottom', 'right'):
                margin = OxmlElement(f'w:{side}')
                margin.set(qn('w:w'), '100')
                margin.set(qn('w:type'), 'dxa')
                tc_mar.append(margin)
            
            # 移除单元格文本中的 Markdown 加粗符号
            clean_text = re.sub(r'\*\*([^*]+)\*\*', r'\1', str(cell_text).strip())
            cell.text = clean_text
            
            # 设置字体
            for para in cell.paragraphs:
                for run in para.runs:
                    set_run_font(run, '宋体', 'Times New Roman', FONT_SIZE['五号'], bold=False)
                para.paragraph_format.line_spacing = Pt(16)
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(0)
                para.paragraph_format.first_line_indent = Pt(0)
                if headers[j] in compact_headers or (tuple(headers) == LEGAL_BASIS_HEADERS and j == 0):
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif tuple(headers) == LEGAL_BASIS_HEADERS:
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # 表头样式：加粗、浅灰底纹；文字保持黑色。
            if i == 0:
                set_cell_shading(cell, 'E7E6E6')
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in para.runs:
                        run.font.bold = True
    if table.rows:
        repeat = OxmlElement('w:tblHeader')
        repeat.set(qn('w:val'), 'true')
        table.rows[0]._tr.get_or_add_trPr().append(repeat)


def split_markdown_row(line):
    """Split one pipe table row while preserving escaped pipes."""
    text = line.strip()[1:-1]
    cells = []
    current = []
    escaped = False
    for char in text:
        if escaped:
            current.append(char)
            escaped = False
        elif char == '\\':
            escaped = True
        elif char == '|':
            cells.append(''.join(current).strip())
            current = []
        else:
            current.append(char)
    if escaped:
        current.append('\\')
    cells.append(''.join(current).strip())
    return cells


def parse_markdown_table(lines, start_idx):
    """解析 Markdown 表格"""
    if start_idx + 1 >= len(lines):
        raise ValueError(f"第{start_idx + 1}行表格缺少分隔行")
    header_line = lines[start_idx].strip()
    separator_line = lines[start_idx + 1].strip()
    if not (header_line.startswith('|') and header_line.endswith('|')):
        raise ValueError(f"第{start_idx + 1}行不是完整表格行")
    headers = split_markdown_row(header_line)
    if not (separator_line.startswith('|') and separator_line.endswith('|')):
        raise ValueError(f"第{start_idx + 2}行表格缺少Markdown分隔行")
    separators = split_markdown_row(separator_line)
    if len(separators) != len(headers) or not all(re.fullmatch(r':?-{3,}:?', cell) for cell in separators):
        raise ValueError(f"第{start_idx + 2}行不是有效的Markdown表格分隔行")

    rows = [headers]
    idx = start_idx + 2
    while idx < len(lines):
        line = lines[idx].strip()
        if not (line.startswith('|') and line.endswith('|')):
            break
        cells = split_markdown_row(line)
        if len(cells) != len(headers):
            raise ValueError(f"第{idx + 1}行表格列数为{len(cells)}，应为{len(headers)}")
        rows.append(cells)
        idx += 1
    return rows, idx


def convert_md_to_docx(md_path, output_path=None):
    """将 Markdown 转换为 Word 文档"""
    md_path = Path(md_path)
    if not md_path.exists():
        print(f"错误：Markdown 文件不存在：{md_path}")
        return False
    
    if output_path:
        output_path = Path(output_path)
    else:
        output_path = md_path.with_suffix('.docx')
    
    # 读取 Markdown 内容
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    
    # 创建文档
    doc = create_document()
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # 空行
        if not line:
            i += 1
            continue
        
        # 分隔线
        if line == '---':
            i += 1
            continue

        # 模板显式分页标记；避免用连续空行或手工换行模拟分页。
        if line == '<!-- PAGEBREAK -->':
            doc.add_page_break()
            i += 1
            continue

        # 引用块/标注说明
        if line.startswith('>'):
            add_paragraph(
                doc,
                line.lstrip('>').strip(),
                first_line_indent=False,
                left_indent_cm=0.37,
                shading='F2F2F2',
            )
            i += 1
            continue
        
        # 主标题（报告标题）
        if line.startswith('# ') and not line.startswith('## '):
            add_title(doc, line[2:].strip(), level=0)
            i += 1
            continue
        
        # 一级标题（一、二、三……）
        if line.startswith('## '):
            add_title(doc, line[3:].strip(), level=1)
            i += 1
            continue
        
        # 二级标题（（一）（二）（三）……）
        if line.startswith('### '):
            add_title(doc, line[4:].strip(), level=2)
            i += 1
            continue
        
        # 三级标题（1. 2. 3. ……）
        if line.startswith('#### '):
            add_title(doc, line[5:].strip(), level=3)
            i += 1
            continue
        
        # 四级标题
        if line.startswith('##### '):
            add_title(doc, line[6:].strip(), level=4)
            i += 1
            continue
        
        # 表格
        if line.startswith('|'):
            try:
                rows, i = parse_markdown_table(lines, i)
            except ValueError as error:
                print(f"错误：{error}")
                return False
            if rows:
                add_table(doc, rows)
            continue

        # 无序列表
        if re.match(r'^[-*]\s+', line):
            add_paragraph(
                doc,
                re.sub(r'^[-*]\s+', '• ', line),
                first_line_indent=False,
                left_indent_cm=0.74,
            )
            i += 1
            continue

        # 有序列表
        if re.match(r'^\d+[.)]\s+', line):
            add_paragraph(
                doc,
                line,
                first_line_indent=False,
                left_indent_cm=0.74,
            )
            i += 1
            continue
        
        # 普通段落
        text = line
        if text.startswith('**') and text.endswith('**'):
            # 整段加粗（如字段名称）
            para = doc.add_paragraph()
            run = para.add_run(text[2:-2])
            set_run_font(run, '宋体', 'Times New Roman', FONT_SIZE['小四'], bold=True)
            set_paragraph_format(para, first_line_indent=True, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
        else:
            add_paragraph(doc, text)
        
        i += 1
    
    # 保存文档
    doc.save(output_path)
    print(f"Word 文档已生成：{output_path}")
    return True


def main():
    """主函数"""
    if len(sys.argv) < 2:
        print("用法：python generate_audit_report.py <markdown文件路径> [--output <输出路径>]")
        print("示例：python generate_audit_report.py ./审计报告.md")
        print("      python generate_audit_report.py ./审计报告.md --output ./输出/审计报告.docx")
        sys.exit(1)
    
    md_path = sys.argv[1]
    output_path = None
    
    # 解析参数
    if '--output' in sys.argv:
        output_idx = sys.argv.index('--output')
        if output_idx + 1 < len(sys.argv):
            output_path = sys.argv[output_idx + 1]
    
    success = convert_md_to_docx(md_path, output_path)
    sys.exit(0 if success else 1)


if __name__ == '__main__':
    main()
